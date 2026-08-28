from __future__ import annotations

import hashlib
import json
import os
from datetime import date
from pathlib import Path

import pytest
from docx import Document
from pydantic import ValidationError

from app.assessment.guidance_bundle import OWNER_ASSESSMENT_BUNDLE
from app.evaluation.live_suite import load_live_evaluation_bundle
from app.evaluation.live_suite_gold import LiveSuiteExpertQualification
from app.evaluation.live_suite_owner_review import (
    HELD_STATUTORY_PROVISIONS,
    RULES_CHECKLIST_FILENAME,
    WORKBOOK_FILENAME,
    _safe_inventory,
    audit_filled_owner_workbook,
    build_owner_return_verification_document,
    build_owner_review_workbook,
    dated_pack_name,
    export_owner_review_workbook,
    prepare_live60_owner_review_pack,
)

PROJECT_ROOT = Path(__file__).resolve().parents[2]
BUNDLE_ROOT = PROJECT_ROOT / "benchmarks/evaluation/live-evaluation-60-v1"


def _workbook_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_owner_review_workbook_covers_all_issues_held_provisions_and_rules(
    tmp_path: Path,
) -> None:
    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    output = tmp_path / WORKBOOK_FILENAME
    export_owner_review_workbook(
        project_root=PROJECT_ROOT,
        output_path=output,
        as_of_date=date(2026, 8, 16),
    )
    text = _workbook_text(output)

    assert "evaluation-only" in text.casefold() or "evaluation only" in text.casefold()
    assert "not eligible for training" in text.casefold()
    assert "/Users/" not in text
    assert "@" not in text
    for case in bundle.registry.cases:
        assert case.case_id in text
        assert case.question in text
        for topic in case.must_cover_issues:
            assert topic in text
    for _provision_id, title, _summary in HELD_STATUTORY_PROVISIONS:
        assert title in text
    for rule in OWNER_ASSESSMENT_BUNDLE.rules:
        assert rule.rule_id in text
    assert "qualified / limited / knowledge_gap" in text
    assert "reviewer:<64-hex>" in text
    assert "owner and the one primary" in text.casefold()
    assert "cannot be the second reviewer" in text.casefold()


def test_owner_review_workbook_does_not_seal_legal_gold() -> None:
    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    document = build_owner_review_workbook(
        bundle,
        as_of_date=date(2026, 8, 16),
        index_build_id="candidate-pending-owner-review",
    )
    assert document.core_properties.author == ""
    with pytest.raises(ValidationError):
        LiveSuiteExpertQualification.model_validate(
            {
                "schema": "legalbot.live-expert-qualification.v1",
                "approval_status": "needs_expert_annotation",
            }
        )


def test_prepare_dated_pack_writes_review_logs_and_empty_answers(
    tmp_path: Path,
) -> None:
    project = tmp_path / "project"
    (project / "benchmarks" / "evaluation").mkdir(parents=True)
    (project / "docs" / "reports").mkdir(parents=True)
    import shutil

    shutil.copytree(
        PROJECT_ROOT / "benchmarks" / "evaluation" / "live-evaluation-60-v1",
        project / "benchmarks" / "evaluation" / "live-evaluation-60-v1",
    )
    shutil.copytree(
        PROJECT_ROOT / "benchmarks" / "evaluation" / "live-evaluation-30-v1",
        project / "benchmarks" / "evaluation" / "live-evaluation-30-v1",
    )
    Document().save(str(project / "docs" / "reports" / RULES_CHECKLIST_FILENAME))

    result = prepare_live60_owner_review_pack(
        project_root=project,
        as_of_date=date(2026, 8, 16),
        overwrite=True,
    )
    pack = project / dated_pack_name(date(2026, 8, 16))
    assert result["pack_name"] == "Live60-2026-08-16"
    assert (pack / "review" / WORKBOOK_FILENAME).is_file()
    assert (pack / "review" / RULES_CHECKLIST_FILENAME).is_file()
    assert (pack / "artifacts" / "expert-qualification-template.json").is_file()
    assert (pack / "answers").is_dir()
    assert not any(pack.joinpath("answers").iterdir())

    template = json.loads(
        (pack / "artifacts" / "expert-qualification-template.json").read_text(encoding="utf-8")
    )
    assert template["approval_status"] == "needs_expert_annotation"
    assert template["seal_sha256"] is None
    assert template["owner_is_primary_reviewer"] is True
    assert template["ai_second_reviewer_forbidden"] is True
    with pytest.raises(ValidationError):
        LiveSuiteExpertQualification.model_validate(template)

    events = (pack / "logs" / "events" / "owner-review-pack-events.jsonl").read_text(
        encoding="utf-8"
    )
    metrics = (pack / "logs" / "metrics" / "live60-pack.jsonl").read_text(encoding="utf-8")
    traces = (pack / "logs" / "traces" / "live60-pack.jsonl").read_text(encoding="utf-8")
    assert "owner_review_workbook_exported" in events
    assert "qualification_template_exported" in events
    assert "no_execution_spans_until_authorized_run" in traces
    assert "owner_review_pack_export_total" in metrics
    assert "/Users/" not in events
    assert "question" not in json.loads(events.splitlines()[0])


def test_run_plan_file_digest_and_object_seal_are_the_same_plan() -> None:
    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    plan_path = BUNDLE_ROOT / "generation-run-plan.json"
    file_sha = hashlib.sha256(plan_path.read_bytes()).hexdigest()
    assert file_sha == bundle.manifest.run_plan_sha256
    assert bundle.run_plan.seal_sha256 != file_sha
    assert bundle.manifest.seal_sha256 != file_sha


def _optional_filled_return_workbook() -> Path:
    raw = os.environ.get("LEGALBOT_OWNER_RETURN_WORKBOOK", "").strip()
    if not raw:
        pytest.skip("set LEGALBOT_OWNER_RETURN_WORKBOOK to audit a local filled workbook")
    workbook = Path(raw)
    if not workbook.is_file():
        pytest.skip("filled RETURN/HOLD workbook is not on this machine")
    return workbook


def test_audit_filled_return_hold_workbook_if_present() -> None:
    workbook = _optional_filled_return_workbook()

    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    audit = audit_filled_owner_workbook(workbook, bundle=bundle)
    assert audit["issue_count"] == 585
    assert audit["knowledge_gap_issue_count"] == 585
    assert audit["qualified_issue_count"] == 0
    assert audit["blank_span_count"] == 585
    assert audit["blank_contrary_count"] == 60
    assert audit["primary_reviewer_role_present"] is False
    assert audit["assessment_keep_count"] == 16
    assert audit["second_review_status"] == "needs_independent_review"


def test_second_review_docx_states_frozen_topics_and_what_to_fill(
    tmp_path: Path,
) -> None:
    workbook = _optional_filled_return_workbook()

    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    audit = audit_filled_owner_workbook(workbook, bundle=bundle)
    inventory = _safe_inventory(audit, bundle=bundle)
    path = tmp_path / "live60-second-review-test.docx"
    build_owner_return_verification_document(bundle=bundle, audit=audit, inventory=inventory).save(
        str(path)
    )
    document = Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            text += "\n" + "\n".join(cell.text for cell in row.cells)
    assert "How to close these knowledge gaps" in text
    assert "Frozen legal topic (the gap)" in text
    assert "What you must fill to close" in text
    assert bundle.registry.cases[0].question in text
    for topic in bundle.registry.cases[0].must_cover_issues:
        assert topic in text
    for topic in bundle.registry.cases[-1].must_cover_issues:
        assert topic in text
    assert text.count("knowledge_gap") >= 585
    assert "/Users/" not in text
