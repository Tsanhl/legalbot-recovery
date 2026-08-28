from __future__ import annotations

import hashlib
import json
import zipfile
from datetime import UTC, datetime
from pathlib import Path

import pytest
from cryptography.fernet import Fernet
from docx import Document
from lxml import etree

from app.crypto import LocalCipher
from app.evaluation.review_docx import (
    ReviewExportError,
    export_live_review_docx,
    load_live_review,
)


def _sha256(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _write_encrypted_json(path: Path, value: dict[str, object], cipher: LocalCipher) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(cipher.encrypt_text(json.dumps(value, ensure_ascii=False, sort_keys=True)))


def _fixture_run(
    tmp_path: Path, *, privacy_passed: bool = True
) -> tuple[Path, LocalCipher, str, str]:
    run_id = "live-review-fixture"
    run_dir = tmp_path / "runs" / run_id
    run_dir.mkdir(parents=True)
    cipher = LocalCipher(Fernet(Fernet.generate_key()))
    run_manifest = {
        "schema": "legalbot.e2e-run-manifest.v1",
        "run_id": run_id,
        "suite_id": "live-evaluation-30-v1",
        "suite_version": "1.0.0",
        "suite_file_sha256": "1" * 64,
        "suite_canonical_sha256": "2" * 64,
        "as_of_date": "2026-08-14",
        "provenance": {
            "git_sha": "a" * 40,
            "git_dirty": False,
            "model_version": "Qwen3.5-9B-4bit",
            "index_build_id": "candidate-test-001",
            "prompt_version": "evidence-first-v2",
            "router_version": "router-v2",
            "classifier_version": "classifier-v2",
            "policy_sha256": "3" * 64,
            "assessment_rules_sha256": "4" * 64,
        },
    }
    run_manifest_bytes = (json.dumps(run_manifest, indent=2, sort_keys=True) + "\n").encode()
    (run_dir / "manifest.json").write_bytes(run_manifest_bytes)

    released_question = (
        "A supplier delivers a bespoke system late and with serious defects. "
        "Advise on breach, termination, causation, remoteness and the limitation clause."
    )
    released_answer = """# Executive answer

The customer has a strong argument that the express deadline was a condition because the contract made time of the essence. The separate defects are also breaches and their seriousness must be assessed before treating them as repudiatory.

## Evidence-led analysis

The loss must be caused by the breach on the balance of probabilities. The disputed commercial contract therefore requires a counterfactual assessment, remoteness analysis and mitigation evidence. The limitation clause must be incorporated, properly construed and tested under the applicable statutory control.

- Identify each contractual promise separately.
- Bind every conclusion to the recorded evidence span.
- State any residual uncertainty instead of inventing a factual finding.

The released conclusion is deliberately concise for this layout fixture."""
    held_question = "HELD QUESTION SENTINEL that must never appear in the review document."
    held_answer = "HELD ANSWER SENTINEL that must never appear in the review document."

    for case_id, question in (("live30-q01", released_question), ("live30-q02", held_question)):
        _write_encrypted_json(
            run_dir / "cases" / case_id / "question.enc",
            {
                "schema": "legalbot.e2e-encrypted-question.v1",
                "run_id": run_id,
                "as_of_date": "2026-08-14",
                "case": {"case_id": case_id, "question": question},
            },
            cipher,
        )
    for case_id, artifact_id, answer in (
        ("live30-q01", "answer-released-v1", released_answer),
        ("live30-q02", "answer-held-v1", held_answer),
    ):
        _write_encrypted_json(
            run_dir / "cases" / case_id / "artifacts" / "answer" / f"{artifact_id}.enc",
            {
                "schema": "legalbot.e2e-sensitive-artifact.v1",
                "run_id": run_id,
                "case_id": case_id,
                "artifact_id": artifact_id,
                "kind": "answer",
                "content": answer,
            },
            cipher,
        )

    released_case = {
        "case_id": "live30-q01",
        "ordinal": 1,
        "case_status": "completed",
        "release_state": "verified_full",
        "released": True,
        "privacy_passed": True,
        "evidence_passed": True,
        "question_sha256": _sha256(released_question),
        "answer_artifact_id": "answer-released-v1",
        "answer_sha256": _sha256(released_answer),
        "subject": "contract law",
        "task_type": "problem",
        "jurisdiction": "England and Wales",
        "as_of_date": "2026-08-14",
        "word_target": 1000,
        "word_count": 128,
        "research_route": "sectioned",
        "drafting_route": "sectioned",
        "assessment_bundle_sha256": "4" * 64,
        "assessment_rule_ids": ["owner.universal.issue_map", "owner.problem.application"],
        "evidence": [
            {
                "evidence_span_id": "ev-contract-001",
                "stable_source_id": "authority-contract-001",
                "legal_locator": "para 42",
                "legal_role": "holding_ratio",
                "identity_state": "verified",
                "support_state": "supported",
                "retrieval_rank": 1,
                "currentness_state": "verified_current",
                "jurisdiction_state": "verified",
            }
        ],
        "rubric": [
            {
                "criterion_id": "issue_spotting",
                "score": 76,
                "status": "advisory",
                "assessment_rule_ids": ["owner.universal.issue_map"],
                "verification_signal": "all_material_issues_mapped",
            }
        ],
        "repairs": [
            {
                "repair_id": "repair-contract-001",
                "section_id": "section-causation",
                "reason_code": "fm7.claim_span_support",
                "status": "completed",
                "attempt_count": 1,
            }
        ],
        "gaps": [
            {
                "gap_id": "gap-contract-001",
                "category": "contrary_authority",
                "severity": "medium",
                "status": "resolved",
                "safe_expected_ids": ["authority-contract-001"],
                "safe_observed_ids": ["authority-contract-001"],
            }
        ],
        "advisory_ai_review": {
            "status": "available",
            "reviewer_execution_mode": "separate_verification_pass_same_model_adapter",
            "model_independent": False,
            "recommendations_only": True,
            "can_decide_or_adopt": False,
            "can_admit_sources": False,
            "can_authorize_gates": False,
            "may_raise_fail_closed_owner_review_hold": True,
            "review_sha256": "9" * 64,
            "recommendation_codes": ["all_claims_recommended_supported"],
            "flagged_claim_count": 0,
            "owner_review_required": False,
            "unavailable_reason_code": None,
        },
        "metrics": [
            {"metric_id": "completion_ms", "value": 9200, "unit": "ms", "gate": "advisory"}
        ],
        "failure_codes": ["fm7.claim_span_support"],
    }
    held_case = {
        "case_id": "live30-q02",
        "ordinal": 2,
        "case_status": "held",
        "release_state": "evidence_failed",
        "released": False,
        "privacy_passed": True,
        "evidence_passed": False,
        "question_sha256": _sha256(held_question),
        "answer_artifact_id": "answer-held-v1",
        "answer_sha256": _sha256(held_answer),
        "subject": "consumer law",
        "task_type": "problem",
        "jurisdiction": "England and Wales",
        "as_of_date": "2026-08-14",
        "word_target": 1000,
        "word_count": None,
        "research_route": "sectioned",
        "drafting_route": "sectioned",
        "assessment_bundle_sha256": None,
        "assessment_rule_ids": [],
        "evidence": [],
        "rubric": [],
        "repairs": [],
        "gaps": [
            {
                "gap_id": "gap-consumer-001",
                "category": "source_needed",
                "severity": "high",
                "status": "source_needed",
                "safe_expected_ids": [],
                "safe_observed_ids": [],
            }
        ],
        "metrics": [],
        "failure_codes": ["fm4.knowledge_gap"],
    }
    review_export = {
        "schema": "legalbot.live-review-export.v1",
        "run_id": run_id,
        "run_manifest_sha256": hashlib.sha256(run_manifest_bytes).hexdigest(),
        "generated_at": datetime.now(UTC).isoformat(),
        "run_status": "running",
        "privacy_report_passed": privacy_passed,
        "purpose": "evaluation_only",
        "eligible_for_training": False,
        "training_export_allowed": False,
        "expected_case_count": 2,
        "cases": [released_case, held_case],
        "aggregate_metrics": [
            {"metric_id": "release_rate", "value": 0.5, "unit": "ratio", "gate": "advisory"}
        ],
        "clusters": [
            {
                "cluster_id": "cluster-evidence-001",
                "category": "evidence_coverage",
                "case_ids": ["live30-q01", "live30-q02"],
                "status": "triaged",
            }
        ],
        "corrections": [
            {
                "correction_id": "correction-001",
                "affected_layer": "claims",
                "case_ids": ["live30-q01"],
                "status": "verified",
                "regression_case_id": "regression-001",
            }
        ],
        "owner_decisions": [
            {
                "decision_id": "decision-001",
                "decision_code": "calibration.review_needed",
                "status": "owner_review_needed",
                "affected_case_ids": ["live30-q01"],
            }
        ],
        "regressions": [
            {
                "regression_case_id": "regression-001",
                "source_issue_id": "gap-contract-001",
                "status": "passing",
                "fixed_version": "candidate-test-002",
                "verification_run_id": run_id,
            }
        ],
    }
    (run_dir / "review-export.json").write_text(
        json.dumps(review_export, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )
    return run_dir, cipher, held_question, held_answer


def test_export_contains_only_release_eligible_plaintext_and_business_sections(
    tmp_path: Path,
) -> None:
    run_dir, cipher, held_question, held_answer = _fixture_run(tmp_path)
    output = tmp_path / "review.docx"
    export_live_review_docx(
        run_dir=run_dir,
        output_path=output,
        cipher=cipher,
        require_complete=False,
    )

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "A supplier delivers a bespoke system late" in text
    assert "The customer has a strong argument" in text
    assert held_question not in text
    assert held_answer not in text
    assert "Thirty-case control matrix" in text
    assert "Evidence and citation controls" in text
    assert "Rubric and activated assessment rules" in text
    assert "Advisory AI evidence review" in text
    assert "Owner substantive review" in text
    assert "recommendations only" in text
    assert "Encrypted review / sign-off references" in table_text
    assert "Cross-case clusters" in text
    assert "Corrections and owner decisions" in text
    assert "Regression register" in text
    assert output.stat().st_mode & 0o777 == 0o600

    with zipfile.ZipFile(output) as package:
        names = set(package.namelist())
        document_xml = package.read("word/document.xml")
        all_xml = b"\n".join(package.read(name) for name in names if name.endswith(".xml"))
        assert b" TOC " in document_xml
        assert b"HELD QUESTION SENTINEL" not in all_xml
        assert b"HELD ANSWER SENTINEL" not in all_xml
        assert "docProps/custom.xml" not in names
        assert b"rsid" not in all_xml
        core = etree.fromstring(package.read("docProps/core.xml"))
        creator = core.xpath("string(//*[local-name()='creator'])")
        modified_by = core.xpath("string(//*[local-name()='lastModifiedBy'])")
        assert creator == ""
        assert modified_by == ""
        assert {item.date_time for item in package.infolist()} == {(1980, 1, 1, 0, 0, 0)}


def test_loader_does_not_open_held_artifact(tmp_path: Path) -> None:
    run_dir, cipher, _, _ = _fixture_run(tmp_path)
    held_path = run_dir / "cases" / "live30-q02" / "artifacts" / "answer" / "answer-held-v1.enc"
    held_path.write_bytes(b"not even a valid encrypted artifact")

    loaded = load_live_review(run_dir=run_dir, cipher=cipher, require_complete=False)

    assert loaded.cases[1].question is None
    assert loaded.cases[1].released_answer is None


def test_export_refuses_failed_run_privacy_report(tmp_path: Path) -> None:
    run_dir, cipher, _, _ = _fixture_run(tmp_path, privacy_passed=False)

    with pytest.raises(ReviewExportError, match="privacy report did not pass"):
        export_live_review_docx(
            run_dir=run_dir,
            output_path=tmp_path / "must-not-exist.docx",
            cipher=cipher,
            require_complete=False,
        )

    assert not (tmp_path / "must-not-exist.docx").exists()


def test_strict_export_requires_completed_reconciled_30_case_run(tmp_path: Path) -> None:
    run_dir, cipher, _, _ = _fixture_run(tmp_path)

    with pytest.raises(ReviewExportError, match="completed, reconciled 30-case"):
        load_live_review(run_dir=run_dir, cipher=cipher, require_complete=True)


def test_manifest_binding_detects_run_manifest_change(tmp_path: Path) -> None:
    run_dir, cipher, _, _ = _fixture_run(tmp_path)
    manifest = json.loads((run_dir / "manifest.json").read_text())
    manifest["as_of_date"] = "2026-08-15"
    (run_dir / "manifest.json").write_text(json.dumps(manifest, sort_keys=True) + "\n")

    with pytest.raises(ReviewExportError, match="not bound"):
        load_live_review(run_dir=run_dir, cipher=cipher, require_complete=False)
