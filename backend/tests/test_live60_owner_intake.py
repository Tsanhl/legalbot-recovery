from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from app.cli import parser
from app.evaluation.live_suite_owner_intake import (
    LEGACY_INTAKE_MIGRATION_SCHEMA,
    OWNER_CONFIRMATION_TOKEN,
    OWNER_INTAKE_DRAFT_SCHEMA,
    audit_legacy_owner_decision_docx,
    build_owner_decision_intake_document,
    migrate_legacy_owner_decision_docx,
    parse_owner_decision_intake,
    write_owner_decision_intake,
)

PROJECT_ROOT = Path(__file__).resolve().parents[2]
AS_OF_DATE = date(2026, 8, 16)


def _filled_workbook(tmp_path: Path) -> Path:
    path = tmp_path / "owner-intake.docx"
    document = build_owner_decision_intake_document(
        project_root=PROJECT_ROOT,
        as_of_date=AS_OF_DATE,
    )
    document.tables[6].rows[7].cells[2].text = OWNER_CONFIRMATION_TOKEN
    document.save(path)
    return path


def _add_table(
    document: Document,
    headers: tuple[str, ...],
    rows: tuple[tuple[str, ...], ...],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        cell.text = value
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values, strict=True):
            cell.text = value


def _legacy_workbook(tmp_path: Path) -> Path:
    path = tmp_path / "legacy.docx"
    document = Document()
    for index in range(16):
        if index == 1:
            _add_table(
                document,
                ("Discarded item", "Why", "Default", "Your decision"),
                (
                    ("draft", "generic", "discard", "no — leave discarded"),
                    ("sealable", "false", "discard", "no — leave discarded"),
                    ("ai", "forbidden", "discard", "no — leave discarded"),
                    ("second", "optional", "not_required", "no reopen"),
                ),
            )
        elif index == 3:
            _add_table(
                document,
                ("Path", "Meaning", "Consequence", "Your decision"),
                (
                    ("A", "all gaps", "no answers", "A — selected"),
                    ("B", "qualify", "later", "not selected"),
                    ("C", "statutes", "no answers", "not selected"),
                ),
            )
        elif index == 6:
            _add_table(
                document,
                ("Decision", "Allowed ticks", "Your tick", "Limit"),
                (
                    ("Leave ACTIVE absent for now", "confirm", "confirm", ""),
                    ("Approve candidate", "hold", "hold", ""),
                    ("Westlaw / scholarship full text", "stay out", "stay out", ""),
                    ("Find Case Law full text", "metadata-only", "metadata-only", ""),
                    ("Elasticsearch", "stay out", "stay out", ""),
                    ("Teaching notes as authority", "never", "never", ""),
                    ("Nearest chunk as gold", "never", "never", ""),
                ),
            )
        elif index == 7:
            _add_table(
                document,
                ("ID", "Provision", "Recorded", "Disposition"),
                tuple(
                    (f"held-{number:02d}", f"Provision {number}", "recorded", "hold")
                    for number in range(1, 5)
                ),
            )
        elif index == 8:
            _add_table(
                document,
                ("Field", "Allowed", "Your value"),
                (
                    ("status", "review choices", "keep pending"),
                    ("defined_source_set_id", "safe ID", "not assigned"),
                    ("defined_source_set_review_method", "method", "not completed"),
                    ("defined_source_set_reviewed_as_of_date", "date", "not completed"),
                    ("bound_contrary_span_count", "count", "0"),
                    ("independent_second_review", "choice", "not_required"),
                    ("means English law has no contrary authority", "false", "false"),
                ),
            )
        elif index == 9:
            _add_table(
                document,
                ("ID", "Decision", "Allowed states", "Your state", "Note"),
                tuple(
                    (
                        f"D-{number:02d}",
                        "decision",
                        "state",
                        ("deferred_later_owner_action" if number in {6, 7, 8, 9} else "pending"),
                        "",
                    )
                    for number in range(1, 16)
                ),
            )
        elif index == 15:
            _add_table(
                document,
                ("Statement", "Owner tick"),
                (
                    (
                        "I am the sole primary qualified England-and-Wales reviewer for this pack.",
                        "PENDING OWNER ATTESTATION",
                    ),
                ),
            )
        else:
            _add_table(document, ("Placeholder",), ((str(index),),))
    document.save(path)
    return path


def test_filled_owner_intake_creates_unsigned_privacy_safe_draft(
    tmp_path: Path,
) -> None:
    workbook = _filled_workbook(tmp_path)
    draft = parse_owner_decision_intake(
        workbook,
        project_root=PROJECT_ROOT,
        as_of_date=AS_OF_DATE,
    )

    assert draft["schema"] == OWNER_INTAKE_DRAFT_SCHEMA
    assert draft["route"]["route"] == "path_b"
    assert draft["route"]["selected_case_count"] == 30
    assert draft["route"]["selected_issue_count"] == 305
    assert len(draft["owner_decisions"]) == 15
    assert draft["owner_decisions"][5]["state"] == "deferred_later_owner_action"
    assert draft["owner_authored_seal"] is False
    assert draft["ready_for_evidence_pack_generation"] is True
    assert draft["ready_for_overlay_seal"] is False
    assert draft["generation_authorised"] is False
    assert draft["writes_active"] is False
    assert draft["writes_previous"] is False
    assert draft["writes_o04"] is False
    assert "/Users/" not in json.dumps(draft)


def test_owner_intake_rejects_route_change_and_premature_owner_gate(
    tmp_path: Path,
) -> None:
    workbook = _filled_workbook(tmp_path)
    document = Document(workbook)
    document.tables[1].rows[1].cells[2].text = "path_a"
    document.save(workbook)
    with pytest.raises(ValueError, match="route must be path_b"):
        parse_owner_decision_intake(
            workbook,
            project_root=PROJECT_ROOT,
            as_of_date=AS_OF_DATE,
        )

    workbook = _filled_workbook(tmp_path)
    document = Document(workbook)
    document.tables[5].rows[6].cells[3].text = "accepted"
    document.save(workbook)
    with pytest.raises(ValueError, match="D-06 must remain deferred_later_owner_action"):
        parse_owner_decision_intake(
            workbook,
            project_root=PROJECT_ROOT,
            as_of_date=AS_OF_DATE,
        )


def test_owner_intake_writes_reviewable_diff_without_applying_a_seal(
    tmp_path: Path,
) -> None:
    workbook = _filled_workbook(tmp_path)
    draft_path = tmp_path / "draft.json"
    diff_path = tmp_path / "diff.md"
    active_path = PROJECT_ROOT / "data" / "indexes" / "ACTIVE.json"
    active_before = active_path.read_bytes() if active_path.is_file() else None
    result = write_owner_decision_intake(
        workbook_path=workbook,
        project_root=PROJECT_ROOT,
        as_of_date=AS_OF_DATE,
        draft_path=draft_path,
        diff_path=diff_path,
    )

    written = json.loads(draft_path.read_text(encoding="utf-8"))
    diff = diff_path.read_text(encoding="utf-8")
    assert result["ready_for_evidence_pack_generation"] is True
    assert result["ready_for_overlay_seal"] is False
    assert written["owner_authored_seal"] is False
    assert "D-06: `unsigned` → `deferred_later_owner_action`" in diff
    assert "Remaining mandatory gates" in diff
    assert str(tmp_path) not in diff
    active_after = active_path.read_bytes() if active_path.is_file() else None
    assert active_after == active_before


def test_legacy_path_a_docx_migrates_compatible_controls_without_a_seal(
    tmp_path: Path,
) -> None:
    legacy = _legacy_workbook(tmp_path)
    audit = audit_legacy_owner_decision_docx(legacy)
    migrated_path = tmp_path / "migrated.docx"
    report_path = tmp_path / "migration.json"
    result = migrate_legacy_owner_decision_docx(
        workbook_path=legacy,
        project_root=PROJECT_ROOT,
        output_path=migrated_path,
        report_path=report_path,
        as_of_date=AS_OF_DATE,
    )

    report = json.loads(report_path.read_text(encoding="utf-8"))
    migrated = Document(migrated_path)
    assert audit["schema"] == LEGACY_INTAKE_MIGRATION_SCHEMA
    assert report["legacy_route"] == "path_a"
    assert report["current_route"] == "path_b"
    assert report["route_superseded_by_later_owner_instruction"] is True
    assert report["legacy_owner_attestation_present"] is False
    assert all(row.cells[3].text == "hold" for row in migrated.tables[3].rows[1:])
    assert migrated.tables[1].rows[1].cells[2].text == "path_b"
    assert migrated.tables[6].rows[7].cells[2].text == ""
    assert result["requires_new_path_b_confirmation"] is True
    assert result["ready_for_overlay_seal"] is False
    assert result["generation_authorised"] is False


def test_cli_registers_local_owner_intake_commands() -> None:
    template = parser().parse_args(["live60-owner-intake-template", "--out", "owner-intake.docx"])
    assert template.command == "live60-owner-intake-template"
    intake = parser().parse_args(
        [
            "live60-owner-intake",
            "--docx",
            "owner-intake.docx",
            "--out",
            "draft.json",
            "--diff",
            "diff.md",
        ]
    )
    assert intake.command == "live60-owner-intake"
    migrated = parser().parse_args(
        [
            "live60-owner-intake-migrate",
            "--legacy-docx",
            "legacy.docx",
            "--out",
            "migrated.docx",
            "--report",
            "migration.json",
        ]
    )
    assert migrated.command == "live60-owner-intake-migrate"
