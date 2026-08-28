from __future__ import annotations

import hashlib
import json
import shutil
import sqlite3
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from app.evaluation.live_suite import load_live_evaluation_bundle
from app.evaluation.live_suite_gold import LiveSuiteExpertQualification
from app.evaluation.live_suite_owner_decisions import (
    KNOWLEDGE_GAP_REASON,
    OWNER_REVIEWER_ROLE,
    apply_owner_ticks,
    assert_mechanical_is_not_gold,
    build_active_promotion_status,
    build_debug_backlog,
    build_issue_decision_document,
    build_issue_decision_pack,
    build_jobs_workers_routes_report,
    build_law_lane_inventory,
    build_mechanical_verification_document,
    build_owner_reviewer_identity,
    export_held_provision_chunks,
    export_owner_decision_artifacts,
    mechanically_verify_held_provisions,
    mint_owner_reviewer_ref,
    official_section_url,
)
from app.evaluation.live_suite_reviewer_identity import reviewer_role_is_forbidden_machine
from app.orchestration.subject_routing_audit import build_subject_routing_audit
from app.types import CASE_PROPOSITION_REVIEWER_ROLES

PROJECT_ROOT = Path(__file__).resolve().parents[2]
BUNDLE_ROOT = PROJECT_ROOT / "benchmarks/evaluation/live-evaluation-60-v1"


def _docx_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_issue_decision_pack_lists_all_frozen_topics() -> None:
    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    pack = build_issue_decision_pack(bundle, as_of_date=date(2026, 8, 16))
    assert pack["issue_count"] == 585
    assert pack["case_count"] == 60
    assert pack["seals_expert_gold"] is False
    assert pack["research_route_counts"] == {"sectioned": 33, "full_enquiry": 27}
    assert pack["selected_research_route_counts"] == {
        "sectioned": 15,
        "full_enquiry": 15,
    }
    assert pack["route_field_used"] == "expected_research_route"
    assert pack["run_identity"]["run_plan_sha_mismatch_is_false_positive"] is True
    q13 = next(case for case in pack["cases"] if case["case_id"] == "live30-q13")
    assert q13["expected_research_route"] == "full_enquiry"
    assert q13["expected_drafting_route"] == "sectioned"
    topics = [issue["topic"] for case in pack["cases"] for issue in case["issues"]]
    expected = [topic for case in bundle.registry.cases for topic in case.must_cover_issues]
    assert topics == expected
    assert all(
        issue["owner_tick"] == "knowledge_gap" for case in pack["cases"] for issue in case["issues"]
    )
    document = build_issue_decision_document(pack)
    text = "\n".join(p.text for p in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            text += "\n" + "\n".join(cell.text for cell in row.cells)
    assert bundle.registry.cases[0].question in text
    assert "knowledge_gap (explicit)" in text
    assert "qualified / limited" not in text
    assert "Research route: full_enquiry" in text
    assert "/Users/" not in text


def test_held_structural_defects_detect_noncontiguous_spans(tmp_path: Path) -> None:
    catalog = tmp_path / "catalog.sqlite3"
    connection = sqlite3.connect(catalog)
    connection.executescript(
        """
        CREATE TABLE documents (
          id TEXT PRIMARY KEY, content_sha256 TEXT NOT NULL,
          source_identity_id TEXT NOT NULL DEFAULT 'doc',
          safe_display_name TEXT NOT NULL DEFAULT 'doc',
          media_type TEXT NOT NULL DEFAULT 'text', status TEXT NOT NULL DEFAULT 'citable',
          lane TEXT, created_at TEXT NOT NULL DEFAULT '2026-08-14',
          updated_at TEXT NOT NULL DEFAULT '2026-08-14'
        );
        CREATE TABLE source_versions (
          id TEXT PRIMARY KEY, document_id TEXT NOT NULL,
          version_sha256 TEXT NOT NULL DEFAULT 'aa',
          canonical_markdown_path TEXT NOT NULL DEFAULT 'x',
          stable_identifier TEXT, currentness_status TEXT NOT NULL DEFAULT 'unknown',
          review_status TEXT NOT NULL DEFAULT 'approved',
          processing_fingerprint TEXT NOT NULL DEFAULT 'test',
          superseded_by TEXT, metadata_json TEXT NOT NULL DEFAULT '{}',
          created_at TEXT NOT NULL DEFAULT '2026-08-14', authority_identity_id TEXT
        );
        CREATE TABLE chunks (
          id TEXT PRIMARY KEY, source_version_id TEXT NOT NULL,
          ordinal INTEGER NOT NULL, locator TEXT NOT NULL,
          text_sha256 TEXT NOT NULL, markdown_text TEXT NOT NULL,
          token_count INTEGER NOT NULL DEFAULT 1,
          metadata_json TEXT NOT NULL DEFAULT '{}', stream TEXT NOT NULL DEFAULT 'body'
        );
        """
    )
    content = "bca5969ce0cd667f59741cf8d1edd2feb9b76512b12a49fafedcf521d89924df"
    connection.execute("INSERT INTO documents(id, content_sha256) VALUES ('doc-1', ?)", (content,))
    connection.execute(
        """
        INSERT INTO source_versions(id, document_id, stable_identifier, authority_identity_id)
        VALUES ('sv-1', 'doc-1', 'ukpga:1980:58:latest-available@2026-08-14', 'ukpga:1980:58')
        """
    )
    bad = (
        "section 14A For the purposes of this section a person’s knowledge includes "
        "knowledge which he might reasonably have been expected to acquire—but a person "
        "shall not be taken by virtue of this subsection to have knowledge of a fact "
        "ascertainable only with the help of expert advice so long as he has taken all "
        "reasonable steps to obtain (and, where appropriate, to act on) that advice."
    )
    connection.execute(
        """
        INSERT INTO chunks(id, source_version_id, ordinal, locator, text_sha256, markdown_text)
        VALUES ('chunk-bad', 'sv-1', 161, 'section 14A', ?, ?)
        """,
        (hashlib.sha256(bad.encode()).hexdigest(), bad),
    )
    connection.commit()
    connection.close()
    from app.evaluation.live_suite_owner_decisions import (
        build_held_span_structural_defects,
        export_held_provision_chunks,
    )

    export = export_held_provision_chunks(catalog)
    defects = build_held_span_structural_defects(export)
    assert defects["defect_count"] >= 1
    assert any(
        "non_contiguous_s14A_10" in code
        for item in defects["defects"]
        for code in item["defect_codes"]
    )
    from app.evaluation.live_suite_held_span_repair import (
        apply_held_span_catalogue_cleanup,
        build_held_span_contiguous_repair,
    )

    repair = build_held_span_contiguous_repair(export)
    assert repair["catalogue_mutated"] is False
    assert repair["seals_expert_gold"] is False
    assert repair["qualified"] is False
    actions = {item["action"] for item in repair["repairs"]}
    assert "replace_spliced_parent_with_chapeau" in actions
    assert "replace_spliced_parent_with_final_proviso" in actions
    chapeau = next(
        item
        for item in repair["repairs"]
        if item["action"] == "replace_spliced_parent_with_chapeau"
    )
    proviso = next(
        item
        for item in repair["repairs"]
        if item["action"] == "replace_spliced_parent_with_final_proviso"
    )
    assert chapeau["markdown_text"].endswith("acquire—")
    assert "but a person" not in chapeau["markdown_text"]
    assert proviso["markdown_text"].startswith("but a person")
    assert "acquire—" not in proviso["markdown_text"]
    assert chapeau["repair_span_id"].startswith("repair-span-")
    assert chapeau["repair_span_id"] != proviso["repair_span_id"]
    cleanup = apply_held_span_catalogue_cleanup(catalog, export=export, repair=repair)
    assert cleanup["parent_bytes_deleted"] is False
    assert "chunk-bad" in cleanup["excluded_parent_chunk_ids"]
    assert len(cleanup["inserted_contiguous_chunks"]) >= 2
    fresh = export_held_provision_chunks(catalog)
    fresh_defects = build_held_span_structural_defects(fresh)
    assert fresh_defects["defect_count"] == 0


def test_owner_reviewer_is_legal_reviewer_and_not_ai() -> None:
    identity = build_owner_reviewer_identity(as_of_date=date(2026, 8, 16))
    assert identity["approval_reviewer_role"] == OWNER_REVIEWER_ROLE
    assert identity["approval_reviewer_role"] == "legal_reviewer"
    assert "expert" not in identity["approval_reviewer_role"]
    assert "england" not in identity["approval_reviewer_role"]
    assert "wales" not in identity["approval_reviewer_role"]
    assert identity["approval_reviewer_role"] in CASE_PROPOSITION_REVIEWER_ROLES
    assert identity["approval_reviewer_ref"].startswith("reviewer:")
    assert len(identity["approval_reviewer_ref"]) == len("reviewer:") + 64
    assert identity["independent_second_review_status"] == "not_required"
    assert identity["independent_second_reviewer_role"] is None
    assert identity["ai_second_reviewer_forbidden"] is True
    assert identity["owner_is_primary_reviewer"] is True
    assert identity["ai_role"] == "mechanical_accuracy_verifier_only"
    assert identity["ai_cannot_seal_gold"] is True
    assert identity["seals_expert_gold"] is False
    assert identity["ai_role_in_reviewer_enum"] is False
    assert "ai" not in {role.casefold() for role in CASE_PROPOSITION_REVIEWER_ROLES}
    assert (
        mint_owner_reviewer_ref(role=OWNER_REVIEWER_ROLE, as_of_date=date(2026, 8, 16))
        == identity["approval_reviewer_ref"]
    )


def test_legal_reviewer_role_is_not_treated_as_machine_identity() -> None:
    assert reviewer_role_is_forbidden_machine("legal_reviewer") is False
    assert reviewer_role_is_forbidden_machine("ai") is True
    assert reviewer_role_is_forbidden_machine("assistant") is True
    assert reviewer_role_is_forbidden_machine("model") is True


def test_apply_owner_ticks_keeps_knowledge_gap_and_does_not_seal() -> None:
    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    identity = build_owner_reviewer_identity(as_of_date=date(2026, 8, 16))
    pack = build_issue_decision_pack(bundle, as_of_date=date(2026, 8, 16))
    mechanical = {"results": [{"held_id": "held-provision-01", "disposition": "hold"}]}
    ticks = apply_owner_ticks(
        bundle=bundle,
        identity=identity,
        issue_pack=pack,
        mechanical=mechanical,
        contrary_authority_status=None,
    )
    assert ticks["knowledge_gap_issue_count"] == 585
    assert ticks["reviewer_policy"]["owner_is_primary_reviewer"] is True
    assert ticks["reviewer_policy"]["second_review"] == "optional"
    assert ticks["reviewer_policy"]["ai_second_reviewer_forbidden"] is True
    assert ticks["qualified_issue_count"] == 0
    assert ticks["overlay_sealable"] is False
    assert ticks["expert_qualification_sealed"] is False
    assert ticks["o04_authorised"] is False
    assert ticks["decision"] == "return_hold"
    assert ticks["cases"][0]["issues"][0]["reason_code"] == KNOWLEDGE_GAP_REASON
    assert "contrary_authority_blank" in ticks["blocking_reason_codes"]
    assert ticks["run_identity"]["registry_canonical_sha256"]
    with pytest.raises(ValueError, match="exact spans"):
        apply_owner_ticks(
            bundle=bundle,
            identity=identity,
            issue_pack=pack,
            mechanical=mechanical,
            contrary_authority_status=None,
            qualified_issue_ids=("live30-q01:issue-01",),
        )


def test_mechanical_verification_is_not_expert_gold() -> None:
    markdown = "An action founded on tort shall not be brought after six years."
    text_sha = hashlib.sha256(markdown.encode()).hexdigest()
    export = {
        "provisions": [
            {
                "held_id": "held-provision-01",
                "title": "Limitation Act 1980 section 2",
                "authority_identity_id": "ukpga:1980:58",
                "legal_locator": "section 2",
                "expected_document_content_sha256": "ab" * 32,
                "chunk_count": 1,
                "chunks": [
                    {
                        "chunk_id": "chunk-test",
                        "source_version_id": "source-version-test",
                        "ordinal": 1,
                        "legal_locator": "section 2",
                        "text_sha256": text_sha,
                        "computed_text_sha256": text_sha,
                        "hash_self_consistent": True,
                        "markdown_text": markdown,
                        "document_content_sha256": "ab" * 32,
                        "stable_identifier": "ukpga:1980:58:latest-available@2026-08-14",
                    }
                ],
            }
        ]
    }
    report = mechanically_verify_held_provisions(
        export,
        fetch_official=lambda _url: b"<root><p>different official text</p></root>",
    )
    assert report["approval_status"] == "mechanical_check_only"
    assert report["expert_approved"] is False
    assert report["qualified_count"] == 0
    assert report["results"][0]["disposition"] == "hold"
    assert report["results"][0]["catalogue_self_consistent"] is True
    assert report["results"][0]["official_normalised_exact_match"] is False
    assert report["results"][0]["difference_class"] == (
        "serialization_or_markup_or_whitespace_or_annotation"
    )
    assert_mechanical_is_not_gold(report)
    with pytest.raises(Exception):  # noqa: B017
        LiveSuiteExpertQualification.model_validate(report)
    document = build_mechanical_verification_document(report)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "mechanical" in text.casefold()


def test_official_section_url_is_deterministic() -> None:
    assert official_section_url("ukpga:1980:58", "section 14A").endswith(
        "/ukpga/1980/58/section/14A/data.xml"
    )


def test_subject_routing_audit_marks_mixed_prefilter() -> None:
    mixed = build_subject_routing_audit(("contract", "tort"))
    assert mixed["mixed"] is True
    assert mixed["filter_mode"] == "mixed_prefilter"
    assert mixed["elasticsearch_used"] is False
    assert "contract" in mixed["recognised_subjects"]
    single = build_subject_routing_audit(("contract",))
    assert single["mixed"] is False
    assert single["filter_mode"] == "single_prefilter"
    empty = build_subject_routing_audit(())
    assert empty["filter_mode"] == "broad_unfiltered"


def test_active_promotion_status_is_read_only_missing(tmp_path: Path) -> None:
    status = build_active_promotion_status(tmp_path)
    assert status["promoted"] is False
    assert status["status"] == "missing"
    assert status["read_only"] is True
    assert status["writes_active"] is False
    assert status["elasticsearch_used"] is False
    assert "no_active_pointer" in status["blocking_reason_codes"]
    assert not (tmp_path / "data" / "indexes" / "ACTIVE.json").exists()


def test_jobs_and_lane_and_debug_reports_are_privacy_safe() -> None:
    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    jobs = build_jobs_workers_routes_report(bundle)
    assert jobs["elasticsearch_used"] is False
    assert jobs["research_worker"] == "separate_process_disabled_first_live"
    assert sum(jobs["task_type_counts"].values()) == 60
    inventory = build_law_lane_inventory(None)
    assert inventory["live_assessment_rule_count"] == 16
    assert inventory["proposed_themes_live"] is False
    assert inventory["teaching_notes_and_ppts"] == "private_teaching_issue_spotting_only"
    backlog = build_debug_backlog()
    assert "o04_then_serial_30" in backlog["items"]


def test_export_writes_pack_without_sealing_gold(tmp_path: Path) -> None:
    project = tmp_path / "project"
    (project / "benchmarks" / "evaluation").mkdir(parents=True)
    shutil.copytree(
        PROJECT_ROOT / "benchmarks" / "evaluation" / "live-evaluation-60-v1",
        project / "benchmarks" / "evaluation" / "live-evaluation-60-v1",
    )
    shutil.copytree(
        PROJECT_ROOT / "benchmarks" / "evaluation" / "live-evaluation-30-v1",
        project / "benchmarks" / "evaluation" / "live-evaluation-30-v1",
    )
    result = export_owner_decision_artifacts(
        project_root=project,
        as_of_date=date(2026, 8, 16),
        catalog_path=tmp_path / "missing.sqlite3",
        fetch_official=None,
        copy_desktop=False,
    )
    pack = project / "Live60-2026-08-16"
    identity = json.loads((pack / "artifacts" / "owner-reviewer-identity.json").read_text())
    ticks = json.loads((pack / "artifacts" / "owner-tick-application.json").read_text())
    mechanical = json.loads((pack / "artifacts" / "mechanical-verification.json").read_text())
    active = json.loads((pack / "artifacts" / "active-promotion-status.json").read_text())
    issues = json.loads((pack / "artifacts" / "issue-decision-pack.json").read_text())
    assert result["overlay_sealable"] is False
    assert result["active_promoted"] is False
    assert identity["independent_second_review_status"] == "not_required"
    assert identity["owner_is_primary_reviewer"] is True
    assert ticks["knowledge_gap_issue_count"] == 585
    assert ticks["contrary_authority_status"] == "reviewed_none"
    assert ticks["overlay_sealable"] is False
    dumped = json.dumps(issues)
    assert '"question"' not in dumped
    assert "question_sha256" in dumped
    assert mechanical["approval_status"] == "mechanical_check_only"
    assert active["promoted"] is False
    assert not (pack / "artifacts" / "expert-qualification.json").exists()
    template = json.loads((pack / "artifacts" / "expert-qualification-template.json").read_text())
    assert template["owner_is_primary_reviewer"] is True
    assert template["ai_role"] == "mechanical_accuracy_verifier_only"
    assert template["ai_second_reviewer_forbidden"] is True
    assert template["seal_sha256"] is None
    gate = json.loads((pack / "artifacts" / "live-run-gate.json").read_text())
    assert gate["after_item_5_live_30_generation_authorised"] is False
    assert gate["suite_verify_and_unit_tests_may_run_now"] is True
    assert result["after_item_5_live_30_generation_authorised"] is False
    accuracy = json.loads((pack / "artifacts" / "user-span-accuracy.json").read_text())
    assert accuracy["exact_match"] is True
    assert accuracy["bound_issue_span_count"] == 0
    remaining = json.loads((pack / "artifacts" / "remaining-gate-attempts.json").read_text())
    assert remaining["fabricated_any_pass"] is False
    assert remaining["any_gate_passed"] is False
    repair = json.loads((pack / "artifacts" / "held-span-contiguous-repair-v2.json").read_text())
    worksheet = json.loads((pack / "artifacts" / "provision-review-worksheet.json").read_text())
    assert repair["catalogue_mutated"] is False
    assert repair["qualified"] is False
    assert worksheet["independent_second_review_status"] == "not_required"
    assert worksheet["ai_cannot_seal_gold"] is True
    text = _docx_text(pack / "review" / "LegalBot-Live60-Issue-Decision-Pack.docx")
    assert "breach" in text
    assert "/Users/" not in text


def test_held_chunk_sql_fixture_round_trip(tmp_path: Path) -> None:
    catalog = tmp_path / "catalog.sqlite3"
    markdown = "Trustee duty of care."
    text_sha = hashlib.sha256(markdown.encode()).hexdigest()
    content_sha = "cd" * 32
    connection = sqlite3.connect(catalog)
    connection.executescript(
        """
        CREATE TABLE documents (
          id TEXT PRIMARY KEY,
          content_sha256 TEXT NOT NULL,
          source_identity_id TEXT NOT NULL DEFAULT 'doc',
          safe_display_name TEXT NOT NULL DEFAULT 'doc',
          media_type TEXT NOT NULL DEFAULT 'text',
          status TEXT NOT NULL DEFAULT 'citable',
          lane TEXT,
          created_at TEXT NOT NULL DEFAULT '2026-08-14',
          updated_at TEXT NOT NULL DEFAULT '2026-08-14'
        );
        CREATE TABLE source_versions (
          id TEXT PRIMARY KEY,
          document_id TEXT NOT NULL,
          version_sha256 TEXT NOT NULL DEFAULT 'aa',
          canonical_markdown_path TEXT NOT NULL DEFAULT 'x',
          stable_identifier TEXT,
          currentness_status TEXT NOT NULL DEFAULT 'unknown',
          review_status TEXT NOT NULL DEFAULT 'approved',
          processing_fingerprint TEXT NOT NULL DEFAULT 'test',
          superseded_by TEXT,
          metadata_json TEXT NOT NULL DEFAULT '{}',
          created_at TEXT NOT NULL DEFAULT '2026-08-14',
          authority_identity_id TEXT
        );
        CREATE TABLE chunks (
          id TEXT PRIMARY KEY,
          source_version_id TEXT NOT NULL,
          ordinal INTEGER NOT NULL,
          locator TEXT NOT NULL,
          text_sha256 TEXT NOT NULL,
          markdown_text TEXT NOT NULL,
          token_count INTEGER NOT NULL DEFAULT 1,
          metadata_json TEXT NOT NULL DEFAULT '{}',
          stream TEXT NOT NULL DEFAULT 'body'
        );
        """
    )
    connection.execute(
        "INSERT INTO documents(id, content_sha256, lane) VALUES ('doc-1', ?, 'primary_authority')",
        (content_sha,),
    )
    connection.execute(
        """
        INSERT INTO source_versions(id, document_id, stable_identifier, authority_identity_id)
        VALUES ('source-version-1', 'doc-1', 'ukpga:2000:29:latest-available@2026-08-14', 'ukpga:2000:29')
        """
    )
    connection.execute(
        """
        INSERT INTO chunks(id, source_version_id, ordinal, locator, text_sha256, markdown_text)
        VALUES ('chunk-1', 'source-version-1', 2, 'section 1', ?, ?)
        """,
        (text_sha, markdown),
    )
    connection.commit()
    connection.close()
    export = export_held_provision_chunks(catalog)
    trustee = next(item for item in export["provisions"] if item["held_id"] == "held-provision-03")
    assert trustee["chunk_count"] == 1
    assert trustee["chunks"][0]["hash_self_consistent"] is True
    assert trustee["status"] == "candidate_held_not_gold"


def test_ipfda_repair_splits_opening_separates_ellipsis_and_excludes_dots() -> None:
    from app.evaluation.live_suite_held_span_repair import (
        build_held_span_contiguous_repair,
    )

    opening = (
        "section 1 Where after the commencement of this Act a person dies "
        "domiciled in England and Wales and is survived by any of the following "
        "persons:—that person may apply to the court for an order under section 2"
    )
    ba = (
        "section 1 any person (not being a person included in paragraph (a) or "
        "(b) above) to whom subsection (1A) ... below applies;"
    )
    dots = "section 1 . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . ."
    export = {
        "provisions": [
            {
                "held_id": "held-provision-04",
                "chunks": [
                    {
                        "chunk_id": "chunk-65e1c2ac95885e5e8e9d65e5ae138b5b6fbcde0a",
                        "ordinal": 1,
                        "markdown_text": opening,
                        "structural_defect_codes": [
                            "non_contiguous_ipfda_s1_1_chapeau_spliced_with_concluding_words"
                        ],
                    },
                    {
                        "chunk_id": "chunk-3fea9d80c6f7fc0011d4db8679b75fe7185c06c7",
                        "ordinal": 4,
                        "markdown_text": ba,
                        "structural_defect_codes": ["editorial_ellipsis_mixed_into_positive_text"],
                    },
                    {
                        "chunk_id": "chunk-0bdcbc97ae11975ac1032cc3c6974aeaad9e43a7",
                        "ordinal": 11,
                        "markdown_text": dots,
                        "structural_defect_codes": ["repealed_or_omitted_editorial_marker"],
                    },
                ],
            }
        ]
    }
    repair = build_held_span_contiguous_repair(export)
    assert repair["repair_span_count"] == 6
    by_action = {item["action"]: item for item in repair["repairs"]}
    assert by_action["replace_spliced_opening_with_chapeau"]["markdown_text"].endswith("persons:—")
    assert by_action["replace_spliced_opening_with_concluding_words"]["markdown_text"].startswith(
        "that person may apply"
    )
    assert (
        by_action["exclude_from_gold_ranking_generation_and_citation"]["gold_eligible_candidate"]
        is False
    )
    assert by_action["exclude_editorial_ellipsis_from_positive_gold_text"]["markdown_text"] == "..."
    assert (
        "(1A) below applies"
        in by_action["bind_official_positive_text_omitting_repealed_or_1B"]["markdown_text"]
    )
    assert (
        "..."
        not in by_action["bind_official_positive_text_omitting_repealed_or_1B"]["markdown_text"]
    )
    assert repair["catalogue_mutated"] is False
    assert repair["qualified"] is False


def test_user_span_accuracy_is_exact_match_fail_closed() -> None:
    from app.evaluation.live_suite_repair_span import REPAIR_SPAN_SCHEMA_V2, repair_span_id_v2
    from app.evaluation.live_suite_span_accuracy import (
        check_user_span_exact_match,
        verify_user_used_spans,
    )

    text = "An action founded on tort shall not be brought after six years."
    text_sha = hashlib.sha256(text.encode()).hexdigest()
    v1 = {
        "schema": "legalbot.live60-held-span-contiguous-repair.v1",
        "repairs": [
            {
                "repair_span_id": "repair-span-" + "a" * 64,
                "required_sublocator": "s 2",
                "markdown_text": text,
                "text_sha256": text_sha,
                "gold_eligible_candidate": True,
            }
        ],
    }
    v1_report = check_user_span_exact_match(
        chunk_id="repair-span-" + "a" * 64,
        content_sha256=text_sha,
        legal_locator="s 2",
        repair=v1,
    )
    assert v1_report["exact_match"] is False
    assert "repair_span_v1_rejected_as_new_gold" in v1_report["mismatch_codes"]
    with pytest.raises(ValueError, match="exact-match"):
        verify_user_used_spans(repair=v1, ticks={"cases": []})

    from app.evaluation.live_suite_repair_span import derivation_manifest_sha256, identity_tuple

    manifest = {"method": "test_fixture", "parent_chunk_id": "chunk-parent"}
    manifest_sha = derivation_manifest_sha256(manifest)
    identity = identity_tuple(
        parent_chunk_id="chunk-parent",
        source_version_id="source-version-1",
        legal_authority_id="ukpga:1980:58",
        official_snapshot_sha256="c" * 64,
        required_sublocator="s 2",
        role="statutory_text_contiguous_candidate",
        markdown_text=text,
        derivation_manifest_sha256=manifest_sha,
        stable_source_id="ukpga-1980-58",
        source_type="legislation",
        jurisdiction="England and Wales",
        legal_locator="s 2",
    )
    span_id = repair_span_id_v2(identity=identity)
    repair = {
        "schema": "legalbot.live60-held-span-contiguous-repair.v2",
        "repairs": [
            {
                "schema": REPAIR_SPAN_SCHEMA_V2,
                "repair_span_id": span_id,
                "parent_chunk_id": "chunk-parent",
                "source_version_id": "source-version-1",
                "legal_authority_id": "ukpga:1980:58",
                "official_snapshot_sha256": "c" * 64,
                "required_sublocator": "s 2",
                "legal_locator": "s 2",
                "stable_source_id": "ukpga-1980-58",
                "source_type": "legislation",
                "jurisdiction": "England and Wales",
                "markdown_text": text,
                "text_sha256": text_sha,
                "role": "statutory_text_contiguous_candidate",
                "derivation_manifest": manifest,
                "derivation_manifest_sha256": manifest_sha,
                "gold_eligible_candidate": True,
            }
        ],
    }
    report = verify_user_used_spans(repair=repair, ticks={"cases": []})
    assert report["exact_match"] is True
    assert report["repair_span_count"] == 1
    assert report["bound_issue_span_count"] == 0
    bad = check_user_span_exact_match(
        chunk_id=span_id,
        content_sha256="b" * 64,
        legal_locator="s 2",
        repair=repair,
    )
    assert bad["exact_match"] is False
    assert "content_sha256_mismatch" in bad["mismatch_codes"]
    with pytest.raises(ValueError, match="100% exact-match"):
        verify_user_used_spans(
            repair={
                "repairs": [
                    {
                        **repair["repairs"][0],
                        "text_sha256": "c" * 64,
                    }
                ]
            }
        )


def test_live_run_gate_blocks_generation_after_item_5(tmp_path: Path) -> None:
    from app.evaluation.live_suite_owner_decisions import build_live_run_gate_report

    gate = build_live_run_gate_report(project_root=tmp_path, official={"applied": True})
    assert gate["reviewer_policy"]["owner_is_primary_reviewer"] is True
    assert gate["reviewer_policy"]["ai_second_reviewer_forbidden"] is True
    assert gate["after_item_5_live_30_generation_authorised"] is False
    assert gate["suite_verify_and_unit_tests_may_run_now"] is True
    assert gate["stage_a_cannot_pass_if_all_585_remain_knowledge_gap"] is True
    assert gate["fabricated_remaining_gates"] is False
    assert "owner_O-04_exact_30_ids" in gate["live_30_generation_requires_after_item_5"]


def test_remaining_live_run_gates_are_attempted_and_refused(tmp_path: Path) -> None:
    from app.evaluation.live_suite_remaining_gates import (
        attempt_remaining_live_run_gates,
        try_seal_overlay_from_owner_ticks,
    )

    ticks = {
        "issue_count": 585,
        "qualified_issue_count": 0,
        "limited_issue_count": 0,
        "knowledge_gap_issue_count": 585,
        "overlay_sealable": False,
        "expert_qualification_sealed": False,
    }
    overlay = try_seal_overlay_from_owner_ticks(ticks=ticks)
    assert overlay["sealed"] is False
    assert overlay["wrote_expert_qualification"] is False
    assert "no_qualified_or_limited_issue_with_exact_spans" in overlay["blocking_reason_codes"]
    report = attempt_remaining_live_run_gates(
        project_root=tmp_path, ticks=ticks, official={"applied": True}
    )
    assert report["fabricated_any_pass"] is False
    assert report["any_gate_passed"] is False
    assert report["generation_authorised"] is False
    assert report["o04_authorised"] is False
    assert report["runtime_blocked_by_path_b_span_gap"] is False
    assert report["live_runtime_separation"]["runtime_blocked_by_path_b_span_gap"] is False
    assert report["live_runtime_separation"]["live60_overlay_status"] == "UNSEALED"
    assert report["live_runtime_separation"]["live60_promotion_status"] == "HOLD"
    assert not (tmp_path / "data" / "indexes" / "ACTIVE.json").exists()
    assert report["attempts"]["owner_O-04_exact_30_ids"]["wrote_o04"] is False
    assert report["attempts"]["owner_promote_ACTIVE"]["wrote_active_pointer"] is False


def test_overlay_ticks_with_gold_still_refuse_to_seal(tmp_path: Path) -> None:
    from app.evaluation.live_suite_remaining_gates import try_seal_overlay_from_owner_ticks

    ticks = {
        "issue_count": 585,
        "qualified_issue_count": 1,
        "limited_issue_count": 0,
        "knowledge_gap_issue_count": 584,
        "overlay_sealable": True,
        "expert_qualification_sealed": False,
    }
    with pytest.raises(ValueError, match="exact-match verified spans"):
        try_seal_overlay_from_owner_ticks(ticks=ticks)
    assert not (tmp_path / "data" / "evaluations" / "expert-qualification.json").exists()


def test_remaining_gates_copy_real_readiness_blockers(tmp_path: Path) -> None:
    from app.evaluation.live_suite_remaining_gates import attempt_remaining_live_run_gates

    report_path = tmp_path / "data" / "reports" / "production-readiness.json"
    report_path.parent.mkdir(parents=True)
    report_path.write_text(
        json.dumps(
            {
                "ready": False,
                "status": "not_ready",
                "blocking_gates": ["source_registry", "owner_promoted_active"],
            }
        ),
        encoding="utf-8",
    )
    ticks = {
        "issue_count": 585,
        "qualified_issue_count": 0,
        "limited_issue_count": 0,
        "knowledge_gap_issue_count": 585,
        "overlay_sealable": False,
        "expert_qualification_sealed": False,
    }
    report = attempt_remaining_live_run_gates(project_root=tmp_path, ticks=ticks)
    readiness = report["attempts"]["readiness_v6_green"]
    assert readiness["passed"] is False
    assert readiness["readiness_report_present"] is True
    assert readiness["blocking_reason_codes"] == [
        "source_registry",
        "owner_promoted_active",
    ]
    assert report["any_gate_passed"] is False
