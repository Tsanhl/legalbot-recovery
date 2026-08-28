from __future__ import annotations

import hashlib
import json
from datetime import UTC, datetime
from pathlib import Path

import pytest
from cryptography.fernet import Fernet
from docx import Document

from app.crypto import LocalCipher
from app.evaluation.live30 import RunProvenance, SensitiveArtifactKind
from app.evaluation.live_suite import load_live_evaluation_bundle
from app.evaluation.live_suite_store import LiveSuiteRunStore
from app.evaluation.review_docx import (
    ReviewExportError,
    ReviewOutputManifest,
    export_live60_review_bundle,
    record_live60_render_gate,
    verify_live60_render_gate,
)

ROOT = Path(__file__).resolve().parents[2]
BUNDLE_ROOT = ROOT / "benchmarks/evaluation/live-evaluation-60-v1"


def _sha(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _live60_review_run(
    tmp_path: Path,
) -> tuple[Path, LocalCipher, tuple[str, ...], tuple[str, ...]]:
    bundle = load_live_evaluation_bundle(BUNDLE_ROOT)
    cipher = LocalCipher(Fernet(Fernet.generate_key()))
    project = tmp_path / "project"
    store = LiveSuiteRunStore(project, cipher)
    manifest = store.create_run(
        run_id="live60-review-test",
        bundle=bundle,
        provenance=RunProvenance(
            git_sha="a" * 40,
            git_dirty=False,
            model_version="qwen-test",
            index_build_id="candidate-test",
            prompt_version="prompt-test",
            router_version="router-test",
            classifier_version="classifier-test",
            policy_sha256="b" * 64,
            assessment_rules_sha256="c" * 64,
        ),
        admitted_at=datetime(2026, 8, 15, 8, 0, tzinfo=UTC),
    )
    run_dir = store.runs_root / manifest.run_id
    run_manifest_bytes = (run_dir / "manifest.json").read_bytes()
    selected = tuple(
        item.case_id for item in bundle.run_plan.cases if item.disposition == "generate_once"
    )
    coverage_only = tuple(
        item.case_id
        for item in bundle.run_plan.cases
        if item.disposition == "coverage_only_not_selected"
    )
    cases: list[dict[str, object]] = []
    for case in bundle.registry.cases:
        if case.case_id in selected:
            answer = f"Released evaluation answer for {case.case_id}."
            artifact_id = f"answer-{case.ordinal:02d}-v1"
            store.store_sensitive_artifact(
                run_id=manifest.run_id,
                case_id=case.case_id,
                kind=SensitiveArtifactKind.ANSWER,
                artifact_id=artifact_id,
                content=answer,
            )
            cases.append(
                {
                    "case_id": case.case_id,
                    "ordinal": case.ordinal,
                    "run_plan_disposition": "generate_once",
                    "run_plan_outcome_count": 1,
                    "coverage_status": "qualified",
                    "case_status": "completed",
                    "release_state": "verified_limited",
                    "released": True,
                    "privacy_passed": True,
                    "evidence_passed": True,
                    "question_sha256": case.question_sha256,
                    "answer_artifact_id": artifact_id,
                    "answer_sha256": _sha(answer),
                    "subject": case.subject,
                    "task_type": case.task_type,
                    "jurisdiction": case.jurisdiction,
                    "as_of_date": manifest.as_of_date,
                    "word_target": case.word_target,
                    "word_count": len(answer.split()),
                    "research_route": case.expected_research_route,
                    "drafting_route": "sectioned",
                    "assessment_bundle_sha256": "c" * 64,
                    "assessment_rule_ids": ["rule.owner.standard"],
                }
            )
        else:
            cases.append(
                {
                    "case_id": case.case_id,
                    "ordinal": case.ordinal,
                    "run_plan_disposition": "coverage_only_not_selected",
                    "run_plan_outcome_count": 0,
                    "coverage_status": "coverage_only_not_selected",
                    "case_status": "completed",
                    "release_state": "not_released",
                    "released": False,
                    "privacy_passed": True,
                    "evidence_passed": False,
                    "question_sha256": case.question_sha256,
                    "subject": case.subject,
                    "task_type": case.task_type,
                    "jurisdiction": case.jurisdiction,
                    "as_of_date": manifest.as_of_date,
                    "word_target": case.word_target,
                    "research_route": case.expected_research_route,
                    "drafting_route": "sectioned",
                }
            )
    review = {
        "schema": "legalbot.live-review-export.v2",
        "run_id": manifest.run_id,
        "run_manifest_sha256": hashlib.sha256(run_manifest_bytes).hexdigest(),
        "generated_at": datetime.now(UTC).isoformat(),
        "run_status": "completed",
        "privacy_report_passed": True,
        "purpose": "evaluation_only",
        "eligible_for_training": False,
        "training_export_allowed": False,
        "expected_case_count": 60,
        "run_plan_id": bundle.run_plan.run_plan_id,
        "run_plan_file_sha256": bundle.manifest.run_plan_sha256,
        "run_plan_seal_sha256": bundle.run_plan.seal_sha256,
        "cases": cases,
    }
    (run_dir / "review-export.json").write_text(
        json.dumps(review, indent=2, sort_keys=True) + "\n", encoding="utf-8"
    )
    return run_dir, cipher, selected, coverage_only


def test_live60_bundle_has_one_control_three_exact_annexes_and_pending_render(
    tmp_path: Path,
) -> None:
    run_dir, cipher, selected, coverage_only = _live60_review_run(tmp_path)
    output_dir = tmp_path / "review-output"

    manifest_path = export_live60_review_bundle(
        run_dir=run_dir, output_dir=output_dir, cipher=cipher
    )
    manifest = ReviewOutputManifest.model_validate_json(manifest_path.read_bytes())

    assert tuple(item.document_id for item in manifest.documents) == (
        "control",
        "annex-a",
        "annex-b",
        "annex-c",
    )
    assert manifest.status == "docx_created_render_pending"
    assert manifest.eligible_for_training is False
    assert manifest.training_export_allowed is False
    control = manifest.documents[0]
    assert control.case_ids == tuple(
        [f"live30-q{number:02d}" for number in range(1, 31)]
        + [f"live60-q{number:02d}" for number in range(31, 61)]
    )
    assert control.released_plaintext_case_ids == ()
    assert set(control.safe_diagnostic_case_ids) == set(selected) | set(coverage_only)
    annex_ids = tuple(
        case_id for document in manifest.documents[1:] for case_id in document.case_ids
    )
    assert annex_ids == selected
    assert all(len(document.case_ids) == 10 for document in manifest.documents[1:])
    assert not (output_dir / "render-gate.json").exists()

    control_doc = Document(output_dir / control.relative_path)
    control_text = "\n".join(paragraph.text for paragraph in control_doc.paragraphs)
    assert "60-case control matrix" in control_text
    assert "Released evaluation answer" not in control_text
    annex_doc = Document(output_dir / manifest.documents[1].relative_path)
    annex_text = "\n".join(paragraph.text for paragraph in annex_doc.paragraphs)
    assert "Released evaluation answer" in annex_text


def test_render_gate_is_separate_immutable_and_detects_document_tampering(
    tmp_path: Path,
) -> None:
    run_dir, cipher, _, _ = _live60_review_run(tmp_path)
    output_dir = tmp_path / "review-output"
    manifest_path = export_live60_review_bundle(
        run_dir=run_dir, output_dir=output_dir, cipher=cipher
    )
    manifest = ReviewOutputManifest.model_validate_json(manifest_path.read_bytes())
    rendered: dict[str, list[Path]] = {}
    for document in manifest.documents:
        page_dir = output_dir / "rendered" / document.document_id
        page_dir.mkdir(parents=True)
        page = page_dir / "page-1.png"
        page.write_bytes(b"\x89PNG\r\n\x1a\nrender-fixture")
        rendered[document.document_id] = [page]

    gate_path = record_live60_render_gate(
        output_dir=output_dir,
        rendered_pages=rendered,
        inspector_ref="reviewer:" + "d" * 64,
    )
    assert gate_path.is_file()
    assert verify_live60_render_gate(output_dir).visual_inspection_passed is True
    with pytest.raises(FileExistsError):
        record_live60_render_gate(
            output_dir=output_dir,
            rendered_pages=rendered,
            inspector_ref="reviewer:" + "d" * 64,
        )

    first = output_dir / manifest.documents[0].relative_path
    first.write_bytes(first.read_bytes() + b"tampered")
    with pytest.raises(ReviewExportError, match="digest changed"):
        verify_live60_render_gate(output_dir)
