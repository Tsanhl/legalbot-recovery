from __future__ import annotations

import hashlib
import json
import sqlite3
from datetime import date
from pathlib import Path

from cryptography.fernet import Fernet
from docx import Document

from app.cli import parser
from app.crypto import LocalCipher
from app.evaluation.live_suite import load_live_evaluation_bundle
from app.evaluation.live_suite_evidence_pack import (
    CASE_CONFIRMATION_TOKEN,
    EVIDENCE_PACK_SCHEMA,
    build_case_evidence_document,
    build_evidence_index_document,
    build_owner_evidence_pack,
    import_owner_evidence_reviews,
)
from app.evaluation.live_suite_path_b import (
    export_review_candidates,
    selected_generation_case_ids,
)
from app.evaluation.live_suite_repair_span import HELD_SPAN_REPAIR_SCHEMA_V2

PROJECT_ROOT = Path(__file__).resolve().parents[2]
AS_OF_DATE = date(2026, 8, 16)


def _catalogue(path: Path) -> None:
    connection = sqlite3.connect(path)
    connection.executescript(
        """
        CREATE TABLE documents (
            id TEXT PRIMARY KEY,
            content_sha256 TEXT NOT NULL,
            source_identity_id TEXT NOT NULL,
            status TEXT NOT NULL,
            lane TEXT,
            jurisdiction TEXT,
            duplicate_of TEXT
        );
        CREATE TABLE source_versions (
            id TEXT PRIMARY KEY,
            document_id TEXT NOT NULL,
            authority_identity_id TEXT,
            title TEXT,
            as_of_date TEXT,
            currentness_status TEXT,
            licence_name TEXT,
            licence_url TEXT,
            review_status TEXT,
            canonical_url TEXT,
            stable_identifier TEXT,
            superseded_by TEXT
        );
        CREATE TABLE chunks (
            id TEXT PRIMARY KEY,
            source_version_id TEXT NOT NULL,
            ordinal INTEGER NOT NULL,
            locator TEXT,
            text_sha256 TEXT,
            markdown_text TEXT,
            stream TEXT
        );
        """
    )
    excerpt = "section 1 A focused statutory test applies to the stated duty."
    digest = hashlib.sha256(excerpt.encode()).hexdigest()
    connection.execute(
        "INSERT INTO documents VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "document-test",
            "a" * 64,
            "source-test-act",
            "citable",
            "primary_authority",
            "England and Wales",
            None,
        ),
    )
    connection.execute(
        "INSERT INTO source_versions VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (
            "source-version-test",
            "document-test",
            "ukpga:2026:1",
            "Test Authority Act 2026",
            "2026-08-16",
            "latest_available_revised_snapshot",
            "Open Government Licence v3.0",
            "https://example.invalid/licence",
            "approved",
            "https://example.invalid/test-act",
            "ukpga:2026:1:latest-available@2026-08-16",
            None,
        ),
    )
    connection.execute(
        "INSERT INTO chunks VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "chunk-test",
            "source-version-test",
            1,
            "section 1",
            digest,
            excerpt,
            "body",
        ),
    )
    connection.commit()
    connection.close()


def _evidence_map(path: Path) -> None:
    bundle = load_live_evaluation_bundle(
        PROJECT_ROOT / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    selected = set(selected_generation_case_ids(bundle))
    issues = []
    for case in bundle.registry.cases:
        if case.case_id not in selected:
            continue
        for number, topic in enumerate(case.must_cover_issues, start=1):
            issue_id = f"issue-{number:02d}"
            candidates = []
            if case.case_id == "live30-q03" and issue_id == "issue-01":
                candidates = [
                    {
                        "source_name": "Test Authority Act 2026",
                        "source_type": "legislation",
                        "rank": "1",
                        "pinpoint": "s 1",
                        "url": "https://example.invalid/test-act",
                        "local_spans": [{"chunk_id": "chunk-test"}],
                    }
                ]
            issues.append(
                {
                    "case_id": case.case_id,
                    "issue_id": issue_id,
                    "topic": topic,
                    "generation_disposition": "generate_once",
                    "candidates": candidates,
                }
            )
    path.write_text(json.dumps({"issues": issues}), encoding="utf-8")


def test_evidence_pack_covers_all_selected_issues_and_keeps_candidates_unsealed(
    tmp_path: Path,
) -> None:
    catalogue = tmp_path / "catalog.sqlite3"
    evidence_map = tmp_path / "evidence-map.json"
    _catalogue(catalogue)
    _evidence_map(evidence_map)
    pack = build_owner_evidence_pack(
        project_root=PROJECT_ROOT,
        catalog_path=catalogue,
        evidence_map_path=evidence_map,
        as_of_date=AS_OF_DATE,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
    )

    assert pack["schema"] == EVIDENCE_PACK_SCHEMA
    assert pack["selected_case_count"] == 30
    assert pack["selected_issue_count"] == 305
    assert pack["candidate_issue_count"] == 1
    assert pack["knowledge_gap_issue_count"] == 304
    assert pack["candidate_count"] == 1
    assert pack["seals_expert_gold"] is False
    assert pack["ready_for_overlay_seal"] is False
    assert pack["generation_authorised"] is False
    q03 = next(case for case in pack["cases"] if case["case_id"] == "live30-q03")
    candidate = q03["issues"][0]["candidates"][0]
    assert candidate["source_type"] == "legislation"
    assert candidate["proposed_legal_role"] == "statutory_text"
    assert candidate["rights_status"] == "approved_full_text"
    assert candidate["excerpt"].startswith("section 1")
    assert candidate["seals_expert_gold"] is False


def test_evidence_documents_are_owner_readable_and_contain_no_question_prose(
    tmp_path: Path,
) -> None:
    catalogue = tmp_path / "catalog.sqlite3"
    evidence_map = tmp_path / "evidence-map.json"
    _catalogue(catalogue)
    _evidence_map(evidence_map)
    pack = build_owner_evidence_pack(
        project_root=PROJECT_ROOT,
        catalog_path=catalogue,
        evidence_map_path=evidence_map,
        as_of_date=AS_OF_DATE,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
    )
    case = next(case for case in pack["cases"] if case["case_id"] == "live30-q03")
    case_path = tmp_path / "case.docx"
    index_path = tmp_path / "index.docx"
    build_case_evidence_document(pack, case).save(case_path)
    build_evidence_index_document(pack).save(index_path)

    case_document = Document(case_path)
    case_text = "\n".join(
        [paragraph.text for paragraph in case_document.paragraphs]
        + [cell.text for table in case_document.tables for row in table.rows for cell in row.cells]
    )
    index_document = Document(index_path)
    index_text = "\n".join(paragraph.text for paragraph in index_document.paragraphs)
    bundle = load_live_evaluation_bundle(
        PROJECT_ROOT / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    frozen = next(item for item in bundle.registry.cases if item.case_id == "live30-q03")
    assert "section 1 A focused statutory test" in case_text
    assert "Owner marking" in case_text
    assert "approve_qualified" in case_text
    assert frozen.question not in case_text
    assert "all 305 issues" in index_text


def test_filled_evidence_documents_import_only_explicit_owner_markings(
    tmp_path: Path,
) -> None:
    catalogue = tmp_path / "catalog.sqlite3"
    evidence_map = tmp_path / "evidence-map.json"
    workbooks = tmp_path / "workbooks"
    workbooks.mkdir()
    _catalogue(catalogue)
    _evidence_map(evidence_map)
    repair = {"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []}
    pack = build_owner_evidence_pack(
        project_root=PROJECT_ROOT,
        catalog_path=catalogue,
        evidence_map_path=evidence_map,
        as_of_date=AS_OF_DATE,
        repair_payload=repair,
    )
    approved_candidate_id = ""
    for case in pack["cases"]:
        document = build_case_evidence_document(pack, case)
        marking_number = 0
        for table in document.tables:
            header = [cell.text for cell in table.rows[0].cells]
            if header == ["Field", "Allowed / instruction", "Owner entry"]:
                marking_number += 1
                if case["case_id"] == "live30-q03" and marking_number == 1:
                    candidate = case["issues"][0]["candidates"][0]
                    approved_candidate_id = candidate["candidate_id"]
                    table.rows[1].cells[2].text = "approve_qualified"
                    table.rows[2].cells[2].text = approved_candidate_id
                    table.rows[4].cells[2].text = f"{approved_candidate_id}=statutory_text"
                else:
                    table.rows[1].cells[2].text = "request_another_candidate"
            if header == ["Field", "Allowed value", "Owner entry"]:
                table.rows[3].cells[2].text = CASE_CONFIRMATION_TOKEN
        document.save(workbooks / f"LegalBot-Live60-Path-B-Evidence-{case['case_id']}.docx")

    review_export = tmp_path / "review-export.json"
    export_review_candidates(
        project_root=PROJECT_ROOT,
        destination=review_export,
        cipher=LocalCipher(Fernet(Fernet.generate_key())),
        as_of_date=AS_OF_DATE,
    )
    imported = import_owner_evidence_reviews(
        project_root=PROJECT_ROOT,
        catalog_path=catalogue,
        evidence_map_path=evidence_map,
        workbook_dir=workbooks,
        review_export_path=review_export,
        as_of_date=AS_OF_DATE,
        repair_payload=repair,
    )

    q03_issue_01 = next(row for row in imported["rows"] if row["row_id"] == "live30-q03:issue-01")
    assert approved_candidate_id
    assert q03_issue_01["status"] == "qualified"
    assert q03_issue_01["exact_gold_spans"][0]["legal_role"] == "statutory_text"
    assert imported["row_count"] == 585
    assert imported["selected_qualified_issue_count"] == 1
    assert imported["selected_knowledge_gap_issue_count"] == 304
    assert imported["selected_evidence_complete"] is False
    assert imported["ready_for_overlay_seal"] is False
    assert imported["generation_authorised"] is False


def test_cli_registers_evidence_pack_command() -> None:
    args = parser().parse_args(["live60-evidence-pack", "--out-dir", "evidence-packs"])
    assert args.command == "live60-evidence-pack"
    evidence_import = parser().parse_args(
        [
            "live60-evidence-import",
            "--workbook-dir",
            "evidence-packs",
            "--review-export",
            "review-export.json",
            "--out",
            "reviewed.json",
        ]
    )
    assert evidence_import.command == "live60-evidence-import"
