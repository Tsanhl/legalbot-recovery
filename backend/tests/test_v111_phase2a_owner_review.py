from __future__ import annotations

import json
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from scripts import build_v111_phase2a_owner_review as owner_review


def _sha(value: str) -> str:
    return owner_review._sha256(value.encode())


def _source_package(root: Path) -> Path:
    root.mkdir(mode=0o700)
    rows = [
        {
            "ordinal": ordinal,
            "row_id": f"live60-q{((ordinal - 1) // 10) + 1:02d}:issue-{((ordinal - 1) % 10) + 1:02d}",
            "case_id": f"live60-q{((ordinal - 1) // 10) + 1:02d}",
            "issue_id": f"issue-{((ordinal - 1) % 10) + 1:02d}",
            "issue_label": f"Issue {ordinal}",
            "issue_label_sha256": _sha(f"issue-label:{ordinal}"),
            "legal_domain": "test",
            "baseline_primary_status": (
                "GOLD_OR_CASE_DEFECT" if ordinal <= 509 else "MATERIAL_CANDIDATE_COVERAGE_GAP"
            ),
            "baseline_official_finding_ids": [],
            "determined_defects": ["MISSING_PROPOSITION_BINDING"],
            "candidate_evidence_candidates": [],
            "remediation_result": "GOLD_SOURCE_VERSION_SPAN_BINDING_INCOMPLETE",
            "technical_status": "BLOCKED_MATERIAL_GAP",
            "row_evidence_sha256": _sha(f"row:{ordinal}"),
        }
        for ordinal in range(1, 586)
    ]
    effects = [
        {
            "ordinal": ordinal,
            "disposition": "OWNER_DECISION_REQUIRED",
            "record_sha256": _sha(f"effect:{ordinal}"),
        }
        for ordinal in range(1, 1897)
    ]
    judgments = [
        {
            "ordinal": ordinal,
            "neutral_citation": f"[2000] UKHL {ordinal}",
            "technical_status": "OWNER_DECISION_REQUIRED",
            "record_sha256": _sha(f"judgment:{ordinal}"),
        }
        for ordinal in range(1, 21)
    ]
    provenance_records = [
        {
            "target_type": "candidate_legislation",
            "target_id": f"source-version-{ordinal}",
            "result": "DOWNLOADED_QUARANTINED",
            "matches_expected_version_sha256": False,
            "sha256": _sha(f"source:{ordinal}"),
        }
        for ordinal in range(1, 66)
    ] + [
        {
            "target_type": "candidate_judgment_source",
            "target_id": f"unavailable-{ordinal}",
            "result": "OFFICIAL_SOURCE_UNAVAILABLE",
            "sha256": _sha(f"unavailable:{ordinal}"),
        }
        for ordinal in range(1, 4)
    ]

    payloads: dict[str, dict[str, Any]] = {
        artifact_id: {"schema": f"test.{artifact_id}.v1", "artifact_sha256": _sha(artifact_id)}
        for artifact_id in owner_review.EXPECTED_ARTIFACT_IDS
    }
    payloads["remediation-matrix-585"].update({"row_count": 585, "rows": rows})
    payloads["legislative-effects-register-1896"].update({"effect_count": 1896, "effects": effects})
    payloads["judgment-later-treatment-register-20"].update(
        {"record_count": 20, "records": judgments}
    )
    payloads["official-source-provenance-register"].update(
        {"record_count": len(provenance_records), "records": provenance_records}
    )
    payloads["corrected-all585-qualification"].update(
        {"blocked_material_gap": 585, "phase2b_allowed": False}
    )
    payloads["final-invariants"].update(
        {
            "terminal_verdict": owner_review.TERMINAL_VERDICT,
            "phase2b_allowed": False,
            "answer_model_invoked": False,
            "development_30_generated": False,
            "promotion_or_live_action": False,
        }
    )

    entries: list[dict[str, Any]] = []
    for ordinal, artifact_id in enumerate(owner_review.EXPECTED_ARTIFACT_IDS, start=1):
        file_name = f"{artifact_id}.json"
        raw = owner_review._canonical_json(payloads[artifact_id])
        (root / file_name).write_bytes(raw)
        entries.append(
            {
                "ordinal": ordinal,
                "artifact_id": artifact_id,
                "file_name": file_name,
                "file_sha256": owner_review._sha256(raw),
                "bytes": len(raw),
            }
        )
    index: dict[str, Any] = {
        "schema": owner_review.PACKAGE_SCHEMA,
        "run_id": "source-package",
        "created_at": "2026-08-24T00:00:00+00:00",
        "phase": "2A",
        "authorizing": False,
        "artifact_count": 18,
        "artifact_order": list(owner_review.EXPECTED_ARTIFACT_IDS),
        "entries": entries,
        "terminal_verdict": owner_review.TERMINAL_VERDICT,
    }
    index["package_digest"] = owner_review._sealed(index)
    (root / "PACKAGE-INDEX.json").write_bytes(owner_review._canonical_json(index))
    return root


def _available_ai(item_sha256: str) -> dict[str, Any]:
    material: dict[str, Any] = {
        "schema": owner_review.ADVISORY_AI_SCHEMA,
        "status": "AVAILABLE",
        "item_sha256": item_sha256,
        "model_id": "verification-model",
        "model_version": "pinned-v1",
        "model_artifact_sha256": _sha("model"),
        "prompt_sha256": _sha("prompt"),
        "configuration_sha256": _sha("configuration"),
        "toolchain_sha256": _sha("toolchain"),
        "reviewer_execution_mode": "separate_verification_pass_same_model_adapter",
        "model_independent": False,
        "pinned": True,
        "logged": True,
        "stateless": True,
        "official_sources_checked_first": True,
        "recommendation": "OWNER_SHOULD_CHECK_EXACT_SPAN",
        "concise_findings": ["candidate span may be relevant"],
        "evidence_reference_sha256s": [_sha("evidence")],
        "unavailable_reason": None,
        "authoritative": False,
        "can_decide_or_adopt": False,
        "can_admit_sources": False,
        "can_authorize_gates": False,
        "may_raise_fail_closed_owner_review_hold": True,
        "owner_decision_applied": False,
        "hidden_reasoning_persisted": False,
    }
    material["review_sha256"] = owner_review._sealed(material)
    return material


def _owner_decision(item: dict[str, Any], ai: dict[str, Any]) -> dict[str, Any]:
    return {
        "schema": owner_review.DECISION_SCHEMA,
        "category": "issue",
        "item_id": item["item_id"],
        "item_sha256": item["item_sha256"],
        "owner_typed_name": "Test Owner",
        "owner_decision_date": "2026-08-24",
        "owner_outcome": "APPROVE_PROPOSITION_BINDINGS",
        "owner_rationale": "I checked the cited version and exact span.",
        "decision_basis_sha256s": [item["source_record_sha256"], _sha("evidence")],
        "advisory_ai_disposition": "USED",
        "advisory_ai_review_sha256": ai["review_sha256"],
        "findings": {
            "proposition": "Test proposition",
            "official_source_version_ids": ["source-version-test"],
            "exact_span_binding_sha256s": [_sha("span")],
            "currentness_conclusion": "Current at the reviewed date",
            "candidate_change_required": False,
        },
    }


def test_builds_complete_non_authorizing_owner_review_inventory(tmp_path: Path) -> None:
    source = _source_package(tmp_path / "source")
    output = tmp_path / "review"

    index = owner_review.build(
        remediation_package_root=source,
        output_root=output,
        run_id="phase2a-owner-review-test",
        recorded_at=datetime(2026, 8, 24, 1, 0, tzinfo=UTC),
        batch_size=100,
    )

    assert index["category_item_counts"] == {
        "issue": 585,
        "legislative_effect": 1896,
        "judgment": 20,
        "source_version": 68,
    }
    assert index["category_batch_counts"] == {
        "issue": 6,
        "legislative_effect": 19,
        "judgment": 1,
        "source_version": 1,
    }
    assert index["artifact_count"] == 29
    assert index["owner_decisions_recorded"] == 0
    assert index["advisory_ai_reviews_recorded"] == 0
    assert index["authorizing"] is False
    assert index["phase2b_allowed"] is False
    assert index["development_30_allowed"] is False
    assert index["answer_model_invoked"] is False
    assert index["source_admitted"] is False
    assert index["candidate_mutated"] is False
    companion = index["companion_document"]
    assert companion["authorizing"] is False
    assert companion["owner_editable"] is True
    companion_path = output / companion["file_name"]
    assert companion_path.is_file()
    document = Document(companion_path)
    companion_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    companion_table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "Phase 2A evidence and hallucination review control" in companion_text
    assert "PHASE 2A SAFELY STOPPED" in companion_table_text
    assert "same configured model adapter" in companion_table_text
    assert "Owner comments" in companion_table_text
    policy = json.loads((output / "owner-review-policy.json").read_bytes())
    assert policy["owner_is_substantive_decision_maker"] is True
    assert policy["advisory_ai_can_decide_or_adopt"] is False
    assert policy["answer_generation_allowed"] is False
    assert owner_review.verify_owner_review_package(output) == index

    first_batch = output / "issue-batch-001.json"
    first_batch.write_bytes(first_batch.read_bytes() + b"\n")
    with pytest.raises(ValueError, match="output_artifact_digest_invalid"):
        owner_review.verify_owner_review_package(output)


def test_source_package_tampering_stops_before_review_output(tmp_path: Path) -> None:
    source = _source_package(tmp_path / "source")
    artifact = source / "remediation-matrix-585.json"
    artifact.write_bytes(artifact.read_bytes() + b"\n")

    with pytest.raises(ValueError, match="artifact_digest_invalid"):
        owner_review.build(
            remediation_package_root=source,
            output_root=tmp_path / "review",
            run_id="phase2a-owner-review-test",
            recorded_at=datetime(2026, 8, 24, 1, 0, tzinfo=UTC),
            batch_size=25,
        )


def test_pinned_ai_review_is_advisory_and_owner_record_is_explicit(tmp_path: Path) -> None:
    source = _source_package(tmp_path / "source")
    output = tmp_path / "review"
    owner_review.build(
        remediation_package_root=source,
        output_root=output,
        run_id="phase2a-owner-review-test",
        recorded_at=datetime(2026, 8, 24, 1, 0, tzinfo=UTC),
        batch_size=100,
    )
    batch = json.loads((output / "issue-batch-001.json").read_bytes())
    item = batch["items"][0]
    ai = _available_ai(item["item_sha256"])
    decision = _owner_decision(item, ai)

    assert owner_review.validate_advisory_ai_review(ai)["authoritative"] is False
    assert (
        owner_review.validate_owner_decision(item=item, decision=decision, advisory_ai_review=ai)[
            "owner_typed_name"
        ]
        == "Test Owner"
    )

    fabricated = dict(ai)
    fabricated["authoritative"] = True
    material = dict(fabricated)
    material.pop("review_sha256")
    fabricated["review_sha256"] = owner_review._sealed(material)
    with pytest.raises(ValueError, match="boundary_invalid"):
        owner_review.validate_advisory_ai_review(fabricated)

    missing_pin = dict(ai)
    missing_pin["model_artifact_sha256"] = None
    material = dict(missing_pin)
    material.pop("review_sha256")
    missing_pin["review_sha256"] = owner_review._sealed(material)
    with pytest.raises(ValueError, match="pin_incomplete"):
        owner_review.validate_advisory_ai_review(missing_pin)

    placeholder_owner = dict(decision)
    placeholder_owner["owner_typed_name"] = "[OWNER FULL NAME]"
    with pytest.raises(ValueError, match="identity_missing"):
        owner_review.validate_owner_decision(
            item=item,
            decision=placeholder_owner,
            advisory_ai_review=ai,
        )


def test_ai_unavailability_requires_owner_acknowledgement_and_sealed_record(
    tmp_path: Path,
) -> None:
    source = _source_package(tmp_path / "source")
    output = tmp_path / "review"
    owner_review.build(
        remediation_package_root=source,
        output_root=output,
        run_id="phase2a-owner-review-test",
        recorded_at=datetime(2026, 8, 24, 1, 0, tzinfo=UTC),
        batch_size=100,
    )
    item = json.loads((output / "issue-batch-001.json").read_bytes())["items"][0]
    unavailable: dict[str, Any] = {
        "schema": owner_review.ADVISORY_AI_SCHEMA,
        "status": "UNAVAILABLE",
        "item_sha256": item["item_sha256"],
        "model_id": None,
        "model_version": None,
        "model_artifact_sha256": None,
        "prompt_sha256": None,
        "configuration_sha256": None,
        "toolchain_sha256": None,
        "reviewer_execution_mode": None,
        "model_independent": False,
        "pinned": False,
        "logged": False,
        "stateless": False,
        "official_sources_checked_first": True,
        "recommendation": None,
        "concise_findings": [],
        "evidence_reference_sha256s": [],
        "unavailable_reason": "No reproducibly pinned reviewer transport is provisioned.",
        "authoritative": False,
        "can_decide_or_adopt": False,
        "can_admit_sources": False,
        "can_authorize_gates": False,
        "may_raise_fail_closed_owner_review_hold": True,
        "owner_decision_applied": False,
        "hidden_reasoning_persisted": False,
    }
    unavailable["review_sha256"] = owner_review._sealed(unavailable)
    decision = _owner_decision(item, _available_ai(item["item_sha256"]))
    decision["advisory_ai_disposition"] = "UNAVAILABLE_OWNER_PROCEEDED_WITH_DETERMINISTIC_EVIDENCE"
    decision["advisory_ai_review_sha256"] = unavailable["review_sha256"]

    verified = owner_review.validate_owner_decision(
        item=item,
        decision=decision,
        advisory_ai_review=unavailable,
    )
    assert verified["advisory_ai_disposition"].startswith("UNAVAILABLE_OWNER")
