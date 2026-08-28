from __future__ import annotations

import hashlib
import json
import sqlite3
from datetime import date
from pathlib import Path

from cryptography.fernet import Fernet
from docx import Document

from app.cli import parser
from app.crypto import LocalCipher
from app.evaluation.live_suite import load_live_evaluation_bundle
from app.evaluation.live_suite_final_check import (
    OWNER_FINAL_CHECK_TOKEN,
    PACK_JSON_NAME,
    SUBSTANTIVE_REVIEW_SCHEMA,
    adapt_substantive_review_for_import,
    build_final_check_case_document,
    build_owner_final_check_pack,
    export_owner_final_check_pack,
    import_owner_final_check,
    is_cross_ref_fragment,
    referenced_operative_locators,
)
from app.evaluation.live_suite_path_b import (
    export_review_candidates,
    selected_generation_case_ids,
)
from app.evaluation.live_suite_repair_span import HELD_SPAN_REPAIR_SCHEMA_V2

PROJECT_ROOT = Path(__file__).resolve().parents[2]
AS_OF_DATE = date(2026, 8, 16)
CRA_S31 = "section 31 section 9 (goods to be of satisfactory quality);"
CRA_S9 = (
    "section 9 Goods are of satisfactory quality if they meet the standard that a "
    "reasonable person would consider satisfactory, taking account of any description "
    "of the goods, the price and all other relevant circumstances."
)
CRA_S59 = (
    "section 59 These definitions apply in this Part (as well as the key definitions "
    "in section 2)—“conditional sales contract” has the meaning given in section 5(3); "
    "“goods” has the meaning given in section 2; “satisfactory quality” is to be read "
    "with section 9 (goods to be of satisfactory quality)."
)
CAPARO = (
    "The law develops incrementally. A duty of care depends on foreseeability, "
    "proximity and whether it is fair, just and reasonable to impose the duty."
)
CPR_351 = (
    "rule 35.1 Expert evidence shall be restricted to that which is reasonably "
    "required to resolve the proceedings."
)
CPR_355 = (
    "rule 35.5 Expert evidence is to be given in a written report unless the court "
    "directs otherwise."
)
CPR_357 = (
    "rule 35.7 Where two or more parties wish to submit expert evidence on a "
    "particular issue, the court may direct that the evidence on that issue is to "
    "be given by a single joint expert."
)


def _digest(text: str) -> str:
    return hashlib.sha256(text.encode()).hexdigest()


def _catalogue(path: Path) -> None:
    connection = sqlite3.connect(path)
    connection.executescript(
        """
        CREATE TABLE documents (
            id TEXT PRIMARY KEY,
            content_sha256 TEXT NOT NULL,
            source_identity_id TEXT NOT NULL,
            status TEXT NOT NULL,
            lane TEXT,
            jurisdiction TEXT,
            duplicate_of TEXT
        );
        CREATE TABLE source_versions (
            id TEXT PRIMARY KEY,
            document_id TEXT NOT NULL,
            authority_identity_id TEXT,
            title TEXT,
            as_of_date TEXT,
            currentness_status TEXT,
            licence_name TEXT,
            licence_url TEXT,
            review_status TEXT,
            canonical_url TEXT,
            stable_identifier TEXT,
            superseded_by TEXT
        );
        CREATE TABLE chunks (
            id TEXT PRIMARY KEY,
            source_version_id TEXT NOT NULL,
            ordinal INTEGER NOT NULL,
            locator TEXT,
            text_sha256 TEXT,
            markdown_text TEXT,
            stream TEXT
        );
        """
    )
    connection.execute(
        "INSERT INTO documents VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "document-cra",
            "a" * 64,
            "ukpga:2015:15",
            "citable",
            "primary_authority",
            "United Kingdom",
            None,
        ),
    )
    connection.execute(
        "INSERT INTO source_versions VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (
            "source-version-cra",
            "document-cra",
            "ukpga:2015:15",
            "Consumer Rights Act 2015",
            "2026-08-16",
            "latest_available_revised_snapshot",
            "Open Government Licence v3.0",
            "https://example.invalid/licence",
            "approved",
            "https://www.legislation.gov.uk/ukpga/2015/15",
            "ukpga:2015:15:latest-available@2026-08-16",
            None,
        ),
    )
    connection.execute(
        "INSERT INTO chunks VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "chunk-cra-s31",
            "source-version-cra",
            1,
            "section 31",
            _digest(CRA_S31),
            CRA_S31,
            "body",
        ),
    )
    connection.execute(
        "INSERT INTO chunks VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "chunk-cra-s9",
            "source-version-cra",
            2,
            "section 9",
            _digest(CRA_S9),
            CRA_S9,
            "body",
        ),
    )
    connection.execute(
        "INSERT INTO chunks VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "chunk-cra-s59",
            "source-version-cra",
            3,
            "section 59",
            _digest(CRA_S59),
            CRA_S59,
            "body",
        ),
    )
    connection.execute(
        "INSERT INTO documents VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "document-caparo",
            "b" * 64,
            "neutral-citation:1990-2-ac-605",
            "citable",
            "primary_authority",
            "England and Wales",
            None,
        ),
    )
    connection.execute(
        "INSERT INTO source_versions VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (
            "source-version-caparo",
            "document-caparo",
            "neutral-citation:1990-2-ac-605",
            "Caparo Industries plc v Dickman [1990] 2 AC 605",
            "1990-02-08",
            "historical",
            "Open Government Licence v3.0",
            "https://example.invalid/licence",
            "approved",
            "https://example.invalid/caparo",
            "neutral-citation:1990-2-ac-605",
            None,
        ),
    )
    connection.execute(
        "INSERT INTO chunks VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "chunk-caparo-1",
            "source-version-caparo",
            1,
            "p 617",
            _digest(CAPARO),
            CAPARO,
            "body",
        ),
    )
    connection.commit()
    connection.close()


def _evidence_map(path: Path) -> None:
    bundle = load_live_evaluation_bundle(
        PROJECT_ROOT / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    selected = set(selected_generation_case_ids(bundle))
    issues = []
    for case in bundle.registry.cases:
        if case.case_id not in selected:
            continue
        for number, topic in enumerate(case.must_cover_issues, start=1):
            issue_id = f"issue-{number:02d}"
            candidates: list[dict[str, object]] = []
            if case.case_id == "live30-q02" and issue_id == "issue-01":
                candidates = [
                    {
                        "source_name": "Consumer Rights Act 2015",
                        "source_type": "legislation",
                        "rank": "1",
                        "pinpoint": "Part 1 and Part 2",
                        "url": "https://www.legislation.gov.uk/ukpga/2015/15",
                        "local_spans": [{"chunk_id": "chunk-cra-s31"}],
                    }
                ]
            if case.case_id == "live30-q03" and issue_id == "issue-01":
                candidates = [
                    {
                        "source_name": "Caparo Industries plc v Dickman [1990] 2 AC 605",
                        "source_type": "case",
                        "rank": "1",
                        "pinpoint": "p 617",
                        "url": "https://example.invalid/caparo",
                        "local_spans": [{"chunk_id": "chunk-caparo-1"}],
                    }
                ]
            issues.append(
                {
                    "case_id": case.case_id,
                    "issue_id": issue_id,
                    "topic": topic,
                    "generation_disposition": "generate_once",
                    "candidates": candidates,
                }
            )
    path.write_text(json.dumps({"issues": issues}), encoding="utf-8")


def _pack(tmp_path: Path) -> dict[str, object]:
    catalogue = tmp_path / "catalog.sqlite3"
    evidence_map = tmp_path / "evidence-map.json"
    _catalogue(catalogue)
    _evidence_map(evidence_map)
    return build_owner_final_check_pack(
        project_root=PROJECT_ROOT,
        catalog_path=catalogue,
        evidence_map_path=evidence_map,
        as_of_date=AS_OF_DATE,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
    )


def test_cross_ref_fragment_detects_cra_section_31_shell() -> None:
    assert is_cross_ref_fragment(CRA_S31) is True
    assert is_cross_ref_fragment(CRA_S9) is False
    assert referenced_operative_locators(CRA_S31) == ("s 9",)


def test_final_check_prefers_operative_section_9_over_section_31_cross_ref(
    tmp_path: Path,
) -> None:
    pack = _pack(tmp_path)
    assert pack["selected_issue_count"] == 305
    assert pack["word_is_import_surface"] is False
    assert pack["seals_expert_gold"] is False
    q02 = next(
        issue
        for case in pack["cases"]
        if case["case_id"] == "live30-q02"
        for issue in case["issues"]
        if issue["issue_id"] == "issue-01"
    )
    assert q02["topic"] == "satisfactory quality"
    assert q02["proposed_action"] == "accept_qualify"
    proposed = next(
        item for item in q02["candidates"] if item["candidate_id"] == q02["proposed_candidate_id"]
    )
    assert proposed["legal_locator"] == "section 9"
    assert proposed["content_sha256"] == _digest(CRA_S9)
    assert "reasonable person would consider satisfactory" in proposed["excerpt"]
    assert proposed["cross_ref_fragment"] is False
    assert not any(
        item["legal_locator"] in {"section 31", "section 59"}
        and item["candidate_id"] == q02["proposed_candidate_id"]
        for item in q02["candidates"]
    )


def test_case_law_is_needs_later_treatment_and_cannot_import_as_qualified(
    tmp_path: Path,
) -> None:
    pack = _pack(tmp_path)
    q03 = next(
        issue
        for case in pack["cases"]
        if case["case_id"] == "live30-q03"
        for issue in case["issues"]
        if issue["issue_id"] == "issue-01"
    )
    assert q03["proposed_action"] == "needs_later_treatment"
    pack["owner_confirmation_token"] = OWNER_FINAL_CHECK_TOKEN
    review_export = tmp_path / "review-export.json"
    export_review_candidates(
        project_root=PROJECT_ROOT,
        destination=review_export,
        cipher=LocalCipher(Fernet(Fernet.generate_key())),
        as_of_date=AS_OF_DATE,
    )
    imported = import_owner_final_check(
        project_root=PROJECT_ROOT,
        catalog_path=tmp_path / "catalog.sqlite3",
        pack_path=_write_pack(tmp_path, pack),
        review_export_path=review_export,
        as_of_date=AS_OF_DATE,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
    )
    row = next(item for item in imported["rows"] if item["row_id"] == "live30-q03:issue-01")
    assert row["status"] == "knowledge_gap"
    assert row["exact_gold_spans"] == []

    q03["owner_action"] = "accept_qualify"
    failed = import_owner_final_check(
        project_root=PROJECT_ROOT,
        catalog_path=tmp_path / "catalog.sqlite3",
        pack_path=_write_pack(tmp_path, pack, name="checked-case.json"),
        review_export_path=review_export,
        as_of_date=AS_OF_DATE,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
    )
    failed_row = next(item for item in failed["rows"] if item["row_id"] == "live30-q03:issue-01")
    assert failed_row["status"] == "knowledge_gap"
    assert failed_row["reason_code"] == "exact_span_verification_failed"
    assert failed_row["exact_gold_spans"] == []
    assert failed["selected_qualified_issue_count"] == 1


def test_json_import_qualifies_statute_and_ignores_word_layout(
    tmp_path: Path,
) -> None:
    pack = _pack(tmp_path)
    pack["owner_confirmation_token"] = OWNER_FINAL_CHECK_TOKEN
    q02_case = next(case for case in pack["cases"] if case["case_id"] == "live30-q02")
    stray = tmp_path / "LegalBot-Live60-Path-B-Final-Check-live30-q02.docx"
    document = build_final_check_case_document(pack, q02_case)
    document.add_table(rows=2, cols=2)
    document.save(stray)
    extra_text = "\n".join(
        [paragraph.text for paragraph in Document(stray).paragraphs]
        + [
            cell.text
            for table in Document(stray).tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    assert "Display only" in extra_text
    assert "satisfactory quality" in extra_text

    review_export = tmp_path / "review-export.json"
    export_review_candidates(
        project_root=PROJECT_ROOT,
        destination=review_export,
        cipher=LocalCipher(Fernet(Fernet.generate_key())),
        as_of_date=AS_OF_DATE,
    )
    imported = import_owner_final_check(
        project_root=PROJECT_ROOT,
        catalog_path=tmp_path / "catalog.sqlite3",
        pack_path=_write_pack(tmp_path, pack),
        review_export_path=review_export,
        as_of_date=AS_OF_DATE,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
    )
    qualified = next(item for item in imported["rows"] if item["row_id"] == "live30-q02:issue-01")
    assert qualified["status"] == "qualified"
    assert qualified["exact_gold_spans"][0]["legal_locator"] == "section 9"
    assert qualified["exact_gold_spans"][0]["content_sha256"] == _digest(CRA_S9)
    assert imported["selected_qualified_issue_count"] == 1
    assert imported["ready_for_overlay_seal"] is False
    assert imported["word_is_import_surface"] is False
    assert imported["writes_active"] is False

    try:
        import_owner_final_check(
            project_root=PROJECT_ROOT,
            catalog_path=tmp_path / "catalog.sqlite3",
            pack_path=stray,
            review_export_path=review_export,
            as_of_date=AS_OF_DATE,
            repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
        )
    except ValueError as exc:
        assert "JSON only" in str(exc)
    else:
        raise AssertionError("DOCX must not be an import surface")


def test_hash_mismatch_fails_closed(tmp_path: Path) -> None:
    pack = _pack(tmp_path)
    pack["owner_confirmation_token"] = OWNER_FINAL_CHECK_TOKEN
    q02 = next(
        issue
        for case in pack["cases"]
        if case["case_id"] == "live30-q02"
        for issue in case["issues"]
        if issue["issue_id"] == "issue-01"
    )
    for item in q02["candidates"]:
        if item["candidate_id"] == q02["proposed_candidate_id"]:
            item["content_sha256"] = "c" * 64
            break
    else:
        raise AssertionError("proposed candidate missing")
    review_export = tmp_path / "review-export.json"
    export_review_candidates(
        project_root=PROJECT_ROOT,
        destination=review_export,
        cipher=LocalCipher(Fernet(Fernet.generate_key())),
        as_of_date=AS_OF_DATE,
    )
    imported = import_owner_final_check(
        project_root=PROJECT_ROOT,
        catalog_path=tmp_path / "catalog.sqlite3",
        pack_path=_write_pack(tmp_path, pack),
        review_export_path=review_export,
        as_of_date=AS_OF_DATE,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
    )
    row = next(item for item in imported["rows"] if item["row_id"] == "live30-q02:issue-01")
    assert row["status"] == "knowledge_gap"
    assert row["reason_code"] == "exact_span_verification_failed"
    assert row["exact_gold_spans"] == []
    assert imported["selected_qualified_issue_count"] == 0
    assert imported["ready_for_overlay_seal"] is False


def test_export_writes_json_and_display_docx(tmp_path: Path) -> None:
    catalogue = tmp_path / "catalog.sqlite3"
    evidence_map = tmp_path / "evidence-map.json"
    _catalogue(catalogue)
    _evidence_map(evidence_map)
    result = export_owner_final_check_pack(
        project_root=PROJECT_ROOT,
        catalog_path=catalogue,
        evidence_map_path=evidence_map,
        output_dir=tmp_path / "final-check",
        as_of_date=AS_OF_DATE,
        overwrite=True,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
    )
    pack_path = tmp_path / "final-check" / PACK_JSON_NAME
    assert pack_path.is_file()
    assert (tmp_path / "final-check" / "LegalBot-Live60-Path-B-Final-Check-Index.docx").is_file()
    assert result["word_is_import_surface"] is False
    assert result["writes_o04"] is False
    payload = json.loads(pack_path.read_text(encoding="utf-8"))
    assert payload["bucket_counts"]["accept_qualify"] >= 1
    assert payload["owner_confirmation_token"] == ""


def test_cli_registers_final_check_commands() -> None:
    pack = parser().parse_args(["live60-final-check-pack", "--out-dir", "final-check"])
    assert pack.command == "live60-final-check-pack"
    imported = parser().parse_args(
        [
            "live60-final-check-import",
            "--pack",
            "final-check.json",
            "--review-export",
            "review-export.json",
            "--out",
            "reviewed.json",
            "--confirm",
            OWNER_FINAL_CHECK_TOKEN,
        ]
    )
    assert imported.command == "live60-final-check-import"
    assert imported.confirm == OWNER_FINAL_CHECK_TOKEN


def test_draft_import_binds_owner_selected_section_9_not_proposed_section_34(
    tmp_path: Path,
) -> None:
    pack = _pack(tmp_path)
    q02 = _issue(pack, "live30-q02", "issue-01")
    section_9 = next(item for item in q02["candidates"] if item["legal_locator"] == "section 9")
    decoy = dict(section_9)
    decoy.update(
        {
            "candidate_id": "candidate-section-34-decoy",
            "legal_locator": "section 34",
            "chunk_id": "chunk-cra-s31",
            "content_sha256": _digest(CRA_S31),
            "excerpt": CRA_S31,
            "cross_ref_fragment": True,
        }
    )
    q02["candidates"] = [decoy, *q02["candidates"]]
    q02["proposed_candidate_id"] = decoy["candidate_id"]
    q02["proposed_action"] = "accept_qualify"
    q02["owner_action"] = "approve_qualified"
    q02["owner_selected_candidate_ids"] = [section_9["candidate_id"]]
    q02["owner_legal_role"] = "statutory_text"
    q02["owner_later_treatment"] = "not_applicable_to_non_case_candidate"
    q02["owner_exact_proposition"] = (
        "Every contract to supply goods includes a term that the quality of the goods "
        "is satisfactory."
    )
    pack["schema"] = SUBSTANTIVE_REVIEW_SCHEMA
    pack["owner_confirmation_token"] = ""
    pack["reviewer_role"] = ""
    pack["reviewer_ref"] = ""
    adapted = adapt_substantive_review_for_import(pack, confirmation_token=OWNER_FINAL_CHECK_TOKEN)
    adapted_issue = _issue(adapted, "live30-q02", "issue-01")
    assert adapted_issue["owner_action"] == "accept_qualify"
    assert adapted_issue["approved_candidate_ids"] == [section_9["candidate_id"]]
    assert adapted_issue["owner_later_treatment"] == ""
    assert adapted_issue["owner_exact_proposition"] == ""
    imported = _import_pack(tmp_path, pack, confirmation_token=OWNER_FINAL_CHECK_TOKEN)
    qualified = next(item for item in imported["rows"] if item["row_id"] == "live30-q02:issue-01")
    assert qualified["status"] == "qualified"
    assert [span["legal_locator"] for span in qualified["exact_gold_spans"]] == ["section 9"]
    assert qualified["exact_gold_spans"][0]["content_sha256"] == _digest(CRA_S9)
    assert imported["selected_qualified_issue_count"] == 1
    assert imported["ready_for_overlay_seal"] is False


def test_draft_import_binds_all_selected_cpr_spans(tmp_path: Path) -> None:
    pack = _pack(tmp_path)
    _add_cpr_chunks(tmp_path / "catalog.sqlite3")
    issue = _issue(pack, "live60-q59", "issue-12")
    template = _legislation_template(pack)
    candidates = [
        {
            **template,
            "candidate_id": f"candidate-cpr-{locator.replace(' ', '-')}",
            "source_name": "Civil Procedure Rules",
            "authority_identity": "uk-cpr",
            "source_version_id": "source-version-cpr",
            "stable_source_id": "uk-cpr",
            "official_snapshot_sha256": "c" * 64,
            "chunk_id": chunk_id,
            "legal_locator": locator,
            "content_sha256": _digest(text),
            "excerpt": text,
            "cross_ref_fragment": False,
        }
        for chunk_id, locator, text in (
            ("chunk-cpr-351", "rule 35.1", CPR_351),
            ("chunk-cpr-355", "rule 35.5", CPR_355),
            ("chunk-cpr-357", "rule 35.7", CPR_357),
        )
    ]
    issue["candidates"] = candidates
    issue["proposed_candidate_id"] = candidates[0]["candidate_id"]
    issue["proposed_action"] = "accept_qualify"
    issue["owner_action"] = "approve_qualified"
    issue["owner_selected_candidate_ids"] = [item["candidate_id"] for item in candidates]
    issue["owner_legal_role"] = "statutory_text"
    issue["owner_later_treatment"] = "not_applicable_to_non_case_candidate"
    issue["owner_exact_proposition"] = (
        "Expert evidence is restricted to what is reasonably required to resolve the proceedings."
    )
    pack["schema"] = SUBSTANTIVE_REVIEW_SCHEMA
    imported = _import_pack(tmp_path, pack, confirmation_token=OWNER_FINAL_CHECK_TOKEN)
    qualified = next(item for item in imported["rows"] if item["row_id"] == "live60-q59:issue-12")
    assert qualified["status"] == "qualified"
    assert [span["legal_locator"] for span in qualified["exact_gold_spans"]] == [
        "rule 35.1",
        "rule 35.5",
        "rule 35.7",
    ]
    assert imported["selected_qualified_issue_count"] == 2
    assert imported["ready_for_overlay_seal"] is False


def test_qualified_current_case_without_limiting_authority_stays_gap(
    tmp_path: Path,
) -> None:
    pack = _pack(tmp_path)
    q03 = _issue(pack, "live30-q03", "issue-01")
    q03["owner_action"] = "accept_qualify"
    q03["owner_later_treatment"] = "qualified_current"
    q03["owner_legal_role"] = "holding_ratio"
    q03["owner_exact_proposition"] = CAPARO
    imported = _import_pack(tmp_path, pack)
    row = next(item for item in imported["rows"] if item["row_id"] == "live30-q03:issue-01")
    assert row["status"] == "knowledge_gap"
    assert row["reason_code"] == "exact_span_verification_failed"
    assert row["exact_gold_spans"] == []


def test_confirmed_current_case_qualifies_when_hash_matches(tmp_path: Path) -> None:
    pack = _pack(tmp_path)
    q03 = _issue(pack, "live30-q03", "issue-01")
    q03["owner_action"] = "accept_qualify"
    q03["owner_later_treatment"] = "confirmed_current"
    q03["owner_legal_role"] = "holding_ratio"
    q03["owner_exact_proposition"] = CAPARO
    imported = _import_pack(tmp_path, pack)
    row = next(item for item in imported["rows"] if item["row_id"] == "live30-q03:issue-01")
    assert row["status"] == "qualified"
    assert row["exact_gold_spans"][0]["content_sha256"] == _digest(CAPARO)
    assert row["exact_gold_spans"][0]["legal_role"] == "holding_ratio"


def _issue(pack: dict[str, object], case_id: str, issue_id: str) -> dict[str, object]:
    return next(
        issue
        for case in pack["cases"]
        if case["case_id"] == case_id
        for issue in case["issues"]
        if issue["issue_id"] == issue_id
    )


def _legislation_template(pack: dict[str, object]) -> dict[str, object]:
    q02 = _issue(pack, "live30-q02", "issue-01")
    return dict(next(item for item in q02["candidates"] if item["legal_locator"] == "section 9"))


def _add_cpr_chunks(path: Path) -> None:
    connection = sqlite3.connect(path)
    connection.execute(
        "INSERT INTO documents VALUES (?, ?, ?, ?, ?, ?, ?)",
        (
            "document-cpr",
            "c" * 64,
            "uk-cpr",
            "citable",
            "primary_authority",
            "England and Wales",
            None,
        ),
    )
    connection.execute(
        "INSERT INTO source_versions VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (
            "source-version-cpr",
            "document-cpr",
            "uk-cpr",
            "Civil Procedure Rules",
            "2026-08-16",
            "latest_available_revised_snapshot",
            "Open Government Licence v3.0",
            "https://example.invalid/licence",
            "approved",
            "https://example.invalid/cpr",
            "uk-cpr:latest-available@2026-08-16",
            None,
        ),
    )
    for chunk_id, locator, text in (
        ("chunk-cpr-351", "rule 35.1", CPR_351),
        ("chunk-cpr-355", "rule 35.5", CPR_355),
        ("chunk-cpr-357", "rule 35.7", CPR_357),
    ):
        connection.execute(
            "INSERT INTO chunks VALUES (?, ?, ?, ?, ?, ?, ?)",
            (chunk_id, "source-version-cpr", 1, locator, _digest(text), text, "body"),
        )
    connection.commit()
    connection.close()


def _import_pack(
    tmp_path: Path,
    pack: dict[str, object],
    *,
    confirmation_token: str | None = None,
    name: str = "pack.json",
) -> dict[str, object]:
    pack = dict(pack)
    if confirmation_token is None:
        pack["owner_confirmation_token"] = OWNER_FINAL_CHECK_TOKEN
    review_export = tmp_path / "review-export.json"
    if not review_export.is_file():
        export_review_candidates(
            project_root=PROJECT_ROOT,
            destination=review_export,
            cipher=LocalCipher(Fernet(Fernet.generate_key())),
            as_of_date=AS_OF_DATE,
        )
    return import_owner_final_check(
        project_root=PROJECT_ROOT,
        catalog_path=tmp_path / "catalog.sqlite3",
        pack_path=_write_pack(tmp_path, pack, name=name),
        review_export_path=review_export,
        as_of_date=AS_OF_DATE,
        repair_payload={"schema": HELD_SPAN_REPAIR_SCHEMA_V2, "repairs": []},
        confirmation_token=confirmation_token,
    )


def _write_pack(tmp_path: Path, pack: dict[str, object], *, name: str = "pack.json") -> Path:
    path = tmp_path / name
    path.write_text(json.dumps(pack), encoding="utf-8")
    return path
