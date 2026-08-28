from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document
from lxml import etree
from scripts import render_v111_phase2a_post_r110_owner_docx as renderer


def _document_text(path: Path) -> str:
    doc = Document(path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(paragraphs)


def test_docx_contains_complete_exact_owner_packet(tmp_path: Path) -> None:
    output = tmp_path / "r112"
    manifest = renderer.build(output)
    path = output / renderer.DOCX_NAME
    text = _document_text(path)

    assert renderer.BATCH_CONTENT_SHA256 in text
    assert "26 source-mapping dispositions" in text
    assert "I APPROVE THIS EXACT DIGEST-BOUND PHASE-2A BATCH" in text
    assert text.count("Mapping ") >= 26
    for source_id in (
        "neutral-citation:[2021] UKSC 3",
        "neutral-citation:[2025] UKSC 22",
        "neutral-citation:[2025] EWHC 38 (Ch)",
        "neutral-citation:[2012] EWHC 1257 (Ch)",
        "uksi:2006:246",
    ):
        assert source_id in text
    assert manifest["status"] == "DOCX_CREATED_VISUAL_QA_REQUIRED"
    assert manifest["visual_qa_completed"] is False
    assert manifest["source_admission_authorized"] is False
    assert manifest["phase2b_authorized"] is False


def test_docx_uses_exact_table_geometry_and_no_fixed_rows(tmp_path: Path) -> None:
    output = tmp_path / "r112"
    renderer.build(output)
    path = output / renderer.DOCX_NAME
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    def attr(name: str) -> str:
        return f"{{{ns['w']}}}{name}"

    with ZipFile(path) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    tables = root.xpath(".//w:tbl", namespaces=ns)
    assert len(tables) == 2
    for table in tables:
        table_width = table.xpath("./w:tblPr/w:tblW", namespaces=ns)[0]
        table_indent = table.xpath("./w:tblPr/w:tblInd", namespaces=ns)[0]
        grid = [
            int(cell.get(attr("w"))) for cell in table.xpath("./w:tblGrid/w:gridCol", namespaces=ns)
        ]
        assert table_width.get(attr("type")) == "dxa"
        assert int(table_width.get(attr("w"))) == renderer.CONTENT_WIDTH_DXA
        assert table_indent.get(attr("type")) == "dxa"
        assert int(table_indent.get(attr("w"))) == renderer.TABLE_INDENT_DXA
        assert sum(grid) == renderer.CONTENT_WIDTH_DXA
        for row in table.xpath("./w:tr", namespaces=ns):
            widths = [
                int(cell.get(attr("w"))) for cell in row.xpath("./w:tc/w:tcPr/w:tcW", namespaces=ns)
            ]
            assert widths == grid
            assert not row.xpath("./w:trPr/w:trHeight", namespaces=ns)


def test_docx_build_manifest_is_sealed(tmp_path: Path) -> None:
    output = tmp_path / "r112"
    manifest = renderer.build(output)
    persisted = json.loads((output / renderer.BUILD_MANIFEST_NAME).read_bytes())

    assert persisted == manifest
    material = dict(persisted)
    supplied = material.pop("manifest_content_sha256")
    assert supplied == renderer._sealed(material)
    assert persisted["docx_file_sha256"] == renderer._sha256_file(output / renderer.DOCX_NAME)


def test_docx_builder_is_create_only(tmp_path: Path) -> None:
    output = tmp_path / "r112"
    renderer.build(output)
    with pytest.raises(ValueError, match="phase2a_r112_output_already_exists"):
        renderer.build(output)
