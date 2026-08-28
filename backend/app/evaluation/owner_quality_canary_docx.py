"""Create-only answer-only owner-quality DOCX and real render receipt."""

from __future__ import annotations

import hashlib
import io
import json
import os
import re
import stat
import struct
import subprocess
import zipfile
import zlib
from collections.abc import Sequence
from contextlib import suppress
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal, Self
from uuid import uuid4

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pydantic import BaseModel, ConfigDict, Field, field_validator, model_validator
from pypdf import PdfReader

from ..privacy import contains_absolute_private_path
from ..text_metrics import word_count
from .canary_review_workspace import CanaryReviewWorkspace
from .live30 import assert_safe_evaluation_payload
from .live_suite import sealed_sha256
from .owner_quality_canary_authorization import OwnerDecisionRequired
from .owner_quality_canary_projection import OwnerCanaryFinalReviewPackage
from .secure_artifact_io import read_file_at, set_file_mode_at, write_create_only_at

OWNER_QUALITY_DOCX_CONTROL_SCHEMA = "legalbot.owner-quality-canary-docx-control.v1"
OWNER_QUALITY_DOCX_RENDER_RECEIPT_SCHEMA = "legalbot.owner-quality-canary-docx-render-receipt.v2"
OWNER_QUALITY_DOCX_INSPECTION_SCHEMA = "legalbot.owner-quality-canary-docx-inspection.v1"

APPROVED_RENDERER_BUNDLE_VERSION = "26.819.11345"
APPROVED_RENDERER_ID = "documents-render-docx-26.819.11345"
APPROVED_RENDERER_SCRIPT_SHA256 = "d8fe979f76e11215e146e53484bb4cb4e5f3906b58debed6844171073b187286"
APPROVED_RENDERER_PYTHON_SHA256 = "71720f1fc66989ebd691e81c96111b47ae6ff3f1a478666084d1cacbf0fccbf2"
APPROVED_RENDERER_COMMAND = (
    "./approved-python-3.12",
    "{approved-renderer-script-fd}",
    "render-input.docx",
    "--output_dir",
    ".",
    "--emit_pdf",
)
APPROVED_RENDER_TIMEOUT_SECONDS = 600

DOCX_PRESET = "compact_reference_guide"
DOCX_HEADER_TEMPLATE = "memo_masthead"
DOCX_ANNEX_LAYOUT = "three_annexes_of_ten"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 0
TABLE_CELL_MARGINS_DXA = (80, 120, 80, 120)
TABLE_HORIZONTAL_MARGIN_DXA = TABLE_CELL_MARGINS_DXA[1] + TABLE_CELL_MARGINS_DXA[3]
TABLE_WIDTH_DXA = CONTENT_WIDTH_DXA - TABLE_HORIZONTAL_MARGIN_DXA
# Keep the grid inside the usable width after accounting for both horizontal
# cell margins. LibreOffice can otherwise relocate a first-column paragraph to
# the physical page edge while paginating dense tables.

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5C6773"
TABLE_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "B8C4D1"

_SHA256 = re.compile(r"^[0-9a-f]{64}$")
_PAGE_PNG = re.compile(r"^page-([1-9][0-9]*)\.png$")
_NUMBERED = re.compile(r"^\s*([0-9]+)[.)]\s+(.+)$")


class OwnerQualityCanaryDocxControl(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True, populate_by_name=True)

    schema_name: Literal["legalbot.owner-quality-canary-docx-control.v1"] = Field(
        default="legalbot.owner-quality-canary-docx-control.v1", alias="schema"
    )
    document_id: str = Field(pattern=r"^owner-canary-review-docx-[0-9a-f]{16}$")
    workspace_seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    final_package_seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    run_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    lane: Literal["development", "blind_holdout"]
    candidate_build_id: str = Field(pattern=r"^[A-Za-z0-9][A-Za-z0-9._:-]{2,127}$")
    candidate_manifest_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    case_count: Literal[30]
    case_ids: tuple[str, ...]
    answer_sha256s: tuple[str, ...]
    projection_receipt_seal_sha256s: tuple[str, ...]
    annex_case_ids: tuple[tuple[str, ...], tuple[str, ...], tuple[str, ...]]
    annex_layout: Literal["three_annexes_of_ten"]
    design_preset: Literal["compact_reference_guide"]
    header_template: Literal["memo_masthead"]
    answer_only: Literal[True]
    plaintext_questions_included: Literal[False]
    held_plaintext_included: Literal[False]
    full_released_answers_included: Literal[True]
    explicit_table_geometry: Literal[True]
    metadata_scrubbed: Literal[True]
    deterministic_zip_normalized: Literal[True]
    create_only: Literal[True]
    render_required: Literal[True]
    document_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    document_byte_count: int = Field(ge=1)
    seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")

    @field_validator("case_ids", "answer_sha256s", "projection_receipt_seal_sha256s")
    @classmethod
    def ordered_identities_are_exact(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if len(values) != 30:
            raise ValueError("owner-quality DOCX control requires exactly 30 identities")
        return values

    @model_validator(mode="after")
    def control_is_exact_and_sealed(self) -> Self:
        if (
            any(len(group) != 10 for group in self.annex_case_ids)
            or tuple(case for group in self.annex_case_ids for case in group) != self.case_ids
            or len(set(self.case_ids)) != 30
            or len(set(self.projection_receipt_seal_sha256s)) != 30
        ):
            raise ValueError("owner-quality DOCX annex or receipt layout is inconsistent")
        if self.seal_sha256 != sealed_sha256(self.model_dump(mode="json", by_alias=True)):
            raise ValueError("owner-quality DOCX control seal does not match")
        return self


class OwnerQualityCanaryRenderedPage(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    ordinal: int = Field(ge=1, le=10_000)
    png_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    width_px: int = Field(ge=500, le=20_000)
    height_px: int = Field(ge=500, le=30_000)


class OwnerQualityCanaryDocxRenderReceipt(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True, populate_by_name=True)

    schema_name: Literal["legalbot.owner-quality-canary-docx-render-receipt.v2"] = Field(
        default="legalbot.owner-quality-canary-docx-render-receipt.v2", alias="schema"
    )
    document_id: str = Field(pattern=r"^owner-canary-review-docx-[0-9a-f]{16}$")
    docx_control_seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    document_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    document_byte_count: int = Field(ge=1)
    renderer_id: Literal["documents-render-docx-26.819.11345"]
    renderer_bundle_version: Literal["26.819.11345"]
    renderer_script_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    renderer_python_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    renderer_command: tuple[str, ...]
    render_output_directory_name: str = Field(pattern=r"^render-output-[0-9a-f]{16}$")
    rendered_at: datetime
    rendered_pdf_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    rendered_pdf_byte_count: int = Field(ge=1)
    page_count: int = Field(ge=1, le=10_000)
    pages: tuple[OwnerQualityCanaryRenderedPage, ...]
    answer_only_verified: Literal[True]
    technical_render_passed: Literal[True]
    owner_visual_inspection_required: Literal[True]
    trusted_owner_visual_inspection_verified: Literal[False]
    authorizes_owner_review_companion: Literal[False]
    create_only: Literal[True]
    seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")

    @field_validator("rendered_at")
    @classmethod
    def rendered_at_is_timezone_aware(cls, value: datetime) -> datetime:
        if value.tzinfo is None or value.utcoffset() is None:
            raise ValueError("DOCX render timestamp must be timezone-aware")
        return value

    @model_validator(mode="after")
    def receipt_covers_every_page_and_is_sealed(self) -> Self:
        if (
            self.page_count != len(self.pages)
            or tuple(page.ordinal for page in self.pages) != tuple(range(1, self.page_count + 1))
            or self.renderer_script_sha256 != APPROVED_RENDERER_SCRIPT_SHA256
            or self.renderer_python_sha256 != APPROVED_RENDERER_PYTHON_SHA256
            or self.renderer_command != APPROVED_RENDERER_COMMAND
        ):
            raise ValueError("DOCX render receipt does not cover every page exactly once")
        if self.seal_sha256 != sealed_sha256(self.model_dump(mode="json", by_alias=True)):
            raise ValueError("DOCX render receipt seal does not match")
        return self


class OwnerQualityCanaryDocxInspectionAttestation(BaseModel):
    """Owner-signed visual inspection, deliberately separate from rendering."""

    model_config = ConfigDict(extra="forbid", frozen=True, populate_by_name=True)

    schema_name: Literal["legalbot.owner-quality-canary-docx-inspection.v1"] = Field(
        default="legalbot.owner-quality-canary-docx-inspection.v1", alias="schema"
    )
    document_id: str = Field(pattern=r"^owner-canary-review-docx-[0-9a-f]{16}$")
    docx_control_seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    render_receipt_seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    rendered_page_set_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    owner_ref: str = Field(pattern=r"^owner:[0-9a-f]{64}$")
    inspected_at: datetime
    inspected_page_count: int = Field(ge=1, le=10_000)
    all_pages_inspected_at_full_size: Literal[True]
    visual_inspection_passed: Literal[True]
    no_clipped_text: Literal[True]
    no_overlapping_objects: Literal[True]
    no_broken_tables: Literal[True]
    no_missing_glyphs: Literal[True]
    signature_algorithm: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,63}$")
    signature_payload_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    signature: str = Field(min_length=16, max_length=4096)
    trusted_owner_signature_verified: Literal[True]
    authorizes_owner_review_companion: Literal[True]
    create_only: Literal[True]
    seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")

    @field_validator("inspected_at")
    @classmethod
    def inspected_at_is_timezone_aware(cls, value: datetime) -> datetime:
        if value.tzinfo is None or value.utcoffset() is None:
            raise ValueError("DOCX inspection timestamp must be timezone-aware")
        return value

    @model_validator(mode="after")
    def inspection_is_sealed(self) -> Self:
        if self.seal_sha256 != sealed_sha256(self.model_dump(mode="json", by_alias=True)):
            raise ValueError("DOCX inspection attestation seal does not match")
        return self


def _validated_package(
    value: OwnerCanaryFinalReviewPackage,
) -> OwnerCanaryFinalReviewPackage:
    return OwnerCanaryFinalReviewPackage.model_validate(
        value.model_dump(mode="json", by_alias=True)
    )


def _set_run_font(
    run: Any,
    *,
    size: float,
    color: str = "000000",
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in zip(("top", "start", "bottom", "end"), TABLE_CELL_MARGINS_DXA, strict=True):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: Sequence[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError("DOCX table grid must preserve the explicit horizontal margin headroom")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_look = tbl_pr.first_child_found_in("w:tblLook")
    if table_look is None:
        table_look = OxmlElement("w:tblLook")
        tbl_pr.append(table_look)
    for name, value in (
        ("firstColumn", "0"),
        ("firstRow", "0"),
        ("lastColumn", "0"),
        ("lastRow", "0"),
        ("noHBand", "1"),
        ("noVBand", "1"),
        ("val", "0000"),
    ):
        table_look.set(qn(f"w:{name}"), value)
    cell_spacing = tbl_pr.first_child_found_in("w:tblCellSpacing")
    if cell_spacing is None:
        cell_spacing = OxmlElement("w:tblCellSpacing")
        tbl_pr.append(cell_spacing)
    cell_spacing.set(qn("w:w"), "0")
    cell_spacing.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        if len(row.cells) != len(widths):
            raise ValueError("DOCX table row does not match explicit column geometry")
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)
        cant_split.set(qn("w:val"), "1")
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph_format = paragraph.paragraph_format
                paragraph_format.left_indent = Pt(0)
                paragraph_format.right_indent = Pt(0)
                paragraph_format.first_line_indent = Pt(0)
                paragraph_format.keep_together = True
                paragraph_format.keep_with_next = False
                paragraph_format.widow_control = False
                p_pr = paragraph._p.get_or_add_pPr()
                keep_lines = p_pr.find(qn("w:keepLines"))
                keep_next = p_pr.find(qn("w:keepNext"))
                widow_control = p_pr.find(qn("w:widowControl"))
                if keep_lines is None or keep_next is None or widow_control is None:
                    raise ValueError("DOCX table paragraph pagination controls are missing")
                keep_lines.set(qn("w:val"), "1")
                keep_next.set(qn("w:val"), "0")
                widow_control.set(qn("w:val"), "0")


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_text(
    cell: Any,
    text: str,
    *,
    size: float = 9.0,
    bold: bool = False,
    color: str = "000000",
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)


def _paragraph_border_bottom(paragraph: Any, *, color: str, size: int = 12) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def _configure_styles(document: DocumentType) -> tuple[int, int]:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if "Owner Review Callout" not in styles:
        callout = styles.add_style("Owner Review Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Owner Review Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(INK)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.left_indent = Pt(8)
    callout.paragraph_format.right_indent = Pt(8)
    callout.paragraph_format.line_spacing = 1.25
    return _add_numbering(document, bullet=True), _add_numbering(document, bullet=False)


def _add_numbering(document: DocumentType, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId"), "0")) for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"), "0")) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "271")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph: Any, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_ref)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def _configure_page(document: DocumentType, run_id: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run("OWNER QUALITY CANARY · ANSWER-ONLY")
    _set_run_font(left, size=8.5, color=MUTED, bold=True)
    paragraph.add_run("\t")
    right = paragraph.add_run(run_id)
    _set_run_font(right, size=8.5, color=MUTED)
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    label = footer_paragraph.add_run("Owner review · Page ")
    _set_run_font(label, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    result_run = OxmlElement("w:r")
    result_text = OxmlElement("w:t")
    result_text.text = "1"
    result_run.append(result_text)
    field.append(result_run)
    footer_paragraph._p.append(field)


def _add_page_break(document: DocumentType) -> None:
    paragraph = document.add_paragraph()
    run = OxmlElement("w:r")
    page_break = OxmlElement("w:br")
    page_break.set(qn("w:type"), "page")
    run.append(page_break)
    paragraph._p.append(run)


def _add_title_block(document: DocumentType, package: OwnerCanaryFinalReviewPackage) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("OWNER QUALITY CANARY REVIEW")
    _set_run_font(run, size=23, color="000000", bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        f"{package.lane.replace('_', ' ').title()} · 30 released answers · Owner-only"
    )
    _set_run_font(run, size=14, color="373737")
    metadata = document.add_table(rows=5, cols=2)
    metadata.style = "Table Grid"
    values = (
        ("Run", package.run_id),
        ("Candidate", package.candidate_build_id),
        ("Candidate manifest", package.candidate_manifest_sha256),
        ("Final package", package.seal_sha256),
        ("Review scope", "Answer-only · 3 annexes × 10 cases"),
    )
    for row, (label, value) in zip(metadata.rows, values, strict=True):
        _shade_cell(row.cells[0], TABLE_FILL)
        _set_cell_text(row.cells[0], label, size=9.5, bold=True, color=INK)
        _set_cell_text(row.cells[1], value, size=9.0)
    _set_table_geometry(metadata, (2700, 6420))
    callout = document.add_paragraph(style="Owner Review Callout")
    p_pr = callout._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_FILL)
    p_pr.append(shading)
    run = callout.add_run("Review boundary. ")
    _set_run_font(run, size=10.5, color=INK, bold=True)
    run = callout.add_run(
        "This pack contains the complete gate-passed released answers only. "
        "Questions, held drafts, and detailed knowledge-gap prose are intentionally excluded."
    )
    _set_run_font(run, size=10.5, color=INK)
    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(4)
    rule.paragraph_format.space_after = Pt(4)
    _paragraph_border_bottom(rule, color=INK, size=14)


def _add_case_metadata(document: DocumentType, receipt: Any) -> None:
    def digest_prefix(value: str) -> str:
        # Break opportunities prevent a proportional-font digest from forcing
        # LibreOffice to widen a fixed cell during pagination.
        prefix = value[:16]
        return " ".join(prefix[index : index + 4] for index in range(0, len(prefix), 4)) + " …"

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    columns = (
        "Words\nEvidence\nStandards",
        "\n".join(
            (
                str(receipt.word_count),
                digest_prefix(receipt.evidence_bundle_seal_sha256),
                digest_prefix(receipt.standards_report_seal_sha256),
            )
        ),
        "Answer SHA\nAI review\nReceipt",
        "\n".join(
            (
                digest_prefix(receipt.answer_sha256),
                digest_prefix(receipt.ai_review_seal_sha256),
                digest_prefix(receipt.seal_sha256),
            )
        ),
    )
    for index, value in enumerate(columns):
        if index % 2 == 0:
            _shade_cell(table.rows[0].cells[index], TABLE_FILL)
            _set_cell_text(table.rows[0].cells[index], value, size=8.5, bold=True, color=INK)
        else:
            _set_cell_text(table.rows[0].cells[index], value, size=8.5)
    _set_table_geometry(table, (1200, 3360, 1200, 3360))


def _add_answer_markdown(
    document: DocumentType, content: str, *, bullet_num_id: int, decimal_num_id: int
) -> None:
    blocks = content.splitlines()
    if not blocks:
        blocks = [content]
    for raw in blocks:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            document.add_paragraph(line[4:].strip(), style="Heading 3")
        elif line.startswith("## "):
            document.add_paragraph(line[3:].strip(), style="Heading 3")
        elif line.startswith("# "):
            document.add_paragraph(line[2:].strip(), style="Heading 3")
        elif line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(line[2:].strip())
            _apply_numbering(paragraph, bullet_num_id)
        elif match := _NUMBERED.match(line):
            paragraph = document.add_paragraph(match.group(2))
            _apply_numbering(paragraph, decimal_num_id)
        else:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.widow_control = True


def _add_owner_review_form(document: DocumentType) -> None:
    heading = document.add_paragraph("Owner review", style="Heading 3")
    heading.paragraph_format.space_before = Pt(8)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _shade_cell(table.rows[0].cells[0], TABLE_FILL)
    _set_cell_text(
        table.rows[0].cells[0], "Decision\nFeedback / date", size=9, bold=True, color=INK
    )
    _set_cell_text(
        table.rows[0].cells[1],
        "☐ Pass    ☐ Revise    ☐ Reject\n________________________________",
        size=9,
    )
    _set_table_geometry(table, (2000, 7120))


def _deterministic_docx_bytes(document: DocumentType) -> bytes:
    raw = io.BytesIO()
    document.save(raw)
    source = zipfile.ZipFile(io.BytesIO(raw.getvalue()), "r")
    output = io.BytesIO()
    with (
        source,
        zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target,
    ):
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = 0
            target.writestr(info, source.read(name))
    return output.getvalue()


def export_owner_quality_canary_docx(
    *,
    workspace: CanaryReviewWorkspace,
    package: OwnerCanaryFinalReviewPackage,
) -> tuple[Path, Path, OwnerQualityCanaryDocxControl]:
    """Create the deterministic three-by-ten answer-only annex and sealed control."""

    package = _validated_package(package)
    if (
        workspace.manifest.seal_sha256 != package.workspace_seal_sha256
        or workspace.manifest.run_id != package.run_id
        or workspace.manifest.lane != package.lane
        or workspace.manifest.expected_case_ids != package.case_ids
        or not package.answer_only
        or package.plaintext_questions_included
    ):
        raise ValueError("DOCX export workspace differs from the finalized answer-only package")
    persisted_package = OwnerCanaryFinalReviewPackage.model_validate_json(
        workspace.read_private_bytes("safe-metrics", "final-review-package.json")
    )
    if persisted_package != package:
        raise ValueError("DOCX export package differs from its persisted final package")
    document_id = f"owner-canary-review-docx-{package.seal_sha256[:16]}"
    existing_review_files = workspace.list_private_directory("review-docx")
    if f"{document_id}.docx" in existing_review_files or "docx-control.json" in (
        existing_review_files
    ):
        raise FileExistsError("owner-quality DOCX export is create-only")

    answers: dict[str, str] = {}
    for case_id, expected_sha in zip(package.case_ids, package.answer_sha256s, strict=True):
        encoded = workspace.read_private_bytes("cases", case_id, "released-answer.md")
        try:
            content = encoded.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError("owner-quality DOCX released answer is not UTF-8") from exc
        if (
            hashlib.sha256(encoded).hexdigest() != expected_sha
            or not content.strip()
            or contains_absolute_private_path(content)
        ):
            raise ValueError("owner-quality DOCX released answer failed identity or privacy")
        receipt = package.projection_receipts[package.case_ids.index(case_id)]
        if word_count(content) != receipt.word_count:
            raise ValueError("owner-quality DOCX answer word count differs from its receipt")
        answers[case_id] = content

    document = Document()
    bullet_num_id, decimal_num_id = _configure_styles(document)
    _configure_page(document, package.run_id)
    core = document.core_properties
    core.title = "Owner Quality Canary Review"
    core.subject = "Answer-only owner review annex"
    core.author = ""
    core.last_modified_by = ""
    core.comments = ""
    core.keywords = ""
    core.category = "Evaluation"
    core.created = datetime(2000, 1, 1, tzinfo=UTC)
    core.modified = datetime(2000, 1, 1, tzinfo=UTC)
    _add_title_block(document, package)
    _add_page_break(document)
    annex_names = ("Annex A", "Annex B", "Annex C")
    annex_groups = (
        package.case_ids[:10],
        package.case_ids[10:20],
        package.case_ids[20:30],
    )
    by_case = {item.case_id: item for item in package.projection_receipts}
    sequence = 0
    for annex_index, (annex_name, group) in enumerate(zip(annex_names, annex_groups, strict=True)):
        if annex_index:
            _add_page_break(document)
        heading = document.add_paragraph(
            f"{annex_name} · Cases {annex_index * 10 + 1:02d}–{annex_index * 10 + 10:02d}",
            style="Heading 1",
        )
        heading.paragraph_format.keep_with_next = True
        for case_index, case_id in enumerate(group):
            if case_index:
                # Do not let the renderer synthesize a page boundary between a
                # case heading and its table; that path can misplace cell text.
                _add_page_break(document)
            sequence += 1
            case_heading = document.add_paragraph(
                f"Case {sequence:02d} · {case_id}", style="Heading 2"
            )
            case_heading.paragraph_format.keep_with_next = True
            _add_case_metadata(document, by_case[case_id])
            answer_heading = document.add_paragraph("Released answer", style="Heading 3")
            answer_heading.paragraph_format.keep_with_next = True
            _add_answer_markdown(
                document,
                answers[case_id],
                bullet_num_id=bullet_num_id,
                decimal_num_id=decimal_num_id,
            )
            _add_owner_review_form(document)

    document_bytes = _deterministic_docx_bytes(document)
    docx_path = workspace.write_private_bytes(
        "review-docx", f"{document_id}.docx", payload=document_bytes
    )
    annex_case_ids = (annex_groups[0], annex_groups[1], annex_groups[2])
    control_material: dict[str, Any] = {
        "schema": OWNER_QUALITY_DOCX_CONTROL_SCHEMA,
        "document_id": document_id,
        "workspace_seal_sha256": workspace.manifest.seal_sha256,
        "final_package_seal_sha256": package.seal_sha256,
        "run_id": package.run_id,
        "lane": package.lane,
        "candidate_build_id": package.candidate_build_id,
        "candidate_manifest_sha256": package.candidate_manifest_sha256,
        "case_count": 30,
        "case_ids": list(package.case_ids),
        "answer_sha256s": list(package.answer_sha256s),
        "projection_receipt_seal_sha256s": list(package.projection_receipt_seal_sha256s),
        "annex_case_ids": [list(group) for group in annex_case_ids],
        "annex_layout": DOCX_ANNEX_LAYOUT,
        "design_preset": DOCX_PRESET,
        "header_template": DOCX_HEADER_TEMPLATE,
        "answer_only": True,
        "plaintext_questions_included": False,
        "held_plaintext_included": False,
        "full_released_answers_included": True,
        "explicit_table_geometry": True,
        "metadata_scrubbed": True,
        "deterministic_zip_normalized": True,
        "create_only": True,
        "render_required": True,
        "document_sha256": hashlib.sha256(document_bytes).hexdigest(),
        "document_byte_count": len(document_bytes),
    }
    control_material["seal_sha256"] = sealed_sha256(control_material)
    control = OwnerQualityCanaryDocxControl.model_validate(control_material)
    assert_safe_evaluation_payload(control.model_dump(mode="json", by_alias=True))
    control_path = workspace.write_private_bytes(
        "review-docx",
        "docx-control.json",
        payload=(
            json.dumps(
                control.model_dump(mode="json", by_alias=True),
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
            )
            + "\n"
        ).encode("utf-8"),
    )
    return docx_path, control_path, control


def _validated_png_bytes(data: bytes) -> tuple[str, int, int]:
    if not data.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ValueError("DOCX render page is not a PNG")
    offset = 8
    width = height = 0
    saw_iend = False
    while offset + 12 <= len(data):
        length = struct.unpack(">I", data[offset : offset + 4])[0]
        chunk_type = data[offset + 4 : offset + 8]
        end = offset + 12 + length
        if end > len(data):
            raise ValueError("DOCX render PNG is truncated")
        chunk_data = data[offset + 8 : offset + 8 + length]
        expected_crc = struct.unpack(">I", data[offset + 8 + length : end])[0]
        if zlib.crc32(chunk_type + chunk_data) & 0xFFFFFFFF != expected_crc:
            raise ValueError("DOCX render PNG checksum is invalid")
        if chunk_type == b"IHDR":
            if length != 13 or width or height:
                raise ValueError("DOCX render PNG header is invalid")
            width, height = struct.unpack(">II", chunk_data[:8])
        if chunk_type == b"IEND":
            saw_iend = True
            if end != len(data):
                raise ValueError("DOCX render PNG has trailing bytes")
            break
        offset = end
    if not saw_iend or width < 500 or height < 500:
        raise ValueError("DOCX render PNG is incomplete or too small")
    return hashlib.sha256(data).hexdigest(), width, height


def _approved_renderer_script_path() -> Path:
    return (
        Path.home()
        / ".codex/plugins/cache/openai-primary-runtime/documents"
        / APPROVED_RENDERER_BUNDLE_VERSION
        / "skills/documents/render_docx.py"
    )


def _approved_renderer_python_path() -> Path:
    return (
        Path.home()
        / ".cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3.12"
    )


def _open_verified_regular_file(path: Path, *, expected_sha256: str) -> tuple[int, bytes]:
    flags = os.O_RDONLY | os.O_NOFOLLOW | int(getattr(os, "O_CLOEXEC", 0))
    try:
        descriptor = os.open(path, flags)
    except OSError as exc:
        raise RuntimeError("approved_docx_renderer_missing") from exc
    try:
        metadata = os.fstat(descriptor)
        if not stat.S_ISREG(metadata.st_mode) or metadata.st_size < 1:
            raise RuntimeError("approved_docx_renderer_identity_mismatch")
        chunks: list[bytes] = []
        while True:
            block = os.read(descriptor, 1024 * 1024)
            if not block:
                break
            chunks.append(block)
        data = b"".join(chunks)
        final = os.fstat(descriptor)
        if (
            len(data) != metadata.st_size
            or final.st_dev != metadata.st_dev
            or final.st_ino != metadata.st_ino
            or final.st_size != metadata.st_size
            or final.st_mtime_ns != metadata.st_mtime_ns
            or final.st_ctime_ns != metadata.st_ctime_ns
            or hashlib.sha256(data).hexdigest() != expected_sha256
        ):
            raise RuntimeError("approved_docx_renderer_identity_mismatch")
        os.lseek(descriptor, 0, os.SEEK_SET)
        return descriptor, data
    except Exception:
        os.close(descriptor)
        raise


def _approved_renderer_environment() -> dict[str, str]:
    dependencies = _approved_renderer_python_path().parents[2]
    python_home = dependencies / "python"
    runtime_override = dependencies / "bin" / "override"
    return {
        "PATH": f"{runtime_override}:/usr/bin:/bin:/usr/sbin:/sbin",
        "PYTHONHOME": os.fspath(python_home),
        "PYTHONNOUSERSITE": "1",
        "PYTHONSAFEPATH": "1",
        "TMPDIR": "/private/tmp",
        "TEMP": "/private/tmp",
        "TMP": "/private/tmp",
    }


def _rendered_page_set_sha256(receipt: OwnerQualityCanaryDocxRenderReceipt) -> str:
    return sealed_sha256(
        {
            "schema": "legalbot.owner-quality-canary-rendered-page-set.v1",
            "document_sha256": receipt.document_sha256,
            "rendered_pdf_sha256": receipt.rendered_pdf_sha256,
            "pages": [page.model_dump(mode="json") for page in receipt.pages],
        }
    )


def _verify_render_output(
    *, workspace: CanaryReviewWorkspace, receipt: OwnerQualityCanaryDocxRenderReceipt
) -> None:
    output = receipt.render_output_directory_name
    expected_names = {
        "render-input.pdf",
        *(f"page-{ordinal}.png" for ordinal in range(1, receipt.page_count + 1)),
    }
    if set(workspace.list_private_directory("review-docx", output)) != expected_names:
        raise ValueError("DOCX render output is not the exact retained page set")
    pdf = workspace.read_private_bytes("review-docx", output, "render-input.pdf")
    if (
        hashlib.sha256(pdf).hexdigest() != receipt.rendered_pdf_sha256
        or len(pdf) != receipt.rendered_pdf_byte_count
    ):
        raise ValueError("DOCX rendered PDF differs from its technical receipt")
    for expected, page in zip(range(1, receipt.page_count + 1), receipt.pages, strict=True):
        raw = workspace.read_private_bytes("review-docx", output, f"page-{expected}.png")
        digest, width, height = _validated_png_bytes(raw)
        if (page.ordinal, page.png_sha256, page.width_px, page.height_px) != (
            expected,
            digest,
            width,
            height,
        ):
            raise ValueError("DOCX rendered page differs from its technical receipt")


def record_owner_quality_canary_docx_render(
    *,
    workspace: CanaryReviewWorkspace,
    control: OwnerQualityCanaryDocxControl,
    rendered_at: datetime,
) -> tuple[Path, OwnerQualityCanaryDocxRenderReceipt]:
    """Own the approved render and seal its exact PDF/PNG output bytes."""

    control = OwnerQualityCanaryDocxControl.model_validate(
        control.model_dump(mode="json", by_alias=True)
    )
    if "render-receipt.json" in workspace.list_private_directory("review-docx"):
        raise FileExistsError("DOCX render receipt is create-only")
    persisted_control = OwnerQualityCanaryDocxControl.model_validate_json(
        workspace.read_private_bytes("review-docx", "docx-control.json")
    )
    document_bytes = workspace.read_private_bytes("review-docx", f"{control.document_id}.docx")
    if (
        workspace.manifest.seal_sha256 != control.workspace_seal_sha256
        or persisted_control != control
        or hashlib.sha256(document_bytes).hexdigest() != control.document_sha256
        or len(document_bytes) != control.document_byte_count
    ):
        raise ValueError("DOCX render inputs differ from the sealed control")

    output_name = f"render-output-{uuid4().hex[:16]}"
    workspace.create_private_directory("review-docx", output_name, exist_ok=False)
    renderer_script_fd, _renderer_script_bytes = _open_verified_regular_file(
        _approved_renderer_script_path(),
        expected_sha256=APPROVED_RENDERER_SCRIPT_SHA256,
    )
    renderer_python_fd, renderer_python_bytes = _open_verified_regular_file(
        _approved_renderer_python_path(),
        expected_sha256=APPROVED_RENDERER_PYTHON_SHA256,
    )
    pages: list[OwnerQualityCanaryRenderedPage] = []
    pdf_bytes = b""
    try:
        with workspace.open_private_directory("review-docx", output_name) as output_fd:
            write_create_only_at(output_fd, "render-input.docx", document_bytes)
            write_create_only_at(output_fd, "approved-python-3.12", renderer_python_bytes)
            set_file_mode_at(output_fd, "approved-python-3.12", 0o700)

            def _enter_private_render_directory() -> None:
                os.fchdir(output_fd)

            try:
                command = (
                    "./approved-python-3.12",
                    f"/dev/fd/{renderer_script_fd}",
                    "render-input.docx",
                    "--output_dir",
                    ".",
                    "--emit_pdf",
                )
                try:
                    completed = subprocess.run(
                        command,
                        check=False,
                        capture_output=True,
                        timeout=APPROVED_RENDER_TIMEOUT_SECONDS,
                        pass_fds=(renderer_script_fd, output_fd),
                        preexec_fn=_enter_private_render_directory,
                        env=_approved_renderer_environment(),
                    )
                except (OSError, subprocess.TimeoutExpired) as exc:
                    raise RuntimeError("owner_quality_docx_render_failed") from exc
                if completed.returncode != 0:
                    raise RuntimeError("owner_quality_docx_render_failed")
            finally:
                for staging_name in ("render-input.docx", "approved-python-3.12"):
                    with suppress(FileNotFoundError):
                        os.unlink(staging_name, dir_fd=output_fd)
            names = tuple(sorted(str(name) for name in os.listdir(output_fd)))
            page_ordinals: dict[str, int] = {}
            for name in names:
                match = _PAGE_PNG.fullmatch(name)
                if match is not None:
                    page_ordinals[name] = int(match.group(1))
            ordered = tuple(sorted(page_ordinals, key=page_ordinals.__getitem__))
            if (
                not ordered
                or ordered != tuple(f"page-{number}.png" for number in range(1, len(ordered) + 1))
                or set(names) != {"render-input.pdf", *ordered}
            ):
                raise ValueError("DOCX renderer did not create one exact contiguous page set")
            for name in names:
                set_file_mode_at(output_fd, name, 0o600)
            pdf_bytes = read_file_at(output_fd, "render-input.pdf", required_mode=0o600)
            try:
                pdf_page_count = len(PdfReader(io.BytesIO(pdf_bytes), strict=True).pages)
            except Exception as exc:
                raise ValueError("DOCX rendered PDF is invalid") from exc
            if pdf_page_count != len(ordered):
                raise ValueError("DOCX rendered PDF/page image counts differ")
            for ordinal, name in enumerate(ordered, start=1):
                digest, width, height = _validated_png_bytes(
                    read_file_at(output_fd, name, required_mode=0o600)
                )
                pages.append(
                    OwnerQualityCanaryRenderedPage(
                        ordinal=ordinal,
                        png_sha256=digest,
                        width_px=width,
                        height_px=height,
                    )
                )
    finally:
        os.close(renderer_python_fd)
        os.close(renderer_script_fd)

    material: dict[str, Any] = {
        "schema": OWNER_QUALITY_DOCX_RENDER_RECEIPT_SCHEMA,
        "document_id": control.document_id,
        "docx_control_seal_sha256": control.seal_sha256,
        "document_sha256": control.document_sha256,
        "document_byte_count": len(document_bytes),
        "renderer_id": APPROVED_RENDERER_ID,
        "renderer_bundle_version": APPROVED_RENDERER_BUNDLE_VERSION,
        "renderer_script_sha256": APPROVED_RENDERER_SCRIPT_SHA256,
        "renderer_python_sha256": APPROVED_RENDERER_PYTHON_SHA256,
        "renderer_command": list(APPROVED_RENDERER_COMMAND),
        "render_output_directory_name": output_name,
        "rendered_at": rendered_at.isoformat().replace("+00:00", "Z"),
        "rendered_pdf_sha256": hashlib.sha256(pdf_bytes).hexdigest(),
        "rendered_pdf_byte_count": len(pdf_bytes),
        "page_count": len(pages),
        "pages": [page.model_dump(mode="json") for page in pages],
        "answer_only_verified": True,
        "technical_render_passed": True,
        "owner_visual_inspection_required": True,
        "trusted_owner_visual_inspection_verified": False,
        "authorizes_owner_review_companion": False,
        "create_only": True,
    }
    material["seal_sha256"] = sealed_sha256(material)
    receipt = OwnerQualityCanaryDocxRenderReceipt.model_validate(material)
    assert_safe_evaluation_payload(receipt.model_dump(mode="json", by_alias=True))
    receipt_path = workspace.write_private_bytes(
        "review-docx",
        "render-receipt.json",
        payload=(
            json.dumps(
                receipt.model_dump(mode="json", by_alias=True),
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
            )
            + "\n"
        ).encode("utf-8"),
    )
    _verify_render_output(workspace=workspace, receipt=receipt)
    return receipt_path, receipt


def _owner_inspection_signature_payload(
    *,
    control: OwnerQualityCanaryDocxControl,
    receipt: OwnerQualityCanaryDocxRenderReceipt,
    owner_ref: str,
    inspected_at: datetime,
    inspected_page_count: int,
    all_pages_inspected_at_full_size: bool,
    visual_inspection_passed: bool,
    no_clipped_text: bool,
    no_overlapping_objects: bool,
    no_broken_tables: bool,
    no_missing_glyphs: bool,
) -> dict[str, Any]:
    return {
        "schema": "legalbot.owner-quality-canary-docx-inspection-signature-payload.v1",
        "document_id": control.document_id,
        "docx_control_seal_sha256": control.seal_sha256,
        "render_receipt_seal_sha256": receipt.seal_sha256,
        "rendered_page_set_sha256": _rendered_page_set_sha256(receipt),
        "owner_ref": owner_ref,
        "inspected_at": inspected_at.isoformat().replace("+00:00", "Z"),
        "inspected_page_count": inspected_page_count,
        "all_pages_inspected_at_full_size": all_pages_inspected_at_full_size,
        "visual_inspection_passed": visual_inspection_passed,
        "no_clipped_text": no_clipped_text,
        "no_overlapping_objects": no_overlapping_objects,
        "no_broken_tables": no_broken_tables,
        "no_missing_glyphs": no_missing_glyphs,
    }


def _verify_trusted_owner_docx_inspection_signature(
    *,
    owner_ref: str,
    signature_algorithm: str,
    signature_payload_sha256: str,
    signature: str,
) -> None:
    del owner_ref, signature_algorithm, signature_payload_sha256, signature
    raise OwnerDecisionRequired("trusted_owner_docx_inspection_signature_verifier_missing")


def record_owner_quality_canary_docx_inspection(
    *,
    workspace: CanaryReviewWorkspace,
    control: OwnerQualityCanaryDocxControl,
    receipt: OwnerQualityCanaryDocxRenderReceipt,
    owner_ref: str,
    inspected_at: datetime,
    inspected_page_count: int,
    all_pages_inspected_at_full_size: Literal[True],
    visual_inspection_passed: Literal[True],
    no_clipped_text: Literal[True],
    no_overlapping_objects: Literal[True],
    no_broken_tables: Literal[True],
    no_missing_glyphs: Literal[True],
    signature_algorithm: str,
    signature: str,
) -> tuple[Path, OwnerQualityCanaryDocxInspectionAttestation]:
    """Persist inspection only after a trusted owner signature verifies."""

    control = OwnerQualityCanaryDocxControl.model_validate(
        control.model_dump(mode="json", by_alias=True)
    )
    receipt = OwnerQualityCanaryDocxRenderReceipt.model_validate(
        receipt.model_dump(mode="json", by_alias=True)
    )
    if (
        receipt.docx_control_seal_sha256 != control.seal_sha256
        or inspected_page_count != receipt.page_count
    ):
        raise ValueError("DOCX owner inspection differs from the technical render")
    _verify_render_output(workspace=workspace, receipt=receipt)
    signature_payload = _owner_inspection_signature_payload(
        control=control,
        receipt=receipt,
        owner_ref=owner_ref,
        inspected_at=inspected_at,
        inspected_page_count=inspected_page_count,
        all_pages_inspected_at_full_size=all_pages_inspected_at_full_size,
        visual_inspection_passed=visual_inspection_passed,
        no_clipped_text=no_clipped_text,
        no_overlapping_objects=no_overlapping_objects,
        no_broken_tables=no_broken_tables,
        no_missing_glyphs=no_missing_glyphs,
    )
    signature_payload_sha256 = sealed_sha256(signature_payload)
    _verify_trusted_owner_docx_inspection_signature(
        owner_ref=owner_ref,
        signature_algorithm=signature_algorithm,
        signature_payload_sha256=signature_payload_sha256,
        signature=signature,
    )
    material: dict[str, Any] = {
        "schema": OWNER_QUALITY_DOCX_INSPECTION_SCHEMA,
        **{key: value for key, value in signature_payload.items() if key != "schema"},
        "signature_algorithm": signature_algorithm,
        "signature_payload_sha256": signature_payload_sha256,
        "signature": signature,
        "trusted_owner_signature_verified": True,
        "authorizes_owner_review_companion": True,
        "create_only": True,
    }
    material["seal_sha256"] = sealed_sha256(material)
    attestation = OwnerQualityCanaryDocxInspectionAttestation.model_validate(material)
    assert_safe_evaluation_payload(attestation.model_dump(mode="json", by_alias=True))
    path = workspace.write_private_bytes(
        "review-docx",
        "owner-inspection.json",
        payload=(
            json.dumps(
                attestation.model_dump(mode="json", by_alias=True),
                ensure_ascii=False,
                indent=2,
                sort_keys=True,
            )
            + "\n"
        ).encode("utf-8"),
    )
    return path, attestation


def require_trusted_owner_quality_canary_docx_inspection(
    *,
    workspace: CanaryReviewWorkspace,
    control: OwnerQualityCanaryDocxControl,
    receipt: OwnerQualityCanaryDocxRenderReceipt,
) -> OwnerQualityCanaryDocxInspectionAttestation:
    """Reload and cryptographically reverify the exact persisted inspection."""

    try:
        raw = workspace.read_private_bytes("review-docx", "owner-inspection.json")
    except FileNotFoundError as exc:
        raise OwnerDecisionRequired(
            "trusted_owner_docx_inspection_signature_verifier_missing"
        ) from exc
    attestation = OwnerQualityCanaryDocxInspectionAttestation.model_validate_json(raw)
    if (
        attestation.document_id != control.document_id
        or attestation.docx_control_seal_sha256 != control.seal_sha256
        or attestation.render_receipt_seal_sha256 != receipt.seal_sha256
        or attestation.rendered_page_set_sha256 != _rendered_page_set_sha256(receipt)
        or attestation.inspected_page_count != receipt.page_count
    ):
        raise ValueError("DOCX owner inspection is bound to different rendered bytes")
    signature_payload = _owner_inspection_signature_payload(
        control=control,
        receipt=receipt,
        owner_ref=attestation.owner_ref,
        inspected_at=attestation.inspected_at,
        inspected_page_count=attestation.inspected_page_count,
        all_pages_inspected_at_full_size=attestation.all_pages_inspected_at_full_size,
        visual_inspection_passed=attestation.visual_inspection_passed,
        no_clipped_text=attestation.no_clipped_text,
        no_overlapping_objects=attestation.no_overlapping_objects,
        no_broken_tables=attestation.no_broken_tables,
        no_missing_glyphs=attestation.no_missing_glyphs,
    )
    if sealed_sha256(signature_payload) != attestation.signature_payload_sha256:
        raise ValueError("DOCX owner inspection signature payload differs")
    _verify_render_output(workspace=workspace, receipt=receipt)
    _verify_trusted_owner_docx_inspection_signature(
        owner_ref=attestation.owner_ref,
        signature_algorithm=attestation.signature_algorithm,
        signature_payload_sha256=attestation.signature_payload_sha256,
        signature=attestation.signature,
    )
    return attestation
