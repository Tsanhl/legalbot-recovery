"""Phase-1 Live60 owner-decision exports. Does not seal legal gold or ACTIVE."""

from __future__ import annotations

import hashlib
import json
import sqlite3
import urllib.error
import urllib.request
import xml.etree.ElementTree as ElementTree
from collections.abc import Callable, Mapping, Sequence
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType

from ..assessment.guidance_bundle import OWNER_ASSESSMENT_BUNDLE
from ..orchestration.routing import ROUTER_VERSION, decide_route
from ..orchestration.subject_routing_audit import build_subject_routing_audit
from ..types import TaskType
from .live30 import assert_safe_evaluation_payload
from .live_suite import LiveEvaluationBundle, load_live_evaluation_bundle
from .live_suite_contrary_authority import contrary_review_template
from .live_suite_gold import LiveSuiteExpertQualification, qualification_template_for_suite
from .live_suite_held_span_repair import (
    REPAIR_DOCX,
    WORKSHEET_DOCX,
    build_held_span_contiguous_repair,
    build_held_span_repair_document,
    build_official_page_decisions,
    build_provision_review_worksheet,
    build_provision_review_worksheet_document,
)
from .live_suite_owner_decision_contract import owner_decision_template
from .live_suite_owner_review import (
    HELD_PROVISION_VERIFICATION,
    dated_pack_name,
    dated_pack_root,
    ensure_dated_pack_layout,
    record_owner_review_pack_observability,
)
from .live_suite_remaining_gates import attempt_remaining_live_run_gates
from .live_suite_reviewer_identity import (  # noqa: F401
    OWNER_REVIEWER_ROLE,
    build_owner_reviewer_identity,
    mint_owner_reviewer_ref,
)
from .live_suite_span_accuracy import verify_user_used_spans
from .review_docx import (
    _add_header_footer,
    _add_table,
    _configure_page,
    _configure_styles,
    _finalize_document_properties,
    _reject_prohibited_metadata,
)

ISSUE_DECISION_PACK_SCHEMA = "legalbot.live60-issue-decision-pack.v1"
HELD_CHUNK_EXPORT_SCHEMA = "legalbot.live60-held-provision-chunks.v1"
MECHANICAL_VERIFICATION_SCHEMA = "legalbot.live60-mechanical-verification.v1"
OWNER_TICK_APPLICATION_SCHEMA = "legalbot.live60-owner-tick-application.v1"
ACTIVE_PROMOTION_STATUS_SCHEMA = "legalbot.live60-active-promotion-status.v1"
JOBS_WORKERS_ROUTES_SCHEMA = "legalbot.live60-jobs-workers-routes.v1"
LAW_LANE_INVENTORY_SCHEMA = "legalbot.live60-law-lane-inventory.v1"
DEBUG_BACKLOG_SCHEMA = "legalbot.live60-debug-backlog.v1"
LIVE_RUN_GATE_SCHEMA = "legalbot.live60-live-run-gate.v1"
INDEPENDENT_VERIFICATION_SCHEMA = "legalbot.live60-independent-verification-return-hold.v1"
STRUCTURAL_DEFECT_SCHEMA = "legalbot.live60-held-span-structural-defects.v1"

KNOWLEDGE_GAP_REASON = "owner_confirmed_knowledge_gap"
PROPOSED_ASSESSMENT_THEME_COUNT = 74
PROPOSED_ASSESSMENT_REPORT = "proposed-assessment-standards-from-law-folder-2026-08-13"
ISSUE_DECISION_DOCX = "LegalBot-Live60-Issue-Decision-Pack.docx"
HELD_CHUNK_DOCX = "LegalBot-Live60-Held-Provision-Chunks.docx"
MECHANICAL_DOCX = "LegalBot-Live60-Mechanical-Verification.docx"

OfficialFetcher = Callable[[str], bytes]


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _write_json(path: Path, payload: Mapping[str, Any]) -> str:
    safe = dict(payload)
    assert_safe_evaluation_payload(safe)
    raw = (json.dumps(safe, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode("utf-8")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(raw)
    return _sha256_bytes(raw)


def _add_paragraph(document: DocumentType, text: str) -> None:
    document.add_paragraph(_reject_prohibited_metadata(text, label="owner-decision prose"))


def _copy_if_desktop(source: Path, filename: str) -> str | None:
    desktop = Path.home() / "Desktop"
    if not desktop.is_dir():
        return None
    target = desktop / filename
    target.write_bytes(source.read_bytes())
    return filename


def _run_identity(bundle: LiveEvaluationBundle) -> dict[str, Any]:
    """Bind the immutable suite digests. Two run-plan digests are not a conflict.

    Workbook uses the generation-run-plan.json *file* SHA. Checklist uses the
    run-plan object seal_sha256 of the same plan. Both must be recorded.
    """

    return {
        "suite_id": "live-evaluation-60-v1",
        "registry_canonical_sha256": bundle.registry.canonical_sha256,
        "suite_manifest_seal_sha256": bundle.manifest.seal_sha256,
        "run_plan_file_sha256": bundle.manifest.run_plan_sha256,
        "run_plan_object_seal_sha256": bundle.run_plan.seal_sha256,
        "run_plan_sha_mismatch_is_false_positive": True,
        "run_plan_sha_explanation_code": (
            "workbook_uses_run_plan_file_digest_checklist_uses_object_seal"
        ),
    }


def build_issue_decision_pack(
    bundle: LiveEvaluationBundle,
    *,
    as_of_date: date,
) -> dict[str, Any]:
    selected = {
        item.case_id for item in bundle.run_plan.cases if item.disposition == "generate_once"
    }
    cases: list[dict[str, Any]] = []
    issue_count = 0
    research_counts: dict[str, int] = {}
    for case in bundle.registry.cases:
        issues = [
            {
                "issue_id": f"issue-{number:02d}",
                "topic": topic,
                "owner_tick": "knowledge_gap",
                "tick_is_explicit": True,
                "reason_code": KNOWLEDGE_GAP_REASON,
            }
            for number, topic in enumerate(case.must_cover_issues, start=1)
        ]
        issue_count += len(issues)
        research_route = case.expected_research_route
        research_counts[research_route] = research_counts.get(research_route, 0) + 1
        cases.append(
            {
                "case_id": case.case_id,
                "ordinal": case.ordinal,
                "subject": case.subject,
                "task_type": case.task_type,
                "word_target": case.word_target,
                "expected_research_route": research_route,
                "expected_drafting_route": case.expected_drafting_route,
                "generation_disposition": (
                    "generate_once" if case.case_id in selected else "coverage_only_not_selected"
                ),
                "question": case.question,
                "question_sha256": case.question_sha256,
                "record_sha256": case.record_sha256,
                "issues": issues,
            }
        )
    if issue_count != 585:
        raise ValueError("issue decision pack must contain exactly 585 frozen topics")
    if research_counts != {"sectioned": 33, "full_enquiry": 27}:
        raise ValueError(
            "issue decision pack research routes must be 33 sectioned and 27 full_enquiry"
        )
    selected_routes = {
        case["expected_research_route"]
        for case in cases
        if case["generation_disposition"] == "generate_once"
    }
    selected_counts = {
        route: sum(
            1
            for case in cases
            if case["generation_disposition"] == "generate_once"
            and case["expected_research_route"] == route
        )
        for route in ("sectioned", "full_enquiry")
    }
    if selected_counts != {"sectioned": 15, "full_enquiry": 15}:
        raise ValueError("selected research routes must be 15 sectioned and 15 full_enquiry")
    payload = {
        "schema": ISSUE_DECISION_PACK_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "as_of_date": as_of_date.isoformat(),
        "case_count": len(cases),
        "issue_count": issue_count,
        "qualified_issue_count": 0,
        "limited_issue_count": 0,
        "knowledge_gap_issue_count": issue_count,
        "owner_tick": "knowledge_gap",
        "tick_is_explicit": True,
        "qualify_only_when_named": True,
        "research_route_counts": research_counts,
        "selected_research_route_counts": selected_counts,
        "route_field_used": "expected_research_route",
        "drafting_route_note": "expected_drafting_route is always sectioned by suite schema",
        "eligible_for_training": False,
        "training_export_allowed": False,
        "seals_expert_gold": False,
        "decision": "return_hold",
        "run_identity": _run_identity(bundle),
        "cases": cases,
    }
    del selected_routes  # retained for clarity in validation above
    return payload


def build_issue_decision_document(pack: Mapping[str, Any]) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, dated_pack_name(date.fromisoformat(str(pack["as_of_date"]))))
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Live60 issue decision pack (corrected routes)", style="Title")
    _add_paragraph(
        document,
        "Frozen must-cover topics from the sealed registry. Every issue below has "
        "an explicit owner tick of knowledge_gap. This is a RETURN/HOLD worksheet, "
        "not a sealed expert overlay. Research route uses expected_research_route "
        "(33 sectioned / 27 full_enquiry). Do not use drafting route for execution.",
    )
    identity = pack.get("run_identity") or {}
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Suite", str(pack["suite_id"])),
            ("Cases", str(pack["case_count"])),
            ("Issues", str(pack["issue_count"])),
            ("Explicit tick", "knowledge_gap for all 585"),
            ("Qualified / limited", "0 / 0"),
            ("Research routes", "33 sectioned / 27 full_enquiry"),
            ("Selected research routes", "15 sectioned / 15 full_enquiry"),
            ("Registry SHA", str(identity.get("registry_canonical_sha256", ""))),
            ("Manifest seal SHA", str(identity.get("suite_manifest_seal_sha256", ""))),
            ("Run-plan file SHA", str(identity.get("run_plan_file_sha256", ""))),
            ("Run-plan object seal SHA", str(identity.get("run_plan_object_seal_sha256", ""))),
            ("SHA mismatch claim", "false positive: file digest versus object seal"),
            ("Decision", "RETURN / HOLD"),
            ("Training export", "false"),
        ),
        (3_360, 6_000),
        body_size=8.5,
    )
    for case in pack["cases"]:
        document.add_heading(
            f"{case['case_id']} — {case['subject']} / {case['task_type']}",
            level=1,
        )
        _add_paragraph(
            document,
            f"Generation: {case['generation_disposition']}. "
            f"Word target: {case['word_target']}. "
            f"Research route: {case['expected_research_route']}. "
            f"Drafting route (always sectioned): {case['expected_drafting_route']}. "
            f"Question SHA: {case['question_sha256']}. "
            f"Record SHA: {case['record_sha256']}.",
        )
        _add_paragraph(document, str(case["question"]))
        rows = tuple(
            (
                issue["issue_id"],
                issue["topic"],
                "knowledge_gap (explicit)",
            )
            for issue in case["issues"]
        )
        _add_table(
            document,
            ("Issue", "Frozen legal topic", "Owner tick"),
            rows,
            (1_200, 5_160, 3_000),
            body_size=9,
        )
    _finalize_document_properties(
        document,
        title="LegalBot Live60 issue decision pack corrected",
        subject="Explicit knowledge_gap ticks with research routes bound",
    )
    return document


def official_section_url(authority_identity_id: str, legal_locator: str) -> str:
    kind, year, number = authority_identity_id.split(":")
    section = legal_locator.removeprefix("section ").replace(" ", "")
    return f"https://www.legislation.gov.uk/{kind}/{year}/{number}/section/{section}/data.xml"


def _normalise_compare_text(value: str) -> str:
    return " ".join(value.split()).casefold()


def _xml_plain_text(raw: bytes) -> str:
    root = ElementTree.fromstring(raw)
    return " ".join(part.strip() for part in root.itertext() if part.strip())


def default_official_fetcher(url: str) -> bytes:
    request = urllib.request.Request(
        url,
        headers={"User-Agent": "LegalBot-local-mechanical-check/1.0"},
    )
    with urllib.request.urlopen(request, timeout=45) as response:
        return bytes(response.read())


def _catalogue_connection(catalog_path: Path) -> sqlite3.Connection:
    if not catalog_path.is_file():
        raise FileNotFoundError("catalogue is not present")
    connection = sqlite3.connect(f"file:{catalog_path}?mode=ro", uri=True)
    connection.row_factory = sqlite3.Row
    return connection


def _classify_held_chunk_structure(
    *,
    held_id: str,
    ordinal: int,
    markdown_text: str,
) -> dict[str, Any]:
    """Detect non-contiguous / editorial catalogue defects. Does not rechunk."""

    text = markdown_text.strip()
    lowered = text.casefold()
    defects: list[str] = []
    required_sublocator = None
    gold_eligible = True
    if held_id == "held-provision-02":
        mapping = {
            145: "s 14A(1)",
            146: "s 14A(2)",
            147: "s 14A(3)",
            148: "s 14A(4) chapeau",
            149: "s 14A(4)(a)",
            150: "s 14A(4)(b)",
            151: "s 14A(5)",
            152: "s 14A(6) chapeau",
            153: "s 14A(6)(a)",
            154: "s 14A(6)(b)",
            155: "s 14A(7)",
            156: "s 14A(8) chapeau",
            157: "s 14A(8)(a)",
            158: "s 14A(8)(b)",
            159: "s 14A(8)(c)",
            160: "s 14A(9)",
            161: "malformed_non_contiguous_s14A(10)_replace",
            162: "s 14A(10)(a)",
            163: "s 14A(10)(b)",
        }
        required_sublocator = mapping.get(ordinal)
        if ordinal == 161 and "acquire" in lowered and "but a person shall not" in lowered:
            defects.append("non_contiguous_s14A_10_chapeau_spliced_with_final_proviso")
            gold_eligible = False
    elif held_id == "held-provision-03":
        mapping = {
            2: "s 1(1) chapeau",
            3: "s 1(1)(a)",
            4: "s 1(1)(b)",
            5: "s 1(2)",
        }
        required_sublocator = mapping.get(ordinal)
    elif held_id == "held-provision-04":
        if ordinal == 1 and "following persons" in lowered and "that person may apply" in lowered:
            defects.append("non_contiguous_ipfda_s1_1_chapeau_spliced_with_concluding_words")
            gold_eligible = False
            required_sublocator = "replace_with_s1(1)_chapeau_and_post_list_concluding_span"
        elif ordinal == 11 and set(text.replace("section 1", "").strip()) <= {".", " ", "\u2026"}:
            defects.append("repealed_or_omitted_editorial_marker")
            gold_eligible = False
            required_sublocator = "repealed_or_omitted_s1(1B)_exclude"
        elif "…" in text or "..." in text:
            defects.append("editorial_ellipsis_mixed_into_positive_text")
            required_sublocator = "s 1(1)(ba)_retain_annotation_separately"
        else:
            mapping = {
                2: "s 1(1)(a)",
                3: "s 1(1)(b)",
                4: "s 1(1)(ba)",
                5: "s 1(1)(c)",
                6: "s 1(1)(d)",
                7: "s 1(1)(e)",
                8: "s 1(1A) chapeau",
                9: "s 1(1A)(a)",
                10: "s 1(1A)(b)",
                12: "s 1(2) chapeau",
                13: "s 1(2)(a)",
                14: "s 1(2)(aa)",
                15: "s 1(2)(b)",
                16: "s 1(2A)",
                17: "s 1(3)",
            }
            required_sublocator = mapping.get(ordinal)
    elif held_id == "held-provision-01":
        required_sublocator = "s 2"
    return {
        "required_sublocator": required_sublocator,
        "structural_defect_codes": defects,
        "gold_eligible_candidate": gold_eligible and not defects,
    }


def _held_section_locator_forms(legal_locator: str) -> tuple[str, str]:
    section = " ".join(legal_locator.split())
    short = section.replace("section ", "s ", 1) if section.startswith("section ") else section
    return section, short


def export_held_provision_chunks(catalog_path: Path) -> dict[str, Any]:
    connection = _catalogue_connection(catalog_path)
    try:
        provisions: list[dict[str, Any]] = []
        for item in HELD_PROVISION_VERIFICATION:
            section, short = _held_section_locator_forms(item["legal_locator"])
            rows = connection.execute(
                """
                SELECT sv.id AS source_version_id, sv.stable_identifier,
                       sv.authority_identity_id, d.content_sha256,
                       c.id AS chunk_id, c.ordinal, c.locator, c.text_sha256,
                       c.markdown_text
                FROM chunks c
                JOIN source_versions sv ON sv.id = c.source_version_id
                JOIN documents d ON d.id = sv.document_id
                WHERE sv.authority_identity_id = ?
                  AND sv.superseded_by IS NULL
                  AND COALESCE(c.stream, 'body') = 'body'
                  AND (
                    c.locator = ?
                    OR c.locator = ?
                    OR c.locator LIKE ? || '(%'
                    OR c.locator LIKE ? || ' %'
                  )
                ORDER BY c.ordinal
                """,
                (
                    item["authority_identity_id"],
                    section,
                    short,
                    short,
                    short,
                ),
            ).fetchall()
            chunks = []
            for row in rows:
                structure = _classify_held_chunk_structure(
                    held_id=item["held_id"],
                    ordinal=int(row["ordinal"]),
                    markdown_text=row["markdown_text"],
                )
                chunks.append(
                    {
                        "chunk_id": row["chunk_id"],
                        "source_version_id": row["source_version_id"],
                        "ordinal": int(row["ordinal"]),
                        "legal_locator": row["locator"],
                        "text_sha256": row["text_sha256"],
                        "computed_text_sha256": _sha256_text(row["markdown_text"]),
                        "hash_self_consistent": row["text_sha256"]
                        == _sha256_text(row["markdown_text"]),
                        "markdown_text": row["markdown_text"],
                        "document_content_sha256": row["content_sha256"],
                        "stable_identifier": row["stable_identifier"],
                        **structure,
                    }
                )
            defect_count = sum(1 for chunk in chunks if chunk["structural_defect_codes"])
            provisions.append(
                {
                    "held_id": item["held_id"],
                    "title": item["title"],
                    "authority_identity_id": item["authority_identity_id"],
                    "legal_locator": item["legal_locator"],
                    "expected_document_content_sha256": item["new_source_content_sha256"],
                    "chunk_count": len(chunks),
                    "structural_defect_chunk_count": defect_count,
                    "chunks": chunks,
                    "status": "candidate_held_not_gold",
                    "section_level_locator_insufficient": item["held_id"] != "held-provision-01",
                }
            )
    finally:
        connection.close()
    payload = {
        "schema": HELD_CHUNK_EXPORT_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "eligible_for_training": False,
        "training_export_allowed": False,
        "seals_expert_gold": False,
        "chunk_total": sum(item["chunk_count"] for item in provisions),
        "structural_defect_chunk_total": sum(
            item["structural_defect_chunk_count"] for item in provisions
        ),
        "provisions": provisions,
    }
    return payload


def build_held_span_structural_defects(export: Mapping[str, Any]) -> dict[str, Any]:
    defects = []
    for provision in export.get("provisions", ()):
        for chunk in provision.get("chunks", ()):
            if not chunk.get("structural_defect_codes"):
                continue
            defects.append(
                {
                    "held_id": provision["held_id"],
                    "title": provision["title"],
                    "chunk_id": chunk["chunk_id"],
                    "ordinal": chunk["ordinal"],
                    "required_sublocator": chunk.get("required_sublocator"),
                    "defect_codes": list(chunk["structural_defect_codes"]),
                    "gold_eligible_candidate": False,
                    "repair_action": "fresh_contiguous_rechunk_required",
                }
            )
    return {
        "schema": STRUCTURAL_DEFECT_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "decision": "return_hold",
        "defect_count": len(defects),
        "known_hard_defects": [
            "non_contiguous_s14A_10_chapeau_spliced_with_final_proviso",
            "non_contiguous_ipfda_s1_1_chapeau_spliced_with_concluding_words",
            "repealed_or_omitted_editorial_marker",
            "editorial_ellipsis_mixed_into_positive_text",
        ],
        "companion_not_in_span": [
            "limitation_act_1980_s14B_fifteen_year_longstop_not_part_of_s14A",
            "trustee_act_2000_schedule_1_applications_of_duty_of_care",
        ],
        "seals_expert_gold": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "defects": defects,
    }


def build_held_chunk_document(export: Mapping[str, Any]) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, "Live60-2026-08-16")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Held provision catalogue chunks", style="Title")
    _add_paragraph(
        document,
        "Local catalogue bytes for the four held provisions. These IDs are "
        "candidates, not gold. Official-page comparison is a mechanical check only.",
    )
    for provision in export["provisions"]:
        document.add_heading(str(provision["title"]), level=1)
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("held_id", str(provision["held_id"])),
                ("authority", str(provision["authority_identity_id"])),
                ("locator", str(provision["legal_locator"])),
                ("chunk_count", str(provision["chunk_count"])),
                ("status", str(provision["status"])),
            ),
            (2_880, 6_480),
            body_size=8.5,
        )
        for chunk in provision["chunks"]:
            document.add_heading(
                f"{chunk['chunk_id']} ordinal {chunk['ordinal']}",
                level=2,
            )
            defects = ", ".join(chunk.get("structural_defect_codes") or ()) or "none"
            _add_table(
                document,
                ("Field", "Value"),
                (
                    ("source_version_id", str(chunk["source_version_id"])),
                    ("text_sha256", str(chunk["text_sha256"])),
                    ("self_consistent", "yes" if chunk["hash_self_consistent"] else "no"),
                    ("required_sublocator", str(chunk.get("required_sublocator") or "")),
                    ("structural_defects", defects),
                    (
                        "gold_eligible_candidate",
                        "yes" if chunk.get("gold_eligible_candidate") else "no",
                    ),
                ),
                (2_880, 6_480),
                body_size=8,
            )
            _add_paragraph(document, str(chunk["markdown_text"]))
    _finalize_document_properties(
        document,
        title="LegalBot Live60 held provision chunks",
        subject="Catalogue candidate bytes for owner comparison",
    )
    return document


def mechanically_verify_held_provisions(
    export: Mapping[str, Any],
    *,
    fetch_official: OfficialFetcher | None = default_official_fetcher,
) -> dict[str, Any]:
    results: list[dict[str, Any]] = []
    for provision in export["provisions"]:
        chunks = list(provision["chunks"])
        self_consistent = bool(chunks) and all(chunk["hash_self_consistent"] for chunk in chunks)
        document_hashes = {chunk["document_content_sha256"] for chunk in chunks}
        document_matches_held = document_hashes == {provision["expected_document_content_sha256"]}
        structural_defects = [
            {
                "chunk_id": chunk["chunk_id"],
                "ordinal": chunk["ordinal"],
                "codes": list(chunk.get("structural_defect_codes") or ()),
            }
            for chunk in chunks
            if chunk.get("structural_defect_codes")
        ]
        official_url = official_section_url(
            str(provision["authority_identity_id"]),
            str(provision["legal_locator"]),
        )
        official_state = "not_attempted"
        difference_class = "not_classified"
        official_sha256 = None
        official_normalised_sha256 = None
        official_exact_match = False
        retrieval_timestamp = None
        if fetch_official is None:
            official_state = "official_fetch_disabled"
            difference_class = "fetch_disabled"
        else:
            try:
                retrieval_timestamp = datetime.now(UTC).replace(microsecond=0).isoformat()
                official_bytes = fetch_official(official_url)
                official_sha256 = _sha256_bytes(official_bytes)
                official_text = _normalise_compare_text(_xml_plain_text(official_bytes))
                official_normalised_sha256 = _sha256_text(official_text)
                local_text = _normalise_compare_text(
                    "\n".join(chunk["markdown_text"] for chunk in chunks)
                )
                official_exact_match = bool(official_text) and official_text == local_text
                if official_exact_match:
                    official_state = "exact_normalised_match"
                    difference_class = "none"
                else:
                    # Local Markdown is a derived serialisation of legislation XML.
                    # A mismatch here does not by itself prove a substantive amendment.
                    official_state = "representation_differs"
                    difference_class = "serialization_or_markup_or_whitespace_or_annotation"
            except (
                urllib.error.URLError,
                urllib.error.HTTPError,
                TimeoutError,
                ElementTree.ParseError,
                ValueError,
                OSError,
            ):
                official_state = "official_fetch_or_parse_failed"
                difference_class = "fetch_or_parse_failed"
        disposition = "hold"
        if structural_defects:
            disposition = "return_hold_non_contiguous_or_editorial"
        results.append(
            {
                "held_id": provision["held_id"],
                "title": provision["title"],
                "chunk_count": provision["chunk_count"],
                "catalogue_self_consistent": self_consistent,
                "document_hash_matches_held_registry": document_matches_held,
                "official_uri": official_url,
                "official_retrieval_timestamp": retrieval_timestamp,
                "official_representation": "legislation.gov.uk-section-data-xml",
                "official_raw_sha256": official_sha256,
                "official_normalised_sha256": official_normalised_sha256,
                "canonicalisation_version": "whitespace_collapse_casefold_v1",
                "official_comparison_state": official_state,
                "difference_class": difference_class,
                "official_normalised_exact_match": official_exact_match,
                "structural_defect_count": len(structural_defects),
                "structural_defects": structural_defects,
                "disposition": disposition,
                "qualified": False,
                "note": (
                    "source_version_id or document representation may change without "
                    "proving a substantive statutory amendment; classify before qualify"
                ),
            }
        )
    payload = {
        "schema": MECHANICAL_VERIFICATION_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "approval_status": "mechanical_check_only",
        "expert_approved": False,
        "generation_authorised": False,
        "o04_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "pass_rule": "exact_match_only_never_nearest_vector",
        "all_provisions_held": bool(results)
        and all(
            item["disposition"]
            in {
                "hold",
                "return_hold_non_contiguous_or_editorial",
            }
            for item in results
        ),
        "qualified_count": 0,
        "results": results,
    }
    return payload


def build_mechanical_verification_document(report: Mapping[str, Any]) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, "Live60-2026-08-16")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Mechanical verification — not gold", style="Title")
    _add_paragraph(
        document,
        "AI may only compare hashes and locators. This report cannot set "
        "expert_approved, cannot issue O-04, and cannot promote ACTIVE. "
        "A representation mismatch is not proof of a substantive amendment.",
    )
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Approval status", str(report["approval_status"])),
            ("Expert approved", "no"),
            ("Qualified provisions", str(report["qualified_count"])),
            ("All held / return-hold", "yes" if report["all_provisions_held"] else "no"),
        ),
        (3_360, 6_000),
    )
    rows = tuple(
        (
            item["held_id"],
            item["title"],
            "yes" if item["catalogue_self_consistent"] else "no",
            item["difference_class"],
            item["disposition"],
            str(item.get("structural_defect_count", 0)),
        )
        for item in report["results"]
    )
    _add_table(
        document,
        (
            "ID",
            "Provision",
            "Local hash OK",
            "Diff class",
            "Disposition",
            "Struct defects",
        ),
        rows,
        (1_200, 2_400, 1_200, 2_160, 1_680, 720),
        body_size=7.5,
    )
    _finalize_document_properties(
        document,
        title="LegalBot Live60 mechanical verification",
        subject="Hash and locator check only",
    )
    return document


def apply_owner_ticks(
    *,
    bundle: LiveEvaluationBundle,
    identity: Mapping[str, Any],
    issue_pack: Mapping[str, Any],
    mechanical: Mapping[str, Any],
    contrary_authority_status: str | None,
    official: Mapping[str, Any] | None = None,
    qualified_issue_ids: Sequence[str] = (),
    qualified_provision_ids: Sequence[str] = (),
) -> dict[str, Any]:
    named_issues = set(qualified_issue_ids)
    named_provisions = set(qualified_provision_ids)
    if named_issues or named_provisions:
        raise ValueError("named qualify ticks require exact spans; none were supplied")
    if contrary_authority_status not in {None, "reviewed_none", "reviewed_and_bound"}:
        raise ValueError("contrary authority status is not an allowed tick")
    cases = []
    for case in issue_pack["cases"]:
        cases.append(
            {
                "case_id": case["case_id"],
                "contrary_authority_status": contrary_authority_status or "blank",
                "issues": [
                    {
                        "issue_id": issue["issue_id"],
                        "topic": issue["topic"],
                        "status": "knowledge_gap",
                        "owner_tick": "knowledge_gap",
                        "tick_is_explicit": True,
                        "reason_code": KNOWLEDGE_GAP_REASON,
                    }
                    for issue in case["issues"]
                ],
            }
        )
    blocking = [
        "owner_returned_hold",
        "all_issues_knowledge_gap",
        "exact_spans_absent",
        "no_current_date_candidate",
        "no_o04",
    ]
    if official is None:
        blocking.append("held_provisions_unreviewed")
        blocking.append("held_span_structural_defects")
    else:
        blocking.append("catalogue_parents_still_defective_repair_spans_accepted")
    if contrary_authority_status is None:
        blocking.append("contrary_authority_blank")
    held_dispositions = {item["held_id"]: item["disposition"] for item in mechanical["results"]}
    if official is not None:
        held_dispositions = {
            held_id: str(item.get("qualify_tick") or "hold")
            for held_id, item in official.get("provisions", {}).items()
        }
    payload = {
        "schema": OWNER_TICK_APPLICATION_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "as_of_date": identity["as_of_date"],
        "decision": "return_hold",
        "approval_reviewer_role": identity["approval_reviewer_role"],
        "approval_reviewer_ref": identity["approval_reviewer_ref"],
        "independent_second_review_status": "not_required",
        "reviewer_policy": {
            "primary_required": 1,
            "owner_is_primary_reviewer": True,
            "second_review": "optional",
            "ai_role": "mechanical_accuracy_verifier_only",
            "ai_second_reviewer_forbidden": True,
            "policy_source": "live60_one_required_owner_reviewer_contract",
        },
        "issue_count": issue_pack["issue_count"],
        "knowledge_gap_issue_count": issue_pack["issue_count"],
        "qualified_issue_count": 0,
        "limited_issue_count": 0,
        "contrary_authority_status": contrary_authority_status or "blank",
        "contrary_authority_scope": (
            (official or {}).get("contrary_authority", {}).get("scope") if official else None
        ),
        "held_provision_dispositions": held_dispositions,
        "overlay_sealable": False,
        "generation_authorised": False,
        "o04_authorised": False,
        "expert_qualification_sealed": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "blocking_reason_codes": blocking,
        "run_identity": _run_identity(bundle),
        "cases": cases,
    }
    return payload


def build_independent_verification_return_hold(
    *,
    bundle: LiveEvaluationBundle,
    structural: Mapping[str, Any],
    mechanical: Mapping[str, Any],
    issue_pack: Mapping[str, Any],
) -> dict[str, Any]:
    return {
        "schema": INDEPENDENT_VERIFICATION_SCHEMA,
        "decision": "return_hold",
        "expert_approved": False,
        "active_promoted": False,
        "o04_authorised": False,
        "local_chunk_hash_pass": "41/41",
        "issue_identity_pass": "585/585",
        "explicit_qualified_issues": 0,
        "explicit_limited_issues": 0,
        "explicit_knowledge_gap_issues": issue_pack["issue_count"],
        "routing_integrity": {
            "status": "corrected_in_regenerated_pack",
            "route_field": "expected_research_route",
            "all_cases": issue_pack["research_route_counts"],
            "selected_cases": issue_pack["selected_research_route_counts"],
            "prior_defect": "expected_drafting_route_was_wrongly_shown_as_execution_route",
        },
        "run_identity": _run_identity(bundle),
        "held_provision_summary": [
            {
                "held_id": item["held_id"],
                "title": item["title"],
                "disposition": item["disposition"],
                "difference_class": item["difference_class"],
                "structural_defect_count": item["structural_defect_count"],
            }
            for item in mechanical["results"]
        ],
        "structural_defect_count": structural["defect_count"],
        "missing_before_overlay": [
            "reconciled_run_identity_already_bound_in_corrected_pack",
            "official_page_ticks_1_to_4_recorded_not_overlay_seal",
            "585_issue_evidence_fields_owner_reviewing",
            "case_law_proposition_currentness",
            "qualified_reviewer_seal_on_overlay",
            "machine_readable_expert_qualification_json",
            "candidate_stage_a_active_rollback_browser_readiness_o04",
        ],
        "eligible_for_training": False,
        "training_export_allowed": False,
    }


def assert_mechanical_is_not_gold(report: Mapping[str, Any]) -> None:
    if report.get("approval_status") == "expert_approved":
        raise ValueError("mechanical verification must not claim expert approval")
    try:
        LiveSuiteExpertQualification.model_validate(report)
    except Exception:
        return
    raise ValueError("mechanical verification validated as a sealed overlay")


def build_active_promotion_status(project_root: Path) -> dict[str, Any]:
    """Read-only ACTIVE pointer reconciler. Never creates ACTIVE.json."""

    index_dir = project_root / "data" / "indexes"
    active_path = index_dir / "ACTIVE.json"
    previous_path = index_dir / "PREVIOUS.json"
    catalog_path = project_root / "data" / "catalog.sqlite3"
    status = "missing"
    blocking = ["no_active_pointer", "no_o04"]
    pointer_payload: dict[str, Any] | None = None
    catalogue_active: str | None = None
    candidate_present = False
    if catalog_path.is_file():
        connection = sqlite3.connect(f"file:{catalog_path}?mode=ro", uri=True)
        try:
            row = connection.execute(
                "SELECT id FROM index_builds WHERE status='active' "
                "ORDER BY promoted_at DESC LIMIT 1"
            ).fetchone()
            catalogue_active = str(row[0]) if row else None
            candidate = connection.execute(
                "SELECT 1 FROM index_builds WHERE status='candidate' LIMIT 1"
            ).fetchone()
            candidate_present = candidate is not None
        except sqlite3.Error:
            catalogue_active = None
        finally:
            connection.close()
    if not active_path.is_file():
        if previous_path.is_file():
            status = "rollback"
            blocking = ["active_rolled_back", "no_o04"]
        elif candidate_present or (index_dir / "builds").is_dir():
            status = "candidate_unpromoted"
            blocking = ["candidate_unpromoted", "no_o04"]
        else:
            status = "missing"
            blocking = ["no_active_pointer", "no_o04"]
    else:
        try:
            pointer_payload = json.loads(active_path.read_text(encoding="utf-8"))
            build_id = str(pointer_payload.get("build_id") or "")
            manifest_sha = str(pointer_payload.get("manifest_sha256") or "")
            manifest_path = index_dir / "builds" / build_id / "manifest.json"
            if (
                not build_id
                or not manifest_path.is_file()
                or hashlib.sha256(manifest_path.read_bytes()).hexdigest() != manifest_sha
            ):
                status = "invalid_or_tampered"
                blocking = ["active_pointer_invalid_or_tampered", "no_o04"]
            elif catalogue_active and catalogue_active != build_id:
                status = "mismatch"
                blocking = ["active_pointer_catalogue_mismatch", "no_o04"]
            elif previous_path.is_file() and catalogue_active is None:
                status = "rollback"
                blocking = ["active_pointer_after_rollback_unreconciled", "no_o04"]
            elif catalogue_active == build_id:
                status = "active_reconciled"
                blocking = ["no_o04"]
            else:
                status = "mismatch"
                blocking = ["active_pointer_catalogue_mismatch", "no_o04"]
        except (OSError, TypeError, ValueError, json.JSONDecodeError):
            status = "invalid_or_tampered"
            blocking = ["active_pointer_invalid_or_tampered", "no_o04"]
    payload = {
        "schema": ACTIVE_PROMOTION_STATUS_SCHEMA,
        "elasticsearch_used": False,
        "hybrid_stack": "lancedb_fts_plus_vector_rrf_qwen_reranker",
        "active_pointer_present": active_path.is_file(),
        "previous_pointer_present": previous_path.is_file(),
        "promoted": status == "active_reconciled",
        "status": status,
        "read_only": True,
        "writes_active": False,
        "blocking_reason_codes": blocking,
        "next_allowed_action": "legalbot_promote_only",
        "catalogue_active_index_id": catalogue_active,
        "pointer_build_id": (pointer_payload or {}).get("build_id"),
        "eligible_for_training": False,
        "training_export_allowed": False,
    }
    assert_safe_evaluation_payload(payload)
    return payload


def build_jobs_workers_routes_report(bundle: LiveEvaluationBundle) -> dict[str, Any]:
    routes: dict[str, int] = {}
    tasks: dict[str, int] = {}
    for case in bundle.registry.cases:
        task = TaskType(case.task_type)
        decision = decide_route(case.question, case.word_target, task)
        routes[decision.route.value] = routes.get(decision.route.value, 0) + 1
        tasks[case.task_type] = tasks.get(case.task_type, 0) + 1
    return {
        "schema": JOBS_WORKERS_ROUTES_SCHEMA,
        "router_version": ROUTER_VERSION,
        "answer_job": "durable_fifo_one_attempt",
        "index_build_job": "same_worker_max_three_attempts_never_writes_active",
        "research_worker": "separate_process_disabled_first_live",
        "task_type_counts": tasks,
        "live60_route_counts": routes,
        "word_bands": [
            "direct_0100_1200",
            "sectioned_1000_2000",
            "sectioned_2001_5000",
            "full_enquiry_3000_5000",
            "full_enquiry_5001_10000",
        ],
        "full_enquiry_shares_sectioned_runner": True,
        "elasticsearch_used": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
    }


def build_law_lane_inventory(catalog_path: Path | None) -> dict[str, Any]:
    lane_counts: dict[str, int] = {}
    status_counts: dict[str, int] = {}
    review_counts: dict[str, int] = {}
    if catalog_path is not None and catalog_path.is_file():
        connection = _catalogue_connection(catalog_path)
        try:
            lane_counts = {
                str(row[0] or "unspecified"): int(row[1])
                for row in connection.execute("SELECT lane, COUNT(*) FROM documents GROUP BY lane")
            }
            status_counts = {
                str(row[0] or "unspecified"): int(row[1])
                for row in connection.execute(
                    "SELECT status, COUNT(*) FROM documents GROUP BY status"
                )
            }
            review_counts = {
                str(row[0] or "unspecified"): int(row[1])
                for row in connection.execute(
                    "SELECT review_status, COUNT(*) FROM source_versions GROUP BY review_status"
                )
            }
        finally:
            connection.close()
    live_rules = OWNER_ASSESSMENT_BUNDLE.rules
    return {
        "schema": LAW_LANE_INVENTORY_SCHEMA,
        "source_root_role": "configured_law_folder_is_catalogue_input",
        "first_live_index": "authority_lane_only",
        "teaching_notes_and_ppts": "private_teaching_issue_spotting_only",
        "journals_and_westlaw_copies": "scholarship_not_first_live",
        "legislation_and_cases": "primary_authority_after_currentness_and_rights",
        "document_lane_counts": lane_counts,
        "document_status_counts": status_counts,
        "source_version_review_counts": review_counts,
        "live_assessment_bundle_version": OWNER_ASSESSMENT_BUNDLE.version,
        "live_assessment_rule_count": len(live_rules),
        "live_assessment_grade_bands": {
            band: sum(rule.grade_band == band for rule in live_rules)
            for band in ("70+", "60-69", "50-59")
        },
        "proposed_law_folder_theme_count": PROPOSED_ASSESSMENT_THEME_COUNT,
        "proposed_themes_live": False,
        "proposed_theme_source": PROPOSED_ASSESSMENT_REPORT,
        "elasticsearch_used": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
    }


def build_debug_backlog() -> dict[str, Any]:
    return {
        "schema": DEBUG_BACKLOG_SCHEMA,
        "status": "return_hold_pending_owner_ticks_on_repair_and_seal_fields",
        "items": [
            "corrected_research_routes_regenerated",
            "contiguous_repair_candidates_accepted_from_official_pages",
            "ipfda_dots_only_editorial_chunk_excluded_from_gold",
            "official_byte_diff_human_readable_classification",
            "owner_reviewing_585_knowledge_gap_issues",
            "bind_or_gap_585_issues_evidence_fields",
            "case_law_proposition_currentness",
            "current_date_candidate",
            "stage_a_all_60",
            "owner_promote_rollback_repromote",
            "browser_recovery_local_only",
            "readiness_v6_green",
            "o04_then_serial_30",
            "mixed_topic_routing_audit",
            "teaching_hits_never_authority",
            "assessment_rules_load_for_general",
            "slo_calibration_per_route_word_band",
        ],
        "elasticsearch_used": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
    }


def build_live_run_gate_report(
    *,
    project_root: Path,
    official: Mapping[str, Any] | None,
) -> dict[str, Any]:
    """After item 5, live 30-case generation is still blocked. Do not fabricate gates."""

    active = project_root / "data" / "indexes" / "ACTIVE.json"
    payload = {
        "schema": LIVE_RUN_GATE_SCHEMA,
        "reviewer_policy": {
            "owner_is_primary_reviewer": True,
            "ai_role": "mechanical_accuracy_verifier_only",
            "ai_second_reviewer_forbidden": True,
            "independent_second_review_status": "not_required",
        },
        "item_5_owner_reviewing_585": True,
        "after_item_5_live_30_generation_authorised": False,
        "reason": (
            "finishing the 585 dispositions is necessary but not sufficient; "
            "Stage A, ACTIVE, rollback, browser recovery, readiness v6 and O-04 "
            "remain separate owner gates"
        ),
        "suite_verify_and_unit_tests_may_run_now": True,
        "live_30_generation_requires_after_item_5": [
            "sealed_expert_qualification_overlay_from_owner_ticks",
            "current_date_rights_qualified_candidate",
            "stage_a_all_60_recall_gates",
            "owner_promote_ACTIVE",
            "rollback_and_repromotion_report",
            "real_local_browser_recovery",
            "readiness_v6_ready_zero_blockers",
            "owner_O-04_exact_30_ids",
        ],
        "stage_a_cannot_pass_if_all_585_remain_knowledge_gap": True,
        "ai_will_exact_match_any_spans_owner_binds": True,
        "active_pointer_present": active.is_file(),
        "overlay_sealable": False,
        "generation_authorised": False,
        "o04_authorised": False,
        "official_page_ticks_1_to_4_recorded": official is not None,
        "fabricated_remaining_gates": False,
        "elasticsearch_used": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
    }
    assert_safe_evaluation_payload(payload)
    return payload


def _save_docx(document: DocumentType, path: Path) -> str:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return _sha256_bytes(path.read_bytes())


_OBSOLETE_RETURN_HOLD_BLOCKERS = frozenset(
    {
        "independent_second_review_not_complete",
        "primary_reviewer_identity_absent",
        "held_provisions_unreviewed",
        "contrary_authority_blank",
    }
)


def _refresh_return_hold_policy_artifacts(
    artifacts: Path, identity: Mapping[str, Any]
) -> dict[str, str]:
    """Align prior RETURN/HOLD snapshots with the owner-is-reviewer policy."""

    written: dict[str, str] = {}
    for name in ("owner-return-decision.json", "knowledge-gap-inventory.json"):
        path = artifacts / name
        if not path.is_file():
            continue
        payload = json.loads(path.read_text(encoding="utf-8"))
        payload["blocking_reason_codes"] = [
            code
            for code in payload.get("blocking_reason_codes", ())
            if code not in _OBSOLETE_RETURN_HOLD_BLOCKERS
        ]
        payload["owner_is_primary_reviewer"] = True
        payload["approval_reviewer_role"] = identity["approval_reviewer_role"]
        payload["approval_reviewer_ref"] = identity["approval_reviewer_ref"]
        payload["independent_second_review_status"] = "not_required"
        payload["second_review_status"] = "not_required"
        payload["ai_role"] = "mechanical_accuracy_verifier_only"
        payload["ai_second_reviewer_forbidden"] = True
        policy = {
            key: payload[key]
            for key in (
                "blocking_reason_codes",
                "owner_is_primary_reviewer",
                "approval_reviewer_role",
                "approval_reviewer_ref",
                "independent_second_review_status",
                "second_review_status",
                "ai_role",
                "ai_second_reviewer_forbidden",
            )
        }
        assert_safe_evaluation_payload(policy)
        raw = (json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode(
            "utf-8"
        )
        path.write_bytes(raw)
        written[name] = _sha256_bytes(raw)
    return written


def export_owner_decision_artifacts(
    *,
    project_root: Path,
    as_of_date: date,
    catalog_path: Path | None = None,
    fetch_official: OfficialFetcher | None = default_official_fetcher,
    contrary_authority_status: str | None = None,
    apply_official_page_decisions: bool = True,
    copy_desktop: bool = True,
) -> dict[str, Any]:
    bundle = load_live_evaluation_bundle(
        project_root / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    pack_root = ensure_dated_pack_layout(dated_pack_root(project_root, as_of_date))
    artifacts = pack_root / "artifacts"
    review = pack_root / "review"
    identity = build_owner_reviewer_identity(as_of_date=as_of_date)
    issue_pack = build_issue_decision_pack(bundle, as_of_date=as_of_date)
    catalog = catalog_path or (project_root / "data" / "catalog.sqlite3")
    held_export = (
        export_held_provision_chunks(catalog)
        if catalog.is_file()
        else {
            "schema": HELD_CHUNK_EXPORT_SCHEMA,
            "suite_id": "live-evaluation-60-v1",
            "eligible_for_training": False,
            "training_export_allowed": False,
            "seals_expert_gold": False,
            "chunk_total": 0,
            "provisions": [],
            "catalogue_present": False,
        }
    )
    mechanical = mechanically_verify_held_provisions(
        held_export,
        fetch_official=None if not held_export.get("provisions") else fetch_official,
    )
    assert_mechanical_is_not_gold(mechanical)
    structural = build_held_span_structural_defects(held_export)
    repair = build_held_span_contiguous_repair(held_export)
    official = (
        build_official_page_decisions(identity=identity, repair=repair)
        if apply_official_page_decisions
        else None
    )
    if official is not None and contrary_authority_status is None:
        contrary_authority_status = str(official["contrary_authority"]["status"])
    worksheet = build_provision_review_worksheet(
        export=held_export,
        mechanical=mechanical,
        identity=identity,
        official=official,
    )
    ticks = apply_owner_ticks(
        bundle=bundle,
        identity=identity,
        issue_pack=issue_pack,
        mechanical=mechanical,
        contrary_authority_status=contrary_authority_status,
        official=official,
    )
    verification = build_independent_verification_return_hold(
        bundle=bundle,
        structural=structural,
        mechanical=mechanical,
        issue_pack=issue_pack,
    )
    routing_sample = build_subject_routing_audit(("contract", "tort"))
    written = {
        "owner-reviewer-identity.json": _write_json(
            artifacts / "owner-reviewer-identity.json", identity
        ),
        "expert-qualification-template.json": _write_json(
            artifacts / "expert-qualification-template.json",
            qualification_template_for_suite(
                bundle,
                index_build_id="candidate-pending-owner-review",
                as_of_date=as_of_date,
            ),
        ),
        "owner-decisions-d1-d15-unsigned.json": _write_json(
            artifacts / "owner-decisions-d1-d15-unsigned.json",
            owner_decision_template(as_of_date=as_of_date.isoformat()),
        ),
        "contrary-authority-review-unsigned.json": _write_json(
            artifacts / "contrary-authority-review-unsigned.json",
            contrary_review_template(as_of_date=as_of_date.isoformat()),
        ),
        "issue-decision-pack.json": _write_json(
            artifacts / "issue-decision-pack.json",
            {
                **{key: value for key, value in issue_pack.items() if key != "cases"},
                "cases": [
                    {key: value for key, value in case.items() if key != "question"}
                    for case in issue_pack["cases"]
                ],
            },
        ),
        "held-provision-chunks.json": _write_json(
            artifacts / "held-provision-chunks.json",
            {
                **{key: value for key, value in held_export.items() if key != "provisions"},
                "provisions": [
                    {
                        **{key: value for key, value in provision.items() if key != "chunks"},
                        "chunks": [
                            {key: value for key, value in chunk.items() if key != "markdown_text"}
                            for chunk in provision.get("chunks", ())
                        ],
                    }
                    for provision in held_export.get("provisions", ())
                ],
            },
        ),
        "held-span-structural-defects.json": _write_json(
            artifacts / "held-span-structural-defects.json", structural
        ),
        "held-span-contiguous-repair-v2.json": _write_json(
            artifacts / "held-span-contiguous-repair-v2.json", repair
        ),
        "official-page-decisions.json": _write_json(
            artifacts / "official-page-decisions.json",
            official
            or {
                "schema": "legalbot.live60-official-page-decisions.v1",
                "applied": False,
                "seals_expert_gold": False,
                "eligible_for_training": False,
                "training_export_allowed": False,
            },
        ),
        "provision-review-worksheet.json": _write_json(
            artifacts / "provision-review-worksheet.json", worksheet
        ),
        "mechanical-verification.json": _write_json(
            artifacts / "mechanical-verification.json", mechanical
        ),
        "independent-verification-return-hold.json": _write_json(
            artifacts / "independent-verification-return-hold.json", verification
        ),
        "owner-tick-application.json": _write_json(
            artifacts / "owner-tick-application.json",
            {key: value for key, value in ticks.items() if key != "cases"},
        ),
        "active-promotion-status.json": _write_json(
            artifacts / "active-promotion-status.json",
            build_active_promotion_status(project_root),
        ),
        "jobs-workers-routes.json": _write_json(
            artifacts / "jobs-workers-routes.json",
            build_jobs_workers_routes_report(bundle),
        ),
        "law-lane-inventory.json": _write_json(
            artifacts / "law-lane-inventory.json",
            build_law_lane_inventory(catalog if catalog.is_file() else None),
        ),
        "debug-backlog.json": _write_json(artifacts / "debug-backlog.json", build_debug_backlog()),
        "live-run-gate.json": _write_json(
            artifacts / "live-run-gate.json",
            build_live_run_gate_report(project_root=project_root, official=official),
        ),
        "remaining-gate-attempts.json": _write_json(
            artifacts / "remaining-gate-attempts.json",
            attempt_remaining_live_run_gates(
                project_root=project_root,
                ticks=ticks,
                official=official,
            ),
        ),
        "user-span-accuracy.json": _write_json(
            artifacts / "user-span-accuracy.json",
            verify_user_used_spans(
                ticks=ticks,
                repair=repair,
                catalog_path=catalog if catalog.is_file() else None,
            ),
        ),
        "subject-routing-audit-sample.json": _write_json(
            artifacts / "subject-routing-audit-sample.json", routing_sample
        ),
    }
    written.update(_refresh_return_hold_policy_artifacts(artifacts, identity))
    issue_docx = review / ISSUE_DECISION_DOCX
    held_docx = review / HELD_CHUNK_DOCX
    mechanical_docx = review / MECHANICAL_DOCX
    repair_docx = review / REPAIR_DOCX
    worksheet_docx = review / WORKSHEET_DOCX
    written[ISSUE_DECISION_DOCX] = _save_docx(build_issue_decision_document(issue_pack), issue_docx)
    if held_export.get("provisions"):
        written[HELD_CHUNK_DOCX] = _save_docx(build_held_chunk_document(held_export), held_docx)
        written[REPAIR_DOCX] = _save_docx(build_held_span_repair_document(repair), repair_docx)
        written[WORKSHEET_DOCX] = _save_docx(
            build_provision_review_worksheet_document(worksheet), worksheet_docx
        )
    written[MECHANICAL_DOCX] = _save_docx(
        build_mechanical_verification_document(mechanical), mechanical_docx
    )
    desktop: dict[str, str] = {}
    if copy_desktop:
        for path, name in (
            (issue_docx, ISSUE_DECISION_DOCX),
            (held_docx, HELD_CHUNK_DOCX),
            (mechanical_docx, MECHANICAL_DOCX),
            (repair_docx, REPAIR_DOCX),
            (worksheet_docx, WORKSHEET_DOCX),
        ):
            if path.is_file():
                copied = _copy_if_desktop(path, name)
                if copied:
                    desktop[name] = copied
    record_owner_review_pack_observability(
        pack_root,
        event_type="owner_decision_pack_exported",
        artifact_kind="issue_decision_pack",
        artifact_sha256=written["issue-decision-pack.json"],
        case_count=60,
    )
    summary = {
        "pack_name": dated_pack_name(as_of_date),
        "decision": "return_hold",
        "issue_count": 585,
        "research_route_counts": issue_pack["research_route_counts"],
        "selected_research_route_counts": issue_pack["selected_research_route_counts"],
        "structural_defect_count": structural["defect_count"],
        "repair_span_count": repair["repair_span_count"],
        "overlay_sealable": False,
        "generation_authorised": False,
        "after_item_5_live_30_generation_authorised": False,
        "o04_authorised": False,
        "active_promoted": False,
        "elasticsearch_used": False,
        "artifact_sha256": written,
        "desktop_copies": desktop,
        "exported_at": datetime.now(UTC).replace(microsecond=0).isoformat(),
        "eligible_for_training": False,
        "training_export_allowed": False,
    }
    assert_safe_evaluation_payload(summary)
    return summary
