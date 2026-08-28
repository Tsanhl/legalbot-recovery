"""Locator-aware Path-B final-check pack: exact official quotes, JSON import.

Word files are display-only. Gold is still the owner-accepted catalogue or v2
repair hash. This module never writes ACTIVE.json, PREVIOUS.json or O-04, and
it never invents later treatment for case law.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import sqlite3
from collections import Counter
from collections.abc import Mapping, Sequence
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType

from .live30 import assert_safe_evaluation_payload
from .live_suite import load_live_evaluation_bundle, sealed_sha256
from .live_suite_evidence_pack import (
    EXPECTED_SELECTED_ISSUES,
    _candidate_exclusion,
    _canonical_duplicate_row,
    _case_review,
    _catalogue_candidate,
    _chunks_for_source,
    _eligible_source_versions,
    _lexical_chunk_score,
    _owner_gold_span,
    _repair_candidate,
    _source_type,
    _title_match_score,
)
from .live_suite_held_span_repair import build_held_span_contiguous_repair
from .live_suite_owner_decisions import export_held_provision_chunks
from .live_suite_owner_review import HELD_STATUTORY_PROVISIONS
from .live_suite_path_b import (
    REVIEW_EXPORT_SCHEMA,
    REVIEW_IMPORT_SCHEMA,
    load_default_v2_repair,
    selected_generation_case_ids,
)
from .live_suite_reviewer_identity import build_owner_reviewer_identity
from .review_docx import (
    _add_header_footer,
    _add_page_break,
    _add_table,
    _add_toc,
    _configure_page,
    _configure_styles,
    _finalize_document_properties,
    _reject_prohibited_metadata,
)

FINAL_CHECK_SCHEMA = "legalbot.live60-owner-final-check.v1"
SUBSTANTIVE_REVIEW_SCHEMA = "legalbot.live60-path-b-substantive-owner-review-draft.v2"
FINAL_CHECK_MANIFEST_SCHEMA = "legalbot.live60-owner-final-check-manifest.v1"
FINAL_CHECK_PACK_RESULT_SCHEMA = "legalbot.live60-owner-final-check-pack-result.v1"
FINAL_CHECK_IMPORT_RESULT_SCHEMA = "legalbot.live60-owner-final-check-import-result.v1"
OWNER_FINAL_CHECK_TOKEN = "CONFIRM_OWNER_FINAL_CHECK"
EXPECTED_SELECTED_CASES = 30
PACK_JSON_NAME = "LegalBot-Live60-Path-B-Final-Check.json"
INDEX_DOCX_NAME = "LegalBot-Live60-Path-B-Final-Check-Index.docx"
PROPOSED_ACTIONS = frozenset({"accept_qualify", "cannot_fill_keep_gap", "needs_later_treatment"})
OWNER_ACTIONS = frozenset({"", "accept_qualify", "reject_keep_gap"})
CASE_LATER_TREATMENT = frozenset({"confirmed_current", "qualified_current"})
CASE_LEGAL_ROLES = frozenset({"holding_ratio", "binding_legal_rule"})
_CATALOGUE_ROW_SQL = """
SELECT c.id AS chunk_id, c.source_version_id, c.ordinal, c.locator,
       c.text_sha256, c.markdown_text, c.stream,
       sv.authority_identity_id, sv.title, sv.as_of_date,
       sv.currentness_status, sv.licence_name, sv.licence_url,
       sv.review_status, sv.canonical_url, sv.stable_identifier,
       sv.superseded_by,
       d.id AS document_id, d.content_sha256 AS official_snapshot_sha256,
       d.source_identity_id AS stable_source_id, d.status AS document_status,
       d.lane, d.jurisdiction, d.duplicate_of
FROM chunks c
JOIN source_versions sv ON sv.id = c.source_version_id
JOIN documents d ON d.id = sv.document_id
WHERE c.id = ?
"""
_PARENTHETICAL_SECTION = re.compile(
    r"\bsection\s+(\d+[A-Za-z]*)\s+\(([^)]{8,120})\)",
    re.IGNORECASE,
)
_SECTION_CHAIN = re.compile(
    r"\bsection\s+\d+[A-Za-z]*\s+section\s+\d+[A-Za-z]*",
    re.IGNORECASE,
)
_PROVISION = re.compile(
    r"\b(sections|section|ss|s|regulations|regulation|regs|reg|"
    r"articles|article|arts|art|rules|rule)\s*\.?\s*(\d+[A-Za-z]*)"
    r"(?:\s*\(([^)]+)\))?(?:\s*\(([^)]+)\))?",
    re.IGNORECASE,
)
_CONTINUATION = re.compile(
    r"(?:and|,|;)\s+(?:(?P<kind>ss|sections?|s|arts?|articles?|regs?|regulations?)\s+)?"
    r"(?P<num>\d+[A-Za-z]*)(?:\s*\((?P<para>[^)]+)\))?",
    re.IGNORECASE,
)
_KIND = {
    "sections": "s",
    "section": "s",
    "ss": "s",
    "s": "s",
    "regulations": "reg",
    "regulation": "reg",
    "regs": "reg",
    "reg": "reg",
    "articles": "art",
    "article": "art",
    "arts": "art",
    "art": "art",
    "rules": "r",
    "rule": "r",
}
_NON_OPERATIVE = (
    "repealed by",
    "substituted by",
    "omitted (",
    "coming into force",
    "transitional provision",
    "after subsection",
    "insert—",
    "insert-",
)
_DEFINITION_MARKERS = (
    "these definitions apply",
    "has the meaning given",
    "key definitions",
    "for the purposes of this part",
)
_PROVIDED_SECTIONS = re.compile(
    r"(?:provided by|described in|meaning given in|see)\s+sections?\s+"
    r"([0-9A-Za-z,\s]+(?:and\s+[0-9A-Za-z]+)?)",
    re.IGNORECASE,
)
_OPERATIVE_MARKERS = (
    " shall ",
    " must ",
    " is to be ",
    " if they ",
    " if it ",
    " if the ",
    " owes ",
    " may not ",
)


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _collapsed(value: str) -> str:
    return " ".join(value.split())


def _write_json(path: Path, value: Mapping[str, Any]) -> bytes:
    raw = (json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode("utf-8")
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(raw)
    os.chmod(path, 0o600)
    return raw


def locator_key(value: str) -> str:
    """Collapse a locator or pinpoint to a comparable provision key."""

    cleaned = value.casefold().replace("\u00a0", " ")
    cleaned = cleaned.replace("section", "s").replace("regulation", "reg")
    cleaned = re.sub(r"\barticle\b", "art", cleaned)
    cleaned = re.sub(r"\brule\b", "r", cleaned)
    cleaned = cleaned.replace(".", " ")
    return " ".join(cleaned.split())


def provision_keys(value: str) -> tuple[str, ...]:
    keys: list[str] = []
    seen: set[str] = set()
    last_kind = "s"

    def add(kind: str, number: str, *parens: str) -> None:
        key = f"{kind} {number.casefold()}"
        for para in parens:
            if para:
                key = f"{key}({_collapsed(para).casefold()})"
        if key not in seen:
            seen.add(key)
            keys.append(key)

    for match in _PROVISION.finditer(value):
        kind = _KIND.get(match.group(1).casefold(), match.group(1).casefold())
        last_kind = kind
        add(kind, match.group(2), match.group(3) or "", match.group(4) or "")
    for match in _CONTINUATION.finditer(value):
        kind = last_kind
        named = match.group("kind")
        if named:
            kind = _KIND.get(named.casefold(), last_kind)
        add(kind, match.group("num"), match.group("para") or "")
    return tuple(keys)


def locator_matches_target(locator: str, target: str) -> bool:
    loc = locator_key(locator)
    tgt = locator_key(target)
    if not loc or not tgt:
        return False
    return loc == tgt or loc.startswith(f"{tgt}(") or loc.startswith(f"{tgt} ")


def is_cross_ref_fragment(excerpt: str) -> bool:
    """True when the printed chunk is a cross-reference shell, not the rule."""

    collapsed = _collapsed(excerpt)
    if not collapsed or len(collapsed) > 400:
        return False
    lowered = f" {collapsed.casefold()} "
    if any(marker in lowered for marker in _OPERATIVE_MARKERS):
        return False
    parenthetical = bool(_PARENTHETICAL_SECTION.search(collapsed))
    chained = bool(_SECTION_CHAIN.search(collapsed))
    if parenthetical and chained:
        return True
    return parenthetical and collapsed.rstrip().endswith((",", ";", ")"))


def referenced_operative_locators(excerpt: str) -> tuple[str, ...]:
    """Section numbers named as the target of a cross-reference fragment."""

    found: list[str] = []
    seen: set[str] = set()

    def add(number: str) -> None:
        key = f"s {number.casefold()}"
        if key not in seen:
            seen.add(key)
            found.append(key)

    for match in _PARENTHETICAL_SECTION.finditer(excerpt):
        add(match.group(1))
    for match in _PROVIDED_SECTIONS.finditer(excerpt):
        for number in re.findall(r"\d+[A-Za-z]*", match.group(1)):
            add(number)
    return tuple(found)


def operative_rank(
    topic: str,
    candidate: Mapping[str, Any],
    *,
    referenced: Sequence[str],
    pinpoint: str,
) -> float:
    excerpt = str(candidate.get("excerpt") or "")
    locator = str(candidate.get("legal_locator") or "")
    score = _lexical_chunk_score(topic, excerpt, locator)
    collapsed = _collapsed(excerpt)
    lowered = collapsed.casefold()
    if is_cross_ref_fragment(excerpt):
        score -= 80.0
    if any(marker in lowered for marker in _DEFINITION_MARKERS):
        score -= 70.0
    if any(marker in lowered for marker in _NON_OPERATIVE):
        score -= 40.0
    if any(locator_matches_target(locator, target) for target in referenced):
        score += 60.0
    pin_keys = provision_keys(pinpoint)
    if any(locator_matches_target(locator, key) for key in pin_keys):
        score += 40.0
    topic_cf = _collapsed(topic).casefold()
    if topic_cf and topic_cf in lowered:
        score += 15.0
        index = lowered.find(topic_cf)
        if 0 <= index < 160:
            score += 25.0
    if any(marker in lowered for marker in _OPERATIVE_MARKERS):
        score += 10.0
    return score


def _matching_versions(
    source: Mapping[str, Any],
    source_versions: Sequence[Any],
) -> list[tuple[float, Any]]:
    source_name = str(source.get("source_name") or "")
    source_url = str(source.get("url") or "").rstrip("/")
    matches: list[tuple[float, Any]] = []
    for version in source_versions:
        title_score = _title_match_score(source_name, str(version["title"] or ""))
        url_match = bool(
            source_url and str(version["canonical_url"] or "").rstrip("/") == source_url
        )
        if title_score < 0.7 and not url_match:
            continue
        matches.append((1.0 if url_match else title_score, version))
    matches.sort(key=lambda item: (-item[0], str(item[1]["source_version_id"])))
    return matches[:3]


def _repairs_for_source(
    source: Mapping[str, Any],
    repair: Mapping[str, Any],
) -> tuple[Mapping[str, Any], ...]:
    name = str(source.get("source_name") or "").casefold()
    pinpoint = str(source.get("pinpoint") or "").casefold()
    selected: list[Mapping[str, Any]] = []
    for item in repair.get("repairs", ()):
        if item.get("gold_eligible_candidate") is not True:
            continue
        authority = str(item.get("legal_authority_id") or "")
        locator = str(item.get("legal_locator") or "").casefold()
        if "limitation act 1980" in name:
            if authority != "ukpga:1980:58":
                continue
            if "14a" in pinpoint:
                if locator.startswith("s 14a"):
                    selected.append(item)
            elif (re.search(r"\bs\.?\s*2\b", pinpoint) or "section 2" in pinpoint) and (
                locator == "s 2" or locator.startswith("s 2(")
            ):
                selected.append(item)
        elif "trustee act 2000" in name:
            if authority == "ukpga:2000:29" and locator.startswith("s 1"):
                selected.append(item)
        elif "inheritance (provision for family and dependants) act 1975" in name:
            if authority == "ukpga:1975:63" and locator.startswith("s 1"):
                selected.append(item)
    return tuple(selected)


def _held_id_for_repair(item: Mapping[str, Any]) -> str | None:
    authority = str(item.get("legal_authority_id") or "")
    locator = str(item.get("legal_locator") or "").casefold()
    if authority == "ukpga:1980:58":
        if locator.startswith("s 14a"):
            return "held-provision-02"
        if locator == "s 2" or locator.startswith("s 2("):
            return "held-provision-01"
        return None
    if authority == "ukpga:2000:29":
        return "held-provision-03"
    if authority == "ukpga:1975:63":
        return "held-provision-04"
    return None


def _eligible_row(
    *,
    connection: Any,
    row: Any,
    chunk_id: str,
    source_type: str,
    as_of_date: date,
    exclusions: Counter[str],
) -> Any | None:
    resolved = _canonical_duplicate_row(connection, row) if row is not None else None
    reason = _candidate_exclusion(
        resolved,
        chunk_id=chunk_id,
        source_type=source_type,
        as_of_date=as_of_date,
    )
    if reason:
        exclusions[reason] += 1
        return None
    return resolved


def _append_candidate(
    collected: dict[str, dict[str, Any]],
    candidate: Mapping[str, Any],
) -> None:
    candidate_id = str(candidate["candidate_id"])
    if candidate_id not in collected:
        collected[candidate_id] = dict(candidate)


def collect_operative_candidates(
    *,
    connection: Any,
    source_versions: Sequence[Any],
    source_cache: dict[str, list[Any]],
    case_id: str,
    issue_id: str,
    topic: str,
    sources: Sequence[Mapping[str, Any]],
    as_of_date: date,
    repair: Mapping[str, Any],
) -> tuple[list[dict[str, Any]], Counter[str]]:
    """Return ranked exact catalogue/repair candidates, operative provision first."""

    collected: dict[str, dict[str, Any]] = {}
    exclusions: Counter[str] = Counter()
    referenced: list[str] = []
    pinpoint = ""

    def add_row(
        row: Any,
        source: Mapping[str, Any],
        origin: str,
    ) -> None:
        candidate = _catalogue_candidate(
            row=row,
            case_id=case_id,
            issue_id=issue_id,
            source_name=str(source.get("source_name") or ""),
            source_kind=str(source.get("source_type") or ""),
            candidate_origin=origin,
            map_rank=str(source.get("rank") or ""),
            map_pinpoint=str(source.get("pinpoint") or ""),
        )
        _append_candidate(collected, candidate)
        referenced.extend(referenced_operative_locators(str(candidate["excerpt"])))

    for source in sources:
        pinpoint = pinpoint or str(source.get("pinpoint") or "")
        source_type = _source_type(str(source.get("source_type") or ""))
        for span in source.get("local_spans", ()):
            original_id = str(span.get("chunk_id") or "")
            row = connection.execute(_CATALOGUE_ROW_SQL, (original_id,)).fetchone()
            eligible = _eligible_row(
                connection=connection,
                row=row,
                chunk_id=original_id,
                source_type=source_type,
                as_of_date=as_of_date,
                exclusions=exclusions,
            )
            if eligible is not None:
                add_row(eligible, source, "resource_map_exact_locator")
        for repair_item in _repairs_for_source(source, repair):
            _append_candidate(
                collected,
                _repair_candidate(
                    item=repair_item,
                    case_id=case_id,
                    issue_id=issue_id,
                    source_name=str(source.get("source_name") or ""),
                    map_rank=str(source.get("rank") or ""),
                    map_pinpoint=str(source.get("pinpoint") or ""),
                ),
            )
        for _title_score, version in _matching_versions(source, source_versions):
            for row in _chunks_for_source(
                connection,
                str(version["source_version_id"]),
                cache=source_cache,
            ):
                eligible = _eligible_row(
                    connection=connection,
                    row=row,
                    chunk_id=str(row["chunk_id"]),
                    source_type=source_type,
                    as_of_date=as_of_date,
                    exclusions=exclusions,
                )
                if eligible is None:
                    continue
                add_row(eligible, source, "operative_locator_search")

    if referenced:
        wanted = tuple(dict.fromkeys(referenced))
        for source in sources:
            source_type = _source_type(str(source.get("source_type") or ""))
            for _title_score, version in _matching_versions(source, source_versions):
                for row in _chunks_for_source(
                    connection,
                    str(version["source_version_id"]),
                    cache=source_cache,
                ):
                    if not any(
                        locator_matches_target(str(row["locator"] or ""), target)
                        for target in wanted
                    ):
                        continue
                    eligible = _eligible_row(
                        connection=connection,
                        row=row,
                        chunk_id=str(row["chunk_id"]),
                        source_type=source_type,
                        as_of_date=as_of_date,
                        exclusions=exclusions,
                    )
                    if eligible is not None:
                        add_row(eligible, source, "operative_locator_follow")

    pin_keys = provision_keys(pinpoint)
    filtered: list[dict[str, Any]] = []
    for item in collected.values():
        lexical = _lexical_chunk_score(
            topic, str(item.get("excerpt") or ""), str(item.get("legal_locator") or "")
        )
        locator_hit = any(
            locator_matches_target(str(item.get("legal_locator") or ""), target)
            for target in referenced
        )
        pin_hit = any(
            locator_matches_target(str(item.get("legal_locator") or ""), key) for key in pin_keys
        )
        mapped = item.get("candidate_origin") == "resource_map_exact_locator"
        if lexical <= 0 and not locator_hit and not pin_hit and not mapped:
            continue
        filtered.append(item)
    ranked = sorted(
        filtered,
        key=lambda item: (
            -operative_rank(
                topic,
                item,
                referenced=referenced,
                pinpoint=pinpoint,
            ),
            str(item["candidate_id"]),
        ),
    )
    if not ranked and not exclusions:
        exclusions["no_catalogue_candidate"] += 1
    return ranked[:3], exclusions


def _public_candidate(candidate: Mapping[str, Any]) -> dict[str, Any]:
    excerpt = _collapsed(str(candidate.get("excerpt") or ""))
    payload = {
        "candidate_id": candidate["candidate_id"],
        "candidate_origin": candidate["candidate_origin"],
        "source_name": candidate["source_name"],
        "source_type": candidate["source_type"],
        "map_pinpoint": candidate.get("map_pinpoint") or "",
        "authority_identity": candidate.get("authority_identity") or "",
        "jurisdiction": candidate.get("jurisdiction") or "",
        "source_version_id": candidate["source_version_id"],
        "stable_source_id": candidate["stable_source_id"],
        "official_snapshot_sha256": candidate.get("official_snapshot_sha256") or "",
        "chunk_id": candidate["chunk_id"],
        "legal_locator": candidate["legal_locator"],
        "content_sha256": candidate["content_sha256"],
        "excerpt": excerpt,
        "proposed_legal_role": candidate["proposed_legal_role"],
        "catalogue_currentness": candidate.get("catalogue_currentness") or "",
        "later_treatment_requirement": candidate.get("later_treatment_requirement") or "",
        "cross_ref_fragment": is_cross_ref_fragment(str(candidate.get("excerpt") or "")),
        "parent_chunk_id": candidate.get("parent_chunk_id") or "",
        "derivation_manifest_sha256": candidate.get("derivation_manifest_sha256") or "",
        "seals_expert_gold": False,
    }
    return payload


def _proposed_action(candidate: Mapping[str, Any] | None) -> str:
    if candidate is None:
        return "cannot_fill_keep_gap"
    if candidate["source_type"] == "case":
        return "needs_later_treatment"
    if candidate.get("cross_ref_fragment") is True:
        return "cannot_fill_keep_gap"
    return "accept_qualify"


def _cannot_fill_reason(
    candidate: Mapping[str, Any] | None,
    exclusions: Mapping[str, int],
    action: str,
) -> str:
    if action == "needs_later_treatment":
        return "case_later_treatment_required"
    if action == "accept_qualify":
        return ""
    if candidate is not None and candidate.get("cross_ref_fragment") is True:
        return "cross_ref_fragment_only"
    if exclusions:
        return sorted(exclusions, key=lambda key: (-exclusions[key], key))[0]
    return "no_catalogue_candidate"


def _load_repair(
    *,
    project_root: Path,
    catalog_path: Path,
    repair_payload: Mapping[str, Any] | None,
) -> dict[str, Any]:
    if repair_payload is not None:
        return dict(repair_payload)
    default = load_default_v2_repair(project_root)
    if default is not None:
        return default
    return build_held_span_contiguous_repair(export_held_provision_chunks(catalog_path))


def _held_statute_rows(repair: Mapping[str, Any]) -> list[dict[str, Any]]:
    by_held: dict[str, list[dict[str, Any]]] = {
        held_id: [] for held_id, _title, _note in HELD_STATUTORY_PROVISIONS
    }
    for item in repair.get("repairs", ()):
        if item.get("gold_eligible_candidate") is not True:
            continue
        held_id = _held_id_for_repair(item)
        if held_id is None or held_id not in by_held:
            continue
        by_held[held_id].append(
            {
                "repair_span_id": item["repair_span_id"],
                "legal_locator": item["legal_locator"],
                "content_sha256": item["text_sha256"],
                "excerpt": _collapsed(str(item.get("markdown_text") or "")),
                "legal_authority_id": item["legal_authority_id"],
                "parent_chunk_id": item["parent_chunk_id"],
                "source_version_id": item["source_version_id"],
                "gold_eligible_candidate": True,
            }
        )
    rows = []
    for held_id, title, note in HELD_STATUTORY_PROVISIONS:
        candidates = by_held[held_id]
        action = "accept_qualify" if candidates else "cannot_fill_keep_gap"
        rows.append(
            {
                "held_id": held_id,
                "provision": title,
                "note": note,
                "gold_eligible_repair_count": len(candidates),
                "proposed_action": action,
                "owner_action": "",
                "cannot_fill_reason": (
                    "" if candidates else "no_accepted_v2_repair_span_in_this_pack"
                ),
                "candidates": candidates,
                "seals_expert_gold": False,
            }
        )
    return rows


def build_owner_final_check_pack(
    *,
    project_root: Path,
    catalog_path: Path,
    evidence_map_path: Path,
    as_of_date: date,
    repair_payload: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Prepare 305 issue rows with the best exact official quote, unsealed."""

    if not catalog_path.is_file():
        raise FileNotFoundError("catalogue is not present")
    bundle = load_live_evaluation_bundle(
        project_root / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    evidence_map_bytes = evidence_map_path.read_bytes()
    evidence_map = json.loads(evidence_map_bytes)
    by_row = {
        f"{item['case_id']}:{item['issue_id']}": item
        for item in evidence_map.get("issues", ())
        if item.get("generation_disposition") == "generate_once"
    }
    selected_ids = selected_generation_case_ids(bundle)
    expected_rows = [
        f"{case.case_id}:issue-{number:02d}"
        for case in bundle.registry.cases
        if case.case_id in selected_ids
        for number in range(1, len(case.must_cover_issues) + 1)
    ]
    if len(expected_rows) != EXPECTED_SELECTED_ISSUES or set(by_row) != set(expected_rows):
        raise ValueError("candidate evidence map does not contain the selected 305 issues")

    repair = _load_repair(
        project_root=project_root,
        catalog_path=catalog_path,
        repair_payload=repair_payload,
    )
    connection = sqlite3.connect(f"file:{catalog_path}?mode=ro", uri=True)
    connection.row_factory = sqlite3.Row
    source_versions = _eligible_source_versions(connection)
    source_cache: dict[str, list[Any]] = {}
    cases: list[dict[str, Any]] = []
    counts: Counter[str] = Counter()
    try:
        for case in bundle.registry.cases:
            if case.case_id not in selected_ids:
                continue
            issues: list[dict[str, Any]] = []
            for number, topic in enumerate(case.must_cover_issues, start=1):
                issue_id = f"issue-{number:02d}"
                mapped = by_row[f"{case.case_id}:{issue_id}"]
                ranked, exclusions = collect_operative_candidates(
                    connection=connection,
                    source_versions=source_versions,
                    source_cache=source_cache,
                    case_id=case.case_id,
                    issue_id=issue_id,
                    topic=topic,
                    sources=tuple(mapped.get("candidates") or ()),
                    as_of_date=as_of_date,
                    repair=repair,
                )
                public = [_public_candidate(item) for item in ranked]
                proposed = public[0] if public else None
                action = _proposed_action(proposed)
                counts[action] += 1
                issues.append(
                    {
                        "row_id": f"{case.case_id}:{issue_id}",
                        "case_id": case.case_id,
                        "issue_id": issue_id,
                        "topic": topic,
                        "topic_sha256": _sha256_text(topic),
                        "proposed_action": action,
                        "owner_action": "",
                        "owner_later_treatment": "",
                        "owner_legal_role": "",
                        "owner_exact_proposition": "",
                        "cannot_fill_reason": _cannot_fill_reason(proposed, exclusions, action),
                        "excluded_reason_counts": dict(sorted(exclusions.items())),
                        "proposed_candidate_id": (
                            proposed["candidate_id"] if proposed is not None else ""
                        ),
                        "candidates": public,
                        "seals_expert_gold": False,
                    }
                )
            cases.append(
                {
                    "case_id": case.case_id,
                    "subject": case.subject,
                    "task_type": case.task_type,
                    "expected_research_route": case.expected_research_route,
                    "question_sha256": case.question_sha256,
                    "record_sha256": case.record_sha256,
                    "issues": issues,
                }
            )
    finally:
        connection.close()

    issue_count = sum(len(case["issues"]) for case in cases)
    if len(cases) != EXPECTED_SELECTED_CASES or issue_count != EXPECTED_SELECTED_ISSUES:
        raise ValueError("final-check pack does not cover the selected 30/305")
    payload = {
        "schema": FINAL_CHECK_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "as_of_date": as_of_date.isoformat(),
        "suite_registry_canonical_sha256": bundle.registry.canonical_sha256,
        "run_plan_sha256": bundle.manifest.run_plan_sha256,
        "evidence_map_sha256": _sha256_bytes(evidence_map_bytes),
        "status": "ai_prepared_owner_final_check_required",
        "word_is_gold": False,
        "word_is_import_surface": False,
        "owner_confirmation_token": "",
        "reviewer_role": "",
        "reviewer_ref": "",
        "selected_case_count": len(cases),
        "selected_issue_count": issue_count,
        "bucket_counts": {
            "accept_qualify": int(counts["accept_qualify"]),
            "needs_later_treatment": int(counts["needs_later_treatment"]),
            "cannot_fill_keep_gap": int(counts["cannot_fill_keep_gap"]),
        },
        "held_statutes": _held_statute_rows(repair),
        "cases": cases,
        "ai_role": "mechanical_accuracy_verifier_only",
        "ai_second_reviewer_forbidden": True,
        "owner_is_primary_reviewer": True,
        "seals_expert_gold": False,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }
    assert_safe_evaluation_payload(
        {key: value for key, value in payload.items() if key not in {"cases", "held_statutes"}}
    )
    return payload


def _add_safe_paragraph(document: DocumentType, value: str, *, style: str = "Normal") -> None:
    document.add_paragraph(
        _reject_prohibited_metadata(value, label="final-check prose"),
        style=style,
    )


def build_final_check_index_document(pack: Mapping[str, Any]) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, f"Live60-{pack['as_of_date']}")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Path-B owner final-check pack", style="Title")
    _add_safe_paragraph(
        document,
        "This Word file is display-only. Import reads the companion JSON. "
        "Quotes are catalogue or accepted v2 repair text. Confirm the JSON with "
        f"{OWNER_FINAL_CHECK_TOKEN}. Case-law rows stay gaps until later treatment "
        "is recorded. This pack does not seal gold, write ACTIVE, or issue O-04.",
    )
    buckets = pack["bucket_counts"]
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("As-of date", str(pack["as_of_date"])),
            ("Selected cases", str(pack["selected_case_count"])),
            ("Selected issues", str(pack["selected_issue_count"])),
            ("Proposed accept_qualify", str(buckets["accept_qualify"])),
            ("Needs later treatment", str(buckets["needs_later_treatment"])),
            ("Cannot fill / keep gap", str(buckets["cannot_fill_keep_gap"])),
            ("Word is gold", "false"),
            ("Word is import surface", "false"),
            ("Overlay sealable", "false"),
            ("Confirmation token", OWNER_FINAL_CHECK_TOKEN),
        ),
        (3_120, 6_240),
    )
    _add_toc(document)
    document.add_heading("Held statutes", level=1)
    for item in pack.get("held_statutes", ()):
        document.add_heading(str(item["held_id"]), level=2)
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("Provision", item["provision"]),
                ("Proposed action", item["proposed_action"]),
                ("Gold-eligible repairs", str(item["gold_eligible_repair_count"])),
                ("Cannot-fill reason", item["cannot_fill_reason"] or "none"),
            ),
            (3_120, 6_240),
        )
        if item["candidates"]:
            _add_safe_paragraph(document, item["candidates"][0]["excerpt"])
    _add_page_break(document)
    document.add_heading("Cases", level=1)
    _add_table(
        document,
        ("Case ID", "Subject", "Accept", "Later treatment", "Gap"),
        tuple(
            (
                case["case_id"],
                case["subject"],
                str(sum(issue["proposed_action"] == "accept_qualify" for issue in case["issues"])),
                str(
                    sum(
                        issue["proposed_action"] == "needs_later_treatment"
                        for issue in case["issues"]
                    )
                ),
                str(
                    sum(
                        issue["proposed_action"] == "cannot_fill_keep_gap"
                        for issue in case["issues"]
                    )
                ),
            )
            for case in pack["cases"]
        ),
        (2_000, 2_400, 1_400, 1_780, 1_780),
    )
    _finalize_document_properties(
        document,
        title="Live60 Path-B owner final-check index",
        subject="Display-only index of extracted official quotes",
    )
    return document


def build_final_check_case_document(
    pack: Mapping[str, Any],
    case: Mapping[str, Any],
) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, f"Live60-{pack['as_of_date']}")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph(
        f"Path-B final-check · {case['case_id']}",
        style="Title",
    )
    _add_safe_paragraph(
        document,
        "Display only. To accept these hashes, set owner_confirmation_token in the "
        f"JSON to {OWNER_FINAL_CHECK_TOKEN}. To reject one quote, set that row's "
        "owner_action to reject_keep_gap. Do not treat this Word file as gold.",
    )
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Case ID", case["case_id"]),
            ("Subject", case["subject"]),
            ("Task", case["task_type"]),
            ("Route", case["expected_research_route"]),
            ("Issue count", str(len(case["issues"]))),
        ),
        (3_120, 6_240),
    )
    _add_toc(document)
    _add_page_break(document)
    for issue in case["issues"]:
        document.add_heading(f"{issue['issue_id']} · {issue['topic']}", level=1)
        proposed = next(
            (
                item
                for item in issue["candidates"]
                if item["candidate_id"] == issue["proposed_candidate_id"]
            ),
            None,
        )
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("Row ID", issue["row_id"]),
                ("Proposed action", issue["proposed_action"]),
                ("Cannot-fill reason", issue["cannot_fill_reason"] or "none"),
                ("Candidate ID", issue["proposed_candidate_id"] or "none"),
                (
                    "Locator",
                    proposed["legal_locator"] if proposed is not None else "none",
                ),
                (
                    "Content SHA-256",
                    proposed["content_sha256"] if proposed is not None else "none",
                ),
                (
                    "Authority",
                    proposed["source_name"] if proposed is not None else "none",
                ),
            ),
            (3_120, 6_240),
        )
        if proposed is None:
            _add_safe_paragraph(
                document,
                "No exact eligible official excerpt is in the catalogue for this issue.",
            )
        else:
            _add_safe_paragraph(document, proposed["excerpt"])
            if issue["proposed_action"] == "needs_later_treatment":
                _add_safe_paragraph(
                    document,
                    "Case-law quote printed only. Record confirmed_current or "
                    "qualified_current plus a legal role in the JSON before this "
                    "row can import as qualified.",
                )
    _finalize_document_properties(
        document,
        title=f"Live60 Path-B final-check {case['case_id']}",
        subject="Display-only extracted official quotes",
    )
    return document


def _safe_manifest(
    pack: Mapping[str, Any], artifacts: Sequence[Mapping[str, Any]]
) -> dict[str, Any]:
    payload = {
        "schema": FINAL_CHECK_MANIFEST_SCHEMA,
        "suite_id": pack["suite_id"],
        "as_of_date": pack["as_of_date"],
        "suite_registry_canonical_sha256": pack["suite_registry_canonical_sha256"],
        "run_plan_sha256": pack["run_plan_sha256"],
        "evidence_map_sha256": pack["evidence_map_sha256"],
        "selected_case_count": pack["selected_case_count"],
        "selected_issue_count": pack["selected_issue_count"],
        "bucket_counts": pack["bucket_counts"],
        "pack_artifact_name": PACK_JSON_NAME,
        "index_artifact_name": INDEX_DOCX_NAME,
        "artifacts": list(artifacts),
        "word_is_import_surface": False,
        "owner_review_required": True,
        "seals_expert_gold": False,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }
    assert_safe_evaluation_payload(payload)
    return payload


def export_owner_final_check_pack(
    *,
    project_root: Path,
    catalog_path: Path,
    evidence_map_path: Path,
    output_dir: Path,
    as_of_date: date,
    overwrite: bool = False,
    repair_payload: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Write the JSON import surface plus display-only Word copies."""

    destination = output_dir.resolve()
    if destination.exists() and any(destination.iterdir()) and not overwrite:
        raise FileExistsError("owner final-check directory is not empty")
    destination.mkdir(parents=True, exist_ok=True)
    pack = build_owner_final_check_pack(
        project_root=project_root,
        catalog_path=catalog_path,
        evidence_map_path=evidence_map_path,
        as_of_date=as_of_date,
        repair_payload=repair_payload,
    )
    artifacts: list[dict[str, Any]] = []
    json_path = destination / PACK_JSON_NAME
    json_bytes = _write_json(json_path, pack)
    artifacts.append(
        {
            "artifact_name": PACK_JSON_NAME,
            "artifact_sha256": _sha256_bytes(json_bytes),
            "kind": "json_import_surface",
        }
    )
    index_path = destination / INDEX_DOCX_NAME
    build_final_check_index_document(pack).save(str(index_path))
    artifacts.append(
        {
            "artifact_name": INDEX_DOCX_NAME,
            "artifact_sha256": _sha256_bytes(index_path.read_bytes()),
            "kind": "docx_display_only",
        }
    )
    for case in pack["cases"]:
        name = f"LegalBot-Live60-Path-B-Final-Check-{case['case_id']}.docx"
        path = destination / name
        build_final_check_case_document(pack, case).save(str(path))
        artifacts.append(
            {
                "artifact_name": name,
                "artifact_sha256": _sha256_bytes(path.read_bytes()),
                "kind": "docx_display_only",
            }
        )
    manifest = _safe_manifest(pack, artifacts)
    manifest_bytes = _write_json(destination / "manifest.json", manifest)
    return {
        "schema": FINAL_CHECK_PACK_RESULT_SCHEMA,
        "output_dir": str(destination),
        "pack_sha256": artifacts[0]["artifact_sha256"],
        "manifest_sha256": _sha256_bytes(manifest_bytes),
        "selected_issue_count": pack["selected_issue_count"],
        "bucket_counts": pack["bucket_counts"],
        "word_is_import_surface": False,
        "seals_expert_gold": False,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "writes_active": False,
        "writes_o04": False,
    }


def _resolved_owner_action(issue: Mapping[str, Any]) -> str:
    owner_action = str(issue.get("owner_action") or "").strip()
    if owner_action in {"accept_qualify", "approve_qualified"}:
        return "accept_qualify"
    if owner_action in {"reject_keep_gap", "request_another_candidate"}:
        return "reject_keep_gap"
    if owner_action:
        raise ValueError(f"{issue['row_id']} has an invalid owner_action")
    proposed = str(issue.get("proposed_action") or "")
    if proposed not in PROPOSED_ACTIONS:
        raise ValueError(f"{issue['row_id']} has an invalid proposed_action")
    if proposed == "accept_qualify":
        return "accept_qualify"
    return "reject_keep_gap"


def _selected_candidate_ids(issue: Mapping[str, Any]) -> tuple[str, ...]:
    selected = tuple(
        str(item)
        for item in issue.get("approved_candidate_ids")
        or issue.get("owner_selected_candidate_ids")
        or ()
        if str(item).strip()
    )
    if selected:
        return selected
    candidate_id = str(
        issue.get("proposed_candidate_id") or issue.get("owner_selected_candidate_id") or ""
    ).strip()
    if candidate_id:
        return (candidate_id,)
    return ()


def adapt_substantive_review_for_import(
    draft: Mapping[str, Any],
    *,
    confirmation_token: str,
) -> dict[str, Any]:
    """Map a completed substantive-review draft onto the JSON import surface."""

    if draft.get("schema") not in {FINAL_CHECK_SCHEMA, SUBSTANTIVE_REVIEW_SCHEMA}:
        raise ValueError("substantive review uses an unsupported schema")
    if confirmation_token != OWNER_FINAL_CHECK_TOKEN:
        raise ValueError("owner final-check confirmation token is missing")
    cases = []
    for case in draft.get("cases", ()):
        issues = []
        for issue in case.get("issues", ()):
            action = _resolved_owner_action(issue)
            selected = _selected_candidate_ids(issue) if action == "accept_qualify" else ()
            source_types = {
                str(item.get("source_type") or "")
                for item in issue.get("candidates") or ()
                if item.get("candidate_id") in set(selected)
            }
            later_treatment = str(issue.get("owner_later_treatment") or "")
            proposition = str(issue.get("owner_exact_proposition") or "")
            if "case" not in source_types:
                later_treatment = ""
                proposition = ""
            reason = str(issue.get("cannot_fill_reason") or "")
            if action != "accept_qualify" and not reason:
                reason = (
                    "owner_requested_another_candidate"
                    if str(issue.get("owner_action") or "") == "request_another_candidate"
                    else "owner_rejected_or_unfilled_final_check"
                )
            issues.append(
                {
                    "row_id": issue["row_id"],
                    "case_id": issue.get("case_id") or case["case_id"],
                    "issue_id": issue["issue_id"],
                    "topic": issue.get("topic") or "",
                    "topic_sha256": issue.get("topic_sha256") or "",
                    "proposed_action": issue.get("proposed_action") or "",
                    "owner_action": action,
                    "owner_later_treatment": later_treatment if "case" in source_types else "",
                    "owner_legal_role": (
                        str(issue.get("owner_legal_role") or "")
                        if action == "accept_qualify"
                        else ""
                    ),
                    "owner_exact_proposition": proposition if "case" in source_types else "",
                    "cannot_fill_reason": reason,
                    "excluded_reason_counts": issue.get("excluded_reason_counts") or {},
                    "proposed_candidate_id": selected[0] if selected else "",
                    "approved_candidate_ids": list(selected),
                    "candidates": list(issue.get("candidates") or ()),
                    "seals_expert_gold": False,
                }
            )
        cases.append(
            {
                "case_id": case["case_id"],
                "subject": case.get("subject") or "",
                "task_type": case.get("task_type") or "",
                "expected_research_route": case.get("expected_research_route") or "",
                "question_sha256": case.get("question_sha256") or "",
                "record_sha256": case.get("record_sha256") or "",
                "issues": issues,
            }
        )
    return {
        "schema": FINAL_CHECK_SCHEMA,
        "suite_id": draft.get("suite_id") or "live-evaluation-60-v1",
        "as_of_date": draft["as_of_date"],
        "suite_registry_canonical_sha256": draft.get("suite_registry_canonical_sha256") or "",
        "run_plan_sha256": draft.get("run_plan_sha256") or "",
        "evidence_map_sha256": draft.get("evidence_map_sha256") or "",
        "status": "owner_adopted_substantive_review",
        "word_is_gold": False,
        "word_is_import_surface": False,
        "owner_confirmation_token": OWNER_FINAL_CHECK_TOKEN,
        "reviewer_role": draft.get("reviewer_role") or "",
        "reviewer_ref": draft.get("reviewer_ref") or "",
        "selected_case_count": draft.get("selected_case_count") or len(cases),
        "selected_issue_count": draft.get("selected_issue_count")
        or sum(len(case["issues"]) for case in cases),
        "bucket_counts": draft.get("bucket_counts") or draft.get("decision_counts") or {},
        "held_statutes": draft.get("held_statutes") or draft.get("held_statutes_original") or [],
        "cases": cases,
        "ai_role": "mechanical_accuracy_verifier_only",
        "ai_second_reviewer_forbidden": True,
        "owner_is_primary_reviewer": True,
        "seals_expert_gold": False,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }


def _parse_final_check_issue(
    *,
    issue: Mapping[str, Any],
    reviewer_role: str,
    reviewer_ref: str,
    as_of_date: date,
    catalog_path: Path,
    repair: Mapping[str, Any],
) -> dict[str, Any]:
    row_id = str(issue["row_id"])
    action = _resolved_owner_action(issue)
    candidates = {str(item["candidate_id"]): item for item in issue.get("candidates") or ()}
    if action != "accept_qualify":
        return {
            "status": "knowledge_gap",
            "reason_code": (
                issue.get("cannot_fill_reason") or "owner_rejected_or_unfilled_final_check"
            ),
            "exact_gold_spans": [],
            "contrary_authority_status": "unresolved",
        }
    selected_ids = _selected_candidate_ids(issue)
    later_treatment = str(issue.get("owner_later_treatment") or "").strip()
    if later_treatment == "not_applicable_to_non_case_candidate":
        later_treatment = ""
    legal_role = str(issue.get("owner_legal_role") or "").strip()
    proposition = _collapsed(str(issue.get("owner_exact_proposition") or ""))
    spans: list[dict[str, Any]] = []
    try:
        if not selected_ids or any(candidate_id not in candidates for candidate_id in selected_ids):
            raise ValueError(f"{row_id} accept_qualify has no proposed candidate")
        for candidate_id in selected_ids:
            candidate = candidates[candidate_id]
            source_type = str(candidate["source_type"])
            if source_type == "case":
                if later_treatment not in CASE_LATER_TREATMENT:
                    raise ValueError(
                        f"{row_id} case-law qualification has unresolved later treatment"
                    )
                role = legal_role
                if role not in CASE_LEGAL_ROLES:
                    raise ValueError(f"{row_id} case-law qualification needs a material legal role")
                bound_proposition = proposition or _collapsed(str(candidate.get("excerpt") or ""))
                if not bound_proposition:
                    raise ValueError(f"{row_id} case-law qualification needs an exact proposition")
                review = _case_review(
                    candidate=candidate,
                    proposition_hash=_sha256_text(bound_proposition),
                    legal_role=role,
                    later_treatment=later_treatment,
                    contrary_ids=(),
                    reviewer_role=reviewer_role,
                    reviewer_ref=reviewer_ref,
                    review_scope="ordinary",
                    second_review_status="not_required",
                    second_reviewer_ref=None,
                    as_of_date=as_of_date,
                )
                spans.append(
                    _owner_gold_span(
                        candidate=candidate,
                        issue_id=str(issue["issue_id"]),
                        legal_role=role,
                        proposition_hash=_sha256_text(bound_proposition),
                        currentness_review=review,
                        catalog_path=catalog_path,
                        repair=repair,
                    )
                )
                continue
            if candidate.get("cross_ref_fragment") is True:
                raise ValueError(f"{row_id} cannot qualify a cross-reference fragment")
            if source_type == "legislation":
                role = legal_role or "statutory_text"
                if role != "statutory_text":
                    raise ValueError(f"{row_id} legislation candidate must be statutory_text")
            else:
                role = "secondary_commentary"
            if later_treatment or proposition:
                raise ValueError(
                    f"{row_id} non-case qualification must leave case-review fields blank"
                )
            spans.append(
                _owner_gold_span(
                    candidate=candidate,
                    issue_id=str(issue["issue_id"]),
                    legal_role=role,
                    proposition_hash=None,
                    currentness_review=None,
                    catalog_path=catalog_path,
                    repair=repair,
                )
            )
    except ValueError:
        return {
            "status": "knowledge_gap",
            "reason_code": "exact_span_verification_failed",
            "exact_gold_spans": [],
            "contrary_authority_status": "unresolved",
        }
    return {
        "status": "qualified",
        "reason_code": None,
        "exact_gold_spans": spans,
        "contrary_authority_status": "unresolved",
    }


def import_owner_final_check(
    *,
    project_root: Path,
    catalog_path: Path,
    pack_path: Path,
    review_export_path: Path,
    as_of_date: date,
    repair_payload: Mapping[str, Any] | None = None,
    confirmation_token: str | None = None,
) -> dict[str, Any]:
    """Bind owner-accepted JSON hashes. Word files are ignored."""

    if pack_path.suffix.casefold() == ".docx":
        raise ValueError("final-check import reads JSON only; Word is display-only")
    pack = json.loads(pack_path.read_text(encoding="utf-8"))
    token = str(pack.get("owner_confirmation_token") or confirmation_token or "")
    if pack.get("schema") == SUBSTANTIVE_REVIEW_SCHEMA:
        pack = adapt_substantive_review_for_import(
            pack,
            confirmation_token=token,
        )
    elif token == OWNER_FINAL_CHECK_TOKEN and not str(pack.get("owner_confirmation_token") or ""):
        pack["owner_confirmation_token"] = token
    if pack.get("schema") != FINAL_CHECK_SCHEMA:
        raise ValueError("final-check import requires the owner final-check JSON")
    if str(pack.get("as_of_date")) != as_of_date.isoformat():
        raise ValueError("final-check pack has a different as-of date")
    if str(pack.get("owner_confirmation_token") or "") != OWNER_FINAL_CHECK_TOKEN:
        raise ValueError("owner final-check confirmation token is missing")

    export = json.loads(review_export_path.read_text(encoding="utf-8"))
    if export.get("schema") != REVIEW_EXPORT_SCHEMA:
        raise ValueError("final-check import requires a Path-B review export")
    expected_export = dict(export)
    expected_export.pop("export_sha256", None)
    if sealed_sha256(expected_export) != export.get("export_sha256"):
        raise ValueError("Path-B review export seal does not match its contents")
    if str(export.get("as_of_date")) != as_of_date.isoformat():
        raise ValueError("Path-B review export has a different as-of date")

    identity = build_owner_reviewer_identity(as_of_date=as_of_date)
    reviewer_role = str(pack.get("reviewer_role") or "") or str(identity["approval_reviewer_role"])
    reviewer_ref = str(pack.get("reviewer_ref") or "") or str(identity["approval_reviewer_ref"])
    if reviewer_role != identity["approval_reviewer_role"]:
        raise ValueError("final-check reviewer role is invalid")
    if reviewer_ref != identity["approval_reviewer_ref"]:
        raise ValueError("final-check reviewer ref is invalid")

    repair = _load_repair(
        project_root=project_root,
        catalog_path=catalog_path,
        repair_payload=repair_payload,
    )
    owner_rows: dict[str, dict[str, Any]] = {}
    for case in pack.get("cases", ()):
        for issue in case.get("issues", ()):
            owner_rows[str(issue["row_id"])] = _parse_final_check_issue(
                issue=issue,
                reviewer_role=reviewer_role,
                reviewer_ref=reviewer_ref,
                as_of_date=as_of_date,
                catalog_path=catalog_path,
                repair=repair,
            )
    if len(owner_rows) != EXPECTED_SELECTED_ISSUES:
        raise ValueError("final-check JSON does not disposition all 305 selected issues")

    selected_ids = {str(case["case_id"]) for case in pack.get("cases", ())}
    reviewed_rows = []
    for source in export.get("rows", ()):
        row_id = str(source["row_id"])
        if source["case_id"] in selected_ids:
            marking = owner_rows[row_id]
        else:
            marking = {
                "status": "knowledge_gap",
                "reason_code": "coverage_only_explicit_knowledge_gap",
                "exact_gold_spans": [],
                "contrary_authority_status": "unresolved",
            }
        reviewed_rows.append(
            {
                "row_id": row_id,
                "case_id": source["case_id"],
                "issue_id": source["issue_id"],
                "status": marking["status"],
                "reason_code": marking["reason_code"],
                "exact_gold_spans": marking["exact_gold_spans"],
                "contrary_authority_status": marking["contrary_authority_status"],
            }
        )
    if len(reviewed_rows) != 585:
        raise ValueError("final-check import does not disposition all 585 issues")
    qualified = sum(
        row["status"] == "qualified" and row["case_id"] in selected_ids for row in reviewed_rows
    )
    payload = {
        "schema": REVIEW_IMPORT_SCHEMA,
        "export_sha256": export["export_sha256"],
        "suite_id": export["suite_id"],
        "as_of_date": export["as_of_date"],
        "suite_registry_canonical_sha256": export["suite_registry_canonical_sha256"],
        "run_plan_sha256": export["run_plan_sha256"],
        "reviewer_identity": "owner_primary_reviewer",
        "row_count": len(reviewed_rows),
        "selected_qualified_issue_count": qualified,
        "selected_knowledge_gap_issue_count": EXPECTED_SELECTED_ISSUES - qualified,
        "rows": reviewed_rows,
        "owner_confirmation_present": True,
        "import_surface": "json",
        "word_is_import_surface": False,
        "held_statute_owner_actions": [
            {
                "held_id": item["held_id"],
                "proposed_action": item["proposed_action"],
                "owner_action": item.get("owner_action") or "",
            }
            for item in pack.get("held_statutes", ())
        ],
        "seals_expert_gold": False,
        "selected_evidence_complete": qualified == EXPECTED_SELECTED_ISSUES,
        "ready_for_overlay_reconstruction": True,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }
    assert_safe_evaluation_payload({key: value for key, value in payload.items() if key != "rows"})
    return payload


def write_owner_final_check_reviews(
    *,
    project_root: Path,
    catalog_path: Path,
    pack_path: Path,
    review_export_path: Path,
    output_path: Path,
    as_of_date: date,
    overwrite: bool = False,
    repair_payload: Mapping[str, Any] | None = None,
    confirmation_token: str | None = None,
) -> dict[str, Any]:
    destination = output_path.resolve()
    if destination.exists() and not overwrite:
        raise FileExistsError("reviewed final-check rows already exist")
    payload = import_owner_final_check(
        project_root=project_root,
        catalog_path=catalog_path,
        pack_path=pack_path,
        review_export_path=review_export_path,
        as_of_date=as_of_date,
        repair_payload=repair_payload,
        confirmation_token=confirmation_token,
    )
    raw = _write_json(destination, payload)
    return {
        "schema": FINAL_CHECK_IMPORT_RESULT_SCHEMA,
        "row_count": payload["row_count"],
        "selected_qualified_issue_count": payload["selected_qualified_issue_count"],
        "selected_knowledge_gap_issue_count": payload["selected_knowledge_gap_issue_count"],
        "selected_evidence_complete": payload["selected_evidence_complete"],
        "ready_for_overlay_seal": payload["ready_for_overlay_seal"],
        "reviewed_rows_sha256": _sha256_bytes(raw),
        "import_surface": "json",
        "word_is_import_surface": False,
        "seals_expert_gold": False,
        "generation_authorised": False,
        "writes_active": False,
        "writes_o04": False,
    }
