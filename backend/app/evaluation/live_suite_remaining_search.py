"""Owner search questions for remaining Path-B selected issues.

This is a worksheet, not gold. Word is display-only. It does not import
spans, seal an overlay, write ACTIVE, or issue O-04.
"""

from __future__ import annotations

import json
from collections import Counter
from collections.abc import Mapping, Sequence
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType

from .live30 import assert_safe_evaluation_payload
from .live_suite import load_live_evaluation_bundle
from .live_suite_path_b import REVIEW_IMPORT_SCHEMA, selected_generation_case_ids
from .review_docx import (
    _add_header_footer,
    _add_table,
    _add_toc,
    _configure_page,
    _configure_styles,
    _finalize_document_properties,
    _reject_prohibited_metadata,
)

REMAINING_SEARCH_SCHEMA = "legalbot.live60-remaining-selected-search.v1"
REMAINING_SEARCH_RESULT_SCHEMA = "legalbot.live60-remaining-selected-search-result.v1"
LEGISLATION_SITE = "legislation.gov.uk"
CASELAW_SITE = "caselaw.nationalarchives.gov.uk"

SUBJECT_HOME: dict[str, tuple[str, str]] = {
    "consumer": (LEGISLATION_SITE, "Consumer Rights Act 2015"),
    "tort": (CASELAW_SITE, "duty of care"),
    "criminal evidence": (LEGISLATION_SITE, "Police and Criminal Evidence Act 1984"),
    "employment and equality": (LEGISLATION_SITE, "Equality Act 2010"),
    "equity and trusts": (CASELAW_SITE, "trust certainty"),
    "intellectual property": (LEGISLATION_SITE, "Copyright, Designs and Patents Act 1988"),
    "wills and succession": (LEGISLATION_SITE, "Wills Act 1837"),
    "corporate governance": (LEGISLATION_SITE, "Companies Act 2006"),
    "medical": (LEGISLATION_SITE, "Mental Capacity Act 2005"),
    "public procurement and administrative": (LEGISLATION_SITE, "Procurement Act 2023"),
    "data protection and privacy": (LEGISLATION_SITE, "UK GDPR"),
    "legal ethics and artificial intelligence": (CASELAW_SITE, "solicitor duty to the court"),
    "construction and commercial": (CASELAW_SITE, "contract interpretation"),
    "constitutional and administrative": (CASELAW_SITE, "judicial review"),
    "corporate fraud regulation and litigation": (LEGISLATION_SITE, "Fraud Act 2006"),
    "agency and commercial contracts": (CASELAW_SITE, "agency actual apparent authority"),
    "defamation": (LEGISLATION_SITE, "Defamation Act 2013"),
    "residential tenancies and housing": (LEGISLATION_SITE, "Housing Act 1988"),
    "sale of goods and retention of title": (LEGISLATION_SITE, "Sale of Goods Act 1979"),
    "planning and judicial review": (LEGISLATION_SITE, "Town and Country Planning Act 1990"),
    "legal professional privilege": (CASELAW_SITE, "legal professional privilege"),
    "international arbitration": (LEGISLATION_SITE, "Arbitration Act 1996"),
    "education": (LEGISLATION_SITE, "Education Act 1996"),
    "cybercrime": (LEGISLATION_SITE, "Computer Misuse Act 1990"),
    "shipping and carriage of goods": (LEGISLATION_SITE, "Carriage of Goods by Sea Act 1971"),
    "privacy media and confidential information": (CASELAW_SITE, "misuse of private information"),
    "sports law and arbitration": (LEGISLATION_SITE, "Arbitration Act 1996"),
    "surveillance and national security": (LEGISLATION_SITE, "Investigatory Powers Act 2016"),
    "collective redress": (LEGISLATION_SITE, "Civil Procedure Rules"),
    "crypto exchange collapse": (LEGISLATION_SITE, "Insolvency Act 1986"),
}

# Look-here first: search targets only, not gold. Keyed by "subject|topic".
LOOK_HERE: dict[str, tuple[str, str]] = {
    "consumer|pre-contract statements": (LEGISLATION_SITE, "Consumer Rights Act 2015 section 50"),
    "consumer|statutory remedies": (LEGISLATION_SITE, "Consumer Rights Act 2015 sections 19 to 24"),
    "consumer|repair and replacement": (
        LEGISLATION_SITE,
        "Consumer Rights Act 2015 sections 23 and 43",
    ),
    "consumer|rejection": (LEGISLATION_SITE, "Consumer Rights Act 2015 sections 20, 22 and 24"),
    "consumer|restriction of consumer rights": (
        LEGISLATION_SITE,
        "Consumer Rights Act 2015 sections 31 and 47",
    ),
    "tort|incremental duty development": (
        CASELAW_SITE,
        "Robinson v Chief Constable of West Yorkshire Police [2018] UKSC 4",
    ),
    "tort|assumption of responsibility": (
        CASELAW_SITE,
        "Hedley Byrne & Co Ltd v Heller & Partners Ltd [1964] AC 465",
    ),
    "tort|omissions": (CASELAW_SITE, "Stovin v Wise [1996] AC 923"),
    "tort|public-authority liability": (
        CASELAW_SITE,
        "Robinson v Chief Constable of West Yorkshire Police [2018] UKSC 4",
    ),
    "tort|pure economic loss": (
        CASELAW_SITE,
        "Spartan Steel & Alloys Ltd v Martin & Co (Contractors) Ltd [1973] 1 QB 27",
    ),
    "tort|foreseeability and proximity": (
        CASELAW_SITE,
        "Caparo Industries plc v Dickman [1990] 2 AC 605",
    ),
    "tort|fairness and precedent": (
        CASELAW_SITE,
        "Robinson v Chief Constable of West Yorkshire Police [2018] UKSC 4",
    ),
    "criminal evidence|confessions": (
        LEGISLATION_SITE,
        "Police and Criminal Evidence Act 1984 section 76",
    ),
    "criminal evidence|oppression and unreliability": (
        LEGISLATION_SITE,
        "Police and Criminal Evidence Act 1984 section 76",
    ),
    "criminal evidence|improperly obtained evidence": (
        LEGISLATION_SITE,
        "Police and Criminal Evidence Act 1984 section 78",
    ),
    "criminal evidence|fairness discretion": (
        LEGISLATION_SITE,
        "Police and Criminal Evidence Act 1984 section 78",
    ),
    "criminal evidence|identification evidence": (CASELAW_SITE, "R v Turnbull [1977] QB 224"),
    "criminal evidence|hearsay": (
        LEGISLATION_SITE,
        "Criminal Justice Act 2003 sections 114 to 118",
    ),
    "criminal evidence|fact-finding and public confidence": (
        LEGISLATION_SITE,
        "Criminal Justice Act 2003 section 114",
    ),
    "employment and equality|unfair dismissal and redundancy": (
        LEGISLATION_SITE,
        "Employment Rights Act 1996 sections 94, 98 and 139",
    ),
    "employment and equality|direct discrimination": (
        LEGISLATION_SITE,
        "Equality Act 2010 section 13",
    ),
    "employment and equality|indirect discrimination": (
        LEGISLATION_SITE,
        "Equality Act 2010 section 19",
    ),
    "employment and equality|disability discrimination": (
        LEGISLATION_SITE,
        "Equality Act 2010 sections 6, 15 and 20",
    ),
    "employment and equality|reasonable adjustments": (
        LEGISLATION_SITE,
        "Equality Act 2010 section 20",
    ),
    "employment and equality|victimisation": (LEGISLATION_SITE, "Equality Act 2010 section 27"),
    "employment and equality|justification": (LEGISLATION_SITE, "Equality Act 2010 section 19"),
    "employment and equality|evidential burdens": (
        LEGISLATION_SITE,
        "Equality Act 2010 section 136",
    ),
    "employment and equality|remedies": (LEGISLATION_SITE, "Equality Act 2010 section 124"),
    "equity and trusts|certainty": (CASELAW_SITE, "Knight v Knight (1840) 3 Beav 148"),
    "equity and trusts|constitution and imperfect gifts": (
        CASELAW_SITE,
        "Milroy v Lord (1862) 4 De G F & J 264",
    ),
    "equity and trusts|testamentary trusts": (LEGISLATION_SITE, "Wills Act 1837 section 9"),
    "equity and trusts|fiduciary duties": (
        CASELAW_SITE,
        "Bristol and West Building Society v Mothew [1998] Ch 1",
    ),
    "equity and trusts|breach of trust": (LEGISLATION_SITE, "Trustee Act 1925 section 61"),
    "equity and trusts|assignment of equitable interests": (
        LEGISLATION_SITE,
        "Law of Property Act 1925 section 53",
    ),
    "equity and trusts|tracing mixed funds": (CASELAW_SITE, "Foskett v McKeown [2001] 1 AC 102"),
    "equity and trusts|proprietary and personal remedies": (
        CASELAW_SITE,
        "Foskett v McKeown [2001] 1 AC 102",
    ),
    "equity and trusts|third-party liability": (
        CASELAW_SITE,
        "Royal Brunei Airlines Sdn Bhd v Tan [1995] 2 AC 378",
    ),
    "intellectual property|employee-created works": (
        LEGISLATION_SITE,
        "Copyright, Designs and Patents Act 1988 section 11",
    ),
    "intellectual property|originality and substantial copying": (
        LEGISLATION_SITE,
        "Copyright, Designs and Patents Act 1988 sections 1 and 16",
    ),
    "intellectual property|computer programs": (
        LEGISLATION_SITE,
        "Copyright, Designs and Patents Act 1988 sections 3 and 50B",
    ),
    "intellectual property|confidential information": (
        CASELAW_SITE,
        "Coco v AN Clark (Engineers) Ltd [1968] FSR 415",
    ),
    "intellectual property|licensing": (
        LEGISLATION_SITE,
        "Copyright, Designs and Patents Act 1988 section 90",
    ),
    "intellectual property|remedies": (
        LEGISLATION_SITE,
        "Copyright, Designs and Patents Act 1988 sections 96 and 97",
    ),
    "wills and succession|execution formalities": (LEGISLATION_SITE, "Wills Act 1837 section 9"),
    "wills and succession|testamentary capacity": (
        CASELAW_SITE,
        "Banks v Goodfellow (1870) LR 5 QB 549",
    ),
    "wills and succession|knowledge and approval": (
        CASELAW_SITE,
        "Gill v Woodall [2010] EWCA Civ 1430",
    ),
    "wills and succession|undue influence": (
        CASELAW_SITE,
        "Edwards v Edwards [2007] EWHC 1119 (Ch)",
    ),
    "wills and succession|revocation and interpretation": (
        LEGISLATION_SITE,
        "Wills Act 1837 sections 18 to 20",
    ),
    "wills and succession|intestacy": (
        LEGISLATION_SITE,
        "Administration of Estates Act 1925 section 46",
    ),
    "wills and succession|reasonable financial provision": (
        LEGISLATION_SITE,
        "Inheritance (Provision for Family and Dependants) Act 1975 section 1",
    ),
    "corporate governance|proper purposes": (LEGISLATION_SITE, "Companies Act 2006 section 171"),
    "corporate governance|stakeholder considerations": (
        LEGISLATION_SITE,
        "Companies Act 2006 section 172",
    ),
    "corporate governance|creditor interests": (LEGISLATION_SITE, "Companies Act 2006 section 172"),
    "corporate governance|ratification and derivative claims": (
        LEGISLATION_SITE,
        "Companies Act 2006 sections 239 and 260",
    ),
    "corporate governance|enforcement limits": (LEGISLATION_SITE, "Companies Act 2006 section 170"),
    "medical|capacity": (LEGISLATION_SITE, "Mental Capacity Act 2005 sections 1 to 3"),
    "medical|consent": (CASELAW_SITE, "Montgomery v Lanarkshire Health Board [2015] UKSC 11"),
    "medical|best interests and emergency treatment": (
        LEGISLATION_SITE,
        "Mental Capacity Act 2005 sections 4 and 5",
    ),
    "medical|risk disclosure": (
        CASELAW_SITE,
        "Montgomery v Lanarkshire Health Board [2015] UKSC 11",
    ),
    "medical|clinical negligence": (
        CASELAW_SITE,
        "Bolam v Friern Hospital Management Committee [1957] 1 WLR 582",
    ),
    "medical|causation": (CASELAW_SITE, "Chester v Afshar [2004] UKHL 41"),
    "medical|confidentiality": (CASELAW_SITE, "W v Egdell [1990] Ch 359"),
    "data protection and privacy|lawful processing and transparency": (
        LEGISLATION_SITE,
        "UK GDPR Articles 5 and 6",
    ),
    "data protection and privacy|special-category inferences": (
        LEGISLATION_SITE,
        "UK GDPR Article 9",
    ),
    "data protection and privacy|profiling and automated decisions": (
        LEGISLATION_SITE,
        "UK GDPR Article 22",
    ),
    "data protection and privacy|security": (LEGISLATION_SITE, "UK GDPR Article 32"),
    "data protection and privacy|breach notification": (
        LEGISLATION_SITE,
        "UK GDPR Articles 33 and 34",
    ),
    "data protection and privacy|access and explanation rights": (
        LEGISLATION_SITE,
        "UK GDPR Articles 15 and 22",
    ),
    "data protection and privacy|compensation and enforcement": (
        LEGISLATION_SITE,
        "Data Protection Act 2018 section 168",
    ),
    "construction and commercial|liquidated damages": (
        CASELAW_SITE,
        "Triple Point Technology Inc v PTT Public Company Ltd [2021] UKSC 29",
    ),
    "defamation|serious harm": (LEGISLATION_SITE, "Defamation Act 2013 section 1"),
    "defamation|truth": (LEGISLATION_SITE, "Defamation Act 2013 section 2"),
    "defamation|honest opinion": (LEGISLATION_SITE, "Defamation Act 2013 section 3"),
    "defamation|publication on matter of public interest": (
        LEGISLATION_SITE,
        "Defamation Act 2013 section 4",
    ),
    "sale of goods and retention of title|passing of property": (
        LEGISLATION_SITE,
        "Sale of Goods Act 1979 sections 17 to 19",
    ),
    "sale of goods and retention of title|retention of title": (
        CASELAW_SITE,
        "Aluminium Industrie Vaassen BV v Romalpa Aluminium Ltd [1976] 1 WLR 676",
    ),
    "legal professional privilege|legal advice privilege": (
        CASELAW_SITE,
        "Three Rivers District Council v Bank of England (No 6) [2004] UKHL 48",
    ),
    "legal professional privilege|litigation privilege": (
        CASELAW_SITE,
        "Waugh v British Railways Board [1980] AC 521",
    ),
    "international arbitration|stay of legal proceedings": (
        LEGISLATION_SITE,
        "Arbitration Act 1996 section 9",
    ),
    "international arbitration|separability": (LEGISLATION_SITE, "Arbitration Act 1996 section 7"),
    "cybercrime|unauthorised access": (LEGISLATION_SITE, "Computer Misuse Act 1990 section 1"),
    "shipping and carriage of goods|seaworthiness": (
        LEGISLATION_SITE,
        "Carriage of Goods by Sea Act 1971 Article III",
    ),
    "privacy media and confidential information|misuse of private information": (
        CASELAW_SITE,
        "Campbell v MGN Ltd [2004] UKHL 22",
    ),
    "privacy media and confidential information|breach of confidence": (
        CASELAW_SITE,
        "Coco v AN Clark (Engineers) Ltd [1968] FSR 415",
    ),
    "surveillance and national security|targeted and bulk interception": (
        LEGISLATION_SITE,
        "Investigatory Powers Act 2016 Parts 2 and 6",
    ),
    "surveillance and national security|communications data": (
        LEGISLATION_SITE,
        "Investigatory Powers Act 2016 Part 3",
    ),
    "surveillance and national security|equipment interference": (
        LEGISLATION_SITE,
        "Investigatory Powers Act 2016 Part 5",
    ),
    "collective redress|group litigation orders": (
        LEGISLATION_SITE,
        "Civil Procedure Rules Part 19",
    ),
    "collective redress|representative proceedings": (
        LEGISLATION_SITE,
        "Civil Procedure Rules rule 19.8",
    ),
    "collective redress|competition collective proceedings": (
        LEGISLATION_SITE,
        "Competition Act 1998 section 47B",
    ),
    "crypto exchange collapse|trust or contractual debt": (
        CASELAW_SITE,
        "Re Lehman Brothers International (Europe) [2012] UKSC 6",
    ),
    "crypto exchange collapse|transactions at an undervalue and preferences": (
        LEGISLATION_SITE,
        "Insolvency Act 1986 sections 238 and 239",
    ),
    "crypto exchange collapse|wrongful and fraudulent trading": (
        LEGISLATION_SITE,
        "Insolvency Act 1986 sections 213 and 214",
    ),
}


def _collapsed(value: str) -> str:
    return " ".join(value.split())


def _safe(value: str, *, label: str) -> str:
    return _reject_prohibited_metadata(_collapsed(value), label=label)


def _source_names(issue: Mapping[str, Any]) -> tuple[str, ...]:
    names: list[str] = []
    seen: set[str] = set()
    for item in issue.get("candidates") or ():
        name = _collapsed(str(item.get("source_name") or ""))
        if name and name not in seen:
            seen.add(name)
            names.append(name)
    return tuple(names)


def _rejected_locators(issue: Mapping[str, Any]) -> tuple[str, ...]:
    locators: list[str] = []
    seen: set[str] = set()
    for item in issue.get("candidates") or ():
        locator = _collapsed(str(item.get("legal_locator") or ""))
        if locator and locator not in seen:
            seen.add(locator)
            locators.append(locator)
    return tuple(locators)


def _kind_and_open(
    *,
    subject: str,
    topic: str,
    sources: Sequence[str],
) -> tuple[str, str]:
    override = LOOK_HERE.get(f"{subject}|{topic}")
    if override is not None:
        return override
    site, home = SUBJECT_HOME.get(
        subject, (LEGISLATION_SITE, "England and Wales primary authority")
    )
    if sources:
        return site, sources[0]
    return site, home


def _search_question(*, subject: str, topic: str, open_name: str, site: str) -> str:
    if site == CASELAW_SITE:
        return (
            f"What is the current UK Supreme Court, House of Lords or other binding "
            f"England-and-Wales paragraph that states the legal rule on {topic} in {subject}?"
        )
    return (
        f"What is the current official England-and-Wales subsection in {open_name} "
        f"that states the legal rule on {topic}?"
    )


def _paste_query(*, site: str, open_name: str, topic: str) -> str:
    if site == CASELAW_SITE:
        return f'site:{site} "{topic}" {open_name}'
    return f'site:{site} {open_name} "{topic}"'


def build_remaining_search_rows(
    *,
    project_root: Path,
    imported: Mapping[str, Any],
    draft: Mapping[str, Any],
) -> dict[str, Any]:
    if imported.get("schema") != REVIEW_IMPORT_SCHEMA:
        raise ValueError("remaining search pack requires imported Path-B rows")
    bundle = load_live_evaluation_bundle(
        project_root / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    selected_ids = set(selected_generation_case_ids(bundle))
    by_issue = {
        str(issue["row_id"]): (case, issue)
        for case in draft.get("cases", ())
        for issue in case.get("issues", ())
    }
    rows: list[dict[str, Any]] = []
    for item in imported.get("rows", ()):
        case_id = str(item["case_id"])
        if case_id not in selected_ids or item.get("status") == "qualified":
            continue
        row_id = str(item["row_id"])
        case, issue = by_issue[row_id]
        subject = _safe(str(case.get("subject") or ""), label="subject")
        topic = _safe(str(issue.get("topic") or ""), label="topic")
        action = str(issue.get("owner_action") or item.get("reason_code") or "knowledge_gap")
        sources = _source_names(issue)
        locators = _rejected_locators(issue)
        site, open_name = _kind_and_open(subject=subject, topic=topic, sources=sources)
        open_name = _safe(open_name, label="search target")
        question = _safe(
            _search_question(subject=subject, topic=topic, open_name=open_name, site=site),
            label="search question",
        )
        extra = ""
        if row_id == "live30-q26:issue-04":
            extra = (
                "This row stayed a gap because qualified_current needs bound "
                "limiting-authority IDs. Either find a confirmed_current paragraph "
                "or keep it as a gap until those IDs exist."
            )
        elif site == CASELAW_SITE:
            extra = (
                "If the answer is case law, copy one exact paragraph and record "
                "later treatment as confirmed_current or qualified_current. "
                "qualified_current also needs limiting-authority IDs."
            )
        elif "Inheritance (Provision for Family and Dependants) Act 1975" in open_name or (
            "Trustee Act 2000 section 1" in open_name
        ):
            extra = (
                "This provision is still on hold. Search it, but do not treat the hold as lifted."
            )
        rows.append(
            {
                "row_id": row_id,
                "case_id": case_id,
                "issue_id": str(issue["issue_id"]),
                "subject": subject,
                "topic": topic,
                "owner_action": action,
                "search_site": site,
                "open_first": open_name,
                "search_question": question,
                "paste_query": _safe(
                    _paste_query(site=site, open_name=open_name, topic=topic),
                    label="paste query",
                ),
                "rejected_locators": list(locators),
                "rejected_sources": list(sources),
                "copy_instruction": (
                    "Copy the operative subsection or paragraph that states the rule. "
                    "Do not copy a heading, contents line, cross-reference shell, or definition dump."
                ),
                "extra_instruction": extra,
                "word_is_gold": False,
                "seals_expert_gold": False,
            }
        )
    selected_issue_ids = {str(issue["row_id"]) for case, issue in by_issue.values()}
    qualified_selected = {
        str(item["row_id"])
        for item in imported.get("rows", ())
        if item.get("status") == "qualified" and str(item["row_id"]) in selected_issue_ids
    }
    if {row["row_id"] for row in rows} != selected_issue_ids - qualified_selected:
        raise ValueError("remaining search pack does not cover every remaining selected issue")
    if not rows:
        raise ValueError("there are no remaining selected issues")
    counts = Counter(row["owner_action"] for row in rows)
    payload = {
        "schema": REMAINING_SEARCH_SCHEMA,
        "as_of_date": imported.get("as_of_date"),
        "selected_remaining_issue_count": len(rows),
        "action_counts": dict(counts),
        "word_is_gold": False,
        "word_is_import_surface": False,
        "seals_expert_gold": False,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "writes_active": False,
        "writes_o04": False,
        "ai_role": "mechanical_accuracy_verifier_only",
        "rows": rows,
    }
    assert_safe_evaluation_payload({key: value for key, value in payload.items() if key != "rows"})
    return payload


def build_remaining_search_document(pack: Mapping[str, Any]) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, f"Live60-{pack['as_of_date']}")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph(
        f"{pack['selected_remaining_issue_count']} remaining selected issues: search questions",
        style="Title",
    )
    document.add_paragraph(
        _safe(
            "This Word file is a search worksheet only. It is not gold, not an import "
            "surface, not an overlay seal, not ACTIVE and not O-04. For each remaining "
            "selected issue there is one question to search on official England-and-Wales "
            "sources. Paste the query, open the named instrument or judgment first, and "
            "copy the operative subsection or paragraph that answers the question.",
            label="cover",
        )
    )
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Remaining selected issues", str(pack["selected_remaining_issue_count"])),
            (
                "Request another candidate",
                str(pack["action_counts"].get("request_another_candidate", 0)),
            ),
            ("Reject / keep gap", str(pack["action_counts"].get("reject_keep_gap", 0))),
            (
                "Approved but failed exact match",
                str(pack["action_counts"].get("approve_qualified", 0)),
            ),
            ("Word is gold", "false"),
            ("Word is import surface", "false"),
            ("Look-here names are gold", "false"),
        ),
        (3_120, 6_240),
    )
    document.add_heading("How to use each question", level=1)
    for line in (
        "Search only legislation.gov.uk or Find Case Law. Do not use commentary as the span.",
        "The named instrument or case is a search start, not a bound gold span.",
        "Copy operative words. Reject headings, contents lists, and cross-reference shells.",
        "Rejected locators listed below already failed. Find a different exact text.",
        "Case-law answers also need later treatment before they can qualify.",
        "After you have the exact official text, it still has to match the local catalogue hash to import.",
    ):
        document.add_paragraph(_safe(line, label="instruction"), style="List Bullet")
    _add_toc(document)
    document.add_heading(
        f"All {pack['selected_remaining_issue_count']} search questions",
        level=1,
    )
    _add_table(
        document,
        ("No.", "Issue", "Search this question"),
        tuple(
            (str(number), row["row_id"], row["search_question"])
            for number, row in enumerate(pack["rows"], start=1)
        ),
        (700, 2_400, 6_260),
        body_size=8,
    )
    current_case = ""
    for number, row in enumerate(pack["rows"], start=1):
        if row["case_id"] != current_case:
            current_case = str(row["case_id"])
            document.add_heading(f"{row['case_id']} · {row['subject']}", level=1)
        document.add_heading(f"Q{number:03d}  {row['issue_id']}  {row['topic']}", level=2)
        document.add_paragraph(_safe(row["search_question"], label="question"))
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("Paste this search", row["paste_query"]),
                ("Search on", row["search_site"]),
                ("Open this first", row["open_first"]),
                ("Why still open", row["owner_action"]),
                (
                    "Do not reuse these locators",
                    "; ".join(row["rejected_locators"]) or "none in the last pack",
                ),
                ("Copy", row["copy_instruction"]),
                ("Extra", row["extra_instruction"] or "none"),
            ),
            (2_400, 6_960),
            body_size=8.5,
        )
    _finalize_document_properties(
        document,
        title="Live60 Path-B remaining selected search questions",
        subject="Owner search worksheet for 293 remaining selected issues",
    )
    return document


def export_remaining_search_pack(
    *,
    project_root: Path,
    imported_path: Path,
    draft_path: Path,
    output_path: Path,
    overwrite: bool = False,
) -> dict[str, Any]:
    destination = output_path.resolve()
    if destination.exists() and not overwrite:
        raise FileExistsError("remaining search document already exists")
    imported = json.loads(imported_path.read_text(encoding="utf-8"))
    draft = json.loads(draft_path.read_text(encoding="utf-8"))
    pack = build_remaining_search_rows(
        project_root=project_root,
        imported=imported,
        draft=draft,
    )
    document = build_remaining_search_document(pack)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    return {
        "schema": REMAINING_SEARCH_RESULT_SCHEMA,
        "selected_remaining_issue_count": pack["selected_remaining_issue_count"],
        "action_counts": pack["action_counts"],
        "output": destination.name,
        "word_is_gold": False,
        "word_is_import_surface": False,
        "seals_expert_gold": False,
        "writes_active": False,
        "writes_o04": False,
    }
