"""Owner-readable Path-B exact-evidence candidate packs.

Candidates come only from approved, rights-qualified authority lanes. They are
not legal gold: the owner must judge relevance, role, currentness, later
treatment and contrary authority before a separate reviewed-row import.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import sqlite3
from collections import Counter
from collections.abc import Mapping, Sequence
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType

from ..types import CasePropositionReview, case_proposition_review_sha256
from .live30 import assert_safe_evaluation_payload
from .live_suite import load_live_evaluation_bundle, sealed_sha256
from .live_suite_gold import LiveGoldSpan
from .live_suite_held_span_repair import build_held_span_contiguous_repair
from .live_suite_owner_decisions import export_held_provision_chunks
from .live_suite_path_b import (
    REVIEW_EXPORT_SCHEMA,
    REVIEW_IMPORT_SCHEMA,
    selected_generation_case_ids,
)
from .live_suite_reviewer_identity import build_owner_reviewer_identity
from .live_suite_tick_draft import SPLICED_PARENT_CHUNK_IDS
from .review_docx import (
    _add_header_footer,
    _add_page_break,
    _add_table,
    _add_toc,
    _configure_page,
    _configure_styles,
    _finalize_document_properties,
    _reject_prohibited_metadata,
)

EVIDENCE_PACK_SCHEMA = "legalbot.live60-owner-evidence-pack.v1"
EVIDENCE_PACK_MANIFEST_SCHEMA = "legalbot.live60-owner-evidence-pack-manifest.v1"
EVIDENCE_PACK_RESULT_SCHEMA = "legalbot.live60-owner-evidence-pack-result.v1"
EVIDENCE_REVIEW_IMPORT_RESULT_SCHEMA = "legalbot.live60-owner-evidence-review-import-result.v1"
EXPECTED_SELECTED_ISSUES = 305
PILOT_CASE_ID = "live30-q03"
CASE_CONFIRMATION_TOKEN = "CONFIRM_CASE_EVIDENCE_REVIEW"

_TOKEN = re.compile(r"[a-z0-9]+")
_SAFE_REVIEW_ID = re.compile(r"^[a-z0-9][a-z0-9._:-]{2,127}$")
_REVIEWER_REF = re.compile(r"^reviewer:[0-9a-f]{64}$")
_TITLE_STOP = frozenset(
    {
        "a",
        "and",
        "co",
        "company",
        "corporation",
        "inc",
        "limited",
        "llc",
        "ltd",
        "of",
        "plc",
        "the",
        "uk",
        "v",
    }
)
_RIGHTS_MARKERS = (
    "open government licence",
    "open supreme court licence",
    "open parliament licence",
    "creative commons attribution",
    "cc by",
)
_ALLOWED_LANES = frozenset({"primary_authority", "official_secondary", "scholarship"})
_ALLOWED_JURISDICTIONS = frozenset({"England and Wales", "United Kingdom"})

_CATALOGUE_ROW_SQL = """
SELECT c.id AS chunk_id, c.source_version_id, c.ordinal, c.locator,
       c.text_sha256, c.markdown_text, c.stream,
       sv.authority_identity_id, sv.title, sv.as_of_date,
       sv.currentness_status, sv.licence_name, sv.licence_url,
       sv.review_status, sv.canonical_url, sv.stable_identifier,
       sv.superseded_by,
       d.id AS document_id, d.content_sha256 AS official_snapshot_sha256,
       d.source_identity_id AS stable_source_id, d.status AS document_status,
       d.lane, d.jurisdiction, d.duplicate_of
FROM chunks c
JOIN source_versions sv ON sv.id = c.source_version_id
JOIN documents d ON d.id = sv.document_id
WHERE c.id = ?
"""


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _normalised_tokens(value: str, *, title: bool = False) -> tuple[str, ...]:
    tokens = tuple(_TOKEN.findall(value.casefold()))
    if title:
        return tuple(token for token in tokens if token not in _TITLE_STOP and not token.isdigit())
    return tokens


def _title_match_score(expected: str, actual: str) -> float:
    wanted = set(_normalised_tokens(expected, title=True))
    found = set(_normalised_tokens(actual, title=True))
    if not wanted or not found:
        return 0.0
    if wanted <= found or found <= wanted:
        return 1.0
    return len(wanted & found) / min(len(wanted), len(found))


def _rights_qualified(licence_name: str | None) -> bool:
    lowered = str(licence_name or "").casefold()
    return any(marker in lowered for marker in _RIGHTS_MARKERS)


def _source_type(value: str) -> str:
    lowered = value.casefold()
    if value == "case":
        return "case"
    if "guidance" in lowered:
        return "official_secondary"
    if "scholar" in lowered:
        return "scholarship"
    return "legislation"


def _proposed_legal_role(source_type: str) -> str:
    if source_type == "case":
        return "owner_must_select_holding_ratio_or_binding_legal_rule"
    if source_type == "official_secondary":
        return "secondary_commentary"
    if source_type == "scholarship":
        return "secondary_commentary"
    return "statutory_text"


def _canonical_duplicate_row(
    connection: sqlite3.Connection, row: sqlite3.Row
) -> sqlite3.Row | None:
    if row["document_status"] != "duplicate" or not row["duplicate_of"]:
        return row
    fetched = connection.execute(
        """
        SELECT c.id AS chunk_id, c.source_version_id, c.ordinal, c.locator,
               c.text_sha256, c.markdown_text, c.stream,
               sv.authority_identity_id, sv.title, sv.as_of_date,
               sv.currentness_status, sv.licence_name, sv.licence_url,
               sv.review_status, sv.canonical_url, sv.stable_identifier,
               sv.superseded_by,
               d.id AS document_id,
               d.content_sha256 AS official_snapshot_sha256,
               d.source_identity_id AS stable_source_id,
               d.status AS document_status, d.lane, d.jurisdiction,
               d.duplicate_of
        FROM documents d
        JOIN source_versions sv ON sv.document_id = d.id
        JOIN chunks c ON c.source_version_id = sv.id
        WHERE d.id = ?
          AND c.text_sha256 = ?
          AND COALESCE(c.stream, 'body') = 'body'
          AND sv.superseded_by IS NULL
        ORDER BY sv.as_of_date DESC, c.ordinal, c.id
        LIMIT 1
        """,
        (row["duplicate_of"], row["text_sha256"]),
    ).fetchone()
    if not isinstance(fetched, sqlite3.Row):
        return None
    return fetched


def _candidate_exclusion(
    row: sqlite3.Row | None,
    *,
    chunk_id: str,
    source_type: str,
    as_of_date: date,
) -> str | None:
    if row is None:
        return "catalogue_chunk_missing"
    if chunk_id in SPLICED_PARENT_CHUNK_IDS:
        return "rejected_spliced_parent"
    if row["document_status"] != "citable":
        return "non_citable_or_duplicate_without_canonical_span"
    if str(row["stream"] or "body") != "body":
        return "non_body_stream"
    if str(row["lane"] or "") not in _ALLOWED_LANES:
        return "non_authority_lane"
    if str(row["jurisdiction"] or "") not in _ALLOWED_JURISDICTIONS:
        return "jurisdiction_out_of_scope"
    if str(row["review_status"] or "") != "approved":
        return "source_not_approved"
    if row["superseded_by"] is not None:
        return "superseded_source_version"
    if not _rights_qualified(row["licence_name"]):
        return "full_text_rights_not_approved"
    currentness = str(row["currentness_status"] or "")
    if source_type == "case":
        if currentness not in {
            "historical",
            "point_in_time",
            "latest_available_revised_snapshot",
        }:
            return "case_currentness_metadata_unusable"
    elif currentness != "latest_available_revised_snapshot" and not (
        currentness == "point_in_time" and str(row["as_of_date"] or "") == as_of_date.isoformat()
    ):
        return "not_current_date_authority"
    excerpt = str(row["markdown_text"] or "")
    if str(row["text_sha256"] or "") != _sha256_text(excerpt):
        return "catalogue_hash_not_self_consistent"
    return None


def _catalogue_candidate(
    *,
    row: sqlite3.Row,
    case_id: str,
    issue_id: str,
    source_name: str,
    source_kind: str,
    candidate_origin: str,
    map_rank: str,
    map_pinpoint: str,
) -> dict[str, Any]:
    normalized_type = _source_type(source_kind)
    candidate_id = "candidate-" + _sha256_text(
        "|".join(
            (
                case_id,
                issue_id,
                str(row["source_version_id"]),
                str(row["chunk_id"]),
                str(row["text_sha256"]),
            )
        )
    )
    return {
        "candidate_id": candidate_id,
        "candidate_origin": candidate_origin,
        "source_name": " ".join(source_name.split()),
        "source_type": normalized_type,
        "map_rank": map_rank,
        "map_pinpoint": " ".join(map_pinpoint.split()),
        "authority_identity": str(row["authority_identity_id"] or ""),
        "jurisdiction": str(row["jurisdiction"]),
        "rights_status": "approved_full_text",
        "licence_name": str(row["licence_name"]),
        "source_version_id": str(row["source_version_id"]),
        "stable_source_id": str(row["stable_source_id"]),
        "official_snapshot_sha256": str(row["official_snapshot_sha256"]),
        "chunk_id": str(row["chunk_id"]),
        "legal_locator": str(row["locator"]),
        "content_sha256": str(row["text_sha256"]),
        "excerpt": str(row["markdown_text"]),
        "proposed_legal_role": _proposed_legal_role(normalized_type),
        "catalogue_currentness": str(row["currentness_status"]),
        "later_treatment_requirement": (
            "owner_proposition_and_later_treatment_review_required"
            if normalized_type == "case"
            else "not_applicable_to_non_case_candidate"
        ),
        "contrary_or_limiting_status": "not_machine_established",
        "canonical_url": str(row["canonical_url"] or ""),
        "seals_expert_gold": False,
    }


def _repair_candidate(
    *,
    item: Mapping[str, Any],
    case_id: str,
    issue_id: str,
    source_name: str,
    map_rank: str,
    map_pinpoint: str,
) -> dict[str, Any]:
    return {
        "candidate_id": (
            "candidate-"
            + _sha256_text(
                "|".join(
                    (
                        case_id,
                        issue_id,
                        str(item["repair_span_id"]),
                        str(item["text_sha256"]),
                    )
                )
            )
        ),
        "candidate_origin": "accepted_v2_repair_span",
        "source_name": " ".join(source_name.split()),
        "source_type": "legislation",
        "map_rank": map_rank,
        "map_pinpoint": " ".join(map_pinpoint.split()),
        "authority_identity": str(item["legal_authority_id"]),
        "jurisdiction": str(item["jurisdiction"]),
        "rights_status": "approved_official_repair_derivation",
        "licence_name": "Open Government Licence v3.0",
        "source_version_id": str(item["source_version_id"]),
        "stable_source_id": str(item["stable_source_id"]),
        "official_snapshot_sha256": str(item["official_snapshot_sha256"]),
        "chunk_id": str(item["repair_span_id"]),
        "parent_chunk_id": str(item["parent_chunk_id"]),
        "legal_locator": str(item["legal_locator"]),
        "content_sha256": str(item["text_sha256"]),
        "excerpt": str(item["markdown_text"]),
        "proposed_legal_role": "statutory_text",
        "catalogue_currentness": "current_official_bytes_repair_candidate",
        "later_treatment_requirement": "not_applicable_to_non_case_candidate",
        "contrary_or_limiting_status": "not_machine_established",
        "canonical_url": "",
        "derivation_manifest_sha256": str(item["derivation_manifest_sha256"]),
        "seals_expert_gold": False,
    }


def _eligible_source_versions(connection: sqlite3.Connection) -> list[sqlite3.Row]:
    return connection.execute(
        """
        SELECT sv.id AS source_version_id, sv.title, sv.canonical_url,
               sv.authority_identity_id, sv.currentness_status, sv.as_of_date,
               sv.licence_name, sv.review_status, sv.superseded_by,
               d.id AS document_id, d.status AS document_status,
               d.lane, d.jurisdiction
        FROM source_versions sv
        JOIN documents d ON d.id = sv.document_id
        WHERE sv.superseded_by IS NULL
          AND sv.review_status = 'approved'
          AND d.status = 'citable'
          AND d.lane IN ('primary_authority', 'official_secondary', 'scholarship')
          AND d.jurisdiction IN ('England and Wales', 'United Kingdom')
        ORDER BY sv.id
        """
    ).fetchall()


def _chunks_for_source(
    connection: sqlite3.Connection,
    source_version_id: str,
    *,
    cache: dict[str, list[sqlite3.Row]],
) -> list[sqlite3.Row]:
    if source_version_id not in cache:
        cache[source_version_id] = connection.execute(
            """
            SELECT c.id AS chunk_id, c.source_version_id, c.ordinal, c.locator,
                   c.text_sha256, c.markdown_text, c.stream,
                   sv.authority_identity_id, sv.title, sv.as_of_date,
                   sv.currentness_status, sv.licence_name, sv.licence_url,
                   sv.review_status, sv.canonical_url, sv.stable_identifier,
                   sv.superseded_by,
                   d.id AS document_id,
                   d.content_sha256 AS official_snapshot_sha256,
                   d.source_identity_id AS stable_source_id,
                   d.status AS document_status, d.lane, d.jurisdiction,
                   d.duplicate_of
            FROM chunks c
            JOIN source_versions sv ON sv.id = c.source_version_id
            JOIN documents d ON d.id = sv.document_id
            WHERE c.source_version_id = ?
              AND COALESCE(c.stream, 'body') = 'body'
            ORDER BY c.ordinal, c.id
            """,
            (source_version_id,),
        ).fetchall()
    return cache[source_version_id]


def _lexical_chunk_score(topic: str, excerpt: str, locator: str) -> float:
    topic_tokens = set(_normalised_tokens(topic))
    if not topic_tokens:
        return 0.0
    searchable = f"{excerpt} {locator}".casefold()
    source_tokens = set(_normalised_tokens(searchable))
    shared = topic_tokens & source_tokens
    exact = " ".join(topic.split()).casefold() in " ".join(searchable.split())
    required = 1 if len(topic_tokens) == 1 else min(3, (len(topic_tokens) + 1) // 2)
    if not exact and len(shared) < required:
        return 0.0
    return (20.0 if exact else 0.0) + len(shared) * 5.0 + (len(shared) / len(topic_tokens))


def _source_discovery_candidates(
    *,
    connection: sqlite3.Connection,
    source_versions: Sequence[sqlite3.Row],
    source_cache: dict[str, list[sqlite3.Row]],
    case_id: str,
    issue_id: str,
    topic: str,
    source: Mapping[str, Any],
    as_of_date: date,
) -> tuple[list[dict[str, Any]], Counter[str]]:
    source_name = str(source.get("source_name") or "")
    source_url = str(source.get("url") or "").rstrip("/")
    matches: list[tuple[float, sqlite3.Row]] = []
    for version in source_versions:
        title_score = _title_match_score(source_name, str(version["title"] or ""))
        url_match = bool(
            source_url and str(version["canonical_url"] or "").rstrip("/") == source_url
        )
        if title_score < 0.7 and not url_match:
            continue
        matches.append((1.0 if url_match else title_score, version))
    matches.sort(key=lambda item: (-item[0], str(item[1]["source_version_id"])))

    exclusions: Counter[str] = Counter()
    scored: list[tuple[float, dict[str, Any]]] = []
    for title_score, version in matches[:3]:
        for row in _chunks_for_source(
            connection,
            str(version["source_version_id"]),
            cache=source_cache,
        ):
            reason = _candidate_exclusion(
                row,
                chunk_id=str(row["chunk_id"]),
                source_type=_source_type(str(source.get("source_type") or "")),
                as_of_date=as_of_date,
            )
            if reason:
                exclusions[reason] += 1
                continue
            score = _lexical_chunk_score(topic, str(row["markdown_text"]), str(row["locator"]))
            if score <= 0:
                continue
            candidate = _catalogue_candidate(
                row=row,
                case_id=case_id,
                issue_id=issue_id,
                source_name=source_name,
                source_kind=str(source.get("source_type") or ""),
                candidate_origin="approved_lane_lexical_discovery",
                map_rank=str(source.get("rank") or ""),
                map_pinpoint=str(source.get("pinpoint") or ""),
            )
            scored.append((score + title_score, candidate))
    scored.sort(key=lambda item: (-item[0], item[1]["candidate_id"]))
    return [item[1] for item in scored[:2]], exclusions


def _repairs_for_source(
    source: Mapping[str, Any],
    repair: Mapping[str, Any],
) -> Sequence[Mapping[str, Any]]:
    name = str(source.get("source_name") or "").casefold()
    pinpoint = str(source.get("pinpoint") or "").casefold()
    if "limitation act 1980" in name and "14a" in pinpoint:
        authority = "ukpga:1980:58"
        locator_prefix = "s 14a"
    elif "inheritance (provision for family and dependants) act 1975" in name:
        authority = "ukpga:1975:63"
        locator_prefix = "s 1"
    else:
        return ()
    return tuple(
        item
        for item in repair.get("repairs", ())
        if item.get("gold_eligible_candidate") is True
        and str(item.get("legal_authority_id") or "") == authority
        and str(item.get("legal_locator") or "").casefold().startswith(locator_prefix)
    )


def build_owner_evidence_pack(
    *,
    project_root: Path,
    catalog_path: Path,
    evidence_map_path: Path,
    as_of_date: date,
    repair_payload: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Build all selected issue rows with exact excerpts kept in memory."""

    if not catalog_path.is_file():
        raise FileNotFoundError("catalogue is not present")
    bundle = load_live_evaluation_bundle(
        project_root / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    evidence_map_bytes = evidence_map_path.read_bytes()
    evidence_map = json.loads(evidence_map_bytes)
    by_row = {
        f"{item['case_id']}:{item['issue_id']}": item
        for item in evidence_map.get("issues", ())
        if item.get("generation_disposition") == "generate_once"
    }
    selected_ids = selected_generation_case_ids(bundle)
    expected_rows = [
        f"{case.case_id}:issue-{number:02d}"
        for case in bundle.registry.cases
        if case.case_id in selected_ids
        for number in range(1, len(case.must_cover_issues) + 1)
    ]
    if len(expected_rows) != EXPECTED_SELECTED_ISSUES or set(by_row) != set(expected_rows):
        raise ValueError("candidate evidence map does not contain the selected 305 issues")

    if repair_payload is None:
        held_export = export_held_provision_chunks(catalog_path)
        repair = build_held_span_contiguous_repair(held_export)
    else:
        repair = dict(repair_payload)
    connection = sqlite3.connect(f"file:{catalog_path}?mode=ro", uri=True)
    connection.row_factory = sqlite3.Row
    source_versions = _eligible_source_versions(connection)
    source_cache: dict[str, list[sqlite3.Row]] = {}
    cases: list[dict[str, Any]] = []
    aggregate_exclusions: Counter[str] = Counter()
    try:
        for case in bundle.registry.cases:
            if case.case_id not in selected_ids:
                continue
            issues: list[dict[str, Any]] = []
            for number, topic in enumerate(case.must_cover_issues, start=1):
                issue_id = f"issue-{number:02d}"
                mapped = by_row[f"{case.case_id}:{issue_id}"]
                candidates: list[dict[str, Any]] = []
                seen: set[str] = set()
                issue_exclusions: Counter[str] = Counter()
                for source in mapped.get("candidates", ()):
                    source_had_eligible_local = False
                    for span in source.get("local_spans", ()):
                        original_id = str(span.get("chunk_id") or "")
                        row = connection.execute(_CATALOGUE_ROW_SQL, (original_id,)).fetchone()
                        row = _canonical_duplicate_row(connection, row) if row is not None else None
                        reason = _candidate_exclusion(
                            row,
                            chunk_id=original_id,
                            source_type=_source_type(str(source.get("source_type") or "")),
                            as_of_date=as_of_date,
                        )
                        if reason:
                            issue_exclusions[reason] += 1
                            continue
                        assert row is not None
                        candidate = _catalogue_candidate(
                            row=row,
                            case_id=case.case_id,
                            issue_id=issue_id,
                            source_name=str(source.get("source_name") or ""),
                            source_kind=str(source.get("source_type") or ""),
                            candidate_origin="resource_map_exact_locator",
                            map_rank=str(source.get("rank") or ""),
                            map_pinpoint=str(source.get("pinpoint") or ""),
                        )
                        if candidate["candidate_id"] not in seen:
                            candidates.append(candidate)
                            seen.add(candidate["candidate_id"])
                            source_had_eligible_local = True
                    for repair_item in _repairs_for_source(source, repair):
                        candidate = _repair_candidate(
                            item=repair_item,
                            case_id=case.case_id,
                            issue_id=issue_id,
                            source_name=str(source.get("source_name") or ""),
                            map_rank=str(source.get("rank") or ""),
                            map_pinpoint=str(source.get("pinpoint") or ""),
                        )
                        if candidate["candidate_id"] not in seen:
                            candidates.append(candidate)
                            seen.add(candidate["candidate_id"])
                            source_had_eligible_local = True
                    if not source_had_eligible_local:
                        discovered, rejected = _source_discovery_candidates(
                            connection=connection,
                            source_versions=source_versions,
                            source_cache=source_cache,
                            case_id=case.case_id,
                            issue_id=issue_id,
                            topic=topic,
                            source=source,
                            as_of_date=as_of_date,
                        )
                        issue_exclusions.update(rejected)
                        for candidate in discovered:
                            if candidate["candidate_id"] not in seen:
                                candidates.append(candidate)
                                seen.add(candidate["candidate_id"])
                if not candidates and not issue_exclusions:
                    issue_exclusions["no_catalogue_candidate"] += 1
                aggregate_exclusions.update(issue_exclusions)
                issues.append(
                    {
                        "issue_id": issue_id,
                        "topic": topic,
                        "topic_sha256": _sha256_text(topic),
                        "candidate_count": len(candidates),
                        "candidate_status": (
                            "owner_review_required"
                            if candidates
                            else "knowledge_gap_no_eligible_candidate"
                        ),
                        "candidates": candidates,
                        "excluded_reason_counts": dict(sorted(issue_exclusions.items())),
                        "owner_decision": "",
                        "approved_candidate_ids": "",
                        "owner_proposition": "",
                        "owner_legal_roles": "",
                        "owner_later_treatment": "",
                        "owner_contrary_or_limiting": "",
                    }
                )
            cases.append(
                {
                    "case_id": case.case_id,
                    "subject": case.subject,
                    "task_type": case.task_type,
                    "expected_research_route": case.expected_research_route,
                    "question_sha256": case.question_sha256,
                    "record_sha256": case.record_sha256,
                    "issues": issues,
                }
            )
    finally:
        connection.close()

    issue_count = sum(len(case["issues"]) for case in cases)
    candidate_issue_count = sum(
        issue["candidate_count"] > 0 for case in cases for issue in case["issues"]
    )
    candidate_count = sum(issue["candidate_count"] for case in cases for issue in case["issues"])
    if len(cases) != 30 or issue_count != EXPECTED_SELECTED_ISSUES:
        raise ValueError("owner evidence pack does not cover the selected 30/305")
    return {
        "schema": EVIDENCE_PACK_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "as_of_date": as_of_date.isoformat(),
        "suite_registry_canonical_sha256": bundle.registry.canonical_sha256,
        "run_plan_sha256": bundle.manifest.run_plan_sha256,
        "evidence_map_sha256": _sha256_bytes(evidence_map_bytes),
        "selected_case_count": len(cases),
        "selected_issue_count": issue_count,
        "candidate_issue_count": candidate_issue_count,
        "knowledge_gap_issue_count": issue_count - candidate_issue_count,
        "candidate_count": candidate_count,
        "excluded_reason_counts": dict(sorted(aggregate_exclusions.items())),
        "repair_schema": repair["schema"],
        "repair_candidate_count": sum(
            item.get("gold_eligible_candidate") is True for item in repair.get("repairs", ())
        ),
        "cases": cases,
        "seals_expert_gold": False,
        "owner_review_required": True,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }


def _add_safe_paragraph(document: DocumentType, value: str, *, style: str = "Normal") -> None:
    document.add_paragraph(
        _reject_prohibited_metadata(value, label="evidence pack prose"),
        style=style,
    )


def _candidate_metadata_rows(
    candidate: Mapping[str, Any],
) -> tuple[tuple[str, str], ...]:
    return (
        ("Candidate ID", candidate["candidate_id"]),
        ("Origin", candidate["candidate_origin"]),
        ("Authority", candidate["source_name"]),
        ("Authority identity", candidate["authority_identity"]),
        ("Jurisdiction", candidate["jurisdiction"]),
        ("Rights status", candidate["rights_status"]),
        ("Licence", candidate["licence_name"]),
        ("Source version", candidate["source_version_id"]),
        ("Stable source ID", candidate["stable_source_id"]),
        ("Chunk / repair span", candidate["chunk_id"]),
        ("Locator", candidate["legal_locator"]),
        ("Content SHA-256", candidate["content_sha256"]),
        ("Proposed legal role", candidate["proposed_legal_role"]),
        ("Currentness", candidate["catalogue_currentness"]),
        ("Later treatment", candidate["later_treatment_requirement"]),
        ("Contrary/limiting", candidate["contrary_or_limiting_status"]),
        ("Resource-map pinpoint", candidate["map_pinpoint"]),
    )


def _owner_marking_rows() -> tuple[tuple[str, str, str], ...]:
    return (
        (
            "Issue decision",
            "approve_qualified / reject_keep_gap / request_another_candidate",
            "",
        ),
        (
            "Approved candidate IDs",
            "comma-separated IDs from above; blank if not approved",
            "",
        ),
        (
            "Exact proposition",
            "write the precise proposition supported; required for case law",
            "",
        ),
        (
            "Legal role per candidate",
            "candidate-id=statutory_text / holding_ratio / binding_legal_rule",
            "",
        ),
        (
            "Later treatment",
            "case law: confirmed_current / qualified_current / not_current / uncertain_hold",
            "",
        ),
        (
            "Contrary / limiting",
            "safe authority IDs, comma-separated; write none if none",
            "",
        ),
        ("Review scope", "case law: ordinary / critical / disputed", ""),
        (
            "Second review status",
            "case law: not_required / confirmed",
            "",
        ),
        (
            "Second reviewer ref",
            "reviewer:<64-hex>; required only when confirmed",
            "",
        ),
    )


def _docx_rows(table: Any) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _protected_pair_table(
    table: Any,
    expected: Sequence[tuple[str, str]],
    *,
    label: str,
) -> None:
    rows = _docx_rows(table)
    if not rows or rows[0] != ["Field", "Value"]:
        raise ValueError(f"{label} table header changed")
    if rows[1:] != [[str(key), str(value)] for key, value in expected]:
        raise ValueError(f"{label} protected values differ from the candidate pack")


def _marking_values(table: Any) -> dict[str, str]:
    rows = _docx_rows(table)
    if not rows or rows[0] != ["Field", "Allowed / instruction", "Owner entry"]:
        raise ValueError("owner marking table header changed")
    expected = _owner_marking_rows()
    if [row[:2] for row in rows[1:]] != [
        [field, instruction] for field, instruction, _entry in expected
    ]:
        raise ValueError("owner marking fields or instructions changed")
    return {row[0]: row[2] for row in rows[1:]}


def _split_values(value: str) -> tuple[str, ...]:
    return tuple(item.strip() for item in re.split(r"[,;\n]+", value) if item.strip())


def _role_assignments(value: str) -> dict[str, str]:
    assignments: dict[str, str] = {}
    for item in _split_values(value):
        if "=" not in item:
            raise ValueError("legal-role entries must use candidate-id=role")
        candidate_id, role = (part.strip() for part in item.split("=", 1))
        if candidate_id in assignments:
            raise ValueError("legal-role entry duplicates a candidate")
        assignments[candidate_id] = role
    return assignments


def _safe_authority_ids(value: str) -> tuple[str, ...]:
    cleaned = value.strip()
    if not cleaned or cleaned.casefold() == "none":
        return ()
    values = _split_values(cleaned)
    if any(not _SAFE_REVIEW_ID.fullmatch(item) for item in values):
        raise ValueError("contrary authority entries must be privacy-safe stable IDs")
    if len(values) != len(set(values)):
        raise ValueError("contrary authority entries are duplicated")
    return values


def _case_review(
    *,
    candidate: Mapping[str, Any],
    proposition_hash: str,
    legal_role: str,
    later_treatment: str,
    contrary_ids: tuple[str, ...],
    reviewer_role: str,
    reviewer_ref: str,
    review_scope: str,
    second_review_status: str,
    second_reviewer_ref: str | None,
    as_of_date: date,
) -> CasePropositionReview:
    payload = {
        "schema": "legalbot.case-proposition-currentness-review.v1",
        "source_version_id": candidate["source_version_id"],
        "chunk_id": candidate["chunk_id"],
        "legal_locator": candidate["legal_locator"],
        "exact_span_sha256": candidate["content_sha256"],
        "proposition_hash": proposition_hash,
        "legal_role": legal_role,
        "later_treatment_reviewed_as_of_date": as_of_date.isoformat(),
        "later_treatment_status": later_treatment,
        "contrary_or_limiting_authority_ids": list(contrary_ids),
        "reviewer_role": reviewer_role,
        "reviewer_ref": reviewer_ref,
        "review_scope": review_scope,
        "second_review_status": second_review_status,
        "second_reviewer_ref": second_reviewer_ref,
    }
    payload["seal_sha256"] = case_proposition_review_sha256(payload)
    return CasePropositionReview.model_validate(payload)


def _owner_gold_span(
    *,
    candidate: Mapping[str, Any],
    issue_id: str,
    legal_role: str,
    proposition_hash: str | None,
    currentness_review: CasePropositionReview | None,
    catalog_path: Path,
    repair: Mapping[str, Any],
) -> dict[str, Any]:
    from .live_suite_span_accuracy import verify_user_span_exact_match

    verify_user_span_exact_match(
        chunk_id=str(candidate["chunk_id"]),
        content_sha256=str(candidate["content_sha256"]),
        legal_locator=str(candidate["legal_locator"]),
        source_version_id=str(candidate["source_version_id"]),
        catalog_path=catalog_path,
        repair=repair,
        legal_authority_id=str(candidate["authority_identity"] or "") or None,
        parent_chunk_id=candidate.get("parent_chunk_id"),
        legal_role=legal_role,
        official_snapshot_sha256=str(candidate.get("official_snapshot_sha256") or "") or None,
        derivation_manifest_sha256=str(candidate.get("derivation_manifest_sha256") or "") or None,
        stable_source_id=str(candidate.get("stable_source_id") or "") or None,
        source_type=str(candidate.get("source_type") or "") or None,
        jurisdiction=str(candidate.get("jurisdiction") or "") or None,
    )
    gold_span_id = "gold-" + _sha256_text(
        "|".join(
            (
                issue_id,
                str(candidate["candidate_id"]),
                legal_role,
                proposition_hash or "",
            )
        )
    )
    span = LiveGoldSpan.model_validate(
        {
            "schema": "legalbot.live-gold-span.v1",
            "gold_span_id": gold_span_id,
            "issue_id": issue_id,
            "stable_source_id": candidate["stable_source_id"],
            "legal_authority_id": candidate["authority_identity"] or None,
            "source_version_id": candidate["source_version_id"],
            "chunk_id": candidate["chunk_id"],
            "legal_locator": candidate["legal_locator"],
            "content_sha256": candidate["content_sha256"],
            "source_type": candidate["source_type"],
            "legal_role": legal_role,
            "proposition_hash": proposition_hash,
            "case_currentness_review": (
                currentness_review.model_dump(mode="json", by_alias=True)
                if currentness_review is not None
                else None
            ),
            "relevance_grade": 3,
            "contrary_or_limiting": False,
        }
    )
    return span.model_dump(mode="json", by_alias=True)


def _parse_issue_marking(
    *,
    issue: Mapping[str, Any],
    values: Mapping[str, str],
    reviewer_role: str,
    reviewer_ref: str,
    as_of_date: date,
    catalog_path: Path,
    repair: Mapping[str, Any],
) -> dict[str, Any]:
    decision = values["Issue decision"].strip()
    allowed_decisions = {
        "approve_qualified",
        "reject_keep_gap",
        "request_another_candidate",
    }
    if decision not in allowed_decisions:
        raise ValueError(f"{issue['issue_id']} requires one explicit allowed issue decision")
    selected_ids = _split_values(values["Approved candidate IDs"])
    roles = _role_assignments(values["Legal role per candidate"])
    candidates = {str(item["candidate_id"]): item for item in issue["candidates"]}
    if any(candidate_id not in candidates for candidate_id in selected_ids):
        raise ValueError(f"{issue['issue_id']} selects an unknown candidate")
    if set(roles) != set(selected_ids):
        raise ValueError(f"{issue['issue_id']} needs exactly one legal role per selected candidate")
    if decision != "approve_qualified":
        if selected_ids or roles:
            raise ValueError(f"{issue['issue_id']} cannot bind candidates to a gap decision")
        reason = (
            "owner_rejected_candidates"
            if decision == "reject_keep_gap"
            else "owner_requested_another_candidate"
        )
        return {
            "status": "knowledge_gap",
            "reason_code": reason,
            "exact_gold_spans": [],
            "contrary_authority_status": "unresolved",
        }
    if not selected_ids:
        raise ValueError(f"{issue['issue_id']} qualification has no selected candidate")

    proposition = " ".join(values["Exact proposition"].split())
    if "/users/" in proposition.casefold() or proposition.casefold().startswith("file:"):
        raise ValueError("owner proposition contains prohibited path metadata")
    contrary_ids = _safe_authority_ids(values["Contrary / limiting"])
    case_selected = any(
        candidates[candidate_id]["source_type"] == "case" for candidate_id in selected_ids
    )
    later_treatment = values["Later treatment"].strip()
    review_scope = values["Review scope"].strip()
    second_review_status = values["Second review status"].strip()
    second_reviewer_ref = values["Second reviewer ref"].strip() or None
    proposition_hash = _sha256_text(proposition) if proposition else None
    if case_selected:
        if proposition_hash is None:
            raise ValueError(
                f"{issue['issue_id']} case-law qualification needs an exact proposition"
            )
        if later_treatment not in {"confirmed_current", "qualified_current"}:
            raise ValueError(
                f"{issue['issue_id']} case-law qualification has unresolved later treatment"
            )
        if review_scope not in {"ordinary", "critical", "disputed"}:
            raise ValueError(f"{issue['issue_id']} has an invalid review scope")
        if second_review_status not in {"not_required", "confirmed"}:
            raise ValueError(f"{issue['issue_id']} has an invalid second-review status")
        if second_reviewer_ref is not None and not _REVIEWER_REF.fullmatch(second_reviewer_ref):
            raise ValueError(f"{issue['issue_id']} has an invalid second reviewer ref")
    elif any(
        (
            proposition,
            later_treatment,
            review_scope,
            second_review_status,
            second_reviewer_ref,
        )
    ):
        raise ValueError(
            f"{issue['issue_id']} non-case qualification must leave case-review fields blank"
        )
    elif contrary_ids:
        raise ValueError(
            f"{issue['issue_id']} non-case contrary authority belongs in the named source review"
        )

    spans = []
    for candidate_id in selected_ids:
        candidate = candidates[candidate_id]
        legal_role = roles[candidate_id]
        source_type = str(candidate["source_type"])
        review = None
        candidate_proposition_hash = None
        if source_type == "case":
            if legal_role not in {"holding_ratio", "binding_legal_rule"}:
                raise ValueError(
                    f"{issue['issue_id']} case candidate has a non-material legal role"
                )
            assert proposition_hash is not None
            review = _case_review(
                candidate=candidate,
                proposition_hash=proposition_hash,
                legal_role=legal_role,
                later_treatment=later_treatment,
                contrary_ids=contrary_ids,
                reviewer_role=reviewer_role,
                reviewer_ref=reviewer_ref,
                review_scope=review_scope,
                second_review_status=second_review_status,
                second_reviewer_ref=second_reviewer_ref,
                as_of_date=as_of_date,
            )
            candidate_proposition_hash = proposition_hash
        elif source_type == "legislation":
            if legal_role != "statutory_text":
                raise ValueError(
                    f"{issue['issue_id']} legislation candidate must be statutory_text"
                )
        elif legal_role != "secondary_commentary":
            raise ValueError(
                f"{issue['issue_id']} secondary candidate must be secondary_commentary"
            )
        spans.append(
            _owner_gold_span(
                candidate=candidate,
                issue_id=str(issue["issue_id"]),
                legal_role=legal_role,
                proposition_hash=candidate_proposition_hash,
                currentness_review=review,
                catalog_path=catalog_path,
                repair=repair,
            )
        )
    return {
        "status": "qualified",
        "reason_code": None,
        "exact_gold_spans": spans,
        "contrary_authority_status": "unresolved",
    }


def _parse_case_evidence_document(
    *,
    workbook_path: Path,
    case: Mapping[str, Any],
    pack: Mapping[str, Any],
    catalog_path: Path,
    repair: Mapping[str, Any],
    as_of_date: date,
) -> dict[str, dict[str, Any]]:
    document = Document(str(workbook_path))
    tables = document.tables
    cursor = 0
    _protected_pair_table(
        tables[cursor],
        (
            ("Case ID", case["case_id"]),
            ("Subject", case["subject"]),
            ("Task", case["task_type"]),
            ("Route", case["expected_research_route"]),
            ("Issue count", str(len(case["issues"]))),
            (
                "Issues with eligible candidate",
                str(sum(issue["candidate_count"] > 0 for issue in case["issues"])),
            ),
            (
                "Issues still gap",
                str(sum(issue["candidate_count"] == 0 for issue in case["issues"])),
            ),
            ("Training export", "false"),
        ),
        label=f"{case['case_id']} cover",
    )
    cursor += 1
    markings: dict[str, dict[str, Any]] = {}
    for issue in case["issues"]:
        _protected_pair_table(
            tables[cursor],
            (
                ("Issue ID", issue["issue_id"]),
                ("Topic", issue["topic"]),
                ("Topic SHA-256", issue["topic_sha256"]),
                ("Eligible candidate count", str(issue["candidate_count"])),
                ("Current status", issue["candidate_status"]),
                (
                    "Contrary/limiting search",
                    "not machine-established; named-source-set owner review required",
                ),
            ),
            label=f"{case['case_id']} {issue['issue_id']}",
        )
        cursor += 1
        if not issue["candidates"] and issue["excluded_reason_counts"]:
            rows = _docx_rows(tables[cursor])
            expected = [
                ["Excluded reason", "Count"],
                *[
                    [reason, str(count)]
                    for reason, count in issue["excluded_reason_counts"].items()
                ],
            ]
            if rows != expected:
                raise ValueError(f"{case['case_id']} {issue['issue_id']} exclusion table changed")
            cursor += 1
        for candidate in issue["candidates"]:
            _protected_pair_table(
                tables[cursor],
                _candidate_metadata_rows(candidate),
                label=f"{case['case_id']} {issue['issue_id']} candidate",
            )
            cursor += 1
        markings[str(issue["issue_id"])] = _marking_values(tables[cursor])
        cursor += 1

    rows = _docx_rows(tables[cursor])
    if not rows or rows[0] != ["Field", "Allowed value", "Owner entry"]:
        raise ValueError(f"{case['case_id']} confirmation table changed")
    if [row[0] for row in rows[1:]] != [
        "reviewer_role",
        "reviewer_ref",
        "confirmation_token",
    ]:
        raise ValueError(f"{case['case_id']} confirmation fields changed")
    signoff = {row[0]: row[2] for row in rows[1:]}
    identity = build_owner_reviewer_identity(as_of_date=as_of_date)
    if signoff["reviewer_role"] != identity["approval_reviewer_role"]:
        raise ValueError(f"{case['case_id']} reviewer role is invalid")
    if signoff["reviewer_ref"] != identity["approval_reviewer_ref"]:
        raise ValueError(f"{case['case_id']} reviewer ref is invalid")
    if signoff["confirmation_token"] != CASE_CONFIRMATION_TOKEN:
        raise ValueError(f"{case['case_id']} owner confirmation is missing")
    cursor += 1
    if cursor != len(tables):
        raise ValueError(f"{case['case_id']} contains unexpected added tables")

    return {
        issue_id: _parse_issue_marking(
            issue=next(item for item in case["issues"] if item["issue_id"] == issue_id),
            values=values,
            reviewer_role=str(identity["approval_reviewer_role"]),
            reviewer_ref=str(identity["approval_reviewer_ref"]),
            as_of_date=as_of_date,
            catalog_path=catalog_path,
            repair=repair,
        )
        for issue_id, values in markings.items()
    }


def import_owner_evidence_reviews(
    *,
    project_root: Path,
    catalog_path: Path,
    evidence_map_path: Path,
    workbook_dir: Path,
    review_export_path: Path,
    as_of_date: date,
    repair_payload: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    """Convert fully marked case DOCXs into rows bound to a sealed Path-B export."""

    export = json.loads(review_export_path.read_text(encoding="utf-8"))
    if export.get("schema") != REVIEW_EXPORT_SCHEMA:
        raise ValueError("evidence review import requires a Path-B review export")
    expected_export = dict(export)
    expected_export.pop("export_sha256", None)
    if sealed_sha256(expected_export) != export.get("export_sha256"):
        raise ValueError("Path-B review export seal does not match its contents")
    if str(export.get("as_of_date")) != as_of_date.isoformat():
        raise ValueError("Path-B review export has a different as-of date")

    pack = build_owner_evidence_pack(
        project_root=project_root,
        catalog_path=catalog_path,
        evidence_map_path=evidence_map_path,
        as_of_date=as_of_date,
        repair_payload=repair_payload,
    )
    repair = (
        dict(repair_payload)
        if repair_payload is not None
        else build_held_span_contiguous_repair(export_held_provision_chunks(catalog_path))
    )
    owner_rows: dict[str, dict[str, Any]] = {}
    for case in pack["cases"]:
        workbook_path = workbook_dir / f"LegalBot-Live60-Path-B-Evidence-{case['case_id']}.docx"
        if not workbook_path.is_file():
            raise ValueError(f"{case['case_id']} filled evidence workbook is missing")
        parsed = _parse_case_evidence_document(
            workbook_path=workbook_path,
            case=case,
            pack=pack,
            catalog_path=catalog_path,
            repair=repair,
            as_of_date=as_of_date,
        )
        for issue_id, marking in parsed.items():
            owner_rows[f"{case['case_id']}:{issue_id}"] = marking

    reviewed_rows = []
    selected_ids = {case["case_id"] for case in pack["cases"]}
    for source in export.get("rows", ()):
        row_id = str(source["row_id"])
        if source["case_id"] in selected_ids:
            if row_id not in owner_rows:
                raise ValueError("filled evidence workbooks omit a selected issue")
            marking = owner_rows[row_id]
        else:
            marking = {
                "status": "knowledge_gap",
                "reason_code": "coverage_only_explicit_knowledge_gap",
                "exact_gold_spans": [],
                "contrary_authority_status": "unresolved",
            }
        reviewed_rows.append(
            {
                "row_id": row_id,
                "case_id": source["case_id"],
                "issue_id": source["issue_id"],
                "status": marking["status"],
                "reason_code": marking["reason_code"],
                "exact_gold_spans": marking["exact_gold_spans"],
                "contrary_authority_status": marking["contrary_authority_status"],
            }
        )
    if len(reviewed_rows) != 585 or len(owner_rows) != EXPECTED_SELECTED_ISSUES:
        raise ValueError("evidence review import does not disposition all 585 issues")
    qualified = sum(row["status"] == "qualified" for row in reviewed_rows)
    payload = {
        "schema": REVIEW_IMPORT_SCHEMA,
        "export_sha256": export["export_sha256"],
        "suite_id": export["suite_id"],
        "as_of_date": export["as_of_date"],
        "suite_registry_canonical_sha256": export["suite_registry_canonical_sha256"],
        "run_plan_sha256": export["run_plan_sha256"],
        "reviewer_identity": "owner_primary_reviewer",
        "row_count": len(reviewed_rows),
        "selected_qualified_issue_count": qualified,
        "selected_knowledge_gap_issue_count": EXPECTED_SELECTED_ISSUES - qualified,
        "rows": reviewed_rows,
        "owner_confirmation_present": True,
        "seals_expert_gold": False,
        "selected_evidence_complete": qualified == EXPECTED_SELECTED_ISSUES,
        "ready_for_overlay_reconstruction": qualified == EXPECTED_SELECTED_ISSUES,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }
    assert_safe_evaluation_payload(payload)
    return payload


def write_owner_evidence_reviews(
    *,
    project_root: Path,
    catalog_path: Path,
    evidence_map_path: Path,
    workbook_dir: Path,
    review_export_path: Path,
    output_path: Path,
    as_of_date: date,
    overwrite: bool = False,
    repair_payload: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    destination = output_path.resolve()
    if destination.exists() and not overwrite:
        raise FileExistsError("reviewed evidence rows already exist")
    payload = import_owner_evidence_reviews(
        project_root=project_root,
        catalog_path=catalog_path,
        evidence_map_path=evidence_map_path,
        workbook_dir=workbook_dir,
        review_export_path=review_export_path,
        as_of_date=as_of_date,
        repair_payload=repair_payload,
    )
    raw = (json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode("utf-8")
    destination.parent.mkdir(parents=True, exist_ok=True)
    destination.write_bytes(raw)
    os.chmod(destination, 0o600)
    return {
        "schema": EVIDENCE_REVIEW_IMPORT_RESULT_SCHEMA,
        "row_count": payload["row_count"],
        "selected_qualified_issue_count": payload["selected_qualified_issue_count"],
        "selected_knowledge_gap_issue_count": payload["selected_knowledge_gap_issue_count"],
        "selected_evidence_complete": payload["selected_evidence_complete"],
        "ready_for_overlay_seal": payload["ready_for_overlay_seal"],
        "reviewed_rows_sha256": _sha256_bytes(raw),
        "seals_expert_gold": False,
        "generation_authorised": False,
        "writes_active": False,
        "writes_o04": False,
    }


def build_case_evidence_document(
    pack: Mapping[str, Any],
    case: Mapping[str, Any],
) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, f"Live60-{pack['as_of_date']}")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph(
        f"Path-B evidence review · {case['case_id']}",
        style="Title",
    )
    _add_safe_paragraph(
        document,
        "Exact excerpts below are candidates only. Approve only authority words "
        "that support the stated issue. Reject irrelevant candidates or request "
        "another candidate. No selection in Word becomes gold automatically.",
    )
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Case ID", case["case_id"]),
            ("Subject", case["subject"]),
            ("Task", case["task_type"]),
            ("Route", case["expected_research_route"]),
            ("Issue count", str(len(case["issues"]))),
            (
                "Issues with eligible candidate",
                str(sum(issue["candidate_count"] > 0 for issue in case["issues"])),
            ),
            (
                "Issues still gap",
                str(sum(issue["candidate_count"] == 0 for issue in case["issues"])),
            ),
            ("Training export", "false"),
        ),
        (3_120, 6_240),
    )
    _add_toc(document)
    _add_page_break(document)

    for issue in case["issues"]:
        document.add_heading(
            f"{issue['issue_id']} · {issue['topic']}",
            level=1,
        )
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("Issue ID", issue["issue_id"]),
                ("Topic", issue["topic"]),
                ("Topic SHA-256", issue["topic_sha256"]),
                ("Eligible candidate count", str(issue["candidate_count"])),
                ("Current status", issue["candidate_status"]),
                (
                    "Contrary/limiting search",
                    "not machine-established; named-source-set owner review required",
                ),
            ),
            (3_120, 6_240),
        )
        if not issue["candidates"]:
            _add_safe_paragraph(
                document,
                "No exact candidate passed authority-lane, jurisdiction, rights, "
                "currentness and hash checks. Keep knowledge_gap or request another "
                "candidate; do not approve a nearest hit.",
            )
            if issue["excluded_reason_counts"]:
                _add_table(
                    document,
                    ("Excluded reason", "Count"),
                    tuple(
                        (reason, str(count))
                        for reason, count in issue["excluded_reason_counts"].items()
                    ),
                    (7_200, 2_160),
                )
        for number, candidate in enumerate(issue["candidates"], start=1):
            document.add_heading(f"Candidate {number}", level=2)
            _add_table(
                document,
                ("Field", "Value"),
                _candidate_metadata_rows(candidate),
                (2_640, 6_720),
                body_size=8.0,
            )
            _add_safe_paragraph(document, "Exact excerpt", style="Review Small")
            _add_safe_paragraph(document, candidate["excerpt"])
        document.add_heading("Owner marking", level=2)
        _add_table(
            document,
            ("Field", "Allowed / instruction", "Owner entry"),
            _owner_marking_rows(),
            (2_160, 4_320, 2_880),
            body_size=8.0,
        )
        _add_page_break(document)

    identity = build_owner_reviewer_identity(as_of_date=date.fromisoformat(str(pack["as_of_date"])))
    document.add_heading("Case review confirmation", level=1)
    _add_table(
        document,
        ("Field", "Allowed value", "Owner entry"),
        (
            (
                "reviewer_role",
                identity["approval_reviewer_role"],
                identity["approval_reviewer_role"],
            ),
            (
                "reviewer_ref",
                "reviewer:<64-hex>",
                identity["approval_reviewer_ref"],
            ),
            (
                "confirmation_token",
                CASE_CONFIRMATION_TOKEN,
                "",
            ),
        ),
        (2_640, 3_600, 3_120),
    )
    _finalize_document_properties(
        document,
        title=f"LegalBot Live60 Path-B evidence review {case['case_id']}",
        subject="Owner review of exact authority candidates; not legal gold",
    )
    return document


def build_evidence_index_document(pack: Mapping[str, Any]) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, f"Live60-{pack['as_of_date']}")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Live60 Path-B evidence-pack index", style="Title")
    _add_safe_paragraph(
        document,
        "This index covers all 305 issues on the frozen 30 selected cases. "
        "Candidate means eligible for owner review, not qualified legal gold.",
    )
    _add_table(
        document,
        ("Metric", "Value"),
        (
            ("Selected cases", str(pack["selected_case_count"])),
            ("Selected issues", str(pack["selected_issue_count"])),
            ("Issues with candidate", str(pack["candidate_issue_count"])),
            ("Issues still knowledge_gap", str(pack["knowledge_gap_issue_count"])),
            ("Exact candidate excerpts", str(pack["candidate_count"])),
            ("Pilot case", PILOT_CASE_ID),
            ("Overlay seal", "not ready"),
            ("Generation", "not authorised"),
        ),
        (4_320, 5_040),
    )
    _add_table(
        document,
        (
            "Case",
            "Subject",
            "Issues",
            "With candidate",
            "Still gap",
            "Candidate excerpts",
        ),
        tuple(
            (
                case["case_id"],
                case["subject"],
                str(len(case["issues"])),
                str(sum(issue["candidate_count"] > 0 for issue in case["issues"])),
                str(sum(issue["candidate_count"] == 0 for issue in case["issues"])),
                str(sum(issue["candidate_count"] for issue in case["issues"])),
            )
            for case in pack["cases"]
        ),
        (1_320, 3_000, 1_080, 1_320, 1_200, 1_440),
        body_size=8.0,
    )
    document.add_heading("How to review", level=1)
    for instruction in (
        "Start with the pilot file for live30-q03, then continue in case order.",
        "For each issue approve exact candidate IDs, reject and keep the gap, or request another candidate.",
        "For case law, write the exact proposition and complete later-treatment review. A judgment paragraph is not qualified merely because it is current catalogue content.",
        "No candidate is a citator result. Contrary or limiting authority remains a separate named-source-set review.",
        "Do not approve teaching notes, foreign law presented as England-and-Wales law, unlicensed full text, spliced parents, or nearest-vector-only hits.",
    ):
        _add_safe_paragraph(document, instruction)
    _finalize_document_properties(
        document,
        title="LegalBot Live60 Path-B evidence-pack index",
        subject="Index of owner-review exact-evidence candidates",
    )
    return document


def _safe_manifest(
    pack: Mapping[str, Any], artifacts: Sequence[Mapping[str, Any]]
) -> dict[str, Any]:
    cases = []
    by_case = {str(item["case_id"]): item for item in artifacts}
    for case in pack["cases"]:
        artifact = by_case[case["case_id"]]
        cases.append(
            {
                "case_id": case["case_id"],
                "issue_count": len(case["issues"]),
                "candidate_issue_count": sum(
                    issue["candidate_count"] > 0 for issue in case["issues"]
                ),
                "knowledge_gap_issue_count": sum(
                    issue["candidate_count"] == 0 for issue in case["issues"]
                ),
                "candidate_count": sum(issue["candidate_count"] for issue in case["issues"]),
                "artifact_name": artifact["artifact_name"],
                "artifact_sha256": artifact["artifact_sha256"],
                "pilot": case["case_id"] == PILOT_CASE_ID,
            }
        )
    payload = {
        "schema": EVIDENCE_PACK_MANIFEST_SCHEMA,
        "suite_id": pack["suite_id"],
        "as_of_date": pack["as_of_date"],
        "suite_registry_canonical_sha256": pack["suite_registry_canonical_sha256"],
        "run_plan_sha256": pack["run_plan_sha256"],
        "evidence_map_sha256": pack["evidence_map_sha256"],
        "selected_case_count": pack["selected_case_count"],
        "selected_issue_count": pack["selected_issue_count"],
        "candidate_issue_count": pack["candidate_issue_count"],
        "knowledge_gap_issue_count": pack["knowledge_gap_issue_count"],
        "candidate_count": pack["candidate_count"],
        "excluded_reason_counts": pack["excluded_reason_counts"],
        "repair_schema": pack["repair_schema"],
        "repair_candidate_count": pack["repair_candidate_count"],
        "index_artifact_name": "LegalBot-Live60-Path-B-Evidence-Pack-Index.docx",
        "index_artifact_sha256": next(
            item["artifact_sha256"] for item in artifacts if item["case_id"] == "index"
        ),
        "cases": cases,
        "owner_review_required": True,
        "seals_expert_gold": False,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }
    assert_safe_evaluation_payload(payload)
    return payload


def export_owner_evidence_pack(
    *,
    project_root: Path,
    catalog_path: Path,
    evidence_map_path: Path,
    output_dir: Path,
    as_of_date: date,
    overwrite: bool = False,
) -> dict[str, Any]:
    """Write one index and 30 local DOCX files plus a prose-free manifest."""

    destination = output_dir.resolve()
    if destination.exists() and any(destination.iterdir()) and not overwrite:
        raise FileExistsError("owner evidence-pack directory is not empty")
    destination.mkdir(parents=True, exist_ok=True)
    pack = build_owner_evidence_pack(
        project_root=project_root,
        catalog_path=catalog_path,
        evidence_map_path=evidence_map_path,
        as_of_date=as_of_date,
    )

    artifacts: list[dict[str, Any]] = []
    index_name = "LegalBot-Live60-Path-B-Evidence-Pack-Index.docx"
    index_path = destination / index_name
    build_evidence_index_document(pack).save(str(index_path))
    os.chmod(index_path, 0o600)
    artifacts.append(
        {
            "case_id": "index",
            "artifact_name": index_name,
            "artifact_sha256": _sha256_bytes(index_path.read_bytes()),
        }
    )
    for case in pack["cases"]:
        artifact_name = f"LegalBot-Live60-Path-B-Evidence-{case['case_id']}.docx"
        artifact_path = destination / artifact_name
        build_case_evidence_document(pack, case).save(str(artifact_path))
        os.chmod(artifact_path, 0o600)
        artifacts.append(
            {
                "case_id": case["case_id"],
                "artifact_name": artifact_name,
                "artifact_sha256": _sha256_bytes(artifact_path.read_bytes()),
            }
        )

    manifest = _safe_manifest(pack, artifacts)
    manifest_path = destination / "manifest.json"
    manifest_bytes = (
        json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True) + "\n"
    ).encode("utf-8")
    manifest_path.write_bytes(manifest_bytes)
    os.chmod(manifest_path, 0o600)
    return {
        "schema": EVIDENCE_PACK_RESULT_SCHEMA,
        "selected_case_count": pack["selected_case_count"],
        "selected_issue_count": pack["selected_issue_count"],
        "candidate_issue_count": pack["candidate_issue_count"],
        "knowledge_gap_issue_count": pack["knowledge_gap_issue_count"],
        "candidate_count": pack["candidate_count"],
        "docx_count": len(artifacts),
        "manifest_sha256": _sha256_bytes(manifest_bytes),
        "pilot_case_id": PILOT_CASE_ID,
        "owner_review_required": True,
        "seals_expert_gold": False,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }
