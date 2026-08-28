"""Local-only DOCX intake for Live60 owner control decisions.

The Word document records workflow choices only. Parsing it creates an unsigned,
privacy-safe draft and a reviewable diff; it never creates legal gold, ACTIVE,
PREVIOUS, O-04, or an owner-authored seal.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
from collections.abc import Mapping, Sequence
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType

from .live30 import assert_safe_evaluation_payload
from .live_suite import LiveEvaluationBundle, load_live_evaluation_bundle
from .live_suite_owner_decision_contract import (
    CONDITIONAL_DECISION_IDS,
    LATER_OWNER_ACTIONS,
    OWNER_DECISION_IDS,
)
from .live_suite_owner_review import HELD_STATUTORY_PROVISIONS
from .live_suite_reviewer_identity import build_owner_reviewer_identity
from .review_docx import (
    _add_header_footer,
    _add_table,
    _configure_page,
    _configure_styles,
    _finalize_document_properties,
)

OWNER_INTAKE_DOCX_SCHEMA = "legalbot.live60-owner-decision-intake-docx.v1"
OWNER_INTAKE_DRAFT_SCHEMA = "legalbot.live60-owner-decision-intake-draft.v1"
OWNER_INTAKE_RESULT_SCHEMA = "legalbot.live60-owner-decision-intake-result.v1"
LEGACY_INTAKE_MIGRATION_SCHEMA = "legalbot.live60-legacy-owner-intake-migration.v1"
OWNER_CONFIRMATION_TOKEN = "CONFIRM_DRAFT_ONLY"
EXPECTED_TABLE_COUNT = 7

_SAFE_ID = re.compile(r"^[a-z0-9][a-z0-9._:-]{2,127}$")
_REVIEWER_REF = re.compile(r"^reviewer:[0-9a-f]{64}$")

ROUTE_ROWS: tuple[tuple[str, str, str], ...] = (
    ("route", "Path B is the active route.", "path_b"),
    ("target", "Prepare every selected case for one answer.", "full_30_answer_run"),
    ("selected_case_count", "Frozen generate-once cases.", "30"),
    ("selected_issue_count", "Issues on the selected cases.", "305"),
    (
        "coverage_only_disposition",
        "The other 280 issues may remain explicit gaps.",
        "explicit_knowledge_gap_allowed",
    ),
)

INDEXING_RIGHT_ROWS: tuple[tuple[str, str, str], ...] = (
    (
        "IDX-01",
        "Candidate authority lane",
        "approved_current_england_wales_sources_only",
    ),
    ("IDX-02", "Westlaw and scholarship full text", "stay_out"),
    ("IDX-03", "Find Case Law judgment text", "metadata_only"),
    ("IDX-04", "Elasticsearch", "stay_out"),
    ("IDX-05", "Teaching and feedback material", "never_independent_authority"),
)

HELD_DEFAULTS: Mapping[str, str] = {
    "held-provision-01": "current_official_bytes_accepted",
    "held-provision-02": "current_official_bytes_accepted_using_repair_spans",
    "held-provision-03": "current_official_bytes_accepted",
    "held-provision-04": "current_official_bytes_accepted_using_repair_spans",
}
HELD_ALLOWED: Mapping[str, frozenset[str]] = {
    "held-provision-01": frozenset({"current_official_bytes_accepted", "hold"}),
    "held-provision-02": frozenset({"current_official_bytes_accepted_using_repair_spans", "hold"}),
    "held-provision-03": frozenset({"current_official_bytes_accepted", "hold"}),
    "held-provision-04": frozenset({"current_official_bytes_accepted_using_repair_spans", "hold"}),
}

CONTRARY_POLICY_ROWS: tuple[tuple[str, str, str], ...] = (
    (
        "contrary-review-scope",
        "Review a named source set before overlay sealing.",
        "named_source_set_required",
    ),
    (
        "contrary-none-meaning",
        "If none is found, describe only the defined source set.",
        "reviewed_none_in_defined_source_set",
    ),
    (
        "contrary-found-treatment",
        "Bind exact contrary or limiting spans when found.",
        "reviewed_and_bound",
    ),
    (
        "critical-disputed-treatment",
        "Critical or disputed propositions need another human.",
        "needs_independent_second_review",
    ),
    (
        "means-no-contrary-english-law",
        "The review must never make this broader claim.",
        "no",
    ),
)

DECISION_ROWS: tuple[tuple[str, str, str], ...] = (
    ("D-01", "Use Path B.", "accepted"),
    ("D-02", "Prepare all 30 frozen selected cases.", "accepted"),
    ("D-03", "Require later-treatment review for every case-law span.", "accepted"),
    ("D-04", "Require the named-source-set contrary review.", "accepted"),
    ("D-05", "Build a current-date candidate only after overlay sealing.", "conditional"),
    ("D-06", "Owner promotes ACTIVE only after Stage A passes.", "deferred_later_owner_action"),
    ("D-07", "Owner observes rollback and re-promotion.", "deferred_later_owner_action"),
    ("D-08", "Owner observes a real loopback browser recovery.", "deferred_later_owner_action"),
    ("D-09", "Owner alone issues O-04 after readiness passes.", "deferred_later_owner_action"),
    ("D-10", "Keep the first run loopback-only.", "accepted"),
    ("D-11", "Keep unapproved licensed full text out.", "accepted"),
    ("D-12", "Keep Find Case Law judgment text metadata-only.", "accepted"),
    ("D-13", "Do not add Elasticsearch.", "accepted"),
    ("D-14", "Keep the frozen 16-rule assessment bundle.", "accepted"),
    ("D-15", "Use accepted repair spans, never rejected parents.", "accepted"),
)


def _bundle(project_root: Path) -> LiveEvaluationBundle:
    return load_live_evaluation_bundle(
        project_root / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )


def _table_rows(table: Any) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _rows_by_id(
    table: Any,
    *,
    expected_ids: Sequence[str],
    width: int,
    label: str,
) -> list[list[str]]:
    rows = _table_rows(table)
    if not rows or any(len(row) != width for row in rows):
        raise ValueError(f"{label} table has an invalid shape")
    body = rows[1:]
    ids = [row[0] for row in body]
    if ids != list(expected_ids):
        raise ValueError(f"{label} rows are missing, duplicated, added, or reordered")
    return body


def _require_exact(value: str, expected: str, *, label: str) -> str:
    if value != expected:
        raise ValueError(f"{label} must be {expected}")
    return value


def _require_safe_id(value: str, *, label: str) -> str:
    if not _SAFE_ID.fullmatch(value):
        raise ValueError(f"{label} is not a privacy-safe identifier")
    return value


def build_owner_decision_intake_document(
    *,
    project_root: Path,
    as_of_date: date,
) -> DocumentType:
    """Build the deterministic Path-B decision form."""

    bundle = _bundle(project_root)
    identity = build_owner_reviewer_identity(as_of_date=as_of_date)
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, f"Live60-{as_of_date.isoformat()}")

    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Live60 Path-B owner decision intake", style="Title")
    document.add_paragraph(
        "This local form records workflow choices only. It does not create legal "
        "gold, seal an overlay, promote ACTIVE, issue O-04, or authorise generation."
    )
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Template schema", OWNER_INTAKE_DOCX_SCHEMA),
            ("Suite", bundle.manifest.suite_id),
            ("As-of date", as_of_date.isoformat()),
            ("Registry SHA-256", bundle.registry.canonical_sha256),
            ("Run-plan SHA-256", bundle.manifest.run_plan_sha256),
            ("Training export", "false"),
        ),
        (3_120, 6_240),
    )

    document.add_heading("1. Route and target", level=1)
    document.add_paragraph(
        "Path B/full-30 is already selected. Do not replace it with Path A or C."
    )
    _add_table(
        document,
        ("Control", "Meaning", "Owner entry"),
        ROUTE_ROWS,
        (2_160, 4_800, 2_400),
    )

    document.add_heading("2. Indexing and source rights", level=1)
    document.add_paragraph(
        "These choices prevent unapproved or non-authority material from becoming gold."
    )
    _add_table(
        document,
        ("ID", "Control", "Owner entry"),
        INDEXING_RIGHT_ROWS,
        (1_200, 4_800, 3_360),
    )

    document.add_heading("3. Four held statutory provisions", level=1)
    document.add_paragraph(
        "The repair choice creates no qualification by itself. Exact accepted spans "
        "are checked again before sealing."
    )
    _add_table(
        document,
        ("ID", "Provision", "Allowed values", "Owner entry"),
        tuple(
            (
                held_id,
                title,
                " / ".join(sorted(HELD_ALLOWED[held_id])),
                HELD_DEFAULTS[held_id],
            )
            for held_id, title, _summary in HELD_STATUTORY_PROVISIONS
        ),
        (1_320, 2_880, 2_640, 2_520),
        body_size=8.0,
    )

    document.add_heading("4. Contrary and limiting authority policy", level=1)
    document.add_paragraph(
        "This is the review policy, not a claim that the source-set review is complete."
    )
    _add_table(
        document,
        ("Control", "Meaning", "Owner entry"),
        CONTRARY_POLICY_ROWS,
        (2_400, 4_560, 2_400),
    )

    document.add_heading("5. D1-D15", level=1)
    document.add_paragraph("D6-D9 are later owner actions. Their rows must remain deferred here.")
    _add_table(
        document,
        ("ID", "Decision", "Required state", "Owner entry"),
        tuple(
            (decision_id, meaning, state, state) for decision_id, meaning, state in DECISION_ROWS
        ),
        (1_080, 4_800, 1_680, 1_800),
        body_size=8.0,
    )

    document.add_heading("6. Sign-off for draft intake", level=1)
    document.add_paragraph(
        f"Type {OWNER_CONFIRMATION_TOKEN} in the last row only after reviewing all "
        "entries. This confirms a draft, not a gold or release seal."
    )
    _add_table(
        document,
        ("Field", "Allowed value", "Owner entry"),
        (
            (
                "reviewer_role",
                identity["approval_reviewer_role"],
                identity["approval_reviewer_role"],
            ),
            (
                "reviewer_ref",
                "reviewer:<64-hex>",
                identity["approval_reviewer_ref"],
            ),
            ("decision_scope", "path_b_full_30", "path_b_full_30"),
            ("word_is_gold_or_owner_seal", "no", "no"),
            ("writes_active", "no", "no"),
            ("writes_o04", "no", "no"),
            ("confirmation_token", OWNER_CONFIRMATION_TOKEN, ""),
        ),
        (2_640, 3_600, 3_120),
    )

    _finalize_document_properties(
        document,
        title="LegalBot Live60 Path-B owner decision intake",
        subject="Local-only unsigned workflow decision intake",
    )
    return document


def export_owner_decision_intake_template(
    *,
    project_root: Path,
    output_path: Path,
    as_of_date: date,
    overwrite: bool = False,
) -> dict[str, Any]:
    destination = output_path.resolve()
    if destination.exists() and not overwrite:
        raise FileExistsError("owner decision intake template already exists")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = build_owner_decision_intake_document(
        project_root=project_root,
        as_of_date=as_of_date,
    )
    document.save(str(destination))
    payload = destination.read_bytes()
    return {
        "schema": OWNER_INTAKE_RESULT_SCHEMA,
        "artifact": "owner_decision_intake_template",
        "sha256": hashlib.sha256(payload).hexdigest(),
        "bytes": len(payload),
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_o04": False,
    }


def audit_legacy_owner_decision_docx(workbook_path: Path) -> dict[str, Any]:
    """Extract controlled decisions from the old 16-table Path-A status pack."""

    document = Document(str(workbook_path))
    if len(document.tables) != 16:
        raise ValueError("legacy owner status pack must contain exactly 16 tables")
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                lowered = cell.text.casefold()
                if (
                    "/users/" in lowered
                    or "\\users\\" in lowered
                    or lowered.startswith("file:")
                    or re.search(
                        r"\b[a-z0-9._%+-]+@[a-z0-9.-]+\.[a-z]{2,}\b",
                        lowered,
                    )
                ):
                    raise ValueError("legacy owner status pack contains private metadata")

    discard_rows = _table_rows(document.tables[1])
    if discard_rows[0][0] != "Discarded item" or any(
        not row[3].casefold().startswith("no") for row in discard_rows[1:]
    ):
        raise ValueError("legacy discarded-draft decisions are incomplete")

    route_rows = _table_rows(document.tables[3])
    if route_rows[0] != ["Path", "Meaning", "Consequence", "Your decision"]:
        raise ValueError("legacy route table is not recognised")
    selected_routes = [
        row[0].casefold()
        for row in route_rows[1:]
        if "selected" in row[3].casefold() and "not selected" not in row[3].casefold()
    ]
    if selected_routes != ["a"]:
        raise ValueError("legacy pack must record its single Path-A selection")

    indexing_rows = _table_rows(document.tables[6])
    indexing = {row[0]: row[2].casefold() for row in indexing_rows[1:]}
    required_indexing = {
        "Westlaw / scholarship full text": "stay out",
        "Find Case Law full text": "metadata-only",
        "Elasticsearch": "stay out",
        "Teaching notes as authority": "never",
        "Nearest chunk as gold": "never",
    }
    for label, expected in required_indexing.items():
        if indexing.get(label) != expected:
            raise ValueError(f"legacy indexing decision changed for {label}")

    held_rows = _table_rows(document.tables[7])
    if held_rows[0][0] != "ID" or len(held_rows) != 5:
        raise ValueError("legacy held-provision table is not recognised")
    held_ids = (
        "held-provision-01",
        "held-provision-02",
        "held-provision-03",
        "held-provision-04",
    )
    held_prefill = []
    for held_id, row in zip(held_ids, held_rows[1:], strict=True):
        raw = row[3].casefold().strip()
        if raw not in {"hold", "knowledge_gap"} and not raw.startswith("accept"):
            raise ValueError(f"legacy {held_id} disposition is unknown")
        held_prefill.append(
            {
                "held_id": held_id,
                "legacy_disposition": "hold" if raw in {"hold", "knowledge_gap"} else "accept",
            }
        )

    contrary_rows = _table_rows(document.tables[8])
    contrary = {row[0]: row[2] for row in contrary_rows[1:]}
    contrary_status = (
        "keep_pending"
        if contrary.get("status", "").casefold().startswith("keep pending")
        else "requires_new_review"
    )

    decision_rows = _table_rows(document.tables[9])
    if [row[0] for row in decision_rows[1:]] != list(OWNER_DECISION_IDS):
        raise ValueError("legacy D1-D15 rows are missing or reordered")
    legacy_states = [{"id": row[0], "state": row[3]} for row in decision_rows[1:]]

    signoff_rows = _table_rows(document.tables[15])
    attestation = next(
        row[1] for row in signoff_rows[1:] if row[0].startswith("I am the sole primary qualified")
    )
    owner_attestation_present = not attestation.casefold().startswith("pending")
    payload = {
        "schema": LEGACY_INTAKE_MIGRATION_SCHEMA,
        "source_document_sha256": hashlib.sha256(workbook_path.read_bytes()).hexdigest(),
        "legacy_route": "path_a",
        "current_route": "path_b",
        "current_target": "full_30_answer_run",
        "route_superseded_by_later_owner_instruction": True,
        "carried_forward_controls": {
            "discard_generic_locator_drafts": True,
            "westlaw_and_scholarship_full_text": "stay_out",
            "find_case_law_full_text": "metadata_only",
            "elasticsearch": "stay_out",
            "teaching_notes_as_authority": "never",
            "nearest_chunk_as_gold": "never",
            "independent_second_review": "not_required",
            "generation_authorised": False,
        },
        "held_provision_prefill": held_prefill,
        "contrary_review_legacy_status": contrary_status,
        "legacy_d1_d15_states": legacy_states,
        "legacy_owner_attestation_present": owner_attestation_present,
        "requires_new_path_b_confirmation": True,
        "owner_authored_seal": False,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_previous": False,
        "writes_o04": False,
    }
    assert_safe_evaluation_payload(payload)
    return payload


def migrate_legacy_owner_decision_docx(
    *,
    workbook_path: Path,
    project_root: Path,
    output_path: Path,
    report_path: Path,
    as_of_date: date,
    overwrite: bool = False,
) -> dict[str, Any]:
    """Create a current Path-B intake form without mutating the legacy DOCX."""

    destination = output_path.resolve()
    report_destination = report_path.resolve()
    if destination == report_destination:
        raise ValueError("migrated DOCX and report destinations must differ")
    for path in (destination, report_destination):
        if path.exists() and not overwrite:
            raise FileExistsError("legacy owner intake migration output exists")
    audit = audit_legacy_owner_decision_docx(workbook_path)
    document = build_owner_decision_intake_document(
        project_root=project_root,
        as_of_date=as_of_date,
    )
    held_by_id = {
        item["held_id"]: item["legacy_disposition"] for item in audit["held_provision_prefill"]
    }
    for row in document.tables[3].rows[1:]:
        held_id = row.cells[0].text
        if held_by_id[held_id] == "hold":
            row.cells[3].text = "hold"
    document.add_heading("Migration note", level=1)
    document.add_paragraph(
        "The earlier Path-A selection is superseded by the later Path-B/full-30 "
        "instruction. Compatible source-rights restrictions remain prefilled. "
        "Earlier held-provision choices remain hold until this current form is "
        "explicitly confirmed. No earlier owner attestation or gold seal was copied."
    )
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    os.chmod(destination, 0o600)

    report_bytes = (json.dumps(audit, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode(
        "utf-8"
    )
    report_destination.parent.mkdir(parents=True, exist_ok=True)
    report_destination.write_bytes(report_bytes)
    os.chmod(report_destination, 0o600)
    migrated_bytes = destination.read_bytes()
    return {
        "schema": OWNER_INTAKE_RESULT_SCHEMA,
        "artifact": "migrated_owner_decision_intake",
        "migrated_docx_sha256": hashlib.sha256(migrated_bytes).hexdigest(),
        "migration_report_sha256": hashlib.sha256(report_bytes).hexdigest(),
        "legacy_route_superseded": True,
        "current_route": "path_b",
        "requires_new_path_b_confirmation": True,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "writes_active": False,
        "writes_previous": False,
        "writes_o04": False,
    }


def parse_owner_decision_intake(
    workbook_path: Path,
    *,
    project_root: Path,
    as_of_date: date,
) -> dict[str, Any]:
    """Parse and validate a filled intake form without applying decisions."""

    bundle = _bundle(project_root)
    document = Document(str(workbook_path))
    if len(document.tables) != EXPECTED_TABLE_COUNT:
        raise ValueError(
            f"owner decision intake must contain exactly {EXPECTED_TABLE_COUNT} tables"
        )

    cover_rows = _rows_by_id(
        document.tables[0],
        expected_ids=(
            "Template schema",
            "Suite",
            "As-of date",
            "Registry SHA-256",
            "Run-plan SHA-256",
            "Training export",
        ),
        width=2,
        label="cover",
    )
    cover = {row[0]: row[1] for row in cover_rows}
    _require_exact(cover["Template schema"], OWNER_INTAKE_DOCX_SCHEMA, label="template schema")
    _require_exact(cover["Suite"], bundle.manifest.suite_id, label="suite")
    _require_exact(cover["As-of date"], as_of_date.isoformat(), label="as-of date")
    _require_exact(
        cover["Registry SHA-256"],
        bundle.registry.canonical_sha256,
        label="registry SHA-256",
    )
    _require_exact(
        cover["Run-plan SHA-256"],
        bundle.manifest.run_plan_sha256,
        label="run-plan SHA-256",
    )
    _require_exact(cover["Training export"], "false", label="training export")

    route_rows = _rows_by_id(
        document.tables[1],
        expected_ids=tuple(row[0] for row in ROUTE_ROWS),
        width=3,
        label="route",
    )
    route = {row[0]: row[2] for row in route_rows}
    for control, _meaning, expected in ROUTE_ROWS:
        _require_exact(route[control], expected, label=control)

    indexing_rows = _rows_by_id(
        document.tables[2],
        expected_ids=tuple(row[0] for row in INDEXING_RIGHT_ROWS),
        width=3,
        label="indexing rights",
    )
    indexing_rights = []
    for row, (control_id, _meaning, expected) in zip(
        indexing_rows, INDEXING_RIGHT_ROWS, strict=True
    ):
        _require_exact(row[2], expected, label=control_id)
        indexing_rights.append({"id": control_id, "decision": expected})

    held_rows = _rows_by_id(
        document.tables[3],
        expected_ids=tuple(item[0] for item in HELD_STATUTORY_PROVISIONS),
        width=4,
        label="held provisions",
    )
    held_provisions = []
    for row, (held_id, title, _summary) in zip(held_rows, HELD_STATUTORY_PROVISIONS, strict=True):
        if row[1] != title:
            raise ValueError(f"{held_id} title differs from the sealed template")
        if row[3] not in HELD_ALLOWED[held_id]:
            raise ValueError(f"{held_id} has an unknown disposition")
        held_provisions.append({"held_id": held_id, "disposition": row[3]})

    contrary_rows = _rows_by_id(
        document.tables[4],
        expected_ids=tuple(row[0] for row in CONTRARY_POLICY_ROWS),
        width=3,
        label="contrary policy",
    )
    contrary_policy = {}
    for row, (control_id, _meaning, expected) in zip(
        contrary_rows, CONTRARY_POLICY_ROWS, strict=True
    ):
        _require_exact(row[2], expected, label=control_id)
        contrary_policy[control_id.replace("-", "_")] = expected

    decision_rows = _rows_by_id(
        document.tables[5],
        expected_ids=OWNER_DECISION_IDS,
        width=4,
        label="D1-D15",
    )
    owner_decisions = []
    decision_defaults = {decision_id: state for decision_id, _meaning, state in DECISION_ROWS}
    for row in decision_rows:
        decision_id = row[0]
        required_state = decision_defaults[decision_id]
        if row[2] != required_state or row[3] != required_state:
            raise ValueError(f"{decision_id} must remain {required_state} in intake")
        owner_decisions.append(
            {
                "id": decision_id,
                "state": required_state,
                "later_owner_action": decision_id in LATER_OWNER_ACTIONS,
                "conditional": decision_id in CONDITIONAL_DECISION_IDS,
                "cannot_be_done_by_ai": True,
            }
        )

    signoff_rows = _rows_by_id(
        document.tables[6],
        expected_ids=(
            "reviewer_role",
            "reviewer_ref",
            "decision_scope",
            "word_is_gold_or_owner_seal",
            "writes_active",
            "writes_o04",
            "confirmation_token",
        ),
        width=3,
        label="sign-off",
    )
    signoff = {row[0]: row[2] for row in signoff_rows}
    identity = build_owner_reviewer_identity(as_of_date=as_of_date)
    _require_exact(
        signoff["reviewer_role"],
        identity["approval_reviewer_role"],
        label="reviewer role",
    )
    if not _REVIEWER_REF.fullmatch(signoff["reviewer_ref"]):
        raise ValueError("reviewer ref must be reviewer:<64-hex>")
    _require_exact(
        signoff["reviewer_ref"],
        identity["approval_reviewer_ref"],
        label="reviewer ref",
    )
    _require_exact(signoff["decision_scope"], "path_b_full_30", label="decision scope")
    _require_exact(
        signoff["word_is_gold_or_owner_seal"],
        "no",
        label="Word gold or seal flag",
    )
    _require_exact(signoff["writes_active"], "no", label="writes ACTIVE")
    _require_exact(signoff["writes_o04"], "no", label="writes O-04")
    _require_exact(
        signoff["confirmation_token"],
        OWNER_CONFIRMATION_TOKEN,
        label="confirmation token",
    )

    source_bytes = workbook_path.read_bytes()
    draft = {
        "schema": OWNER_INTAKE_DRAFT_SCHEMA,
        "suite_id": bundle.manifest.suite_id,
        "as_of_date": as_of_date.isoformat(),
        "source_document_sha256": hashlib.sha256(source_bytes).hexdigest(),
        "suite_registry_canonical_sha256": bundle.registry.canonical_sha256,
        "run_plan_sha256": bundle.manifest.run_plan_sha256,
        "route": {
            "route": route["route"],
            "target": route["target"],
            "selected_case_count": int(route["selected_case_count"]),
            "selected_issue_count": int(route["selected_issue_count"]),
            "coverage_only_disposition": route["coverage_only_disposition"],
        },
        "indexing_rights": indexing_rights,
        "held_provisions": held_provisions,
        "contrary_review_policy": contrary_policy,
        "owner_decisions": owner_decisions,
        "sign_off": {
            "reviewer_role": signoff["reviewer_role"],
            "reviewer_ref": signoff["reviewer_ref"],
            "decision_scope": signoff["decision_scope"],
            "owner_confirmation_present": True,
        },
        "owner_authored_seal": False,
        "unsigned": True,
        "ready_for_evidence_pack_generation": True,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_previous": False,
        "writes_o04": False,
    }
    assert_safe_evaluation_payload(draft)
    return draft


def render_owner_decision_intake_diff(draft: Mapping[str, Any]) -> str:
    """Render a privacy-safe review diff against the unsigned baseline."""

    lines = [
        "# Live60 owner decision intake diff",
        "",
        "This is an unsigned workflow draft. It is not legal gold, an overlay seal, "
        "ACTIVE, PREVIOUS, O-04, or generation authority.",
        "",
        "## Route",
        "",
        f"- route: `path_b` → `{draft['route']['route']}`",
        f"- target: `full_30_answer_run` → `{draft['route']['target']}`",
        f"- selected cases: `30` → `{draft['route']['selected_case_count']}`",
        f"- selected issues: `305` → `{draft['route']['selected_issue_count']}`",
        "",
        "## Indexing and rights",
        "",
    ]
    lines.extend(f"- {item['id']}: `{item['decision']}`" for item in draft["indexing_rights"])
    lines.extend(["", "## Held provisions", ""])
    lines.extend(
        f"- {item['held_id']}: `{item['disposition']}`" for item in draft["held_provisions"]
    )
    lines.extend(["", "## D1-D15 changes from unsigned", ""])
    lines.extend(
        f"- {item['id']}: `unsigned` → `{item['state']}`" for item in draft["owner_decisions"]
    )
    lines.extend(
        [
            "",
            "## Remaining mandatory gates",
            "",
            "- Review and approve exact evidence for all 305 selected-paper issues.",
            "- Complete the named-source-set contrary and limiting authority review.",
            "- Mechanically validate currentness, locators, hashes, roles, and repair spans.",
            "- Create a separate explicit owner-authored confirmation before any seal.",
            "- Pass Stage A before owner promotion; pass readiness before owner O-04.",
            "",
        ]
    )
    return "\n".join(lines)


def write_owner_decision_intake(
    *,
    workbook_path: Path,
    project_root: Path,
    as_of_date: date,
    draft_path: Path,
    diff_path: Path,
    overwrite: bool = False,
) -> dict[str, Any]:
    """Validate once, then write the unsigned draft and human-readable diff."""

    draft_destination = draft_path.resolve()
    diff_destination = diff_path.resolve()
    if draft_destination == diff_destination:
        raise ValueError("draft and diff destinations must differ")
    for destination in (draft_destination, diff_destination):
        if destination.exists() and not overwrite:
            raise FileExistsError("owner decision intake output already exists")

    draft = parse_owner_decision_intake(
        workbook_path,
        project_root=project_root,
        as_of_date=as_of_date,
    )
    draft_bytes = (json.dumps(draft, ensure_ascii=False, indent=2, sort_keys=True) + "\n").encode(
        "utf-8"
    )
    diff_bytes = render_owner_decision_intake_diff(draft).encode("utf-8")
    for destination, payload in (
        (draft_destination, draft_bytes),
        (diff_destination, diff_bytes),
    ):
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_bytes(payload)
        os.chmod(destination, 0o600)

    return {
        "schema": OWNER_INTAKE_RESULT_SCHEMA,
        "artifact": "owner_decision_intake",
        "draft_sha256": hashlib.sha256(draft_bytes).hexdigest(),
        "diff_sha256": hashlib.sha256(diff_bytes).hexdigest(),
        "ready_for_evidence_pack_generation": True,
        "ready_for_overlay_seal": False,
        "generation_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "writes_active": False,
        "writes_previous": False,
        "writes_o04": False,
    }
