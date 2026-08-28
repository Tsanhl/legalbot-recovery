"""Privacy-gated DOCX export for a completed local live-evaluation run.

The exporter consumes two kinds of input only:

* plaintext, schema-validated identifiers and metrics from ``manifest.json`` and
  ``review-export.json``; and
* the exact encrypted question plus an encrypted *released* answer, decrypted
  with :class:`~app.crypto.LocalCipher` only after the safe release gates pass.

It deliberately never reads human-review, issue-detail or knowledge-gap-detail
artifacts.  Held, failed and privacy-failed answer artifacts are not opened and
their plaintext cannot enter the document.  The generated document is an
owner-only evaluation artefact and is not eligible for training.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import tempfile
import zipfile
from collections.abc import Iterable, Mapping, Sequence
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal, Self, cast

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from lxml import etree  # type: ignore[import-untyped]
from pydantic import BaseModel, ConfigDict, Field, field_validator, model_validator

from ..crypto import LocalCipher
from ..legal_roles import REPORT_LEGAL_ROLES

REVIEW_EXPORT_SCHEMA = "legalbot.live-review-export.v1"
REVIEW_EXPORT_SCHEMA_V2 = "legalbot.live-review-export.v2"
LIVE30_EXPECTED_CASE_IDS = tuple(f"live30-q{number:02d}" for number in range(1, 31))
LIVE60_EXPECTED_CASE_IDS = tuple(
    [f"live30-q{number:02d}" for number in range(1, 31)]
    + [f"live60-q{number:02d}" for number in range(31, 61)]
)
EXPECTED_CASE_IDS = LIVE30_EXPECTED_CASE_IDS
RELEASED_STATES = frozenset({"verified_full", "verified_concise", "verified_limited"})
USABLE_WIDTH_DXA = 9_360

_SAFE_ID = re.compile(r"^[a-z0-9][a-z0-9._:-]{2,127}$")
_SAFE_CODE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._:+/@() -]{0,255}$")
_SHA256 = re.compile(r"^[0-9a-f]{64}$")
_PROHIBITED_METADATA = (
    re.compile(r"/Users/", re.IGNORECASE),
    re.compile(r"/home/", re.IGNORECASE),
    re.compile(r"[A-Za-z]:\\\\Users\\\\", re.IGNORECASE),
    re.compile(r"file://", re.IGNORECASE),
    re.compile(r"(?:^|[\s<])[^\s<>@]+@[^\s<>@]+\.[A-Za-z]{2,}(?:$|[\s>])"),
)


class ReviewExportError(ValueError):
    """Raised when a run cannot be exported without weakening privacy or integrity."""


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _sha256_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _safe_code(value: str, *, label: str) -> str:
    cleaned = value.strip()
    if not _SAFE_CODE.fullmatch(cleaned):
        raise ValueError(f"{label} is not safe for a plaintext review manifest")
    return cleaned


def _reject_prohibited_metadata(value: str, *, label: str) -> str:
    if any(pattern.search(value) for pattern in _PROHIBITED_METADATA):
        raise ReviewExportError(f"{label} contains prohibited path or identifier metadata")
    return value


def _validate_safe_json_tree(value: Any, *, location: str = "$") -> None:
    if isinstance(value, str):
        _reject_prohibited_metadata(value, label=location)
        return
    if isinstance(value, Mapping):
        for key, item in value.items():
            if not isinstance(key, str):
                raise ReviewExportError(f"{location} contains a non-string JSON key")
            _validate_safe_json_tree(item, location=f"{location}.{key}")
        return
    if isinstance(value, Sequence) and not isinstance(value, str | bytes | bytearray):
        for index, item in enumerate(value):
            _validate_safe_json_tree(item, location=f"{location}[{index}]")


class SafeMetric(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    metric_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    value: int | float | bool | str | None
    unit: str | None = None
    gate: Literal["pass", "fail", "advisory", "not_scored"] = "advisory"

    @field_validator("unit")
    @classmethod
    def unit_is_safe(cls, value: str | None) -> str | None:
        return None if value is None else _safe_code(value, label="metric unit")

    @field_validator("value")
    @classmethod
    def string_value_is_safe(
        cls, value: int | float | bool | str | None
    ) -> int | float | bool | str | None:
        return _safe_code(value, label="metric value") if isinstance(value, str) else value


class SafeEvidenceRecord(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    evidence_span_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    stable_source_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    legal_locator: str = Field(min_length=1, max_length=255)
    legal_role: str
    identity_state: Literal["verified", "unverified"]
    support_state: Literal["supported", "partial", "contrary", "limiting"]
    retrieval_rank: int | None = Field(default=None, ge=1, le=10_000)
    currentness_state: Literal[
        "verified_current",
        "latest_available_revised_snapshot",
        "historical",
        "not_applicable",
        "unverified",
    ]
    jurisdiction_state: Literal["verified", "not_applicable", "unverified"]

    @field_validator("legal_locator")
    @classmethod
    def locator_is_safe(cls, value: str) -> str:
        cleaned = " ".join(value.split())
        if not cleaned or len(cleaned) > 255:
            raise ValueError("legal locator is empty or too long")
        return _reject_prohibited_metadata(cleaned, label="legal locator")

    @field_validator("legal_role")
    @classmethod
    def role_is_from_canonical_taxonomy(cls, value: str) -> str:
        if value not in REPORT_LEGAL_ROLES:
            raise ValueError("evidence legal role is not in the canonical taxonomy")
        return value


class SafeRubricResult(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    criterion_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    score: float | None = Field(default=None, ge=0, le=100)
    status: Literal["pass", "repair", "hold", "advisory", "not_scored"]
    assessment_rule_ids: tuple[str, ...] = ()
    verification_signal: str | None = None

    @field_validator("assessment_rule_ids")
    @classmethod
    def rule_ids_are_safe_and_unique(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if any(not _SAFE_ID.fullmatch(value) for value in values):
            raise ValueError("rubric assessment rule ID is invalid")
        if len(set(values)) != len(values):
            raise ValueError("rubric assessment rule IDs are duplicated")
        return values

    @field_validator("verification_signal")
    @classmethod
    def verification_signal_is_safe(cls, value: str | None) -> str | None:
        return None if value is None else _safe_code(value, label="verification signal")


class SafeRepairRecord(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    repair_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    section_id: str | None = Field(default=None, pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    reason_code: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    status: Literal["completed", "failed", "skipped", "pending"]
    attempt_count: int = Field(ge=0, le=100)


class SafeGapRecord(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    gap_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    category: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    severity: Literal["critical", "high", "medium", "low"]
    status: Literal[
        "open",
        "triaged",
        "source_needed",
        "metadata_currentness_needed",
        "retrieval_fix_needed",
        "accepted_out_of_scope",
        "resolved",
        "regression_verified",
    ]
    safe_expected_ids: tuple[str, ...] = ()
    safe_observed_ids: tuple[str, ...] = ()

    @field_validator("safe_expected_ids", "safe_observed_ids")
    @classmethod
    def referenced_ids_are_safe(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if any(not _SAFE_ID.fullmatch(value) for value in values):
            raise ValueError("gap contains an invalid safe reference ID")
        return values


class SafeAdvisoryAIReview(BaseModel):
    """Prose-free AI recommendation summary for owner review, never gate authority."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    status: Literal["available", "unavailable", "not_run"]
    reviewer_execution_mode: Literal["separate_verification_pass_same_model_adapter"] = (
        "separate_verification_pass_same_model_adapter"
    )
    model_independent: Literal[False] = False
    recommendations_only: Literal[True] = True
    can_decide_or_adopt: Literal[False] = False
    can_admit_sources: Literal[False] = False
    can_authorize_gates: Literal[False] = False
    may_raise_fail_closed_owner_review_hold: Literal[True] = True
    review_sha256: str | None = Field(default=None, pattern=r"^[0-9a-f]{64}$")
    recommendation_codes: tuple[str, ...] = ()
    flagged_claim_count: int = Field(default=0, ge=0, le=100_000)
    owner_review_required: bool = False
    unavailable_reason_code: str | None = Field(
        default=None, pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$"
    )

    @field_validator("recommendation_codes")
    @classmethod
    def recommendation_codes_are_safe_unique(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if any(not _SAFE_ID.fullmatch(value) for value in values):
            raise ValueError("AI recommendation contains an invalid safe code")
        if len(values) != len(set(values)):
            raise ValueError("AI recommendation codes are duplicated")
        return values

    @model_validator(mode="after")
    def advisory_record_is_consistent(self) -> Self:
        if self.status == "available":
            if self.review_sha256 is None or self.unavailable_reason_code is not None:
                raise ValueError("available AI recommendation lacks its exact review binding")
            if self.owner_review_required != bool(self.flagged_claim_count):
                raise ValueError("AI owner-review flag differs from its claim findings")
        elif (
            self.review_sha256 is not None or self.flagged_claim_count or self.recommendation_codes
        ):
            raise ValueError("unavailable AI recommendation contains fabricated review results")
        elif self.status == "unavailable" and self.unavailable_reason_code is None:
            raise ValueError("unavailable AI recommendation lacks a safe reason code")
        elif self.status == "not_run" and self.unavailable_reason_code is not None:
            raise ValueError("not-run AI recommendation cannot claim an unavailability reason")
        return self


class SafeClusterRecord(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    cluster_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    category: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    case_ids: tuple[str, ...]
    status: Literal["open", "triaged", "resolved", "regression_verified"]

    @field_validator("case_ids")
    @classmethod
    def case_ids_are_valid(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if any(value not in LIVE60_EXPECTED_CASE_IDS for value in values):
            raise ValueError("cluster contains an invalid live-evaluation case ID")
        return values


class SafeCorrectionRecord(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    correction_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    affected_layer: Literal[
        "index",
        "filtering",
        "retrieval",
        "generation",
        "claims",
        "citations",
        "privacy",
        "routing",
        "worker",
        "ui",
    ]
    case_ids: tuple[str, ...]
    status: Literal["proposed", "implemented", "verified", "rejected"]
    regression_case_id: str | None = Field(default=None, pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")

    @field_validator("case_ids")
    @classmethod
    def case_ids_are_valid(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if any(value not in LIVE60_EXPECTED_CASE_IDS for value in values):
            raise ValueError("correction contains an invalid live-evaluation case ID")
        return values


class SafeOwnerDecision(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    decision_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    decision_code: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    status: Literal["owner_review_needed", "approved", "rejected", "deferred"]
    affected_case_ids: tuple[str, ...] = ()

    @field_validator("affected_case_ids")
    @classmethod
    def case_ids_are_valid(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if any(value not in LIVE60_EXPECTED_CASE_IDS for value in values):
            raise ValueError("owner decision contains an invalid live-evaluation case ID")
        return values


class SafeRegressionRecord(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    regression_case_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    source_issue_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    status: Literal["pending", "passing", "failing", "retired"]
    fixed_version: str | None = None
    verification_run_id: str | None = Field(default=None, pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")

    @field_validator("fixed_version")
    @classmethod
    def fixed_version_is_safe(cls, value: str | None) -> str | None:
        return None if value is None else _safe_code(value, label="fixed version")


class LiveReviewCase(BaseModel):
    """Safe per-case release record. It contains no question or answer prose."""

    model_config = ConfigDict(extra="forbid", frozen=True)

    case_id: str = Field(pattern=r"^(?:live30|live60)-q[0-9]{2}$")
    ordinal: int = Field(ge=1, le=60)
    run_plan_disposition: Literal["generate_once", "coverage_only_not_selected"] | None = None
    run_plan_outcome_count: int | None = Field(default=None, ge=0, le=1)
    coverage_status: str | None = Field(default=None, pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    case_status: Literal["completed", "held", "failed", "cancelled"]
    release_state: Literal[
        "verified_full",
        "verified_concise",
        "verified_limited",
        "held",
        "not_released",
        "privacy_failed",
        "evidence_failed",
    ]
    released: bool
    privacy_passed: bool
    evidence_passed: bool
    question_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    answer_artifact_id: str | None = Field(default=None, pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    answer_sha256: str | None = Field(default=None, pattern=r"^[0-9a-f]{64}$")
    subject: str = Field(min_length=2, max_length=120)
    task_type: Literal["essay", "problem", "general"]
    jurisdiction: str
    as_of_date: str
    word_target: int = Field(ge=1_000, le=10_000)
    word_count: int | None = Field(default=None, ge=0, le=100_000)
    research_route: Literal["direct", "sectioned", "full_enquiry"]
    drafting_route: Literal["direct", "sectioned", "full_enquiry"]
    assessment_bundle_sha256: str | None = Field(default=None, pattern=r"^[0-9a-f]{64}$")
    assessment_rule_ids: tuple[str, ...] = ()
    evidence: tuple[SafeEvidenceRecord, ...] = ()
    rubric: tuple[SafeRubricResult, ...] = ()
    repairs: tuple[SafeRepairRecord, ...] = ()
    gaps: tuple[SafeGapRecord, ...] = ()
    advisory_ai_review: SafeAdvisoryAIReview | None = None
    metrics: tuple[SafeMetric, ...] = ()
    failure_codes: tuple[str, ...] = ()

    @field_validator("subject", "jurisdiction", "as_of_date")
    @classmethod
    def labels_are_safe(cls, value: str) -> str:
        return _safe_code(value, label="case label")

    @field_validator("assessment_rule_ids")
    @classmethod
    def assessment_rule_ids_are_safe(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if any(not _SAFE_ID.fullmatch(value) for value in values):
            raise ValueError("case contains an invalid assessment rule ID")
        if len(set(values)) != len(values):
            raise ValueError("case assessment rule IDs are duplicated")
        return values

    @field_validator("failure_codes")
    @classmethod
    def failure_codes_are_safe(cls, values: tuple[str, ...]) -> tuple[str, ...]:
        if any(not _SAFE_ID.fullmatch(value) for value in values):
            raise ValueError("case contains an invalid failure code")
        return values

    @model_validator(mode="after")
    def release_contract_is_consistent(self) -> Self:
        expected_id = (
            f"live30-q{self.ordinal:02d}" if self.ordinal <= 30 else f"live60-q{self.ordinal:02d}"
        )
        if self.case_id != expected_id:
            raise ValueError("case ID and ordinal disagree")
        if self.run_plan_disposition == "coverage_only_not_selected" and self.released:
            raise ValueError("a coverage-only case cannot expose a released answer")
        is_released = self.release_state in RELEASED_STATES
        if is_released != self.released:
            raise ValueError("released flag and release state disagree")
        if is_released:
            if (
                self.case_status != "completed"
                or not self.privacy_passed
                or not self.evidence_passed
            ):
                raise ValueError(
                    "a released case has not passed completion, privacy and evidence gates"
                )
            if not self.answer_artifact_id or not self.answer_sha256:
                raise ValueError("a released case must bind an encrypted answer artifact and hash")
            if self.word_count is None:
                raise ValueError("a released case must record its word count")
            if self.release_state != "verified_limited" and not self.evidence:
                raise ValueError("a released case must report at least one frozen evidence span")
            if not self.assessment_bundle_sha256 or not self.assessment_rule_ids:
                raise ValueError("a released case must report its assessment bundle and rule IDs")
        return self


class LiveReviewExport(BaseModel):
    """Plaintext-safe master review manifest; answer prose is never allowed here."""

    model_config = ConfigDict(extra="forbid", frozen=True, populate_by_name=True)

    schema_name: Literal["legalbot.live-review-export.v1", "legalbot.live-review-export.v2"] = (
        Field(default="legalbot.live-review-export.v1", alias="schema")
    )
    run_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    run_manifest_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    generated_at: datetime
    run_status: Literal["created", "running", "completed", "failed", "cancelled"]
    privacy_report_passed: bool
    purpose: Literal["evaluation_only"] = "evaluation_only"
    eligible_for_training: Literal[False] = False
    training_export_allowed: Literal[False] = False
    expected_case_count: int = Field(default=30, ge=1, le=60)
    run_plan_id: str | None = Field(default=None, pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    run_plan_file_sha256: str | None = Field(default=None, pattern=r"^[0-9a-f]{64}$")
    run_plan_seal_sha256: str | None = Field(default=None, pattern=r"^[0-9a-f]{64}$")
    cases: tuple[LiveReviewCase, ...]
    aggregate_metrics: tuple[SafeMetric, ...] = ()
    clusters: tuple[SafeClusterRecord, ...] = ()
    corrections: tuple[SafeCorrectionRecord, ...] = ()
    owner_decisions: tuple[SafeOwnerDecision, ...] = ()
    regressions: tuple[SafeRegressionRecord, ...] = ()

    @model_validator(mode="after")
    def case_registry_is_unique(self) -> Self:
        case_ids = [case.case_id for case in self.cases]
        ordinals = [case.ordinal for case in self.cases]
        if len(case_ids) != len(set(case_ids)) or len(ordinals) != len(set(ordinals)):
            raise ValueError("review export contains duplicate cases")
        if len(self.cases) > self.expected_case_count:
            raise ValueError("review export contains more cases than declared")
        if self.schema_name == REVIEW_EXPORT_SCHEMA_V2:
            if (
                self.expected_case_count != 60
                or not self.run_plan_id
                or not self.run_plan_file_sha256
                or not self.run_plan_seal_sha256
            ):
                raise ValueError("Live60 review export lacks its run-plan binding")
            for case in self.cases:
                if case.run_plan_disposition is None or case.run_plan_outcome_count is None:
                    raise ValueError("Live60 case lacks its immutable disposition")
        return self


@dataclass(frozen=True, slots=True)
class RunIdentity:
    manifest_schema: str
    run_id: str
    suite_id: str
    suite_version: str
    suite_file_sha256: str
    suite_canonical_sha256: str
    as_of_date: str
    git_sha: str
    git_dirty: bool
    model_version: str | None
    index_build_id: str | None
    prompt_version: str | None
    router_version: str | None
    classifier_version: str | None
    policy_sha256: str | None
    assessment_rules_sha256: str | None
    run_plan_id: str | None = None
    run_plan_file_sha256: str | None = None
    run_plan_seal_sha256: str | None = None


@dataclass(frozen=True, slots=True)
class ReviewContract:
    expected_case_ids: tuple[str, ...]
    selected_case_ids: tuple[str, ...]
    coverage_only_case_ids: tuple[str, ...]
    annexes: Mapping[str, tuple[str, ...]]
    run_plan_id: str | None
    run_plan_file_sha256: str | None
    run_plan_seal_sha256: str | None


@dataclass(frozen=True, slots=True)
class LoadedReviewCase:
    safe: LiveReviewCase
    question: str | None
    released_answer: str | None


@dataclass(frozen=True, slots=True)
class LoadedLiveReview:
    identity: RunIdentity
    review: LiveReviewExport
    cases: tuple[LoadedReviewCase, ...]
    incomplete_fixture: bool
    contract: ReviewContract


def _read_safe_json(path: Path) -> tuple[dict[str, Any], bytes]:
    if not path.is_file():
        raise ReviewExportError(f"required safe manifest is missing: {path.name}")
    raw = path.read_bytes()
    try:
        value = json.loads(raw)
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ReviewExportError(f"safe manifest is invalid JSON: {path.name}") from exc
    if not isinstance(value, dict):
        raise ReviewExportError(f"safe manifest is not an object: {path.name}")
    _validate_safe_json_tree(value)
    return value, raw


def _safe_version(value: Any, *, label: str, optional: bool = False) -> str | None:
    if value is None and optional:
        return None
    if not isinstance(value, str):
        raise ReviewExportError(f"{label} is missing or not a string")
    return _safe_code(value, label=label)


def _run_identity(run_manifest: Mapping[str, Any], *, expected_run_id: str) -> RunIdentity:
    manifest_schema = str(run_manifest.get("schema") or "")
    if manifest_schema not in {
        "legalbot.e2e-run-manifest.v1",
        "legalbot.live-evaluation-run-manifest.v2",
    }:
        raise ReviewExportError("run manifest schema is not supported")
    if run_manifest.get("run_id") != expected_run_id:
        raise ReviewExportError("run manifest identity does not match the review export")
    provenance = run_manifest.get("provenance")
    if not isinstance(provenance, Mapping):
        raise ReviewExportError("run manifest provenance is missing")
    git_sha = str(provenance.get("git_sha", ""))
    if not re.fullmatch(r"[0-9a-f]{7,64}", git_sha):
        raise ReviewExportError("run manifest git identity is invalid")
    suite_file_sha256 = str(
        run_manifest.get("suite_file_sha256", run_manifest.get("suite_registry_file_sha256", ""))
    )
    suite_canonical_sha256 = str(
        run_manifest.get(
            "suite_canonical_sha256",
            run_manifest.get("suite_registry_canonical_sha256", ""),
        )
    )
    if not _SHA256.fullmatch(suite_file_sha256) or not _SHA256.fullmatch(suite_canonical_sha256):
        raise ReviewExportError("run manifest suite hashes are invalid")
    policy_sha256 = provenance.get("policy_sha256")
    rules_sha256 = provenance.get("assessment_rules_sha256")
    if policy_sha256 is not None and not _SHA256.fullmatch(str(policy_sha256)):
        raise ReviewExportError("run policy hash is invalid")
    if rules_sha256 is not None and not _SHA256.fullmatch(str(rules_sha256)):
        raise ReviewExportError("run assessment-rules hash is invalid")
    for key in ("run_plan_file_sha256", "run_plan_seal_sha256"):
        value = run_manifest.get(key)
        if value is not None and not _SHA256.fullmatch(str(value)):
            raise ReviewExportError(f"run manifest {key} is invalid")
    git_dirty = provenance.get("git_dirty")
    if not isinstance(git_dirty, bool):
        raise ReviewExportError("run manifest git-dirty flag is invalid")
    return RunIdentity(
        manifest_schema=manifest_schema,
        run_id=expected_run_id,
        suite_id=str(_safe_version(run_manifest.get("suite_id"), label="suite ID")),
        suite_version=str(_safe_version(run_manifest.get("suite_version"), label="suite version")),
        suite_file_sha256=suite_file_sha256,
        suite_canonical_sha256=suite_canonical_sha256,
        as_of_date=str(_safe_version(run_manifest.get("as_of_date"), label="as-of date")),
        git_sha=git_sha,
        git_dirty=git_dirty,
        model_version=_safe_version(
            provenance.get("model_version"), label="model version", optional=True
        ),
        index_build_id=_safe_version(
            provenance.get("index_build_id"), label="index build ID", optional=True
        ),
        prompt_version=_safe_version(
            provenance.get("prompt_version"), label="prompt version", optional=True
        ),
        router_version=_safe_version(
            provenance.get("router_version"), label="router version", optional=True
        ),
        classifier_version=_safe_version(
            provenance.get("classifier_version"), label="classifier version", optional=True
        ),
        policy_sha256=None if policy_sha256 is None else str(policy_sha256),
        assessment_rules_sha256=None if rules_sha256 is None else str(rules_sha256),
        run_plan_id=_safe_version(
            run_manifest.get("run_plan_id"), label="run-plan ID", optional=True
        ),
        run_plan_file_sha256=(
            str(run_manifest["run_plan_file_sha256"])
            if run_manifest.get("run_plan_file_sha256") is not None
            else None
        ),
        run_plan_seal_sha256=(
            str(run_manifest["run_plan_seal_sha256"])
            if run_manifest.get("run_plan_seal_sha256") is not None
            else None
        ),
    )


def _review_contract(
    *, run_dir: Path, run_manifest: Mapping[str, Any], identity: RunIdentity
) -> ReviewContract:
    if identity.manifest_schema == "legalbot.e2e-run-manifest.v1":
        if identity.suite_id != "live-evaluation-30-v1":
            raise ReviewExportError("legacy review manifest is not the sealed Live30 suite")
        return ReviewContract(
            expected_case_ids=LIVE30_EXPECTED_CASE_IDS,
            selected_case_ids=LIVE30_EXPECTED_CASE_IDS,
            coverage_only_case_ids=(),
            annexes={},
            run_plan_id=None,
            run_plan_file_sha256=None,
            run_plan_seal_sha256=None,
        )

    from .live_suite import LiveGenerationRunPlan, LiveSuiteManifest

    suite_path = run_dir / "suite-manifest.json"
    run_plan_path = run_dir / "generation-run-plan.json"
    try:
        suite = LiveSuiteManifest.model_validate_json(suite_path.read_bytes())
        plan = LiveGenerationRunPlan.model_validate_json(run_plan_path.read_bytes())
    except (OSError, TypeError, ValueError) as exc:
        raise ReviewExportError("Live60 suite or run-plan snapshot is invalid") from exc
    if (
        identity.suite_id != suite.suite_id
        or run_manifest.get("suite_manifest_seal_sha256") != suite.seal_sha256
        or identity.suite_file_sha256 != suite.registry_file_sha256
        or identity.suite_canonical_sha256 != suite.registry_canonical_sha256
        or identity.run_plan_id != plan.run_plan_id
        or identity.run_plan_file_sha256 != suite.run_plan_sha256
        or identity.run_plan_seal_sha256 != plan.seal_sha256
        or _sha256_bytes(run_plan_path.read_bytes()) != suite.run_plan_sha256
        or plan.suite_registry_canonical_sha256 != suite.registry_canonical_sha256
    ):
        raise ReviewExportError("Live60 run snapshots do not match the run manifest")
    selected = tuple(item.case_id for item in plan.cases if item.disposition == "generate_once")
    coverage_only = tuple(
        item.case_id for item in plan.cases if item.disposition == "coverage_only_not_selected"
    )
    if (
        tuple(item.case_id for item in plan.cases) != LIVE60_EXPECTED_CASE_IDS
        or len(selected) != 30
        or len(coverage_only) != 30
        or sum(item.pass_count for item in plan.cases) != 30
        or plan.stability_repeats != 0
    ):
        raise ReviewExportError("Live60 run plan is not the exact single-pass contract")
    return ReviewContract(
        expected_case_ids=LIVE60_EXPECTED_CASE_IDS,
        selected_case_ids=selected,
        coverage_only_case_ids=coverage_only,
        annexes={key: tuple(value) for key, value in plan.annexes.items()},
        run_plan_id=plan.run_plan_id,
        run_plan_file_sha256=suite.run_plan_sha256,
        run_plan_seal_sha256=plan.seal_sha256,
    )


def _decrypt_question(
    *, run_dir: Path, run_id: str, case: LiveReviewCase, cipher: LocalCipher
) -> str:
    path = run_dir / "cases" / case.case_id / "question.enc"
    if not path.is_file():
        raise ReviewExportError(f"encrypted question is missing for {case.case_id}")
    try:
        envelope = json.loads(cipher.decrypt_text(path.read_bytes()))
    except (json.JSONDecodeError, UnicodeDecodeError) as exc:
        raise ReviewExportError(
            f"encrypted question envelope is invalid for {case.case_id}"
        ) from exc
    if not isinstance(envelope, dict):
        raise ReviewExportError(f"encrypted question envelope is not an object for {case.case_id}")
    private_case = envelope.get("case")
    expected_schema = (
        "legalbot.e2e-encrypted-question.v2"
        if case.run_plan_disposition is not None
        else "legalbot.e2e-encrypted-question.v1"
    )
    expected = (expected_schema, run_id, case.case_id)
    observed = (
        envelope.get("schema"),
        envelope.get("run_id"),
        private_case.get("case_id") if isinstance(private_case, dict) else None,
    )
    if observed != expected:
        raise ReviewExportError(f"encrypted question identity check failed for {case.case_id}")
    question = private_case.get("question") if isinstance(private_case, dict) else None
    if not isinstance(question, str) or not question.strip():
        raise ReviewExportError(f"encrypted question is empty for {case.case_id}")
    if _sha256_text(question) != case.question_sha256:
        raise ReviewExportError(f"encrypted question hash check failed for {case.case_id}")
    return _reject_prohibited_metadata(question, label=f"question {case.case_id}")


def _decrypt_released_answer(
    *, run_dir: Path, run_id: str, case: LiveReviewCase, cipher: LocalCipher
) -> str:
    if case.release_state not in RELEASED_STATES:
        raise ReviewExportError("internal error: attempted to decrypt an unreleased answer")
    artifact_id = case.answer_artifact_id
    if artifact_id is None:
        raise ReviewExportError(f"released answer artifact ID is missing for {case.case_id}")
    if case.run_plan_disposition is not None:
        path = run_dir / "cases" / case.case_id / "artifacts" / f"answer-{artifact_id}.enc"
    else:
        path = run_dir / "cases" / case.case_id / "artifacts" / "answer" / f"{artifact_id}.enc"
    if not path.is_file():
        raise ReviewExportError(f"released answer artifact is missing for {case.case_id}")
    try:
        plaintext = cipher.decrypt_text(path.read_bytes())
    except UnicodeDecodeError as exc:
        raise ReviewExportError(f"released answer is invalid for {case.case_id}") from exc
    if case.run_plan_disposition is not None:
        answer = plaintext
        if not answer.strip():
            raise ReviewExportError(f"released answer is empty for {case.case_id}")
        if _sha256_text(answer) != case.answer_sha256:
            raise ReviewExportError(f"released answer hash check failed for {case.case_id}")
        return _reject_prohibited_metadata(answer, label=f"released answer {case.case_id}")
    try:
        envelope = json.loads(plaintext)
    except json.JSONDecodeError as exc:
        raise ReviewExportError(f"released answer envelope is invalid for {case.case_id}") from exc
    expected = (
        "legalbot.e2e-sensitive-artifact.v1",
        run_id,
        case.case_id,
        artifact_id,
        "answer",
    )
    observed = (
        envelope.get("schema") if isinstance(envelope, dict) else None,
        envelope.get("run_id") if isinstance(envelope, dict) else None,
        envelope.get("case_id") if isinstance(envelope, dict) else None,
        envelope.get("artifact_id") if isinstance(envelope, dict) else None,
        envelope.get("kind") if isinstance(envelope, dict) else None,
    )
    if observed != expected:
        raise ReviewExportError(f"released answer identity check failed for {case.case_id}")
    legacy_answer = envelope.get("content") if isinstance(envelope, dict) else None
    if not isinstance(legacy_answer, str) or not legacy_answer.strip():
        raise ReviewExportError(f"released answer is empty for {case.case_id}")
    if _sha256_text(legacy_answer) != case.answer_sha256:
        raise ReviewExportError(f"released answer hash check failed for {case.case_id}")
    return _reject_prohibited_metadata(legacy_answer, label=f"released answer {case.case_id}")


def load_live_review(
    *,
    run_dir: Path,
    cipher: LocalCipher,
    require_complete: bool = True,
) -> LoadedLiveReview:
    """Load an evaluation run without ever opening an unreleased answer artifact."""

    resolved_run_dir = run_dir.resolve()
    run_manifest, run_manifest_bytes = _read_safe_json(resolved_run_dir / "manifest.json")
    review_value, _ = _read_safe_json(resolved_run_dir / "review-export.json")
    try:
        review = LiveReviewExport.model_validate(review_value)
    except Exception as exc:
        raise ReviewExportError("review-export.json failed schema validation") from exc
    if review.run_manifest_sha256 != _sha256_bytes(run_manifest_bytes):
        raise ReviewExportError("review export is not bound to the current run manifest")
    if not review.privacy_report_passed:
        raise ReviewExportError("run privacy report did not pass; no review DOCX may be exported")
    identity = _run_identity(run_manifest, expected_run_id=review.run_id)
    contract = _review_contract(
        run_dir=resolved_run_dir, run_manifest=run_manifest, identity=identity
    )
    is_live60 = identity.manifest_schema == "legalbot.live-evaluation-run-manifest.v2"
    if is_live60:
        if (
            review.schema_name != REVIEW_EXPORT_SCHEMA_V2
            or review.run_plan_id != contract.run_plan_id
            or review.run_plan_file_sha256 != contract.run_plan_file_sha256
            or review.run_plan_seal_sha256 != contract.run_plan_seal_sha256
        ):
            raise ReviewExportError("Live60 review export is not bound to its run plan")
    elif review.schema_name != REVIEW_EXPORT_SCHEMA:
        raise ReviewExportError("legacy Live30 run requires the v1 review schema")

    expected_ids = set(contract.expected_case_ids)
    observed_ids = {case.case_id for case in review.cases}
    incomplete = (
        review.run_status != "completed"
        or review.expected_case_count != len(contract.expected_case_ids)
        or len(review.cases) != len(contract.expected_case_ids)
        or observed_ids != expected_ids
    )
    if require_complete and incomplete:
        if is_live60:
            raise ReviewExportError(
                "strict export requires one completed run reconciled to the exact suite"
            )
        raise ReviewExportError("strict export requires one completed, reconciled 30-case run")

    if is_live60:
        by_id = {case.case_id: case for case in review.cases}
        for case_id in contract.selected_case_ids:
            case = by_id.get(case_id)
            if (
                case is None
                or case.run_plan_disposition != "generate_once"
                or case.run_plan_outcome_count != 1
                or case.case_status not in {"completed", "held", "failed"}
            ):
                raise ReviewExportError("selected Live60 case lacks exactly one terminal outcome")
        for case_id in contract.coverage_only_case_ids:
            case = by_id.get(case_id)
            if (
                case is None
                or case.run_plan_disposition != "coverage_only_not_selected"
                or case.run_plan_outcome_count != 0
                or case.case_status != "completed"
                or case.released
            ):
                raise ReviewExportError(
                    "coverage-only Live60 case has an invalid generation outcome"
                )

    loaded: list[LoadedReviewCase] = []
    for case in sorted(review.cases, key=lambda item: item.ordinal):
        if case.release_state in RELEASED_STATES:
            question = _decrypt_question(
                run_dir=resolved_run_dir, run_id=review.run_id, case=case, cipher=cipher
            )
            answer = _decrypt_released_answer(
                run_dir=resolved_run_dir, run_id=review.run_id, case=case, cipher=cipher
            )
            loaded.append(LoadedReviewCase(case, question, answer))
        else:
            # Deliberately do not read question.enc or any answer artifact here.
            loaded.append(LoadedReviewCase(case, None, None))
    return LoadedLiveReview(identity, review, tuple(loaded), incomplete, contract)


def _set_cell_shading(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell: _Cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != USABLE_WIDTH_DXA:
        raise ValueError("table column widths must sum to the usable page width")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(USABLE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width_dxa))
        grid.append(column)

    margins = properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        properties.append(margins)
    for edge, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width_dxa in zip(row.cells, widths_dxa, strict=True):
            _set_cell_width(cell, width_dxa)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row: _Row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_cell_text(
    cell: _Cell,
    text: str,
    *,
    bold: bool = False,
    size: float = 8.5,
    color: str | None = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_table(
    document: DocumentType,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths_dxa: Sequence[int],
    *,
    body_size: float = 8.5,
) -> Table:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_shading(cell, "F2F4F7")
        _set_cell_text(cell, value, bold=True, size=body_size, color="1F4D78")
    _repeat_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values, strict=True):
            _set_cell_text(cell, value, size=body_size)
    _set_table_geometry(table, widths_dxa)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return cast(Table, table)


def _set_run_font(run: Any, *, size: float, bold: bool = False, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _configure_styles(document: DocumentType) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    specifications = {
        "Title": (24, "1F4D78", 0, 14),
        "Subtitle": (12, "5B6573", 0, 10),
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "1F4D78", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in specifications.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Review Small" not in document.styles:
        style = document.styles.add_style("Review Small", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
    small = document.styles["Review Small"]
    small.font.name = "Calibri"
    small._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    small.font.size = Pt(8.5)
    small.font.color.rgb = RGBColor.from_string("5B6573")
    small.paragraph_format.space_after = Pt(3)
    small.paragraph_format.line_spacing = 1.0


def _configure_page(document: DocumentType) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def _add_field(paragraph: Paragraph, instruction: str, result: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction_node = OxmlElement("w:instrText")
    instruction_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction_node.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend((begin, instruction_node, separate))
    if result:
        result_run = paragraph.add_run(result)
        _set_run_font(result_run, size=9, color="5B6573")
    paragraph.add_run()._r.append(end)


def _add_page_break(document: DocumentType) -> None:
    paragraph = document.add_paragraph()
    break_node = OxmlElement("w:br")
    break_node.set(qn("w:type"), "page")
    paragraph.add_run()._r.append(break_node)


def _add_header_footer(document: DocumentType, run_id: str) -> None:
    section = document.sections[0]
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    _set_cell_text(
        table.cell(0, 0), "LEGALBOT · LIVE EVALUATION", bold=True, size=8, color="1F4D78"
    )
    right = table.cell(0, 1)
    _set_cell_text(right, f"OWNER REVIEW · {run_id}", bold=True, size=8, color="1F4D78")
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_table_geometry(table, (4_680, 4_680))

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Evaluation only · Local controlled copy · Page ")
    _set_run_font(run, size=8, color="5B6573")
    _add_field(paragraph, " PAGE ", "1")


def _add_toc(document: DocumentType) -> None:
    document.add_heading("Contents", level=1)
    paragraph = document.add_paragraph()
    _add_field(
        paragraph,
        ' TOC \\o "1-3" \\h \\z \\u ',
        "Open in Word and update this field to populate the contents list.",
    )
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    document.settings.element.append(update)


def _identity_rows(identity: RunIdentity) -> list[tuple[str, str]]:
    return [
        ("Run", identity.run_id),
        ("Suite", f"{identity.suite_id} · {identity.suite_version}"),
        ("As-of date", identity.as_of_date),
        ("Git", f"{identity.git_sha}{' · dirty' if identity.git_dirty else ' · clean'}"),
        ("Model", identity.model_version or "not recorded"),
        ("Index", identity.index_build_id or "not recorded"),
        (
            "Prompt / router",
            f"{identity.prompt_version or 'not recorded'} / {identity.router_version or 'not recorded'}",
        ),
        ("Policy SHA", identity.policy_sha256 or "not recorded"),
        ("Assessment SHA", identity.assessment_rules_sha256 or "not recorded"),
    ]


def _metric_display(metric: SafeMetric) -> str:
    value = "—" if metric.value is None else str(metric.value)
    return f"{value}{' ' + metric.unit if metric.unit else ''}"


def _add_cover(document: DocumentType, loaded: LoadedLiveReview) -> None:
    paragraph = document.add_paragraph("LEGALBOT", style="Subtitle")
    paragraph.paragraph_format.space_before = Pt(54)
    title_text = (
        "Live60 evaluation\ncontrol owner review"
        if loaded.identity.suite_id == "live-evaluation-60-v1"
        else "Live evaluation\nmaster owner review"
    )
    title = document.add_paragraph(title_text, style="Title")
    title.paragraph_format.keep_with_next = True
    subtitle = (
        "60-case coverage and status record · answers packaged in three annexes"
        if loaded.identity.suite_id == "live-evaluation-60-v1"
        else "30-case release, evidence, quality and defect record"
    )
    if loaded.incomplete_fixture:
        subtitle = "INCOMPLETE FIXTURE · " + subtitle
    document.add_paragraph(subtitle, style="Subtitle")
    status = document.add_paragraph()
    run = status.add_run("CONTROLLED LOCAL DOCUMENT · EVALUATION ONLY · NOT TRAINING DATA")
    _set_run_font(run, size=9, bold=True, color="9C2F2F")
    status.paragraph_format.space_before = Pt(8)
    status.paragraph_format.space_after = Pt(18)
    _add_table(
        document,
        ("Identity", "Recorded value"),
        _identity_rows(loaded.identity),
        (2_180, 7_180),
        body_size=9,
    )
    note = document.add_paragraph(style="Review Small")
    note.add_run(
        "This document contains only safe run metadata and plaintext for answers whose "
        "release, privacy and evidence gates all passed. Held or failed drafts are omitted."
    )
    _add_page_break(document)


def _add_executive_metrics(document: DocumentType, loaded: LoadedLiveReview) -> None:
    document.add_heading("Executive review", level=1)
    released = [case for case in loaded.cases if case.released_answer is not None]
    held = [case for case in loaded.cases if case.released_answer is None]
    evidence_count = sum(len(case.safe.evidence) for case in loaded.cases)
    repair_count = sum(len(case.safe.repairs) for case in loaded.cases)
    gap_count = sum(len(case.safe.gaps) for case in loaded.cases)
    calculated = [
        (
            "Cases reconciled",
            str(len(loaded.cases)),
            ("pass" if len(loaded.cases) == len(loaded.contract.expected_case_ids) else "advisory"),
        ),
        (
            "Released answers",
            str(len(released)),
            "pass" if len(released) == len(loaded.cases) else "advisory",
        ),
        ("Held / omitted plaintext", str(len(held)), "advisory"),
        ("Safe evidence rows", str(evidence_count), "advisory"),
        ("Repair records", str(repair_count), "advisory"),
        ("Knowledge-gap records", str(gap_count), "advisory"),
        ("Run privacy report", "PASS", "pass"),
    ]
    _add_table(
        document,
        ("Measure", "Value", "Gate"),
        calculated,
        (4_960, 2_200, 2_200),
        body_size=9,
    )
    if loaded.review.aggregate_metrics:
        document.add_heading("Recorded aggregate metrics", level=2)
        _add_table(
            document,
            ("Metric", "Value", "Gate"),
            (
                (metric.metric_id, _metric_display(metric), metric.gate)
                for metric in loaded.review.aggregate_metrics
            ),
            (4_960, 2_200, 2_200),
            body_size=9,
        )


def _case_gate_text(case: LiveReviewCase) -> str:
    return f"privacy={'P' if case.privacy_passed else 'F'} · evidence={'P' if case.evidence_passed else 'F'}"


def _add_case_matrix(document: DocumentType, loaded: LoadedLiveReview) -> None:
    matrix_title = (
        "Thirty-case control matrix"
        if loaded.identity.suite_id == "live-evaluation-30-v1"
        else f"{len(loaded.contract.expected_case_ids)}-case control matrix"
    )
    document.add_heading(matrix_title, level=1)
    document.add_paragraph(
        "The matrix is complete only when reconciled to the immutable suite and run plan. "
        "Omitted plaintext remains encrypted."
    )
    rows = []
    for item in loaded.cases:
        case = item.safe
        actual = "—" if case.word_count is None else str(case.word_count)
        rows.append(
            (
                f"{case.ordinal:02d}",
                case.case_id,
                case.subject,
                f"{case.word_target:,}\n{actual}",
                case.research_route,
                case.release_state,
                _case_gate_text(case),
                (
                    f"{case.run_plan_disposition or 'legacy'}\n"
                    f"E {len(case.evidence)} · R {len(case.repairs)} · G {len(case.gaps)}"
                ),
            )
        )
    _add_table(
        document,
        (
            "#",
            "Case",
            "Subject",
            "Target\nactual",
            "Route",
            "Release",
            "Gates",
            "Plan / E-R-G",
        ),
        rows,
        (420, 1_160, 1_550, 900, 1_050, 1_350, 1_620, 1_310),
        body_size=7.5,
    )


def _add_plaintext_blocks(document: DocumentType, text: str, *, answer: bool) -> None:
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if not line:
            document.add_paragraph()
            continue
        heading = re.match(r"^#{1,6}\s+(.+)$", line)
        if heading:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(heading.group(1).strip())
            _set_run_font(run, size=11, bold=True, color="1F4D78")
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.keep_with_next = True
            continue
        bullet = re.match(r"^[-*]\s+(.+)$", line)
        numbered = re.match(r"^\d+[.)]\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(bullet.group(1))
        elif numbered:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(numbered.group(1))
        else:
            paragraph = document.add_paragraph(line)
        if answer:
            paragraph.paragraph_format.widow_control = True


def _add_evidence(document: DocumentType, case: LiveReviewCase) -> None:
    document.add_heading("Evidence and citation controls", level=2)
    if not case.evidence:
        document.add_paragraph("No safe evidence summary recorded.", style="Review Small")
        return
    _add_table(
        document,
        ("Evidence span", "Stable source", "Locator", "Role / support", "Rank", "Currentness"),
        (
            (
                evidence.evidence_span_id,
                evidence.stable_source_id,
                evidence.legal_locator,
                f"{evidence.legal_role} / {evidence.support_state}",
                "—" if evidence.retrieval_rank is None else str(evidence.retrieval_rank),
                evidence.currentness_state,
            )
            for evidence in case.evidence
        ),
        (1_450, 1_760, 1_520, 1_700, 560, 2_370),
        body_size=7.5,
    )


def _add_rubric(document: DocumentType, case: LiveReviewCase) -> None:
    document.add_heading("Rubric and activated assessment rules", level=2)
    document.add_paragraph(
        f"Assessment bundle SHA-256: {case.assessment_bundle_sha256 or 'not recorded'}",
        style="Review Small",
    )
    document.add_paragraph(
        "Applied rule IDs: " + (", ".join(case.assessment_rule_ids) or "none recorded"),
        style="Review Small",
    )
    if case.rubric:
        _add_table(
            document,
            ("Criterion", "Score", "Status", "Rule IDs", "Verification signal"),
            (
                (
                    result.criterion_id,
                    "—" if result.score is None else f"{result.score:g}",
                    result.status,
                    ", ".join(result.assessment_rule_ids) or "—",
                    result.verification_signal or "—",
                )
                for result in case.rubric
            ),
            (2_120, 820, 1_100, 2_520, 2_800),
            body_size=7.5,
        )
    else:
        document.add_paragraph("No automated rubric result recorded.", style="Review Small")
    warning = document.add_paragraph(style="Review Small")
    warning.add_run(
        "Automated academic scoring is review guidance until blind human calibration confirms "
        "that the 70+ classification is calibrated."
    )


def _add_repairs_and_gaps(document: DocumentType, case: LiveReviewCase) -> None:
    document.add_heading("Repairs, failures and knowledge gaps", level=2)
    if case.repairs:
        _add_table(
            document,
            ("Repair", "Section", "Reason", "Status", "Attempts"),
            (
                (
                    repair.repair_id,
                    repair.section_id or "whole answer",
                    repair.reason_code,
                    repair.status,
                    str(repair.attempt_count),
                )
                for repair in case.repairs
            ),
            (1_650, 1_580, 2_700, 1_350, 2_080),
            body_size=8,
        )
    else:
        document.add_paragraph("No repair event recorded.", style="Review Small")
    if case.gaps:
        _add_table(
            document,
            ("Gap", "Category", "Severity", "Status", "Expected / observed safe IDs"),
            (
                (
                    gap.gap_id,
                    gap.category,
                    gap.severity,
                    gap.status,
                    f"{', '.join(gap.safe_expected_ids) or '—'} / {', '.join(gap.safe_observed_ids) or '—'}",
                )
                for gap in case.gaps
            ),
            (1_520, 1_620, 1_050, 2_120, 3_050),
            body_size=8,
        )
    else:
        document.add_paragraph("No knowledge-gap record linked.", style="Review Small")
    if case.failure_codes:
        document.add_paragraph(
            "Failure-mode codes: " + ", ".join(case.failure_codes), style="Review Small"
        )


def _add_advisory_ai_review(document: DocumentType, case: LiveReviewCase) -> None:
    document.add_heading("Advisory AI evidence review", level=2)
    review = case.advisory_ai_review
    if review is None:
        document.add_paragraph(
            "No safe advisory AI review summary was recorded for this case.",
            style="Review Small",
        )
        return
    _add_table(
        document,
        ("Status", "Execution", "Independent", "Flags", "Owner review", "Review SHA / reason"),
        [
            (
                review.status,
                "same model adapter\nseparate pass",
                "NO",
                ", ".join(review.recommendation_codes) or "none",
                "REQUIRED" if review.owner_review_required else "not flagged",
                review.review_sha256 or review.unavailable_reason_code or "not run",
            )
        ],
        (900, 1_560, 900, 2_000, 1_200, 2_800),
        body_size=7.5,
    )
    document.add_paragraph(
        "AI findings are recommendations only. They cannot decide, adopt, admit sources or "
        "authorize a gate. A concern may raise a fail-closed owner-review hold; a positive "
        "recommendation never overrides deterministic checks or owner authority.",
        style="Review Small",
    )


def _add_human_review_fields(document: DocumentType, case: LiveReviewCase) -> None:
    document.add_heading("Owner substantive review", level=2)
    rows = [
        ("Overall legal quality", "[ ] Pass   [ ] Repair   [ ] Hold", ""),
        ("Issue spotting / completeness", "[ ] Pass   [ ] Repair   [ ] Hold", ""),
        ("Rule, authority, citations / OSCOLA", "[ ] Pass   [ ] Repair   [ ] Hold", ""),
        (
            "Application, critical analysis / uncertainty",
            "[ ] Pass   [ ] Repair   [ ] Hold",
            "",
        ),
        ("Calibrated band", "[ ] 70+   [ ] 60–69   [ ] 50–59   [ ] Below 50", ""),
        ("AI recommendation disposition", "[ ] Considered   [ ] Rejected   [ ] Unavailable", ""),
        ("Owner comments", "", ""),
        ("Owner decision", "[ ] Approve   [ ] Repair   [ ] Hold", ""),
        ("Owner typed name / decision date", "", ""),
        ("Encrypted review / sign-off references", "", ""),
    ]
    _add_table(
        document,
        ("Review field", "Decision", "Safe reference / date"),
        rows,
        (3_000, 3_820, 2_540),
        body_size=8.5,
    )


def _add_case_detail(document: DocumentType, item: LoadedReviewCase) -> None:
    case = item.safe
    heading = document.add_heading(f"Case {case.ordinal:02d} — {case.case_id}", level=1)
    heading.paragraph_format.page_break_before = True
    document.add_paragraph(
        f"{case.subject} · {case.task_type} · {case.jurisdiction} · as at {case.as_of_date}",
        style="Subtitle",
    )
    _add_table(
        document,
        ("Target / actual", "Route", "Release", "Privacy", "Evidence"),
        [
            (
                f"{case.word_target:,} / {case.word_count if case.word_count is not None else '—'}",
                f"{case.research_route} / {case.drafting_route}",
                case.release_state,
                "PASS" if case.privacy_passed else "FAIL",
                "PASS" if case.evidence_passed else "FAIL",
            )
        ],
        (1_700, 2_300, 1_900, 1_730, 1_730),
        body_size=8.5,
    )
    if item.question is None or item.released_answer is None:
        document.add_heading("Plaintext intentionally omitted", level=2)
        document.add_paragraph(
            "This case has no export-eligible released answer. Its encrypted question, held "
            "drafts and sensitive issue detail were not opened by the exporter."
        )
    else:
        document.add_heading("Exact submitted question", level=2)
        _add_plaintext_blocks(document, item.question, answer=False)
        document.add_heading("Released answer", level=2)
        _add_plaintext_blocks(document, item.released_answer, answer=True)
    _add_evidence(document, case)
    _add_rubric(document, case)
    _add_repairs_and_gaps(document, case)
    _add_advisory_ai_review(document, case)
    _add_human_review_fields(document, case)


def _add_programme_appendices(document: DocumentType, loaded: LoadedLiveReview) -> None:
    heading = document.add_heading("Cross-case clusters", level=1)
    heading.paragraph_format.page_break_before = True
    if loaded.review.clusters:
        _add_table(
            document,
            ("Cluster", "Category", "Cases", "Status"),
            (
                (cluster.cluster_id, cluster.category, ", ".join(cluster.case_ids), cluster.status)
                for cluster in loaded.review.clusters
            ),
            (2_000, 2_200, 3_800, 1_360),
            body_size=8.5,
        )
    else:
        document.add_paragraph("No cross-case cluster recorded.")

    document.add_heading("Corrections and owner decisions", level=1)
    if loaded.review.corrections:
        _add_table(
            document,
            ("Correction", "Layer", "Cases", "Status", "Regression"),
            (
                (
                    correction.correction_id,
                    correction.affected_layer,
                    ", ".join(correction.case_ids),
                    correction.status,
                    correction.regression_case_id or "—",
                )
                for correction in loaded.review.corrections
            ),
            (1_850, 1_450, 2_800, 1_360, 1_900),
            body_size=8.5,
        )
    else:
        document.add_paragraph("No correction record supplied.")
    if loaded.review.owner_decisions:
        _add_table(
            document,
            ("Decision", "Decision code", "Affected cases", "Status"),
            (
                (
                    decision.decision_id,
                    decision.decision_code,
                    ", ".join(decision.affected_case_ids) or "programme",
                    decision.status,
                )
                for decision in loaded.review.owner_decisions
            ),
            (1_850, 2_850, 2_800, 1_860),
            body_size=8.5,
        )
    else:
        document.add_paragraph("No owner decision record supplied.")

    document.add_heading("Regression register", level=1)
    if loaded.review.regressions:
        _add_table(
            document,
            ("Regression case", "Source issue", "Status", "Fixed version", "Verification run"),
            (
                (
                    regression.regression_case_id,
                    regression.source_issue_id,
                    regression.status,
                    regression.fixed_version or "—",
                    regression.verification_run_id or "—",
                )
                for regression in loaded.review.regressions
            ),
            (2_050, 1_850, 1_300, 1_800, 2_360),
            body_size=8.5,
        )
    else:
        document.add_paragraph("No regression record supplied.")


def _finalize_document_properties(document: DocumentType, *, title: str, subject: str) -> None:
    core = document.core_properties
    core.title = title
    core.subject = subject
    core.author = ""
    core.last_modified_by = ""
    core.keywords = ""
    core.comments = ""
    core.category = ""
    core.created = datetime(2000, 1, 1, tzinfo=UTC)
    core.modified = datetime(2000, 1, 1, tzinfo=UTC)


def _add_annex_cover(
    document: DocumentType,
    loaded: LoadedLiveReview,
    *,
    annex_id: str,
    case_ids: Sequence[str],
) -> None:
    paragraph = document.add_paragraph("LEGALBOT", style="Subtitle")
    paragraph.paragraph_format.space_before = Pt(54)
    title = document.add_paragraph(f"Live60 answer review\nAnnex {annex_id}", style="Title")
    title.paragraph_format.keep_with_next = True
    document.add_paragraph(
        "Ten selected single-pass outcomes · released answers only", style="Subtitle"
    )
    status = document.add_paragraph()
    run = status.add_run("CONTROLLED LOCAL DOCUMENT · EVALUATION ONLY · NOT TRAINING DATA")
    _set_run_font(run, size=9, bold=True, color="9C2F2F")
    status.paragraph_format.space_before = Pt(8)
    status.paragraph_format.space_after = Pt(18)
    _add_table(
        document,
        ("Identity", "Recorded value"),
        [*_identity_rows(loaded.identity), ("Annex cases", ", ".join(case_ids))],
        (2_180, 7_180),
        body_size=9,
    )
    note = document.add_paragraph(style="Review Small")
    note.add_run(
        "A selected case with no export-eligible released answer appears only as a safe "
        "diagnostic. The exporter does not open its encrypted question or held draft."
    )
    _add_page_break(document)


def build_live60_control_document(loaded: LoadedLiveReview) -> DocumentType:
    if loaded.identity.suite_id != "live-evaluation-60-v1":
        raise ReviewExportError("Live60 control document requires a Live60 run")
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, loaded.identity.run_id)
    _add_cover(document, loaded)
    _add_toc(document)
    _add_page_break(document)
    _add_executive_metrics(document, loaded)
    _add_case_matrix(document, loaded)
    _add_programme_appendices(document, loaded)
    _finalize_document_properties(
        document,
        title="LegalBot Live60 control owner review",
        subject="All-60 coverage, status, quality and defect control",
    )
    return document


def build_live60_annex_document(loaded: LoadedLiveReview, *, annex_id: str) -> DocumentType:
    case_ids = loaded.contract.annexes.get(annex_id)
    if case_ids is None or len(case_ids) != 10:
        raise ReviewExportError("annex is absent from the sealed Live60 run plan")
    by_id = {item.safe.case_id: item for item in loaded.cases}
    if any(case_id not in by_id for case_id in case_ids):
        raise ReviewExportError("annex case is absent from the reconciled review")
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, loaded.identity.run_id)
    _add_annex_cover(document, loaded, annex_id=annex_id, case_ids=case_ids)
    _add_toc(document)
    _add_page_break(document)
    for case_id in case_ids:
        _add_case_detail(document, by_id[case_id])
    _finalize_document_properties(
        document,
        title=f"LegalBot Live60 answer review Annex {annex_id}",
        subject="Ten selected single-pass answer outcomes",
    )
    return document


def build_live_review_document(loaded: LoadedLiveReview) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, loaded.identity.run_id)
    _add_cover(document, loaded)
    _add_toc(document)
    _add_page_break(document)
    _add_executive_metrics(document, loaded)
    _add_case_matrix(document, loaded)
    for case in loaded.cases:
        _add_case_detail(document, case)
    _add_programme_appendices(document, loaded)
    _finalize_document_properties(
        document,
        title="LegalBot live evaluation owner review",
        subject="Local evaluation-only quality and defect review",
    )
    return document


def _scrub_package_xml(name: str, payload: bytes) -> bytes | None:
    parser = etree.XMLParser(resolve_entities=False, no_network=True, remove_blank_text=False)
    root = etree.fromstring(payload, parser=parser)
    if name == "[Content_Types].xml":
        for node in root.xpath("//*[local-name()='Override'][@PartName='/docProps/custom.xml']"):
            node.getparent().remove(node)
    elif name == "_rels/.rels":
        for node in root.xpath(
            "//*[local-name()='Relationship'][contains(@Type, '/custom-properties')]"
        ):
            node.getparent().remove(node)
    elif name == "docProps/core.xml":
        for local_name in ("creator", "lastModifiedBy", "keywords", "description", "category"):
            for node in root.xpath(f"//*[local-name()='{local_name}']"):
                node.text = ""
        for local_name in ("created", "modified"):
            for node in root.xpath(f"//*[local-name()='{local_name}']"):
                node.text = "2000-01-01T00:00:00Z"
    elif name == "docProps/app.xml":
        for local_name in ("Company", "Manager"):
            for node in root.xpath(f"//*[local-name()='{local_name}']"):
                node.text = ""
    if name.startswith("word/") and name.endswith(".xml"):
        for element in root.iter():
            for attribute in tuple(element.attrib):
                if etree.QName(attribute).localname.startswith("rsid"):
                    del element.attrib[attribute]
        for node in root.xpath("//*[starts-with(local-name(), 'rsid')]"):
            node.getparent().remove(node)
    return cast(
        bytes,
        etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True),
    )


def scrub_docx_metadata(path: Path) -> None:
    """Remove author/custom properties, revision IDs and variable ZIP timestamps."""

    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{path.stem}.scrub-", suffix=".docx", dir=path.parent
    )
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        with (
            zipfile.ZipFile(path, "r") as source,
            zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as destination,
        ):
            for item in source.infolist():
                if item.filename == "docProps/custom.xml":
                    continue
                payload = source.read(item.filename)
                if item.filename.endswith(".xml") or item.filename.endswith(".rels"):
                    scrubbed = _scrub_package_xml(item.filename, payload)
                    if scrubbed is None:
                        continue
                    payload = scrubbed
                clean_info = zipfile.ZipInfo(item.filename, date_time=(1980, 1, 1, 0, 0, 0))
                clean_info.compress_type = zipfile.ZIP_DEFLATED
                clean_info.external_attr = item.external_attr
                clean_info.create_system = item.create_system
                destination.writestr(clean_info, payload)
        os.replace(temporary, path)
        path.chmod(0o600)
    finally:
        temporary.unlink(missing_ok=True)


class ReviewOutputDocument(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    document_id: Literal["control", "annex-a", "annex-b", "annex-c"]
    kind: Literal["all60_control", "answer_annex"]
    relative_path: str = Field(pattern=r"^[A-Za-z0-9][A-Za-z0-9._/-]{0,255}$")
    sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    case_ids: tuple[str, ...]
    released_plaintext_case_ids: tuple[str, ...]
    safe_diagnostic_case_ids: tuple[str, ...]
    render_gate_status: Literal["pending"] = "pending"

    @model_validator(mode="after")
    def document_scope_is_safe(self) -> Self:
        if self.relative_path.startswith("/") or ".." in self.relative_path:
            raise ValueError("review document path is not safely relative")
        if any(case_id not in LIVE60_EXPECTED_CASE_IDS for case_id in self.case_ids):
            raise ValueError("review document contains an unknown Live60 case")
        if not set(self.released_plaintext_case_ids).issubset(self.case_ids):
            raise ValueError("plaintext case is outside the document scope")
        if not set(self.safe_diagnostic_case_ids).issubset(self.case_ids):
            raise ValueError("diagnostic case is outside the document scope")
        return self


class ReviewOutputManifest(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True, populate_by_name=True)

    schema_name: Literal["legalbot.live-review-output-manifest.v2"] = Field(
        default="legalbot.live-review-output-manifest.v2", alias="schema"
    )
    run_id: str = Field(pattern=r"^[a-z0-9][a-z0-9._:-]{2,127}$")
    suite_id: Literal["live-evaluation-60-v1"]
    suite_registry_canonical_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    run_plan_id: Literal["live60-single-pass-30-v1"]
    run_plan_file_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    run_plan_seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    run_manifest_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    review_export_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    purpose: Literal["evaluation_only"] = "evaluation_only"
    eligible_for_training: Literal[False] = False
    training_export_allowed: Literal[False] = False
    created_at: datetime
    document_count: Literal[4] = 4
    status: Literal["docx_created_render_pending"] = "docx_created_render_pending"
    documents: tuple[ReviewOutputDocument, ...]
    seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")

    @model_validator(mode="after")
    def manifest_is_exact_and_sealed(self) -> Self:
        if tuple(document.document_id for document in self.documents) != (
            "control",
            "annex-a",
            "annex-b",
            "annex-c",
        ):
            raise ValueError("review output must contain one control and three annexes")
        material = self.model_dump(mode="json", by_alias=True)
        material.pop("seal_sha256", None)
        if self.seal_sha256 != _sha256_bytes(
            (
                json.dumps(
                    material,
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                )
                + "\n"
            ).encode("utf-8")
        ):
            raise ValueError("review output manifest seal is invalid")
        return self


def _sealed_model_payload(value: Mapping[str, Any]) -> str:
    material = dict(value)
    material.pop("seal_sha256", None)
    return _sha256_bytes(
        (
            json.dumps(
                material,
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
            )
            + "\n"
        ).encode("utf-8")
    )


def _save_scrubbed_document(document: DocumentType, destination: Path) -> None:
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{destination.stem}.build-", suffix=".docx", dir=destination.parent
    )
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        document.save(str(temporary))
        scrub_docx_metadata(temporary)
        os.replace(temporary, destination)
        destination.chmod(0o600)
    finally:
        temporary.unlink(missing_ok=True)


def _write_json_atomic(path: Path, value: Mapping[str, Any]) -> None:
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{path.stem}.write-", suffix=".json", dir=path.parent
    )
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        temporary.write_text(
            json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
        )
        temporary.chmod(0o600)
        os.replace(temporary, path)
        path.chmod(0o600)
    finally:
        temporary.unlink(missing_ok=True)


def export_live60_review_bundle(
    *,
    run_dir: Path,
    output_dir: Path,
    cipher: LocalCipher,
    overwrite: bool = False,
) -> Path:
    """Create one all-60 control and three exact ten-outcome annexes.

    The returned output manifest remains explicitly render-pending.  Visual
    inspection is recorded in a separate immutable render-gate artifact so a
    successful DOCX save can never be mistaken for layout approval.
    """

    loaded = load_live_review(run_dir=run_dir, cipher=cipher, require_complete=True)
    if loaded.identity.suite_id != "live-evaluation-60-v1":
        raise ReviewExportError("bundle export requires the sealed Live60 contract")
    destination = output_dir.resolve()
    destination.mkdir(parents=True, exist_ok=True)
    filenames = {
        "control": "LegalBot-Live60-Control.docx",
        "annex-a": "LegalBot-Live60-Annex-A.docx",
        "annex-b": "LegalBot-Live60-Annex-B.docx",
        "annex-c": "LegalBot-Live60-Annex-C.docx",
    }
    manifest_path = destination / "output-manifest.json"
    render_gate_path = destination / "render-gate.json"
    existing = [destination / filename for filename in filenames.values()]
    existing.extend((manifest_path, render_gate_path))
    if any(path.exists() for path in existing) and not overwrite:
        raise FileExistsError("review output bundle already exists")
    if render_gate_path.exists():
        # A render decision belongs to the previous document digests and must
        # never survive a rebuilt bundle.
        raise FileExistsError("existing render gate must be archived before rebuilding")

    documents = {
        "control": build_live60_control_document(loaded),
        "annex-a": build_live60_annex_document(loaded, annex_id="A"),
        "annex-b": build_live60_annex_document(loaded, annex_id="B"),
        "annex-c": build_live60_annex_document(loaded, annex_id="C"),
    }
    for document_id, document in documents.items():
        _save_scrubbed_document(document, destination / filenames[document_id])

    by_id = {item.safe.case_id: item for item in loaded.cases}
    output_documents: list[dict[str, Any]] = []
    scopes: dict[str, tuple[str, ...]] = {
        "control": loaded.contract.expected_case_ids,
        "annex-a": loaded.contract.annexes["A"],
        "annex-b": loaded.contract.annexes["B"],
        "annex-c": loaded.contract.annexes["C"],
    }
    for document_id in ("control", "annex-a", "annex-b", "annex-c"):
        case_ids = scopes[document_id]
        released_ids = (
            ()
            if document_id == "control"
            else tuple(
                case_id for case_id in case_ids if by_id[case_id].released_answer is not None
            )
        )
        diagnostics = (
            tuple(case_id for case_id in case_ids if by_id[case_id].released_answer is None)
            if document_id != "control"
            else case_ids
        )
        document_path = destination / filenames[document_id]
        output_documents.append(
            {
                "document_id": document_id,
                "kind": ("all60_control" if document_id == "control" else "answer_annex"),
                "relative_path": filenames[document_id],
                "sha256": _sha256_bytes(document_path.read_bytes()),
                "case_ids": list(case_ids),
                "released_plaintext_case_ids": list(released_ids),
                "safe_diagnostic_case_ids": list(diagnostics),
                "render_gate_status": "pending",
            }
        )
    run_manifest_path = run_dir.resolve() / "manifest.json"
    review_export_path = run_dir.resolve() / "review-export.json"
    value: dict[str, Any] = {
        "schema": "legalbot.live-review-output-manifest.v2",
        "run_id": loaded.identity.run_id,
        "suite_id": "live-evaluation-60-v1",
        "suite_registry_canonical_sha256": loaded.identity.suite_canonical_sha256,
        "run_plan_id": loaded.contract.run_plan_id,
        "run_plan_file_sha256": loaded.contract.run_plan_file_sha256,
        "run_plan_seal_sha256": loaded.contract.run_plan_seal_sha256,
        "run_manifest_sha256": _sha256_bytes(run_manifest_path.read_bytes()),
        "review_export_sha256": _sha256_bytes(review_export_path.read_bytes()),
        "purpose": "evaluation_only",
        "eligible_for_training": False,
        "training_export_allowed": False,
        "created_at": datetime.now(UTC).isoformat().replace("+00:00", "Z"),
        "document_count": 4,
        "status": "docx_created_render_pending",
        "documents": output_documents,
    }
    value["seal_sha256"] = _sealed_model_payload(value)
    validated = ReviewOutputManifest.model_validate(value)
    _write_json_atomic(manifest_path, validated.model_dump(mode="json", by_alias=True))
    return manifest_path


class ReviewRenderPage(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    relative_path: str = Field(pattern=r"^[A-Za-z0-9][A-Za-z0-9._/-]{0,255}$")
    sha256: str = Field(pattern=r"^[0-9a-f]{64}$")


class ReviewRenderDocument(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True)

    document_id: Literal["control", "annex-a", "annex-b", "annex-c"]
    document_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    page_count: int = Field(ge=1, le=10_000)
    pages: tuple[ReviewRenderPage, ...]


class ReviewRenderGate(BaseModel):
    model_config = ConfigDict(extra="forbid", frozen=True, populate_by_name=True)

    schema_name: Literal["legalbot.live-review-render-gate.v1"] = Field(
        default="legalbot.live-review-render-gate.v1", alias="schema"
    )
    output_manifest_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")
    purpose: Literal["evaluation_only"] = "evaluation_only"
    eligible_for_training: Literal[False] = False
    training_export_allowed: Literal[False] = False
    visual_inspection_passed: Literal[True]
    inspector_ref: str = Field(pattern=r"^reviewer:[0-9a-f]{64}$")
    inspected_at: datetime
    documents: tuple[ReviewRenderDocument, ...]
    seal_sha256: str = Field(pattern=r"^[0-9a-f]{64}$")

    @model_validator(mode="after")
    def gate_is_complete_and_sealed(self) -> Self:
        if tuple(item.document_id for item in self.documents) != (
            "control",
            "annex-a",
            "annex-b",
            "annex-c",
        ):
            raise ValueError("render gate must cover all four review documents")
        if any(item.page_count != len(item.pages) for item in self.documents):
            raise ValueError("render page count is inconsistent")
        if self.seal_sha256 != _sealed_model_payload(self.model_dump(mode="json", by_alias=True)):
            raise ValueError("render gate seal is invalid")
        return self


def record_live60_render_gate(
    *,
    output_dir: Path,
    rendered_pages: Mapping[str, Sequence[Path]],
    inspector_ref: str,
) -> Path:
    """Record a human-inspected PNG render set without automating approval."""

    destination = output_dir.resolve()
    manifest_path = destination / "output-manifest.json"
    gate_path = destination / "render-gate.json"
    if gate_path.exists():
        raise FileExistsError("render gate is immutable and already exists")
    try:
        manifest = ReviewOutputManifest.model_validate_json(manifest_path.read_bytes())
    except (OSError, TypeError, ValueError) as exc:
        raise ReviewExportError("review output manifest is invalid") from exc
    records: list[dict[str, Any]] = []
    by_id = {item.document_id: item for item in manifest.documents}
    if set(rendered_pages) != set(by_id):
        raise ReviewExportError("render set does not cover all review documents")
    for document_id in ("control", "annex-a", "annex-b", "annex-c"):
        document_record = by_id[document_id]
        document_path = destination / document_record.relative_path
        if _sha256_bytes(document_path.read_bytes()) != document_record.sha256:
            raise ReviewExportError("review document changed before visual inspection")
        pages: list[dict[str, str]] = []
        for page_number, page_path in enumerate(rendered_pages[document_id], start=1):
            resolved = page_path.resolve()
            try:
                relative = resolved.relative_to(destination)
            except ValueError as exc:
                raise ReviewExportError("rendered page escaped the output directory") from exc
            if resolved.suffix.casefold() != ".png" or not resolved.is_file():
                raise ReviewExportError("rendered page is missing or not PNG")
            if resolved.name != f"page-{page_number}.png":
                raise ReviewExportError("rendered pages are not a contiguous ordered set")
            page_bytes = resolved.read_bytes()
            if not page_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
                raise ReviewExportError("rendered page does not have a PNG signature")
            pages.append(
                {
                    "relative_path": relative.as_posix(),
                    "sha256": _sha256_bytes(page_bytes),
                }
            )
        if not pages:
            raise ReviewExportError("each review document needs at least one rendered page")
        records.append(
            {
                "document_id": document_id,
                "document_sha256": document_record.sha256,
                "page_count": len(pages),
                "pages": pages,
            }
        )
    value: dict[str, Any] = {
        "schema": "legalbot.live-review-render-gate.v1",
        "output_manifest_sha256": _sha256_bytes(manifest_path.read_bytes()),
        "purpose": "evaluation_only",
        "eligible_for_training": False,
        "training_export_allowed": False,
        "visual_inspection_passed": True,
        "inspector_ref": inspector_ref,
        "inspected_at": datetime.now(UTC).isoformat().replace("+00:00", "Z"),
        "documents": records,
    }
    value["seal_sha256"] = _sealed_model_payload(value)
    validated = ReviewRenderGate.model_validate(value)
    _write_json_atomic(gate_path, validated.model_dump(mode="json", by_alias=True))
    return gate_path


def verify_live60_render_gate(output_dir: Path) -> ReviewRenderGate:
    """Revalidate document and rendered-page digests before delivery."""

    destination = output_dir.resolve()
    manifest_path = destination / "output-manifest.json"
    gate_path = destination / "render-gate.json"
    try:
        manifest = ReviewOutputManifest.model_validate_json(manifest_path.read_bytes())
        gate = ReviewRenderGate.model_validate_json(gate_path.read_bytes())
    except (OSError, TypeError, ValueError) as exc:
        raise ReviewExportError("review render gate is absent or invalid") from exc
    if gate.output_manifest_sha256 != _sha256_bytes(manifest_path.read_bytes()):
        raise ReviewExportError("render gate is bound to another output manifest")
    documents = {item.document_id: item for item in manifest.documents}
    for rendered in gate.documents:
        document = documents[rendered.document_id]
        document_path = destination / document.relative_path
        if (
            rendered.document_sha256 != document.sha256
            or _sha256_bytes(document_path.read_bytes()) != document.sha256
        ):
            raise ReviewExportError("render-gated document digest changed")
        for page in rendered.pages:
            page_path = (destination / page.relative_path).resolve()
            if not page_path.is_relative_to(destination):
                raise ReviewExportError("render page escaped the output directory")
            if _sha256_bytes(page_path.read_bytes()) != page.sha256:
                raise ReviewExportError("render-gated page digest changed")
    return gate


def export_live_review_docx(
    *,
    run_dir: Path,
    output_path: Path,
    cipher: LocalCipher,
    require_complete: bool = True,
    overwrite: bool = False,
) -> Path:
    """Validate, decrypt release-eligible plaintext, render, scrub and atomically publish."""

    if output_path.suffix.casefold() != ".docx":
        raise ReviewExportError("review output must use the .docx extension")
    output = output_path.resolve()
    if output.exists() and not overwrite:
        raise FileExistsError(f"review output already exists: {output}")
    output.parent.mkdir(parents=True, exist_ok=True)
    loaded = load_live_review(run_dir=run_dir, cipher=cipher, require_complete=require_complete)
    document = build_live_review_document(loaded)
    descriptor, temporary_name = tempfile.mkstemp(
        prefix=f".{output.stem}.build-", suffix=".docx", dir=output.parent
    )
    os.close(descriptor)
    temporary = Path(temporary_name)
    try:
        document.save(str(temporary))
        scrub_docx_metadata(temporary)
        os.replace(temporary, output)
        output.chmod(0o600)
    finally:
        temporary.unlink(missing_ok=True)
    return output
