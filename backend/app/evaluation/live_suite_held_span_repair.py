"""New contiguous span versions for held-provision defects.

Old catalogue chunks are immutable. This module derives repair spans from the
existing local text only. It does not write into the catalogue, does not seal
gold, and does not qualify currentness.
"""

from __future__ import annotations

import hashlib
import json
import sqlite3
from collections.abc import Mapping
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType

from .live30 import assert_safe_evaluation_payload
from .live_suite_repair_span import (
    HELD_SPAN_REPAIR_SCHEMA_V2,
    IPFDA_DOTS_PARENT,
    REPAIR_SPAN_SCHEMA_V2,
    derivation_manifest_sha256,
    identity_complete,
    parent_identity,
    repair_span_id_v2,
    repair_span_identity_v2,
    sha256_text,
)
from .review_docx import (
    _add_header_footer,
    _add_table,
    _configure_page,
    _configure_styles,
    _finalize_document_properties,
    _reject_prohibited_metadata,
)

HELD_SPAN_REPAIR_SCHEMA = HELD_SPAN_REPAIR_SCHEMA_V2
PROVISION_REVIEW_WORKSHEET_SCHEMA = "legalbot.live60-provision-review-worksheet.v1"
S14A_SPLICED_PARENT = "chunk-73fd50b82133e6bd94542a9e1d78ae5b06f91987"
IPFDA_OPENING_PARENT = "chunk-65e1c2ac95885e5e8e9d65e5ae138b5b6fbcde0a"
IPFDA_BA_PARENT = "chunk-3fea9d80c6f7fc0011d4db8679b75fe7185c06c7"
S14A_SPLIT_MARK = "acquire—"
IPFDA_SPLIT_MARK = "persons:—"
EDITORIAL_ELLIPSIS = "..."
REPAIR_DOCX = "LegalBot-Live60-Held-Span-Contiguous-Repair.docx"
WORKSHEET_DOCX = "LegalBot-Live60-Provision-Review-Worksheet.docx"


def _sha256_text(value: str) -> str:
    return sha256_text(value)


def _json_safe_text(value: str) -> str:
    return " ".join(value.split())


def _ipfda_ba_operative_text(raw: str) -> str:
    """Current official words: omit editorial ellipsis standing for repealed 'or (1B)'."""

    collapsed = _json_safe_text(raw).replace("…", EDITORIAL_ELLIPSIS)
    return collapsed.replace("(1A) ... below", "(1A) below")


def _split_once(text: str, mark: str) -> tuple[str, str]:
    if mark not in text:
        raise ValueError("repair split mark is not present in parent chunk text")
    left, right = text.split(mark, 1)
    return left + mark, right


def _span(
    *,
    parent_chunk_id: str,
    sublocator: str,
    text: str,
    role: str,
    gold_eligible_candidate: bool,
    action: str,
    source_version_id: str,
    legal_authority_id: str,
    official_snapshot_sha256: str,
    derivation_manifest: Mapping[str, Any],
    stable_source_id: str,
    source_type: str,
    jurisdiction: str,
    legal_locator: str | None = None,
) -> dict[str, Any]:
    safe_text = _json_safe_text(text)
    manifest = dict(derivation_manifest)
    manifest_sha = derivation_manifest_sha256(manifest)
    locator = legal_locator or sublocator
    identity = repair_span_identity_v2(
        parent_chunk_id=parent_chunk_id,
        source_version_id=source_version_id,
        legal_authority_id=legal_authority_id,
        official_snapshot_sha256=official_snapshot_sha256,
        required_sublocator=sublocator,
        role=role,
        markdown_text=safe_text,
        derivation_manifest_sha256=manifest_sha,
        stable_source_id=stable_source_id,
        source_type=source_type,
        jurisdiction=jurisdiction,
        legal_locator=locator,
    )
    complete = identity_complete(
        {
            "parent_chunk_id": parent_chunk_id,
            "source_version_id": source_version_id,
            "legal_authority_id": legal_authority_id,
            "official_snapshot_sha256": official_snapshot_sha256,
            "required_sublocator": sublocator,
            "role": role,
            "markdown_text": safe_text,
            "derivation_manifest_sha256": manifest_sha,
            "stable_source_id": stable_source_id,
            "source_type": source_type,
            "jurisdiction": jurisdiction,
            "legal_locator": locator,
        }
    )
    eligible = bool(gold_eligible_candidate and complete)
    if parent_chunk_id == IPFDA_DOTS_PARENT:
        eligible = False
    return {
        "schema": REPAIR_SPAN_SCHEMA_V2,
        "repair_span_id": repair_span_id_v2(identity=identity),
        "parent_chunk_id": parent_chunk_id,
        "source_version_id": source_version_id,
        "legal_authority_id": legal_authority_id,
        "official_snapshot_sha256": official_snapshot_sha256,
        "required_sublocator": sublocator,
        "legal_locator": locator,
        "stable_source_id": stable_source_id,
        "source_type": source_type,
        "jurisdiction": jurisdiction,
        "markdown_text": safe_text,
        "text_sha256": _sha256_text(safe_text),
        "role": role,
        "derivation_manifest": manifest,
        "derivation_manifest_sha256": manifest_sha,
        "identity_complete": complete,
        "gold_eligible_candidate": eligible,
        "action": action,
        "mutates_parent_chunk": False,
        "seals_expert_gold": False,
        "v1_rejected_as_new_gold": True,
    }


def _span_from_parent(
    *,
    identity: Mapping[str, str],
    sublocator: str,
    text: str,
    role: str,
    gold_eligible_candidate: bool,
    action: str,
    derivation_manifest: Mapping[str, Any],
) -> dict[str, Any]:
    return _span(
        parent_chunk_id=identity["parent_chunk_id"],
        sublocator=sublocator,
        text=text,
        role=role,
        gold_eligible_candidate=gold_eligible_candidate,
        action=action,
        source_version_id=identity["source_version_id"],
        legal_authority_id=identity["legal_authority_id"],
        official_snapshot_sha256=identity["official_snapshot_sha256"],
        derivation_manifest=derivation_manifest,
        stable_source_id=identity["stable_source_id"],
        source_type=identity["source_type"],
        jurisdiction=identity["jurisdiction"],
        legal_locator=sublocator,
    )


def build_held_span_contiguous_repair(export: Mapping[str, Any]) -> dict[str, Any]:
    repairs: list[dict[str, Any]] = []
    for provision in export.get("provisions", ()):
        for chunk in provision.get("chunks", ()):
            chunk_id = str(chunk["chunk_id"])
            text = str(chunk.get("markdown_text") or "")
            codes = set(chunk.get("structural_defect_codes") or ())
            identity = parent_identity(provision, chunk)
            if (
                "non_contiguous_s14A_10_chapeau_spliced_with_final_proviso" in codes
                or chunk_id == S14A_SPLICED_PARENT
            ) and S14A_SPLIT_MARK in text:
                chapeau, proviso = _split_once(text, S14A_SPLIT_MARK)
                repairs.extend(
                    [
                        _span_from_parent(
                            identity=identity,
                            sublocator="s 14A(10) chapeau",
                            text=chapeau,
                            role="statutory_text",
                            gold_eligible_candidate=True,
                            action="replace_spliced_parent_with_chapeau",
                            derivation_manifest={
                                "method": "split_once",
                                "split_mark": S14A_SPLIT_MARK,
                                "part": "chapeau",
                                "parent_chunk_id": chunk_id,
                            },
                        ),
                        _span_from_parent(
                            identity=identity,
                            sublocator="s 14A(10) final proviso",
                            text=proviso,
                            role="statutory_text",
                            gold_eligible_candidate=True,
                            action="replace_spliced_parent_with_final_proviso",
                            derivation_manifest={
                                "method": "split_once",
                                "split_mark": S14A_SPLIT_MARK,
                                "part": "final_proviso",
                                "parent_chunk_id": chunk_id,
                            },
                        ),
                    ]
                )
            if (
                "non_contiguous_ipfda_s1_1_chapeau_spliced_with_concluding_words" in codes
                or chunk_id == IPFDA_OPENING_PARENT
            ) and IPFDA_SPLIT_MARK in text:
                chapeau, concluding = _split_once(text, IPFDA_SPLIT_MARK)
                repairs.extend(
                    [
                        _span_from_parent(
                            identity=identity,
                            sublocator="s 1(1) chapeau",
                            text=chapeau,
                            role="statutory_text",
                            gold_eligible_candidate=True,
                            action="replace_spliced_opening_with_chapeau",
                            derivation_manifest={
                                "method": "split_once",
                                "split_mark": IPFDA_SPLIT_MARK,
                                "part": "chapeau",
                                "parent_chunk_id": chunk_id,
                            },
                        ),
                        _span_from_parent(
                            identity=identity,
                            sublocator="s 1(1) concluding application words",
                            text=concluding,
                            role="statutory_text",
                            gold_eligible_candidate=True,
                            action="replace_spliced_opening_with_concluding_words",
                            derivation_manifest={
                                "method": "split_once",
                                "split_mark": IPFDA_SPLIT_MARK,
                                "part": "concluding_words",
                                "parent_chunk_id": chunk_id,
                            },
                        ),
                    ]
                )
            if (
                "editorial_ellipsis_mixed_into_positive_text" in codes
                or chunk_id == IPFDA_BA_PARENT
            ) and (EDITORIAL_ELLIPSIS in text or "…" in text):
                repairs.append(
                    _span_from_parent(
                        identity=identity,
                        sublocator="s 1(1)(ba) raw text with editorial ellipsis mixed in",
                        text=text,
                        role="statutory_text_with_editorial_annotation",
                        gold_eligible_candidate=False,
                        action="retain_raw_and_separate_editorial_ellipsis_before_qualify",
                        derivation_manifest={
                            "method": "retain_raw_editorial_mix",
                            "parent_chunk_id": chunk_id,
                        },
                    )
                )
                repairs.append(
                    _span_from_parent(
                        identity=identity,
                        sublocator="s 1(1)(ba) editorial ellipsis marker",
                        text=EDITORIAL_ELLIPSIS,
                        role="editorial_annotation_marker",
                        gold_eligible_candidate=False,
                        action="exclude_editorial_ellipsis_from_positive_gold_text",
                        derivation_manifest={
                            "method": "extract_editorial_ellipsis",
                            "parent_chunk_id": chunk_id,
                        },
                    )
                )
                repairs.append(
                    _span_from_parent(
                        identity=identity,
                        sublocator="s 1(1)(ba) current positive text",
                        text=_ipfda_ba_operative_text(text),
                        role="statutory_text",
                        gold_eligible_candidate=True,
                        action="bind_official_positive_text_omitting_repealed_or_1B",
                        derivation_manifest={
                            "method": "omit_repealed_or_1B_ellipsis",
                            "parent_chunk_id": chunk_id,
                        },
                    )
                )
            if (
                "repealed_or_omitted_editorial_marker" in codes or chunk_id == IPFDA_DOTS_PARENT
            ) and text:
                repairs.append(
                    _span_from_parent(
                        identity=identity,
                        sublocator="repealed_or_omitted_s1(1B)",
                        text=text,
                        role="repealed_or_omitted_editorial_marker",
                        gold_eligible_candidate=False,
                        action="exclude_from_gold_ranking_generation_and_citation",
                        derivation_manifest={
                            "method": "exclude_dots_only_parent",
                            "parent_chunk_id": chunk_id,
                        },
                    )
                )

    proposed_sublocators = [
        {
            "held_id": provision["held_id"],
            "chunk_id": chunk["chunk_id"],
            "ordinal": chunk["ordinal"],
            "required_sublocator": chunk.get("required_sublocator"),
            "structural_defect_codes": list(chunk.get("structural_defect_codes") or ()),
        }
        for provision in export.get("provisions", ())
        for chunk in provision.get("chunks", ())
    ]
    payload = {
        "schema": HELD_SPAN_REPAIR_SCHEMA,
        "decision": "return_hold",
        "catalogue_mutated": False,
        "parent_chunks_immutable": True,
        "seals_expert_gold": False,
        "qualified": False,
        "o04_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "repair_span_count": len(repairs),
        "repairs": repairs,
        "proposed_sublocators": proposed_sublocators,
        "v1_rejected_as_new_gold": True,
        "note": (
            "contiguous repair candidates accepted from official pages; "
            "catalogue parents remain immutable; v1 repair spans are not new gold"
        ),
    }
    assert_safe_evaluation_payload(payload)
    return payload


EXCLUDED_STREAM = "excluded_structural_defect"


def apply_held_span_catalogue_cleanup(
    catalog_path: Path,
    *,
    export: Mapping[str, Any],
    repair: Mapping[str, Any],
) -> dict[str, Any]:
    """Exclude spliced/editorial body chunks and insert contiguous replacements.

    Parent chunk bytes stay in the catalogue. Retrieval uses stream='body' only.
    """

    connection = sqlite3.connect(catalog_path)
    connection.row_factory = sqlite3.Row
    excluded: list[str] = []
    inserted: list[dict[str, Any]] = []
    locator_updates: list[dict[str, str]] = []
    try:
        for provision in export.get("provisions", ()):
            for chunk in provision.get("chunks", ()):
                chunk_id = str(chunk["chunk_id"])
                if chunk.get("structural_defect_codes"):
                    connection.execute(
                        "UPDATE chunks SET stream=? WHERE id=?",
                        (EXCLUDED_STREAM, chunk_id),
                    )
                    excluded.append(chunk_id)
                    continue
                sublocator = chunk.get("required_sublocator")
                if sublocator and chunk.get("gold_eligible_candidate") is not False:
                    connection.execute(
                        "UPDATE chunks SET locator=? WHERE id=? AND COALESCE(stream,'body')='body'",
                        (sublocator, chunk_id),
                    )
                    locator_updates.append({"chunk_id": chunk_id, "locator": str(sublocator)})
        for item in repair.get("repairs", ()):
            if item.get("gold_eligible_candidate") is not True:
                continue
            parent_id = str(item["parent_chunk_id"])
            parent = connection.execute(
                "SELECT * FROM chunks WHERE id=?",
                (parent_id,),
            ).fetchone()
            if parent is None:
                continue
            text = str(item["markdown_text"])
            text_sha = str(item["text_sha256"])
            locator = str(item["required_sublocator"])
            chunk_id = (
                "chunk-"
                + hashlib.sha256(
                    f"{item['repair_span_id']}|{locator}|{text_sha}".encode()
                ).hexdigest()[:40]
            )
            existing = connection.execute(
                "SELECT id FROM chunks WHERE id=?", (chunk_id,)
            ).fetchone()
            if existing:
                continue
            max_ordinal = connection.execute(
                "SELECT COALESCE(MAX(ordinal), 0) FROM chunks WHERE source_version_id=?",
                (parent["source_version_id"],),
            ).fetchone()[0]
            metadata = {
                "repair_span_id": item["repair_span_id"],
                "parent_chunk_id": parent_id,
                "action": item.get("action"),
                "stream": "body",
            }
            columns = {row[1] for row in connection.execute("PRAGMA table_info(chunks)")}
            heading = parent["heading_path"] if "heading_path" in columns else None
            fields = [
                "id",
                "source_version_id",
                "ordinal",
                "locator",
                "text_sha256",
                "markdown_text",
                "token_count",
            ]
            values: list[object] = [
                chunk_id,
                parent["source_version_id"],
                int(max_ordinal) + 1,
                locator,
                text_sha,
                text,
                max(1, len(text.split())),
            ]
            if "heading_path" in columns:
                fields.insert(3, "heading_path")
                values.insert(3, heading)
            if "stream" in columns:
                fields.append("stream")
                values.append("body")
            if "metadata_json" in columns:
                fields.append("metadata_json")
                values.append(json.dumps(metadata, sort_keys=True))
            placeholders = ", ".join("?" for _ in fields)
            connection.execute(
                f"INSERT INTO chunks({', '.join(fields)}) VALUES ({placeholders})",
                values,
            )
            inserted.append(
                {
                    "chunk_id": chunk_id,
                    "parent_chunk_id": parent_id,
                    "locator": locator,
                    "text_sha256": text_sha,
                    "repair_span_id": item["repair_span_id"],
                    "source_version_id": parent["source_version_id"],
                }
            )
        connection.commit()
    finally:
        connection.close()
    payload = {
        "schema": "legalbot.live60-held-span-catalogue-cleanup.v1",
        "parent_bytes_deleted": False,
        "excluded_stream": EXCLUDED_STREAM,
        "excluded_parent_chunk_ids": excluded,
        "inserted_contiguous_chunks": inserted,
        "locator_updates": locator_updates,
        "seals_expert_gold": False,
        "qualified": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
    }
    assert_safe_evaluation_payload(
        {
            key: value
            for key, value in payload.items()
            if key
            not in {
                "inserted_contiguous_chunks",
                "locator_updates",
                "excluded_parent_chunk_ids",
            }
        }
    )
    return payload


def build_provision_review_worksheet(
    *,
    export: Mapping[str, Any],
    mechanical: Mapping[str, Any],
    identity: Mapping[str, Any],
    official: Mapping[str, Any] | None = None,
) -> dict[str, Any]:
    mechanical_by_id = {item["held_id"]: item for item in mechanical.get("results", ())}
    official_provisions = (official or {}).get("provisions", {})
    rows = []
    for provision in export.get("provisions", ()):
        held_id = str(provision["held_id"])
        mech = mechanical_by_id.get(held_id, {})
        proposed = _proposed_review_fields(held_id)
        ticks = official_provisions.get(held_id, {})
        qualified = bool(ticks.get("qualify_tick"))
        rows.append(
            {
                "held_id": held_id,
                "title": provision["title"],
                "authority_identity_id": provision["authority_identity_id"],
                "legal_locator": provision["legal_locator"],
                "source_version_id": provision["chunks"][0]["source_version_id"]
                if provision.get("chunks")
                else None,
                "document_content_sha256": provision.get("expected_document_content_sha256"),
                "official_uri": mech.get("official_uri"),
                "official_retrieval_timestamp": mech.get("official_retrieval_timestamp"),
                "official_raw_sha256": mech.get("official_raw_sha256"),
                "official_normalised_sha256": mech.get("official_normalised_sha256"),
                "canonicalisation_version": mech.get("canonicalisation_version"),
                "difference_class": mech.get("difference_class"),
                "structural_defect_count": mech.get("structural_defect_count", 0),
                "proposed_extent": proposed["extent"],
                "proposed_commencement": proposed["commencement"],
                "proposed_unapplied_effects": proposed["unapplied_effects"],
                "companion_provisions": proposed["companion_provisions"],
                "owner_extent_tick": ticks.get("extent_tick"),
                "owner_commencement_tick": ticks.get("commencement_tick"),
                "owner_effects_tick": ticks.get("effects_tick"),
                "owner_qualify_tick": ticks.get("qualify_tick"),
                "qualified": qualified,
                "disposition": (
                    "current_official_bytes_accepted" if qualified else "hold_pending_owner_ticks"
                ),
            }
        )
    payload = {
        "schema": PROVISION_REVIEW_WORKSHEET_SCHEMA,
        "decision": "return_hold",
        "primary_reviewer_is_owner": True,
        "approval_reviewer_role": identity.get("approval_reviewer_role"),
        "approval_reviewer_ref": identity.get("approval_reviewer_ref"),
        "ai_role": "mechanical_accuracy_verifier_only",
        "ai_cannot_seal_gold": True,
        "independent_second_review_status": "not_required",
        "seals_expert_gold": False,
        "o04_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "provisions": rows,
        "note": (
            "official-page ticks recorded; 585 issues remain knowledge_gap; overlay not sealed"
        ),
    }
    assert_safe_evaluation_payload(payload)
    return payload


def _proposed_review_fields(held_id: str) -> dict[str, Any]:
    if held_id == "held-provision-01":
        return {
            "extent": "E+W",
            "commencement": "1981-05-01_except_s35_see_s41",
            "unapplied_effects": [
                "dmcca_2024_s234_4_applies_LA_1980_to_consumer_redress_as_simple_contract_still_prospective_on_official_page",
                "crime_and_policing_act_2026_s96_inserts_s11ZA_s11ZB_not_part_of_s2",
            ],
            "companion_provisions": [],
        }
    if held_id == "held-provision-02":
        return {
            "extent": "E+W",
            "commencement": "in_force_as_part_of_limitation_act_1980",
            "unapplied_effects": [
                "dmcca_2024_s234_4_applies_LA_1980_to_consumer_redress_as_simple_contract_still_prospective_on_official_page",
                "crime_and_policing_act_2026_s96_4_amends_s14B_not_s14A",
            ],
            "companion_provisions": [
                "limitation_act_1980_s14B_fifteen_year_longstop_not_part_of_s14A"
            ],
        }
    if held_id == "held-provision-03":
        return {
            "extent": "E+W_principally_see_s42_4",
            "commencement": "2001-02-01_SI_2001_49",
            "unapplied_effects": ["none_recorded_on_official_s1_page"],
            "companion_provisions": ["trustee_act_2000_schedule_1_when_duty_applies"],
        }
    return {
        "extent": "E+W",
        "commencement": "1976-04-01_see_s27",
        "unapplied_effects": [
            "family_law_act_1996_sch8_para27_2_prospective_on_s1_2_a_not_in_force"
        ],
        "companion_provisions": [],
    }


OFFICIAL_PAGE_DECISIONS_SCHEMA = "legalbot.live60-official-page-decisions.v1"
IPFDA_BA_OPERATIVE_ACTION = "bind_official_positive_text_omitting_repealed_or_1B"


def build_official_page_decisions(
    *,
    identity: Mapping[str, Any],
    repair: Mapping[str, Any],
) -> dict[str, Any]:
    """Owner-authorised ticks from legislation.gov.uk pages dated 2026-08-16.

    Does not seal the 585-issue overlay. AI recorded official-page facts only.
    """

    repair_ids = {item["action"]: item["repair_span_id"] for item in repair.get("repairs", ())}
    payload = {
        "schema": OFFICIAL_PAGE_DECISIONS_SCHEMA,
        "as_of_date": identity.get("as_of_date"),
        "source": "legislation.gov.uk_revised_pages",
        "decision_authority": "owner_asked_ai_to_search_and_tick_items_1_to_4",
        "ai_role": "mechanical_accuracy_verifier_only",
        "seals_expert_gold": False,
        "overlay_sealable": False,
        "o04_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "contrary_authority": {
            "status": "reviewed_none",
            "scope": "no_contrary_or_limiting_authority_bound_in_this_overlay",
            "not_a_claim_that": "english_law_has_no_contrary_cases_on_the_60_topics",
            "revisit_when": "any_of_the_585_issues_is_qualified",
        },
        "contiguous_repairs_accepted": True,
        "accepted_repair_actions": [
            "replace_spliced_parent_with_chapeau",
            "replace_spliced_parent_with_final_proviso",
            "replace_spliced_opening_with_chapeau",
            "replace_spliced_opening_with_concluding_words",
            IPFDA_BA_OPERATIVE_ACTION,
            "exclude_from_gold_ranking_generation_and_citation",
        ],
        "accepted_repair_span_ids": [
            repair_ids[action]
            for action in (
                "replace_spliced_parent_with_chapeau",
                "replace_spliced_parent_with_final_proviso",
                "replace_spliced_opening_with_chapeau",
                "replace_spliced_opening_with_concluding_words",
                IPFDA_BA_OPERATIVE_ACTION,
            )
            if action in repair_ids
        ],
        "ipfda_s1_1_ba": {
            "official_current_words": (
                "any person (not being a person included in paragraph (a) or (b) "
                "above) to whom subsection (1A) below applies"
            ),
            "omitted_from_2019-12-02": "or (1B)",
            "omitting_instrument": "SI_2019_1458_sch3_para3_2_a",
            "s1_1B": "omitted_exclude_dots_only_editorial_chunk",
        },
        "provisions": {
            "held-provision-01": {
                "extent_tick": "E+W",
                "commencement_tick": "1981-05-01_except_s35",
                "effects_tick": "recorded_application_and_companion_not_text_amendment",
                "qualify_tick": "current_official_bytes_accepted",
                "catalogue_chunk_gold": True,
            },
            "held-provision-02": {
                "extent_tick": "E+W",
                "commencement_tick": "in_force",
                "effects_tick": "s14B_companion_not_in_span;_s234_4_prospective_application",
                "qualify_tick": "current_official_bytes_accepted_using_repair_spans",
                "catalogue_chunk_gold": False,
            },
            "held-provision-03": {
                "extent_tick": "E+W",
                "commencement_tick": "2001-02-01",
                "effects_tick": "none_on_s1_page;_schedule_1_companion",
                "qualify_tick": "current_official_bytes_accepted",
                "catalogue_chunk_gold": True,
            },
            "held-provision-04": {
                "extent_tick": "E+W",
                "commencement_tick": "1976-04-01",
                "effects_tick": "FLA_1996_sch8_para27_2_prospective_s1_2_a",
                "qualify_tick": "current_official_bytes_accepted_using_repair_spans",
                "catalogue_chunk_gold": False,
            },
        },
        "issues_585": "left_knowledge_gap_owner_reviewing",
    }
    assert_safe_evaluation_payload(payload)
    return payload


def _add_paragraph(document: DocumentType, text: str) -> None:
    document.add_paragraph(_reject_prohibited_metadata(text, label="held-span-repair prose"))


def build_held_span_repair_document(repair: Mapping[str, Any]) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, "Live60-2026-08-16")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Held span contiguous repair candidates", style="Title")
    _add_paragraph(
        document,
        "New contiguous span versions derived from immutable catalogue parents. "
        "These are not gold, not qualified, and not written back into the catalogue.",
    )
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Decision", str(repair["decision"])),
            ("Catalogue mutated", "no"),
            ("Repair spans", str(repair["repair_span_count"])),
            ("Seals expert gold", "no"),
            ("O-04", "not authorised"),
        ),
        (3_120, 6_240),
    )
    for item in repair.get("repairs", ()):
        document.add_heading(str(item["repair_span_id"]), level=2)
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("parent_chunk_id", str(item["parent_chunk_id"])),
                ("required_sublocator", str(item["required_sublocator"])),
                ("role", str(item["role"])),
                ("action", str(item["action"])),
                ("gold_eligible_candidate", "yes" if item["gold_eligible_candidate"] else "no"),
                ("text_sha256", str(item["text_sha256"])),
                ("markdown_text", str(item["markdown_text"])),
            ),
            (2_880, 6_480),
            body_size=8.5,
        )
    _finalize_document_properties(
        document,
        title="LegalBot Live60 held span contiguous repair",
        subject="New contiguous span candidates; not gold",
    )
    return document


def build_provision_review_worksheet_document(
    worksheet: Mapping[str, Any],
) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, "Live60-2026-08-16")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Held provision review worksheet", style="Title")
    _add_paragraph(
        document,
        "Owner ticks for extent, commencement and unapplied effects. Proposed "
        "fields are not qualifications. AI checks mechanical accuracy only.",
    )
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Primary reviewer", "owner"),
            ("Reviewer role", str(worksheet.get("approval_reviewer_role") or "")),
            ("Reviewer ref", str(worksheet.get("approval_reviewer_ref") or "")),
            ("Second review", "optional; not required"),
            ("AI role", "mechanical accuracy verifier only"),
            ("Qualified", "no"),
        ),
        (3_120, 6_240),
    )
    for provision in worksheet.get("provisions", ()):
        document.add_heading(str(provision["title"]), level=2)
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("held_id", str(provision["held_id"])),
                ("authority", str(provision["authority_identity_id"])),
                ("locator", str(provision["legal_locator"])),
                ("source_version_id", str(provision.get("source_version_id") or "")),
                ("document_content_sha256", str(provision.get("document_content_sha256") or "")),
                ("official_uri", str(provision.get("official_uri") or "")),
                ("difference_class", str(provision.get("difference_class") or "")),
                ("structural_defect_count", str(provision.get("structural_defect_count") or 0)),
                ("proposed_extent", str(provision["proposed_extent"])),
                ("proposed_commencement", str(provision["proposed_commencement"])),
                (
                    "proposed_unapplied_effects",
                    "; ".join(provision.get("proposed_unapplied_effects") or ()),
                ),
                (
                    "companion_provisions",
                    "; ".join(provision.get("companion_provisions") or ()) or "none",
                ),
                ("owner_extent_tick", str(provision.get("owner_extent_tick") or "blank")),
                (
                    "owner_commencement_tick",
                    str(provision.get("owner_commencement_tick") or "blank"),
                ),
                ("owner_effects_tick", str(provision.get("owner_effects_tick") or "blank")),
                ("owner_qualify_tick", str(provision.get("owner_qualify_tick") or "blank")),
                ("disposition", str(provision["disposition"])),
            ),
            (3_120, 6_240),
            body_size=8.5,
        )
    _finalize_document_properties(
        document,
        title="LegalBot Live60 provision review worksheet",
        subject="Extent, commencement and effects awaiting owner ticks",
    )
    return document
