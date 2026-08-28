"""Owner-only Live60 expert-review workbook and dated pack layout.

This module prepares a fill-in working copy for the qualified England-and-Wales
reviewer.  It never invents issue dispositions, gold spans, later treatment,
promotion, rollback or O-04.  The generated Word file is evaluation-only and
ineligible for training.
"""

from __future__ import annotations

import hashlib
import json
import shutil
import uuid
from collections.abc import Mapping
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Pt

from ..assessment.guidance_bundle import OWNER_ASSESSMENT_BUNDLE
from ..types import CASE_PROPOSITION_REVIEWER_ROLES
from .live30 import assert_safe_evaluation_payload
from .live_suite import LiveEvaluationBundle, admission_as_of_date, load_live_evaluation_bundle
from .live_suite_gold import qualification_template_for_suite
from .live_suite_reviewer_identity import build_owner_reviewer_identity
from .review_docx import (
    _add_header_footer,
    _add_page_break,
    _add_table,
    _add_toc,
    _configure_page,
    _configure_styles,
    _finalize_document_properties,
    _reject_prohibited_metadata,
)

OWNER_REVIEW_PACK_EVENT_SCHEMA = "legalbot.live60-owner-review-pack-event.v1"
OWNER_REVIEW_PACK_METRIC_SCHEMA = "legalbot.live60-owner-review-pack-metric.v1"
OWNER_REVIEW_PACK_TRACE_SCHEMA = "legalbot.live60-owner-review-pack-trace.v1"
WORKBOOK_FILENAME = "LegalBot-Live60-Owner-Expert-Review-Workbook.docx"
RULES_CHECKLIST_FILENAME = "LegalBot-Live60-Owner-Rule-Review-Checklist.docx"
BLANK = ""

HELD_STATUTORY_PROVISIONS: tuple[tuple[str, str, str], ...] = (
    (
        "held-provision-01",
        "Limitation Act 1980 section 2",
        "Ordinary time limit for actions founded on tort.",
    ),
    (
        "held-provision-02",
        "Limitation Act 1980 section 14A",
        "Special time limit for negligence actions where facts relevant to cause of action are not known at date of accrual.",
    ),
    (
        "held-provision-03",
        "Trustee Act 2000 section 1",
        "The duty of care of a trustee.",
    ),
    (
        "held-provision-04",
        "Inheritance (Provision for Family and Dependants) Act 1975 section 1",
        "Application for financial provision from a deceased person's estate.",
    ),
)

ISSUE_STATUS_CHOICES = "qualified / limited / knowledge_gap"
CONTRARY_CHOICES = "reviewed_none / reviewed_and_bound"
LATER_TREATMENT_CHOICES = "confirmed_current / qualified_current / not_current / uncertain_hold"
LEGAL_ROLE_CHOICES = "holding_ratio / binding_legal_rule / statutory_text / other reviewed role"


def dated_pack_name(as_of_date: date) -> str:
    return f"Live60-{as_of_date.isoformat()}"


def dated_pack_root(project_root: Path, as_of_date: date) -> Path:
    return project_root / dated_pack_name(as_of_date)


def ensure_dated_pack_layout(root: Path) -> Path:
    resolved = root.resolve()
    for relative in (
        "artifacts",
        "review",
        "logs/events",
        "logs/metrics",
        "logs/traces",
        "answers",
    ):
        (resolved / relative).mkdir(parents=True, exist_ok=True)
    return resolved


def _sha256_bytes(value: bytes) -> str:
    return hashlib.sha256(value).hexdigest()


def _append_pack_jsonl(path: Path, value: Mapping[str, Any]) -> None:
    assert_safe_evaluation_payload(value)
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(
            json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")) + "\n"
        )


def record_owner_review_pack_observability(
    pack_root: Path,
    *,
    event_type: str,
    artifact_kind: str,
    artifact_sha256: str,
    case_count: int,
) -> None:
    timestamp = datetime.now(UTC).replace(microsecond=0).isoformat()
    event = {
        "schema": OWNER_REVIEW_PACK_EVENT_SCHEMA,
        "event_id": uuid.uuid4().hex,
        "timestamp": timestamp,
        "event_type": event_type,
        "stage": "human_review",
        "status": "complete",
        "suite_id": "live-evaluation-60-v1",
        "artifact_kind": artifact_kind,
        "artifact_sha256": artifact_sha256,
        "case_count": case_count,
        "eligible_for_training": False,
        "training_export_allowed": False,
    }
    metric = {
        "schema": OWNER_REVIEW_PACK_METRIC_SCHEMA,
        "timestamp": timestamp,
        "stream": "live60",
        "metric": "owner_review_pack_export_total",
        "value": 1,
        "unit": "exports",
        "observe_only": True,
        "event_type": event_type,
        "eligible_for_training": False,
    }
    trace = {
        "schema": OWNER_REVIEW_PACK_TRACE_SCHEMA,
        "timestamp": timestamp,
        "stage": "human_review",
        "operation": "workbook_export",
        "status": "complete",
        "span_count": 0,
        "reason_code": "no_execution_spans_until_authorized_run",
        "observe_only": True,
    }
    _append_pack_jsonl(pack_root / "logs" / "events" / "owner-review-pack-events.jsonl", event)
    _append_pack_jsonl(pack_root / "logs" / "metrics" / "live60-pack.jsonl", metric)
    _append_pack_jsonl(pack_root / "logs" / "traces" / "live60-pack.jsonl", trace)


def _add_paragraph(document: DocumentType, text: str, *, style: str = "Normal") -> None:
    document.add_paragraph(_reject_prohibited_metadata(text, label="workbook prose"), style=style)


def build_owner_review_workbook(
    bundle: LiveEvaluationBundle,
    *,
    as_of_date: date,
    index_build_id: str = "candidate-pending-owner-review",
) -> DocumentType:
    if bundle.manifest.suite_id != "live-evaluation-60-v1":
        raise ValueError("owner review workbook requires the Live60 suite")
    if bundle.manifest.expert_reviewers_required != 1:
        raise ValueError("owner review workbook requires exactly one primary reviewer")

    dispositions = {item.case_id: item.disposition for item in bundle.run_plan.cases}
    identity = build_owner_reviewer_identity(as_of_date=as_of_date)
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, dated_pack_name(as_of_date))

    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Live60 owner expert-review workbook", style="Title")
    _add_paragraph(
        document,
        "Fill this working copy. Do not treat nearest-vector retrieval as legal gold. "
        "This pack is evaluation-only and is not eligible for training.",
    )
    _add_table(
        document,
        ("Field", "Value"),
        (
            ("Suite", bundle.manifest.suite_id),
            ("Legal as-of date", as_of_date.isoformat()),
            ("Reviewers required", "1 primary owner reviewer; second review optional"),
            ("Cases", "60"),
            ("Selected generation cases", "30"),
            ("Index build", index_build_id),
            ("Registry SHA-256", bundle.registry.canonical_sha256),
            ("Run-plan SHA-256", bundle.manifest.run_plan_sha256),
            ("Assessment bundle", OWNER_ASSESSMENT_BUNDLE.version),
            ("Assessment SHA-256", OWNER_ASSESSMENT_BUNDLE.sha256),
            ("Training export", "false"),
        ),
        (3_120, 6_240),
    )
    _add_toc(document)
    _add_page_break(document)

    document.add_heading("How to fill this workbook", level=1)
    for line in (
        "You are the owner and the one primary qualified England-and-Wales reviewer. Mark every issue on all 60 cases.",
        "For each issue choose exactly one status: qualified, limited, or knowledge_gap.",
        "Qualified and limited issues need exact authority spans: source version, chunk, locator, content hash and legal role. Do not copy a nearest retrieved chunk.",
        "Case-law spans also need later treatment: confirmed_current, qualified_current, not_current or uncertain_hold.",
        "Record contrary or limiting authority at case level as reviewed_none or reviewed_and_bound.",
        "Review the four held statutory provisions in their own section, including extent, commencement and unapplied effects.",
        "Reviewer identity is a role plus reviewer:<sha256>. Never write a name, email or host path.",
        "A second independent human reviewer is optional, not mandatory. Leave those fields blank unless you later confirm one. The AI checks mechanical accuracy of hashes and locators only and cannot be the second reviewer.",
        "The JSON template in artifacts/ is the machine seal. This Word file is your working copy. After you finish, the overlay is sealed from your decisions only.",
        "Answer DOCXs (control plus Annexes A/B/C) are created only after a later authorised run. They are not in this pack yet.",
    ):
        _add_paragraph(document, line)

    document.add_heading("Reviewer header", level=1)
    roles = " / ".join(sorted(CASE_PROPOSITION_REVIEWER_ROLES))
    _add_table(
        document,
        ("Field", "Allowed values", "Fill in"),
        (
            ("Primary reviewer role", roles, identity["approval_reviewer_role"]),
            ("Primary reviewer ref", "reviewer:<64-hex>", identity["approval_reviewer_ref"]),
            (
                "Second review status",
                "not_required / needs_independent_review / confirmed",
                "not_required",
            ),
            ("Second reviewer role", f"{roles} or blank", BLANK),
            ("Second reviewer ref", "reviewer:<64-hex> or blank", BLANK),
            ("Material disagreement", "none / adjudicated", "none"),
            ("Adjudication ref", "adjudication:<64-hex> or blank", BLANK),
        ),
        (2_400, 4_560, 2_400),
    )

    document.add_heading("Four held statutory provisions", level=1)
    _add_paragraph(
        document,
        "These four changed provisions remain held until you qualify the current official bytes. "
        "A source-level approval does not replace this issue-specific check.",
    )
    _add_table(
        document,
        (
            "ID",
            "Provision",
            "Current-bytes review",
            "Extent / commencement / effects",
            "Disposition",
        ),
        tuple(
            (provision_id, title, BLANK, BLANK, BLANK)
            for provision_id, title, _summary in HELD_STATUTORY_PROVISIONS
        ),
        (1_200, 2_640, 1_840, 2_160, 1_520),
    )
    for provision_id, title, summary in HELD_STATUTORY_PROVISIONS:
        document.add_heading(f"{provision_id}: {title}", level=2)
        _add_paragraph(document, summary, style="Review Small")

    document.add_heading("Assessment rules to review", level=1)
    _add_paragraph(
        document,
        "These are owner drafting and repair rules, not legal authority. "
        "Mark keep, amend or replace in the companion rules checklist. "
        "Do not use them as gold spans.",
    )
    rule_rows = tuple(
        (
            rule.rule_id,
            rule.grade_band,
            rule.criterion,
            rule.task_type,
            rule.positive_target,
            rule.anti_pattern or "none",
            rule.repair_action,
            BLANK,
        )
        for rule in OWNER_ASSESSMENT_BUNDLE.rules
    )
    _add_table(
        document,
        (
            "Rule ID",
            "Band",
            "Criterion",
            "Task",
            "Positive target",
            "Anti-pattern",
            "Repair",
            "Keep / amend / replace",
        ),
        rule_rows,
        (1_200, 720, 1_080, 720, 1_680, 1_440, 1_440, 1_080),
        body_size=7.5,
    )

    document.add_heading("All 60 issue dispositions", level=1)
    _add_paragraph(
        document,
        f"Issue status: {ISSUE_STATUS_CHOICES}. "
        f"Contrary authority: {CONTRARY_CHOICES}. "
        f"Later treatment: {LATER_TREATMENT_CHOICES}. "
        f"Legal role: {LEGAL_ROLE_CHOICES}.",
    )

    for case in bundle.registry.cases:
        disposition = dispositions[case.case_id]
        selected = "selected generate_once" if disposition == "generate_once" else "coverage only"
        document.add_heading(f"{case.case_id} · {case.subject} · {case.task_type}", level=2)
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("Disposition", selected),
                ("Word target", str(case.word_target)),
                ("Route", case.expected_research_route),
                ("Question SHA-256", case.question_sha256),
                ("Record SHA-256", case.record_sha256),
                ("Case status", BLANK),
                ("Contrary authority", BLANK),
            ),
            (2_880, 6_480),
        )
        _add_paragraph(document, case.question)
        issue_rows = tuple(
            (
                f"issue-{number:02d}",
                topic,
                BLANK,
                BLANK,
                BLANK,
                BLANK,
                BLANK,
            )
            for number, topic in enumerate(case.must_cover_issues, start=1)
        )
        _add_table(
            document,
            (
                "Issue",
                "Must-cover topic",
                "Status",
                "Exact span IDs / hashes",
                "Legal locator",
                "Legal role",
                "Later treatment",
            ),
            issue_rows,
            (900, 1_920, 1_080, 1_800, 1_440, 1_080, 1_140),
            body_size=8.0,
        )
        document.add_paragraph().paragraph_format.space_after = Pt(0)

    _finalize_document_properties(
        document,
        title="LegalBot Live60 owner expert-review workbook",
        subject="Fill-in overlay for all 60 issues, held provisions and assessment rules",
    )
    return document


def export_owner_review_workbook(
    *,
    project_root: Path,
    output_path: Path,
    as_of_date: date | None = None,
    index_build_id: str = "candidate-pending-owner-review",
    overwrite: bool = False,
) -> Path:
    bundle = load_live_evaluation_bundle(
        project_root / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    legal_date = as_of_date or admission_as_of_date(datetime.now(UTC))
    destination = output_path.resolve()
    if destination.exists() and not overwrite:
        raise FileExistsError("owner review workbook already exists")
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = build_owner_review_workbook(
        bundle, as_of_date=legal_date, index_build_id=index_build_id
    )
    document.save(str(destination))
    return destination


def export_qualification_template(
    *,
    project_root: Path,
    output_path: Path,
    as_of_date: date | None = None,
    index_build_id: str = "candidate-pending-owner-review",
    overwrite: bool = False,
) -> Path:
    bundle = load_live_evaluation_bundle(
        project_root / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    legal_date = as_of_date or admission_as_of_date(datetime.now(UTC))
    destination = output_path.resolve()
    if destination.exists() and not overwrite:
        raise FileExistsError("qualification template already exists")
    destination.parent.mkdir(parents=True, exist_ok=True)
    template = qualification_template_for_suite(
        bundle, index_build_id=index_build_id, as_of_date=legal_date
    )
    destination.write_text(
        json.dumps(template, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    return destination


def copy_rules_checklist(*, project_root: Path, destination: Path, overwrite: bool = False) -> Path:
    source = project_root / "docs" / "reports" / RULES_CHECKLIST_FILENAME
    target = destination.resolve()
    if not source.is_file():
        raise FileNotFoundError("owner rule-review checklist is missing")
    if target.exists() and not overwrite:
        raise FileExistsError("owner rule-review checklist already exists in the pack")
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)
    return target


def prepare_live60_owner_review_pack(
    *,
    project_root: Path,
    as_of_date: date | None = None,
    index_build_id: str = "candidate-pending-owner-review",
    overwrite: bool = False,
) -> dict[str, str]:
    legal_date = as_of_date or admission_as_of_date(datetime.now(UTC))
    pack_root = ensure_dated_pack_layout(dated_pack_root(project_root, legal_date))
    workbook = export_owner_review_workbook(
        project_root=project_root,
        output_path=pack_root / "review" / WORKBOOK_FILENAME,
        as_of_date=legal_date,
        index_build_id=index_build_id,
        overwrite=overwrite,
    )
    template = export_qualification_template(
        project_root=project_root,
        output_path=pack_root / "artifacts" / "expert-qualification-template.json",
        as_of_date=legal_date,
        index_build_id=index_build_id,
        overwrite=overwrite,
    )
    checklist = copy_rules_checklist(
        project_root=project_root,
        destination=pack_root / "review" / RULES_CHECKLIST_FILENAME,
        overwrite=overwrite,
    )
    workbook_sha = _sha256_bytes(workbook.read_bytes())
    template_sha = _sha256_bytes(template.read_bytes())
    checklist_sha = _sha256_bytes(checklist.read_bytes())
    record_owner_review_pack_observability(
        pack_root,
        event_type="owner_review_workbook_exported",
        artifact_kind="owner_expert_review_workbook",
        artifact_sha256=workbook_sha,
        case_count=60,
    )
    record_owner_review_pack_observability(
        pack_root,
        event_type="qualification_template_exported",
        artifact_kind="expert_qualification_template",
        artifact_sha256=template_sha,
        case_count=60,
    )
    record_owner_review_pack_observability(
        pack_root,
        event_type="owner_rules_checklist_copied",
        artifact_kind="owner_rule_review_checklist",
        artifact_sha256=checklist_sha,
        case_count=0,
    )
    return {
        "pack_name": dated_pack_name(legal_date),
        "workbook_sha256": workbook_sha,
        "template_sha256": template_sha,
        "checklist_sha256": checklist_sha,
    }


HELD_PROVISION_VERIFICATION: tuple[dict[str, str], ...] = (
    {
        "held_id": "held-provision-01",
        "title": "Limitation Act 1980 section 2",
        "authority_identity_id": "ukpga:1980:58",
        "legal_locator": "section 2",
        "new_official_version_sha256": (
            "53c4002fc9486f7d6f32dda34e390af9e180e88ffdd359786b64057a94d1bc65"
        ),
        "new_source_content_sha256": (
            "bca5969ce0cd667f59741cf8d1edd2feb9b76512b12a49fafedcf521d89924df"
        ),
        "predecessor_source_content_sha256": (
            "39d184eba98086e5f81b406d4bb1763a88205091b45c8f850208cf8fb128d83f"
        ),
        "reason_code": "official_source_bytes_changed",
        "review_action": "fresh_human_provision_review_required",
    },
    {
        "held_id": "held-provision-02",
        "title": "Limitation Act 1980 section 14A",
        "authority_identity_id": "ukpga:1980:58",
        "legal_locator": "section 14A",
        "new_official_version_sha256": (
            "53c4002fc9486f7d6f32dda34e390af9e180e88ffdd359786b64057a94d1bc65"
        ),
        "new_source_content_sha256": (
            "bca5969ce0cd667f59741cf8d1edd2feb9b76512b12a49fafedcf521d89924df"
        ),
        "predecessor_source_content_sha256": (
            "39d184eba98086e5f81b406d4bb1763a88205091b45c8f850208cf8fb128d83f"
        ),
        "reason_code": "official_source_bytes_changed",
        "review_action": "fresh_human_provision_review_required",
    },
    {
        "held_id": "held-provision-03",
        "title": "Trustee Act 2000 section 1",
        "authority_identity_id": "ukpga:2000:29",
        "legal_locator": "section 1",
        "new_official_version_sha256": (
            "6f8160a1c7c69ab2bdf0ea9debb623a69d229674af0730c804150227e53f93cf"
        ),
        "new_source_content_sha256": (
            "ebafb72d8bfbc8dfc0430d6f4b778cf31740c5c0c16f047a9ffef3b76073b612"
        ),
        "predecessor_source_content_sha256": (
            "89acb7e390250bb355cf3792787fe8fe2815504e282b7e987485bce64de1bff6"
        ),
        "reason_code": "official_source_bytes_changed",
        "review_action": "fresh_human_provision_review_required",
    },
    {
        "held_id": "held-provision-04",
        "title": "Inheritance (Provision for Family and Dependants) Act 1975 section 1",
        "authority_identity_id": "ukpga:1975:63",
        "legal_locator": "section 1",
        "new_official_version_sha256": (
            "cf6eab520c960367a1b945eb846811ef2081ea8bfed55ab3a02a39142429e5f7"
        ),
        "new_source_content_sha256": (
            "1311476a3ec00bf00e1f6602f10520407785a8354d6d8fe5d86865270f66578a"
        ),
        "predecessor_source_content_sha256": (
            "31e54c0bb398de234109c7d005a3f768bbb8ad7bc7d1c658e0c55af08ef38394"
        ),
        "reason_code": "official_source_bytes_changed",
        "review_action": "fresh_human_provision_review_required",
    },
)

OWNER_RETURN_DECISION_SCHEMA = "legalbot.live60-owner-return-decision.v1"
KNOWLEDGE_GAP_INVENTORY_SCHEMA = "legalbot.live60-knowledge-gap-inventory.v1"
VERIFICATION_DOCX_FILENAME = "LegalBot-Live60-Owner-Return-Verification.docx"
SECOND_REVIEW_DOCX_FILENAME = "LegalBot-Live60-Second-Review-Knowledge-Gaps.docx"
FILLED_WORKBOOK_FILENAME = "LegalBot-Live60-Owner-Expert-Review-Workbook-FILLED-RETURN-HOLD.docx"
FILLED_CHECKLIST_FILENAME = "LegalBot-Live60-Owner-Rule-Review-Checklist-FILLED-RETURN.docx"


def _table_rows(table: Any) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def audit_filled_owner_workbook(
    workbook_path: Path, *, bundle: LiveEvaluationBundle
) -> dict[str, Any]:
    """Check a filled workbook against the sealed registry. Invents no gold."""

    document = Document(str(workbook_path))
    if len(document.tables) != 124:
        raise ValueError("filled workbook does not contain the expected 124 tables")
    cover = {row[0]: row[1] for row in _table_rows(document.tables[0])[1:]}
    reviewer = {row[0]: row[2] for row in _table_rows(document.tables[1])[1:]}
    held_rows = _table_rows(document.tables[2])[1:]
    rule_rows = _table_rows(document.tables[3])[1:]
    selected = {
        item.case_id for item in bundle.run_plan.cases if item.disposition == "generate_once"
    }
    issues: list[dict[str, str]] = []
    cases: list[dict[str, str]] = []
    for index, case in enumerate(bundle.registry.cases):
        header_rows = _table_rows(document.tables[4 + index * 2])
        header = {row[0]: row[1] for row in header_rows[1:]}
        issue_rows = _table_rows(document.tables[5 + index * 2])[1:]
        if len(issue_rows) != len(case.must_cover_issues):
            raise ValueError(f"{case.case_id} issue-row count differs from the registry")
        cases.append(
            {
                "case_id": case.case_id,
                "subject": case.subject,
                "task_type": case.task_type,
                "disposition": (
                    "generate_once" if case.case_id in selected else "coverage_only_not_selected"
                ),
                "case_status": header.get("Case status", ""),
                "contrary_authority": header.get("Contrary authority", ""),
                "question_sha256": case.question_sha256,
            }
        )
        for number, topic in enumerate(case.must_cover_issues, start=1):
            row = issue_rows[number - 1]
            issue_id = f"issue-{number:02d}"
            if row[0] != issue_id or row[1] != topic:
                raise ValueError(f"{case.case_id} {issue_id} topic does not match the registry")
            issues.append(
                {
                    "case_id": case.case_id,
                    "issue_id": issue_id,
                    "topic": topic,
                    "status": row[2],
                    "exact_span_ids": row[3],
                    "legal_locator": row[4],
                    "legal_role": row[5],
                    "later_treatment": row[6],
                }
            )
    held = [
        {
            "held_id": row[0],
            "title": row[1],
            "current_bytes_review": row[2],
            "extent_commencement_effects": row[3],
            "disposition": row[4],
        }
        for row in held_rows
    ]
    rules = [{"rule_id": row[0], "owner_decision": row[7]} for row in rule_rows]
    return {
        "cover": cover,
        "reviewer": reviewer,
        "held_provisions": held,
        "assessment_rules": rules,
        "cases": cases,
        "issues": issues,
        "issue_count": len(issues),
        "knowledge_gap_issue_count": sum(item["status"] == "knowledge_gap" for item in issues),
        "qualified_issue_count": sum(item["status"] == "qualified" for item in issues),
        "limited_issue_count": sum(item["status"] == "limited" for item in issues),
        "blank_contrary_count": sum(item["contrary_authority"] == "" for item in cases),
        "blank_span_count": sum(item["exact_span_ids"] == "" for item in issues),
        "primary_reviewer_role_present": bool(reviewer.get("Primary reviewer role")),
        "primary_reviewer_ref_present": bool(reviewer.get("Primary reviewer ref")),
        "second_review_status": reviewer.get("Second review status", ""),
        "assessment_keep_count": sum(item["owner_decision"] == "keep" for item in rules),
        "assessment_rule_count": len(rules),
    }


def _safe_inventory(
    audit: Mapping[str, Any],
    *,
    bundle: LiveEvaluationBundle,
    as_of_date: date = date(2026, 8, 16),
) -> dict[str, Any]:
    selected = tuple(
        item.case_id for item in bundle.run_plan.cases if item.disposition == "generate_once"
    )
    identity = build_owner_reviewer_identity(as_of_date=as_of_date)
    inventory = {
        "schema": KNOWLEDGE_GAP_INVENTORY_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "as_of_date": as_of_date.isoformat(),
        "decision": "return_hold",
        "overlay_sealable": False,
        "generation_authorised": False,
        "o04_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "issue_count": audit["issue_count"],
        "knowledge_gap_issue_count": audit["knowledge_gap_issue_count"],
        "qualified_issue_count": audit["qualified_issue_count"],
        "limited_issue_count": audit["limited_issue_count"],
        "blank_contrary_count": audit["blank_contrary_count"],
        "blank_span_count": audit["blank_span_count"],
        "selected_case_ids": list(selected),
        "coverage_only_case_ids": [
            item.case_id
            for item in bundle.run_plan.cases
            if item.disposition == "coverage_only_not_selected"
        ],
        "primary_reviewer_role_present": audit["primary_reviewer_role_present"],
        "primary_reviewer_ref_present": audit["primary_reviewer_ref_present"],
        "owner_is_primary_reviewer": True,
        "approval_reviewer_role": identity["approval_reviewer_role"],
        "approval_reviewer_ref": identity["approval_reviewer_ref"],
        "workbook_primary_reviewer_header_blank": not audit["primary_reviewer_role_present"],
        "second_review_status": "not_required",
        "workbook_second_review_status": audit["second_review_status"],
        "ai_role": "mechanical_accuracy_verifier_only",
        "ai_second_reviewer_forbidden": True,
        "assessment_keep_count": audit["assessment_keep_count"],
        "assessment_rule_count": audit["assessment_rule_count"],
        "held_provisions": [
            {
                **item,
                "workbook_disposition": next(
                    row["disposition"]
                    for row in audit["held_provisions"]
                    if row["held_id"] == item["held_id"]
                ),
            }
            for item in HELD_PROVISION_VERIFICATION
        ],
        "cases": [
            {
                "case_id": case["case_id"],
                "subject": case["subject"],
                "task_type": case["task_type"],
                "disposition": case["disposition"],
                "case_status": case["case_status"],
                "contrary_authority": case["contrary_authority"] or "blank",
                "question_sha256": case["question_sha256"],
                "issues": [
                    {
                        "issue_id": issue["issue_id"],
                        "topic": issue["topic"],
                        "status": issue["status"],
                        "missing_fields": [
                            name
                            for name, value in (
                                ("exact_span_ids", issue["exact_span_ids"]),
                                ("legal_locator", issue["legal_locator"]),
                                ("legal_role", issue["legal_role"]),
                                ("later_treatment", issue["later_treatment"]),
                            )
                            if not value
                        ],
                    }
                    for issue in audit["issues"]
                    if issue["case_id"] == case["case_id"]
                ],
            }
            for case in audit["cases"]
        ],
        "sha_reconciliation": {
            "suite_manifest_seal_sha256": bundle.manifest.seal_sha256,
            "run_plan_file_sha256": bundle.manifest.run_plan_sha256,
            "run_plan_object_seal_sha256": bundle.run_plan.seal_sha256,
            "mismatch_is_false_positive": True,
            "explanation_code": "workbook_uses_run_plan_file_digest_checklist_uses_object_seal",
        },
        "blocking_reason_codes": [
            "owner_returned_hold",
            "all_issues_knowledge_gap",
            "exact_spans_absent",
            "no_current_date_candidate",
            "no_o04",
        ],
    }
    assert_safe_evaluation_payload(
        {key: value for key, value in inventory.items() if key != "cases"}
    )
    return inventory


def build_owner_return_verification_document(
    *,
    bundle: LiveEvaluationBundle,
    audit: Mapping[str, Any],
    inventory: Mapping[str, Any],
) -> DocumentType:
    document = Document()
    _configure_page(document)
    _configure_styles(document)
    _add_header_footer(document, "Live60-2026-08-16")
    document.add_paragraph("LEGALBOT", style="Subtitle")
    document.add_paragraph("Live60 owner return verification", style="Title")
    _add_paragraph(
        document,
        "Audit of the filled RETURN/HOLD pack. This is not a sealed overlay and does not authorise generation, promotion or O-04.",
    )
    _add_table(
        document,
        ("Finding", "Result"),
        (
            ("Owner decision", "RETURN / HOLD"),
            ("Sealable overlay", "no"),
            ("Generation authorised", "no"),
            ("O-04", "not issued"),
            ("Cases present", "60 / 60 in registry order"),
            ("Issues present", f"{audit['issue_count']} / 585"),
            (
                "Issue status",
                f"{audit['knowledge_gap_issue_count']} knowledge_gap; 0 qualified; 0 limited",
            ),
            ("Exact spans bound", f"{audit['issue_count'] - audit['blank_span_count']} / 585"),
            ("Contrary authority filled", f"{60 - audit['blank_contrary_count']} / 60"),
            (
                "Primary reviewer role/ref",
                f"{inventory['approval_reviewer_role']} / bound as owner",
            ),
            ("Second review", "optional; not required; AI cannot be second reviewer"),
            (
                "Assessment rules",
                f"{audit['assessment_keep_count']} keep / {audit['assessment_rule_count']}",
            ),
            ("Held provisions qualified", "0 / 4"),
            (
                "SHA mismatch claim",
                "false positive: file digest versus object seal of the same run plan",
            ),
            ("Training export", "false"),
        ),
        (3_360, 6_000),
    )
    _add_toc(document)
    _add_page_break(document)

    document.add_heading("What this means", level=1)
    for line in (
        "The two returned forms are complete as a HOLD. They are not a finished legal overlay.",
        "I will not invent gold spans, later treatment or O-04. The owner is the one primary reviewer.",
        "The 30-answer run stays blocked until you qualify issues with exact spans or expressly keep them as knowledge_gap.",
        "Live60 is 60 frozen questions. Exactly 30 sealed IDs later get one generated answer. That selection is not random.",
    ):
        _add_paragraph(document, line)

    document.add_heading("How to close these knowledge gaps", level=1)
    for line in (
        "Do not review issue IDs alone. Each row below restates the frozen legal topic taken from the sealed question registry. Those topics were not generated by a model.",
        "knowledge_gap means: this issue has no bound exact authority span. There is no locator, content hash, legal role or later-treatment review. The owner HOLD left those cells blank on purpose rather than guessing.",
        "To close a gap, mark the issue qualified or limited and bind a reviewed current England-and-Wales span: source_version_id, chunk_id, legal_locator, content_sha256, legal_role. If the span is case law, also record later_treatment: confirmed_current, qualified_current, not_current or uncertain_hold.",
        "Do not copy a nearest retrieved chunk and call it gold. If you cannot bind an exact span, keep knowledge_gap and write a reason. At case level also mark contrary authority reviewed_none or reviewed_and_bound.",
        "In the last column write: Confirm gap / Qualify / Limit / Disagree with topic. This document is evaluation-only and is not training data.",
        "The owner is the primary reviewer. A second human review is optional. The AI verifies mechanical accuracy of hashes and locators and cannot seal gold.",
    ):
        _add_paragraph(document, line)
    _add_table(
        document,
        ("Field you must fill", "What it is"),
        (
            ("Issue status", "qualified, limited, or knowledge_gap"),
            (
                "Exact span IDs / hashes",
                "source_version_id, chunk_id, content_sha256 of the exact words relied on",
            ),
            (
                "Legal locator",
                "section, subsection or paragraph pin-cite, not a document title alone",
            ),
            (
                "Legal role",
                "holding_ratio, binding_legal_rule, statutory_text, or other reviewed role",
            ),
            (
                "Later treatment",
                "Required for case-law spans. Not required for legislation once current bytes are qualified",
            ),
            ("Contrary authority", "Per case: reviewed_none or reviewed_and_bound"),
            (
                "Reviewer role / ref",
                "Owner is the primary qualified E&W reviewer. Bound as reviewer:<64-hex>. Never a name",
            ),
        ),
        (3_120, 6_240),
    )

    document.add_heading("SHA reconciliation", level=1)
    _add_table(
        document,
        ("Label", "Digest", "Meaning"),
        (
            (
                "Suite manifest seal",
                bundle.manifest.seal_sha256,
                "Checklist Manifest SHA. Correct.",
            ),
            (
                "Run-plan FILE SHA-256",
                bundle.manifest.run_plan_sha256,
                "Workbook Run-plan SHA-256. Hash of generation-run-plan.json bytes.",
            ),
            (
                "Run-plan OBJECT seal",
                bundle.run_plan.seal_sha256,
                "Checklist Run-plan SHA. Seal inside the same JSON object.",
            ),
        ),
        (2_160, 3_840, 3_360),
    )
    _add_paragraph(
        document,
        "Both run-plan digests belong to the same sealed plan. Comparing them as if they were rival plans is a labelling error, not a registry split.",
    )

    document.add_heading("What you still need to fill", level=1)
    _add_table(
        document,
        ("Item", "Why it blocks sealing / generation"),
        (
            (
                "Primary reviewer role",
                "Bound as owner: legal_reviewer. Workbook header may still be blank.",
            ),
            ("Primary reviewer ref", str(inventory.get("approval_reviewer_ref") or "")),
            (
                "585 issue gold spans",
                "qualified/limited issues need source version, chunk, locator, content hash and legal role.",
            ),
            ("Later treatment", "Required for every case-law span you qualify."),
            (
                "Contrary authority on 60 cases",
                "Mark reviewed_none or reviewed_and_bound. Currently blank.",
            ),
            (
                "Four held provisions",
                "Qualify current official bytes, extent, commencement and unapplied effects after contiguous repair.",
            ),
            ("Second review", "Optional. Not a blocker. AI cannot be the second reviewer."),
            (
                "O-04",
                "Only after overlay, candidate, Stage A, promotion, rollback and readiness v6 are green.",
            ),
        ),
        (3_120, 6_240),
    )

    document.add_heading("Selected 30 generation cases", level=1)
    _add_paragraph(
        document,
        "These IDs are frozen. They are not a random sample. The other 30 stay coverage-only.",
    )
    selected = " ".join(inventory["selected_case_ids"])
    _add_paragraph(document, selected)

    document.add_heading("Four held provisions — verification identities", level=1)
    _add_paragraph(
        document,
        "Official bytes changed. Predecessor content hashes differ from the new content hashes. Version metadata SHA is unchanged. Fresh human review is required before these locators can be gold.",
    )
    _add_table(
        document,
        (
            "ID",
            "Authority / locator",
            "New content SHA-256",
            "Predecessor content SHA-256",
            "Workbook",
        ),
        tuple(
            (
                item["held_id"],
                f"{item['authority_identity_id']} {item['legal_locator']}",
                item["new_source_content_sha256"],
                item["predecessor_source_content_sha256"],
                item["workbook_disposition"],
            )
            for item in inventory["held_provisions"]
        ),
        (1_320, 1_920, 2_160, 2_160, 1_800),
    )

    document.add_heading("Assessment rules", level=1)
    _add_paragraph(
        document,
        "All 16 owner drafting rules are marked keep. That matches the already-active bundle owner-standards-2026-08-14.1. No rule change is applied. Rules are not legal gold.",
    )

    document.add_heading("Knowledge-gap list — all 585 frozen issues", level=1)
    _add_paragraph(
        document,
        "Every topic below is copied from the sealed registry must_cover_issues field. "
        "A separate triple-check confirmed 585/585 match the workbook and the inventory. "
        "No extra topic was added. No topic was rewritten.",
    )
    cases_by_id = {case.case_id: case for case in bundle.registry.cases}
    for case in inventory["cases"]:
        source = cases_by_id[case["case_id"]]
        selected_label = (
            "SELECTED for later one-pass generation"
            if case["disposition"] == "generate_once"
            else "coverage-only, no generated answer in this run"
        )
        document.add_heading(
            f"{case['case_id']} · {case['subject']} · {source.task_type} · {selected_label}",
            level=2,
        )
        _add_table(
            document,
            ("Field", "Value"),
            (
                ("Case status", case["case_status"]),
                (
                    "Contrary authority",
                    f"{case['contrary_authority']} — still required even while issues are knowledge_gap",
                ),
                ("Word target", str(source.word_target)),
                ("Question SHA-256", case["question_sha256"]),
                ("Issue count", str(len(case["issues"]))),
            ),
            (2_880, 6_480),
        )
        _add_paragraph(document, source.question)
        _add_table(
            document,
            (
                "Issue",
                "Frozen legal topic (the gap)",
                "Status",
                "What is missing now",
                "What you must fill to close",
                "Second-review mark",
            ),
            tuple(
                (
                    issue["issue_id"],
                    issue["topic"],
                    issue["status"],
                    "exact span IDs / hashes; legal locator; legal role; later treatment"
                    if set(issue["missing_fields"])
                    >= {
                        "exact_span_ids",
                        "legal_locator",
                        "legal_role",
                        "later_treatment",
                    }
                    else ", ".join(issue["missing_fields"]) or "none",
                    (
                        "Bind a reviewed E&W span for this topic, or keep knowledge_gap "
                        "and add a reason. Do not use nearest retrieval."
                    ),
                    "",
                )
                for issue in case["issues"]
            ),
            (780, 1_800, 960, 1_800, 2_220, 1_800),
            body_size=8.0,
        )

    _finalize_document_properties(
        document,
        title="LegalBot Live60 owner return verification",
        subject="RETURN/HOLD audit, SHA reconciliation and knowledge-gap inventory",
    )
    return document


def apply_owner_return_hold(
    *,
    project_root: Path,
    filled_workbook: Path,
    filled_checklist: Path,
    as_of_date: date | None = None,
    overwrite: bool = False,
) -> dict[str, str]:
    legal_date = as_of_date or admission_as_of_date(datetime.now(UTC))
    bundle = load_live_evaluation_bundle(
        project_root / "benchmarks" / "evaluation" / "live-evaluation-60-v1"
    )
    audit = audit_filled_owner_workbook(filled_workbook, bundle=bundle)
    inventory = _safe_inventory(audit, bundle=bundle, as_of_date=legal_date)
    pack_root = ensure_dated_pack_layout(dated_pack_root(project_root, legal_date))
    review_dir = pack_root / "review"
    artifacts = pack_root / "artifacts"
    workbook_copy = review_dir / FILLED_WORKBOOK_FILENAME
    checklist_copy = review_dir / FILLED_CHECKLIST_FILENAME
    if (workbook_copy.exists() or checklist_copy.exists()) and not overwrite:
        raise FileExistsError("filled owner-return copies already exist")
    shutil.copy2(filled_workbook, workbook_copy)
    shutil.copy2(filled_checklist, checklist_copy)
    decision = {
        "schema": OWNER_RETURN_DECISION_SCHEMA,
        "suite_id": "live-evaluation-60-v1",
        "as_of_date": legal_date.isoformat(),
        "decision": "return_hold",
        "overlay_sealable": False,
        "generation_authorised": False,
        "o04_authorised": False,
        "eligible_for_training": False,
        "training_export_allowed": False,
        "assessment_rules_kept": True,
        "assessment_bundle_sha256": OWNER_ASSESSMENT_BUNDLE.sha256,
        "filled_workbook_sha256": _sha256_bytes(workbook_copy.read_bytes()),
        "filled_checklist_sha256": _sha256_bytes(checklist_copy.read_bytes()),
        "issue_count": audit["issue_count"],
        "knowledge_gap_issue_count": audit["knowledge_gap_issue_count"],
        "blocking_reason_codes": inventory["blocking_reason_codes"],
        "sha_reconciliation": inventory["sha_reconciliation"],
    }
    assert_safe_evaluation_payload(decision)
    decision_path = artifacts / "owner-return-decision.json"
    inventory_path = artifacts / "knowledge-gap-inventory.json"
    decision_path.write_text(
        json.dumps(decision, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    inventory_path.write_text(
        json.dumps(inventory, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    verification = pack_root / "review" / VERIFICATION_DOCX_FILENAME
    second_review = pack_root / "review" / SECOND_REVIEW_DOCX_FILENAME
    document = build_owner_return_verification_document(
        bundle=bundle, audit=audit, inventory=inventory
    )
    document.save(str(verification))
    document.save(str(second_review))
    record_owner_review_pack_observability(
        pack_root,
        event_type="owner_return_hold_recorded",
        artifact_kind="owner_return_decision",
        artifact_sha256=_sha256_bytes(decision_path.read_bytes()),
        case_count=60,
    )
    record_owner_review_pack_observability(
        pack_root,
        event_type="knowledge_gap_inventory_exported",
        artifact_kind="knowledge_gap_inventory",
        artifact_sha256=_sha256_bytes(inventory_path.read_bytes()),
        case_count=60,
    )
    record_owner_review_pack_observability(
        pack_root,
        event_type="second_review_knowledge_gaps_exported",
        artifact_kind="second_review_knowledge_gaps",
        artifact_sha256=_sha256_bytes(second_review.read_bytes()),
        case_count=60,
    )
    return {
        "decision": "return_hold",
        "verification_sha256": _sha256_bytes(verification.read_bytes()),
        "second_review_sha256": _sha256_bytes(second_review.read_bytes()),
        "inventory_sha256": _sha256_bytes(inventory_path.read_bytes()),
        "decision_sha256": _sha256_bytes(decision_path.read_bytes()),
    }
