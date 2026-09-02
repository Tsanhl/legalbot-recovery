#!/usr/bin/env python3
"""Build the complete GE visible-training-unseen owner-review DOCX."""

from __future__ import annotations

import hashlib
import json
import os
from collections import defaultdict
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RUN_ID = "LegalBot-GE-2026-09-01-improvement-training-unseen-r3"
RUN_ROOT = ROOT / "data/evaluations/general-enquiries" / RUN_ID
OUTPUT = ROOT / "output/docx/LegalBot-GE-331-Training-and-60-Unseen-Full-Review-r2.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "007F79"
RED = "9B1C1C"
GOLD = "7A5A00"
INK = "1F2937"
MUTED = "5B6573"
GRID = "B8C5D6"
PALE_BLUE = "E8EEF5"
PALE_RED = "FDECEC"
PALE_GOLD = "FFF4D6"
PALE_TEAL = "E8F5F4"
WHITE = "FFFFFF"

TOPIC_LABELS = {
    "administrative-law": "Administrative law",
    "ai-and-data-protection": "AI and data protection",
    "business-and-company-law": "Business and company law",
    "commercial-law": "Commercial law",
    "competition-law": "Competition law",
    "contemporary-biolaw-and-regulation": "Contemporary biolaw and regulation",
    "contract-law": "Contract law",
    "criminal-law": "Criminal law",
    "eu-internal-market-law": "EU internal market law",
    "international-commercial-mediation": "International commercial mediation",
    "land-law": "Land law",
    "law-and-medicine": "Law and medicine",
    "pensions-law": "Pensions law",
    "private-international-law": "Private international law",
    "tort-law": "Tort law",
    "trusts-law": "Trusts law",
    "wills-and-estates": "Wills and estates",
}


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _load_json(path: Path) -> dict[str, Any]:
    value = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(value, dict):
        raise RuntimeError(f"expected JSON object: {path}")
    return value


def _load_jsonl(path: Path) -> list[dict[str, Any]]:
    values = [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line]
    if not all(isinstance(value, dict) for value in values):
        raise RuntimeError(f"invalid JSONL rows: {path}")
    return values


def _load_run() -> tuple[dict[str, Any], list[dict[str, Any]], list[dict[str, Any]]]:
    manifest = _load_json(RUN_ROOT / "RUN-MANIFEST.json")
    visible = _load_jsonl(RUN_ROOT / "visible/RESULTS.jsonl")
    unseen = _load_jsonl(RUN_ROOT / "unseen/RESULTS.jsonl")
    if len(visible) != 331 or [row["ordinal"] for row in visible] != list(range(1, 332)):
        raise RuntimeError("visible result denominator or order differs from 331")
    if len(unseen) != 60 or [row["ordinal"] for row in unseen] != list(range(1, 61)):
        raise RuntimeError("diagnostic unseen denominator or order differs from 60")
    artifacts = {item["path"]: item for item in manifest.get("artifacts", [])}
    for relative in ("visible/RESULTS.jsonl", "unseen/RESULTS.jsonl"):
        item = artifacts.get(relative)
        path = RUN_ROOT / relative
        if not item or item.get("sha256") != _sha256(path):
            raise RuntimeError(f"run artifact identity differs: {relative}")
    return manifest, visible, unseen


def _set_font(run: Any, *, size: float | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _shade(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _cell_margins(
    cell: Any, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, width in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")


def _table_geometry(table: Any, widths: Iterable[int], *, indent: int = 120) -> None:
    widths = tuple(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_pr = table._tbl.tblPr
    width_node = table_pr.first_child_found_in("w:tblW")
    if width_node is None:
        width_node = OxmlElement("w:tblW")
        table_pr.append(width_node)
    width_node.set(qn("w:w"), str(sum(widths)))
    width_node.set(qn("w:type"), "dxa")
    indent_node = table_pr.first_child_found_in("w:tblInd")
    if indent_node is None:
        indent_node = OxmlElement("w:tblInd")
        table_pr.append(indent_node)
    indent_node.set(qn("w:w"), str(indent))
    indent_node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(widths[index]))
            tc_width.set(qn("w:type"), "dxa")
            _cell_margins(cell)


def _repeat_header(row: Any) -> None:
    row_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    row_pr.append(node)


def _field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    visible = OxmlElement("w:t")
    visible.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, visible, end))


def _configure(doc: DocumentType) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 30, NAVY, 0, 10),
        ("Subtitle", 14, MUTED, 0, 8),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("LEGALBOT  /  GENERAL ENQUIRIES  /  OWNER REVIEW")
    run.bold = True
    _set_font(run, size=8, color=NAVY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Non-authorizing review  |  Page ")
    _set_font(run, size=8, color=MUTED)
    _field(footer, "PAGE")


def _add_kv(doc: DocumentType, label: str, value: str, *, color: str = INK) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(label + ": ")
    run.bold = True
    _set_font(run, color=NAVY)
    run = paragraph.add_run(value)
    _set_font(run, color=color)


def _add_cover(doc: DocumentType, manifest: dict[str, Any]) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(108)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("PHASE 2  /  EVALUATION → IMPROVEMENT → UNSEEN")
    run.bold = True
    _set_font(run, size=10, color=TEAL)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("General Enquiries\nFull Owner Review")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "331 visible cases  •  retrieval training  •  60 realistic diagnostic unseen cases"
    )
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.paragraph_format.space_before = Pt(26)
    notice.paragraph_format.space_after = Pt(20)
    run = notice.add_run(
        "Every case shows its answer, evidence, factual result, 70+ result or ineligibility, "
        "and the exact reason for re-evaluation or improvement."
    )
    run.bold = True
    _set_font(run, size=11, color=NAVY)
    card = doc.add_table(rows=5, cols=2)
    _table_geometry(card, (1800, 7560))
    _repeat_header(card.rows[0])
    rows = (
        ("Run", manifest["run_id"]),
        ("Legal cutoff", "28 August 2026"),
        ("Visible", "331 cases; exact visible-pack order preserved"),
        ("Unseen", "60 exposed diagnostic/regression records; 15 families; not sealed Validation"),
        (
            "Authority",
            "Owner review only; not qualified legal gold, promotion, live or sealed validation",
        ),
    )
    for index, (label, value) in enumerate(rows):
        left, right = card.rows[index].cells
        left.text = label
        right.text = value
        _shade(left, NAVY)
        _shade(right, PALE_BLUE)
        for run in left.paragraphs[0].runs:
            run.bold = True
            _set_font(run, size=9, color=WHITE)
        for run in right.paragraphs[0].runs:
            _set_font(run, size=9, color=INK)
    doc.add_page_break()


def _summary_table(
    doc: DocumentType, visible: list[dict[str, Any]], unseen: list[dict[str, Any]]
) -> None:
    doc.add_heading("Results at a glance", level=1)
    table = doc.add_table(rows=1, cols=5)
    _table_geometry(table, (3120, 1560, 1560, 1560, 1560))
    headers = ("Lane", "Cases", "Evidence", "Factual pass", "70+ pass")
    for index, text in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = text
        _shade(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            _set_font(run, size=8.5, color=WHITE)
    _repeat_header(table.rows[0])
    for label, values in (
        ("Visible evaluation / training", visible),
        ("Diagnostic unseen", unseen),
    ):
        cells = table.add_row().cells
        summary = (
            label,
            len(values),
            sum(bool(row.get("evidence")) for row in values),
            sum(row["factual_result"]["outcome"] == "FACTUAL_PASS" for row in values),
            sum(
                row["quality_70_plus"]["outcome"] in {"MEETS_70_STANDARD", "EXCEEDS_70_STANDARD"}
                for row in values
            ),
        )
        for index, value in enumerate(summary):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[index].paragraphs[0].runs:
                _set_font(run, size=9)
        _shade(cells[3], PALE_RED)
        _shade(cells[4], PALE_RED)


def _add_summary(
    doc: DocumentType,
    manifest: dict[str, Any],
    visible: list[dict[str, Any]],
    unseen: list[dict[str, Any]],
) -> None:
    _summary_table(doc, visible, unseen)
    doc.add_heading("What the improvement cycle completed", level=2)
    _add_kv(
        doc,
        "Retrieval repair",
        "Exact provision routing replaced broad subject matches; absent essential authority now fails closed.",
    )
    _add_kv(
        doc,
        "Training",
        "331 retrieval-planner examples were generated and prompt/retrieval tuning completed.",
    )
    _add_kv(
        doc,
        "Weight training",
        "Not run: no qualified legal-gold answer targets exist. Treating AI drafts as gold would train errors.",
        color=RED,
    )
    _add_kv(
        doc,
        "Unseen",
        "The 60 diagnostic cases are exposed regression material, not fresh unseen or sealed Validation. "
        "They comprise 15 scenario families with four modifiers each (60 question records).",
        color=GOLD,
    )
    _add_kv(
        doc,
        "Sealed validation",
        "Not run: no distinct owner-approved external private Validation root was supplied.",
        color=GOLD,
    )
    doc.add_heading("Why every case remains ineligible for 70+", level=2)
    doc.add_paragraph(
        "The factual gate runs before prose quality. The admitted source snapshots were reviewed only "
        "through 14 August 2026, earlier than the 28 August 2026 case cutoff, and the answers have not "
        "received qualified legal review. A 70+ score is therefore prohibited even where the quoted "
        "provision is relevant. This is an ineligibility result, not a zero-quality score."
    )
    doc.add_heading("Evidence coverage by topic", level=2)
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in visible:
        grouped[str(row["topic_id"])].append(row)
    table = doc.add_table(rows=1, cols=4)
    _table_geometry(table, (5040, 1440, 1440, 1440))
    for index, text in enumerate(("Topic", "Cases", "Evidence", "Gaps")):
        cell = table.cell(0, index)
        cell.text = text
        _shade(cell, NAVY)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            _set_font(run, size=8.5, color=WHITE)
    _repeat_header(table.rows[0])
    for topic, rows in grouped.items():
        cells = table.add_row().cells
        evidence_count = sum(bool(row.get("evidence")) for row in rows)
        values = (
            TOPIC_LABELS.get(topic, topic),
            len(rows),
            evidence_count,
            len(rows) - evidence_count,
        )
        for index, value in enumerate(values):
            cells[index].text = str(value)
            for run in cells[index].paragraphs[0].runs:
                _set_font(run, size=8.5)
        if len(rows) - evidence_count:
            _shade(cells[3], PALE_GOLD)
    doc.add_heading("How to review a case", level=2)
    doc.add_paragraph(
        "Read the response, then inspect the evidence and factual checks. If the evidence is absent or "
        "the source-currentness check fails, use Re-evaluate or Improve / train. Approve only after the "
        "answer and cited propositions have received the required legal and currentness review."
    )
    doc.add_page_break()


def _status_table(doc: DocumentType, row: dict[str, Any]) -> None:
    table = doc.add_table(rows=2, cols=3)
    _table_geometry(table, (3120, 3120, 3120))
    _repeat_header(table.rows[0])
    for index, label in enumerate(("FACTUAL RESULT", "70+ RESULT", "EVIDENCE SPANS")):
        cell = table.cell(0, index)
        cell.text = label
        _shade(cell, NAVY)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            _set_font(run, size=8, color=WHITE)
    values = (
        row["factual_result"]["outcome"].replace("_", " "),
        row["quality_70_plus"]["outcome"].replace("_", " "),
        str(len(row.get("evidence") or [])),
    )
    for index, value in enumerate(values):
        cell = table.cell(1, index)
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shade(cell, PALE_BLUE if index == 2 else PALE_RED)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            _set_font(run, size=9, color=NAVY if index == 2 else RED)


ANSWER_HEADINGS = {
    "Initial legal orientation",
    "Planner output",
    "Answer shown to the user",
    "What I would do now",
    "Documents to gather",
    "Questions I need answered",
    "Status",
}


def _answer(doc: DocumentType, text: str) -> None:
    for raw in text.splitlines():
        line = raw.strip().replace("\u200b", "")
        if not line:
            continue
        if line in ANSWER_HEADINGS:
            doc.add_heading(line, level=3)
            continue
        paragraph = doc.add_paragraph()
        if line.startswith("- "):
            paragraph.paragraph_format.left_indent = Inches(0.25)
            line = line[2:]
        if line.startswith("“"):
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.15)
            run = paragraph.add_run(line)
            run.italic = True
            _set_font(run, color=DARK_BLUE)
        else:
            paragraph.add_run(line)


def _evidence(doc: DocumentType, evidence: list[dict[str, Any]]) -> None:
    if not evidence:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(
            "No sufficiently relevant primary-authority passage was admitted for this case."
        )
        run.bold = True
        _set_font(run, color=RED)
        return
    for index, item in enumerate(evidence, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"Evidence {index}: {item['oscola_parenthetical']}")
        run.bold = True
        _set_font(run, color=NAVY)
        quote = doc.add_paragraph()
        quote.paragraph_format.left_indent = Inches(0.25)
        quote.paragraph_format.right_indent = Inches(0.15)
        run = quote.add_run(f"“{item['quote']}”")
        run.italic = True
        _set_font(run, color=DARK_BLUE)
        _add_kv(
            doc,
            "Source identity",
            f"{item['title']} | {item['locator']} | {item['stable_identifier']}",
        )
        _add_kv(
            doc,
            "Currentness / extent",
            f"reviewed as of {item.get('currentness_reviewed_as_of_date') or 'unknown'}; "
            f"status {item.get('currentness_status')}; extent {item.get('provision_extent_status')}; "
            f"full-current-law eligible={item.get('full_current_law_verification_eligible')}",
            color=RED if not item.get("full_current_law_verification_eligible") else TEAL,
        )


def _factual(doc: DocumentType, row: dict[str, Any]) -> None:
    result = row["factual_result"]
    _add_kv(doc, "Outcome", str(result["outcome"]).replace("_", " "), color=RED)
    for name, outcome in result["checks"].items():
        reason = str(result["reasons"].get(name) or "No reason recorded.")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{name.replace('_', ' ').title()} — {outcome}: ")
        run.bold = True
        if outcome == "FAIL":
            color = RED
        elif outcome == "PASS":
            color = TEAL
        elif outcome == "NOT_ASSESSABLE":
            color = GOLD
        else:
            color = MUTED
        _set_font(run, color=color)
        p.add_run(reason)
    diagnostic = result.get("diagnostic_checks") or {}
    if diagnostic:
        doc.add_heading("Split diagnostic checks", level=3)
        for name, payload in diagnostic.items():
            if isinstance(payload, dict):
                outcome = str(payload.get("outcome") or "")
                reason = str(payload.get("reason") or "No reason recorded.")
            else:
                outcome = str(payload)
                reason = ""
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"{name.replace('_', ' ').title()} — {outcome}: ")
            run.bold = True
            if outcome == "FAIL":
                color = RED
            elif outcome == "PASS":
                color = TEAL
            elif outcome == "NOT_ASSESSABLE":
                color = GOLD
            else:
                color = MUTED
            _set_font(run, color=color)
            if reason:
                p.add_run(reason)


def _improvement(doc: DocumentType, row: dict[str, Any]) -> None:
    reasons = list(row.get("improvement_reasons") or [])
    if not reasons:
        doc.add_paragraph("No improvement reason was recorded.")
    for reason in reasons:
        _add_kv(doc, "Reason", str(reason).replace("_", " "), color=RED)
    missing = list(row.get("known_missing_primary_authorities") or [])
    if missing:
        _add_kv(doc, "Known missing primary authority", "; ".join(map(str, missing)), color=GOLD)
    training = row.get("training_eligibility") or {}
    if training:
        _add_kv(
            doc,
            "Retrieval-planner tuning eligible",
            str(training.get("retrieval_planner_tuning_eligible", training.get("retrieval_planner_tuning"))),
        )
        _add_kv(
            doc,
            "Retrieval-planner tuning consumed",
            str(training.get("retrieval_planner_tuning_consumed", False)),
            color=RED if training.get("retrieval_planner_tuning_consumed") else MUTED,
        )
        _add_kv(
            doc, "Answer-weight training", str(training.get("answer_weight_training")), color=RED
        )
        if training.get("reason"):
            _add_kv(
                doc,
                "Weight-training hold",
                str(training["reason"]),
                color=RED,
            )


def _decision(doc: DocumentType) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    labels = ("☐ Approve", "☐ Re-evaluate", "☐ Improve / train", "☐ Hold")
    for index, label in enumerate(labels):
        run = paragraph.add_run(label + ("     " if index != len(labels) - 1 else ""))
        run.bold = True
        _set_font(
            run,
            size=9,
            color=TEAL if index == 0 else GOLD if index in {1, 2} else RED,
        )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.add_run(
        "Owner note: __________________________________________________________________________"
    )


def _case(doc: DocumentType, row: dict[str, Any], *, label: str, ordinal: int) -> None:
    heading = doc.add_heading(f"{label} {ordinal:03d}  ·  {row['case_id']}", level=1)
    heading.paragraph_format.keep_with_next = True
    _add_kv(
        doc,
        "Topic / scenario",
        f"{TOPIC_LABELS.get(str(row['topic_id']), str(row['topic_id']))} | {row['scenario_family_id']}",
    )
    question = doc.add_paragraph()
    question.paragraph_format.space_before = Pt(5)
    question.paragraph_format.space_after = Pt(10)
    run = question.add_run(str(row["question"]))
    run.bold = True
    _set_font(run, size=11.5, color=NAVY)
    _status_table(doc, row)
    doc.add_heading("Answer shown for review", level=2)
    _answer(doc, str(row["answer"]))
    doc.add_heading("Evidence and OSCOLA parentheticals", level=2)
    _evidence(doc, list(row.get("evidence") or []))
    doc.add_heading("Factual result", level=2)
    _factual(doc, row)
    doc.add_heading("70+ result or ineligibility", level=2)
    quality = row["quality_70_plus"]
    _add_kv(doc, "Outcome", str(quality["outcome"]).replace("_", " "), color=RED)
    _add_kv(doc, "Score", "Not scored" if quality.get("score") is None else str(quality["score"]))
    _add_kv(doc, "Reason", str(quality["reason"]), color=RED)
    doc.add_heading("Exact re-evaluation / improvement reason", level=2)
    _improvement(doc, row)
    doc.add_heading("Owner decision", level=2)
    _decision(doc)


def _cases(
    doc: DocumentType, rows: list[dict[str, Any]], *, title: str, label: str, intro: str
) -> None:
    doc.add_heading(title, level=1)
    doc.add_paragraph(intro)
    doc.add_page_break()
    for ordinal, row in enumerate(rows, start=1):
        _case(doc, row, label=label, ordinal=ordinal)
        if ordinal != len(rows):
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    doc.add_page_break()


def _appendix(doc: DocumentType, manifest: dict[str, Any]) -> None:
    doc.add_heading("Appendix  ·  Integrity and release limits", level=1)
    identities = (
        ("Run ID", manifest["run_id"]),
        ("Run content SHA-256", manifest["content_sha256"]),
        ("Visible pack manifest SHA-256", manifest["visible_pack_manifest_sha256"]),
        ("Source manifest SHA-256", manifest["source_manifest_sha256"]),
        ("Source count", str(manifest["source_count"])),
        ("Visible results SHA-256", _sha256(RUN_ROOT / "visible/RESULTS.jsonl")),
        ("Unseen questions SHA-256", _sha256(RUN_ROOT / "unseen/PRIVATE-QUESTIONS.jsonl")),
        ("Unseen results SHA-256", _sha256(RUN_ROOT / "unseen/RESULTS.jsonl")),
    )
    for label, value in identities:
        _add_kv(doc, label, value)
    doc.add_heading("Release limits", level=2)
    doc.add_paragraph(
        "This package is a readable projection of immutable run artifacts. It does not convert the AI "
        "candidate answers into legal gold, does not provide qualified legal review, does not execute "
        "sealed Validation, and does not authorize promotion or live activation. The diagnostic unseen "
        "lane is useful for failure discovery but is not the official sealed Validation lane."
    )
    doc.add_paragraph(
        "No question, source, chunk, embedding, prior run, diagnostic, or evidence artifact was deleted "
        "or overwritten while preparing this package."
    )


def _properties(doc: DocumentType) -> None:
    props = doc.core_properties
    props.title = "LegalBot General Enquiries 331 Training and 60 Unseen Full Review"
    props.subject = (
        "Factual-first owner review of GE evaluation, improvement training and diagnostic unseen"
    )
    props.author = "LegalBot recovery workspace"
    props.keywords = (
        "LegalBot, General Enquiries, evaluation, training, unseen, factual gate, 70+, OSCOLA"
    )
    props.comments = "Non-authorizing owner-review artifact; not legal gold, sealed validation, promotion or live output."


def main() -> int:
    manifest, visible, unseen = _load_run()
    if OUTPUT.exists() or OUTPUT.is_symlink():
        raise FileExistsError(f"refusing to replace existing DOCX: {OUTPUT}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _configure(doc)
    _properties(doc)
    _add_cover(doc, manifest)
    _add_summary(doc, manifest, visible, unseen)
    _cases(
        doc,
        visible,
        title="Part I  ·  331 visible evaluation and training cases",
        label="Visible case",
        intro=(
            "The exact visible denominator and order are preserved. These answers are the strict post-repair "
            "candidate outputs used to create retrieval-planner training examples."
        ),
    )
    _cases(
        doc,
        unseen,
        title="Part II  ·  60 realistic diagnostic unseen cases",
        label="Unseen case",
        intro=(
            "These 60 records are 15 scenario families with four modifiers. After owner review they are "
            "exposed diagnostic/regression material, not fresh unseen and not sealed Validation."
        ),
    )
    _appendix(doc, manifest)
    doc.save(OUTPUT)
    os.chmod(OUTPUT, 0o600)
    print(
        json.dumps(
            {
                "path": str(OUTPUT),
                "sha256": _sha256(OUTPUT),
                "visible_cases": len(visible),
                "unseen_cases": len(unseen),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            },
            indent=2,
            sort_keys=True,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
