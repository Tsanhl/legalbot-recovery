#!/usr/bin/env python3
"""Build create-only owner-review batches from a blocked Phase-2A package.

This command implements the owner-performed internal-research-tool route.  It
does not invoke a model, make a legal judgment, admit a source, mutate a sealed
candidate, qualify an issue, or authorize Phase 2B.  Advisory-AI findings and
owner decisions are separate, sealed records; only the owner record can express
the owner's substantive decision.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import os
import re
import stat
from collections.abc import Mapping, Sequence
from datetime import UTC, date, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell

PACKAGE_SCHEMA = "legalbot.v111-phase2a-remediation-package-index.v1"
REVIEW_INDEX_SCHEMA = "legalbot.v111-phase2a-owner-review-index.v1"
REVIEW_POLICY_SCHEMA = "legalbot.v111-phase2a-owner-review-policy.v1"
DECISION_SCHEMA = "legalbot.v111-phase2a-owner-decision.v1"
ADVISORY_AI_SCHEMA = "legalbot.v111-phase2a-advisory-ai-review.v1"
TERMINAL_VERDICT = "PHASE 2A SAFELY STOPPED — PHASE 2B AND DEVELOPMENT 30 NOT AUTHORIZED"
OWNER_ROUTE = "OWNER_PERFORMED_INTERNAL_RESEARCH_TOOL_WITH_ADVISORY_AI"
EXPECTED_ARTIFACT_IDS = (
    "canonical-registry-snapshot",
    "remediation-matrix-585",
    "gold-case-reconciliation-509",
    "candidate-impact-reconciliation-76",
    "legislative-effects-register-1896",
    "judgment-later-treatment-register-20",
    "official-source-provenance-register",
    "gold-successor-manifest",
    "successor-source-admission-manifest",
    "successor-candidate-decision",
    "retrieval-reattestation",
    "corrected-all585-qualification",
    "cutoff-proposal",
    "material-change-policy",
    "advisory-ai-audit",
    "exact-head-verification",
    "owner-adoption-draft",
    "final-invariants",
)
ALLOWED_OWNER_OUTCOMES: dict[str, frozenset[str]] = {
    "issue": frozenset(
        {
            "APPROVE_PROPOSITION_BINDINGS",
            "APPROVE_NONMATERIAL_NOTE",
            "CONFIRM_MATERIAL_GAP",
            "REQUEST_MORE_EVIDENCE",
        }
    ),
    "legislative_effect": frozenset(
        {
            "APPROVE_EFFECT_DISPOSITION",
            "APPROVE_NONMATERIAL_NOTE",
            "CONFIRM_MATERIAL_GAP",
            "REQUEST_MORE_EVIDENCE",
        }
    ),
    "judgment": frozenset(
        {
            "APPROVE_JUDGMENT_TREATMENT",
            "CONFIRM_MATERIAL_GAP",
            "REQUEST_MORE_EVIDENCE",
        }
    ),
    "source_version": frozenset(
        {
            "APPROVE_SOURCE_VERSION_MATERIALITY",
            "APPROVE_NONMATERIAL_NOTE",
            "CONFIRM_MATERIAL_GAP",
            "REQUEST_MORE_EVIDENCE",
        }
    ),
}
_SAFE_RUN_ID = re.compile(r"^[a-z0-9][a-z0-9._:-]{2,127}$")
_SAFE_FILE = re.compile(r"^[a-zA-Z0-9][a-zA-Z0-9._-]{0,199}\.json$")
_SAFE_DOCX = re.compile(r"^[a-zA-Z0-9][a-zA-Z0-9._-]{0,199}\.docx$")
_SHA256 = re.compile(r"^[0-9a-f]{64}$")


def _canonical_json(value: Any) -> bytes:
    return (
        json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")) + "\n"
    ).encode("utf-8")


def _sha256(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def _sealed(value: Any) -> str:
    return _sha256(_canonical_json(value))


def _load_json(path: Path) -> Any:
    if path.is_symlink() or not path.is_file():
        raise ValueError("phase2a_owner_review_input_must_be_regular_file")
    return json.loads(path.read_bytes())


def _write_exclusive(path: Path, raw: bytes) -> None:
    descriptor = os.open(path, os.O_WRONLY | os.O_CREAT | os.O_EXCL | os.O_NOFOLLOW, 0o600)
    try:
        with os.fdopen(descriptor, "wb") as handle:
            handle.write(raw)
            handle.flush()
            os.fsync(handle.fileno())
    except BaseException:
        path.unlink(missing_ok=True)
        raise


def _set_run_font(
    run: Any,
    *,
    size: float,
    bold: bool = False,
    color: str = "000000",
) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell: _Cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    for existing in properties.findall(qn("w:shd")):
        properties.remove(existing)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_table_geometry(table: Table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != 9_360 or len(widths_dxa) != len(table.columns):
        raise ValueError("phase2a_owner_review_docx_table_geometry_invalid")
    table.autofit = False
    properties = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        for existing in properties.findall(qn(tag)):
            properties.remove(existing)
    width = OxmlElement("w:tblW")
    width.set(qn("w:w"), "9360")
    width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    properties.append(width)
    properties.append(indent)
    properties.append(layout)

    grid = table._tbl.tblGrid
    for existing in list(grid):
        grid.remove(existing)
    for value in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa, strict=True):
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")
            margins = cell_properties.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell_properties.append(margins)
            for side, amount in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    margins.append(node)
                node.set(qn("w:w"), str(amount))
                node.set(qn("w:type"), "dxa")


def _set_cell_text(
    cell: _Cell,
    value: str,
    *,
    size: float = 9,
    bold: bool = False,
    color: str = "000000",
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    _set_run_font(paragraph.add_run(value), size=size, bold=bold, color=color)


def _add_review_table(
    document: DocumentType,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths_dxa: Sequence[int],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers, strict=True):
        _set_cell_text(cell, value, bold=True, color="1F4D78")
        _set_cell_shading(cell, "F2F4F7")
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            _set_cell_text(cell, value)
    _set_table_geometry(table, widths_dxa)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _configure_review_docx(document: DocumentType, *, run_id: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    styles = {
        "Title": (24, "000000", 0, 6),
        "Subtitle": (12, "555555", 0, 12),
        "Heading 1": (16, "2E74B5", 12, 6),
        "Heading 2": (13, "2E74B5", 10, 5),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in styles.items():
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(
        header.add_run(f"LEGALBOT · PHASE 2A OWNER REVIEW · {run_id}"),
        size=8,
        bold=True,
        color="5B6573",
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run_font(
        footer.add_run("Private local review · Non-authorizing · Page "),
        size=8,
        color="5B6573",
    )
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = footer.add_run()
    field_run._r.extend((begin, instruction, separate))
    _set_run_font(footer.add_run("1"), size=8, color="5B6573")
    footer.add_run()._r.append(end)


def _owner_review_companion_docx(
    *,
    run_id: str,
    recorded_at: datetime,
    source_package_digest: str,
    category_counts: Mapping[str, int],
    category_batch_counts: Mapping[str, int],
) -> bytes:
    """Build a readable, non-authorizing companion to the complete JSON logs."""

    document = Document()
    _configure_review_docx(document, run_id=run_id)
    document.add_paragraph("OWNER DECISION MEMO", style="Subtitle")
    document.add_paragraph("Phase 2A evidence and hallucination review control", style="Title")
    document.add_paragraph(
        "Owner-performed internal research-tool review with pinned advisory AI",
        style="Subtitle",
    )
    _add_review_table(
        document,
        ("Control", "Recorded value"),
        (
            ("Run ID", run_id),
            ("Recorded at", recorded_at.astimezone(UTC).isoformat(timespec="seconds")),
            ("Route", OWNER_ROUTE),
            ("Source Phase-2A package digest", source_package_digest),
            ("Authority", "NON-AUTHORIZING OWNER-REVIEW COMPANION"),
        ),
        (2_520, 6_840),
    )
    status = document.add_table(rows=1, cols=1)
    status.style = "Table Grid"
    _set_cell_text(
        status.cell(0, 0),
        "PHASE 2A SAFELY STOPPED. This document does not authorize Phase 2B, a split, "
        "Development 30, source admission, candidate mutation, promotion, validation or live use.",
        size=10,
        bold=True,
        color="7A3E00",
    )
    _set_cell_shading(status.cell(0, 0), "FFF4E5")
    _set_table_geometry(status, (9_360,))

    document.add_heading("Decision boundary", level=1)
    document.add_paragraph(
        "Official-source evidence and deterministic checks come first. Advisory AI may organise "
        "evidence, compare versions, identify discrepancies and recommend a row-level outcome. "
        "Only an explicit owner record may make the substantive decision. Silence is never approval."
    )
    document.add_paragraph(
        "The full record is the sealed JSON folder beside this DOCX. This companion is for reading, "
        "comments and decision preparation; editing it does not alter or approve any machine record."
    )

    document.add_heading("Review inventory", level=1)
    _add_review_table(
        document,
        ("Category", "Items", "JSON batches"),
        tuple(
            (
                category.replace("_", " ").title(),
                str(category_counts[category]),
                str(category_batch_counts[category]),
            )
            for category in ("issue", "legislative_effect", "judgment", "source_version")
        ),
        (4_200, 2_280, 2_880),
    )

    document.add_page_break()
    document.add_heading("Advisory AI authority", level=1)
    _add_review_table(
        document,
        ("Control", "Required treatment"),
        (
            ("Execution mode", "Separate verification pass using the same configured model adapter"),
            ("Model-independent", "No"),
            ("May do", "Recommend, organise, compare, flag discrepancies and cite frozen evidence IDs"),
            ("Cannot do", "Decide, adopt, qualify, admit a source, rebuild, sign or authorize a gate"),
            ("Positive recommendation", "Never overrides deterministic checks or owner authority"),
            ("Concern or uncertainty", "May raise a fail-closed owner-review hold only"),
            ("Persisted content", "Concise findings and evidence references; no hidden reasoning"),
        ),
        (2_520, 6_840),
    )

    document.add_heading("Owner review workflow", level=1)
    _add_review_table(
        document,
        ("Step", "Owner action", "Required record"),
        (
            ("1", "Open the matching JSON batch and locate the item ID", "Exact item SHA-256"),
            ("2", "Check official source/version/span evidence first", "Evidence-reference SHA-256s"),
            ("3", "Read any advisory AI recommendation", "AI review SHA-256 or unavailable record"),
            ("4", "Write your own conclusion and rationale", "Owner outcome, findings and comments"),
            ("5", "Type your name and decision date", "Explicit owner decision record"),
            ("6", "Re-run package validation", "New receipt; no automatic next phase"),
        ),
        (700, 5_220, 3_440),
    )

    document.add_page_break()
    document.add_heading("Knowledge issue and hallucination record", level=1)
    document.add_paragraph(
        "Use one record per issue. For a future authorized internal Development 30, the owner packet "
        "must preserve the exact question, full release-eligible output, claim-level deterministic "
        "quality findings, typed-fact and atomicity failures, evidence IDs, knowledge gaps, advisory "
        "AI recommendation, owner comments and final owner decision. Held or unsafe drafts remain "
        "encrypted and are represented by safe failure references."
    )
    _add_review_table(
        document,
        ("Field", "Owner / reviewer entry"),
        (
            ("Run, case and item ID", ""),
            ("Question SHA-256 / output SHA-256", ""),
            ("Deterministic finding codes", ""),
            ("Knowledge-gap IDs and evidence references", ""),
            ("Advisory AI recommendation and review SHA-256", ""),
            ("Owner comments", "\n\n"),
            ("Owner decision", "[ ] Approve  [ ] Repair  [ ] Hold  [ ] Request more evidence"),
            ("Owner typed name / decision date", ""),
        ),
        (3_240, 6_120),
    )

    document.add_page_break()
    document.add_heading("Phase-2A owner decision form", level=1)
    _add_review_table(
        document,
        ("Decision field", "Owner entry"),
        (
            ("Batch file and item ID", ""),
            ("Decision basis SHA-256s", ""),
            ("Outcome", "[ ] Approve  [ ] Nonmaterial note  [ ] Material gap  [ ] More evidence"),
            ("Substantive findings", "\n\n"),
            ("Owner rationale and comments", "\n\n\n"),
            ("Advisory AI disposition", "[ ] Used  [ ] Unavailable; deterministic evidence used"),
            ("Owner typed name", ""),
            ("Decision date (YYYY-MM-DD)", ""),
        ),
        (3_240, 6_120),
    )
    document.add_paragraph(
        "STOP VERDICT: PHASE 2A OWNER REVIEW REQUIRED. PHASE 2B AND DEVELOPMENT 30 ARE NOT AUTHORIZED.",
    ).runs[0].bold = True

    core = document.core_properties
    core.title = "LegalBot Phase 2A owner review control"
    core.subject = "Non-authorizing evidence, hallucination and owner-decision companion"
    core.author = ""
    core.last_modified_by = ""
    core.keywords = ""
    core.comments = ""
    core.created = datetime(2000, 1, 1, tzinfo=UTC)
    core.modified = datetime(2000, 1, 1, tzinfo=UTC)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _artifact(schema: str, payload: Mapping[str, Any]) -> dict[str, Any]:
    artifact = {"schema": schema, **payload}
    artifact["artifact_sha256"] = _sealed(artifact)
    return artifact


def _require_sequence(value: Any, *, code: str) -> list[Any]:
    if not isinstance(value, list):
        raise ValueError(code)
    return value


def verify_remediation_package(package_root: Path) -> dict[str, Any]:
    """Replay the blocked package before producing owner-review material."""

    if package_root.is_symlink() or not package_root.is_dir():
        raise ValueError("phase2a_owner_review_package_root_invalid")
    index_path = package_root / "PACKAGE-INDEX.json"
    index = _load_json(index_path)
    if not isinstance(index, dict):
        raise ValueError("phase2a_owner_review_package_index_invalid")
    supplied_digest = str(index.get("package_digest") or "")
    digest_material = dict(index)
    digest_material.pop("package_digest", None)
    if (
        index.get("schema") != PACKAGE_SCHEMA
        or index.get("phase") != "2A"
        or index.get("authorizing") is not False
        or index.get("terminal_verdict") != TERMINAL_VERDICT
        or supplied_digest != _sealed(digest_material)
    ):
        raise ValueError("phase2a_owner_review_package_binding_invalid")

    entries = _require_sequence(
        index.get("entries"), code="phase2a_owner_review_package_entries_invalid"
    )
    if (
        len(entries) != len(EXPECTED_ARTIFACT_IDS)
        or tuple(index.get("artifact_order") or ()) != EXPECTED_ARTIFACT_IDS
        or int(index.get("artifact_count") or 0) != len(EXPECTED_ARTIFACT_IDS)
    ):
        raise ValueError("phase2a_owner_review_package_inventory_invalid")

    artifacts: dict[str, dict[str, Any]] = {}
    for ordinal, (expected_id, entry) in enumerate(
        zip(EXPECTED_ARTIFACT_IDS, entries, strict=True), start=1
    ):
        if not isinstance(entry, dict):
            raise ValueError("phase2a_owner_review_package_entry_invalid")
        file_name = str(entry.get("file_name") or "")
        if (
            entry.get("ordinal") != ordinal
            or entry.get("artifact_id") != expected_id
            or not _SAFE_FILE.fullmatch(file_name)
            or file_name != f"{expected_id}.json"
        ):
            raise ValueError("phase2a_owner_review_package_entry_binding_invalid")
        path = package_root / file_name
        raw = path.read_bytes() if path.is_file() and not path.is_symlink() else b""
        if _sha256(raw) != entry.get("file_sha256") or len(raw) != entry.get("bytes"):
            raise ValueError("phase2a_owner_review_package_artifact_digest_invalid")
        value = json.loads(raw)
        if not isinstance(value, dict):
            raise ValueError("phase2a_owner_review_package_artifact_invalid")
        artifacts[expected_id] = value

    remediation = artifacts["remediation-matrix-585"]
    effects = artifacts["legislative-effects-register-1896"]
    judgments = artifacts["judgment-later-treatment-register-20"]
    provenance = artifacts["official-source-provenance-register"]
    qualification = artifacts["corrected-all585-qualification"]
    invariants = artifacts["final-invariants"]
    if (
        remediation.get("row_count") != 585
        or len(_require_sequence(remediation.get("rows"), code="phase2a_rows_invalid")) != 585
        or effects.get("effect_count") != 1896
        or len(_require_sequence(effects.get("effects"), code="phase2a_effects_invalid")) != 1896
        or judgments.get("record_count") != 20
        or len(_require_sequence(judgments.get("records"), code="phase2a_judgments_invalid")) != 20
        or qualification.get("blocked_material_gap") != 585
        or qualification.get("phase2b_allowed") is not False
        or invariants.get("terminal_verdict") != TERMINAL_VERDICT
        or invariants.get("phase2b_allowed") is not False
        or invariants.get("answer_model_invoked") is not False
        or invariants.get("development_30_generated") is not False
        or invariants.get("promotion_or_live_action") is not False
    ):
        raise ValueError("phase2a_owner_review_fail_closed_invariants_invalid")

    records = _require_sequence(provenance.get("records"), code="phase2a_provenance_invalid")
    return {
        "index": index,
        "index_file_sha256": _sha256(index_path.read_bytes()),
        "artifacts": artifacts,
        "provenance_records": records,
    }


def verify_owner_review_package(package_root: Path) -> dict[str, Any]:
    """Replay every output seal, file digest, count, and phase boundary."""

    if package_root.is_symlink() or not package_root.is_dir():
        raise ValueError("phase2a_owner_review_output_root_invalid")
    index = _load_json(package_root / "OWNER-REVIEW-INDEX.json")
    if not isinstance(index, dict):
        raise ValueError("phase2a_owner_review_index_invalid")
    supplied_digest = str(index.get("package_digest") or "")
    digest_material = dict(index)
    digest_material.pop("package_digest", None)
    if (
        index.get("schema") != REVIEW_INDEX_SCHEMA
        or index.get("route") != OWNER_ROUTE
        or index.get("authorizing") is not False
        or index.get("phase2b_allowed") is not False
        or index.get("development_30_allowed") is not False
        or index.get("answer_model_invoked") is not False
        or index.get("source_admitted") is not False
        or index.get("candidate_mutated") is not False
        or index.get("terminal_verdict") != TERMINAL_VERDICT
        or supplied_digest != _sealed(digest_material)
    ):
        raise ValueError("phase2a_owner_review_index_binding_invalid")

    entries = _require_sequence(
        index.get("entries"), code="phase2a_owner_review_output_entries_invalid"
    )
    if len(entries) != index.get("artifact_count"):
        raise ValueError("phase2a_owner_review_output_inventory_invalid")
    companion = index.get("companion_document")
    if companion is not None:
        if not isinstance(companion, dict):
            raise ValueError("phase2a_owner_review_companion_invalid")
        file_name = str(companion.get("file_name") or "")
        path = package_root / file_name
        raw = path.read_bytes() if path.is_file() and not path.is_symlink() else b""
        if (
            set(companion)
            != {
                "file_name",
                "file_sha256",
                "bytes",
                "role",
                "authorizing",
                "owner_editable",
            }
            or not _SAFE_DOCX.fullmatch(file_name)
            or companion.get("role") != "READABLE_OWNER_REVIEW_COMPANION"
            or companion.get("authorizing") is not False
            or companion.get("owner_editable") is not True
            or _sha256(raw) != companion.get("file_sha256")
            or len(raw) != companion.get("bytes")
        ):
            raise ValueError("phase2a_owner_review_companion_binding_invalid")
    category_counts: dict[str, int] = {
        "issue": 0,
        "legislative_effect": 0,
        "judgment": 0,
        "source_version": 0,
    }
    category_batch_counts = dict.fromkeys(category_counts, 0)
    for ordinal, entry in enumerate(entries, start=1):
        if not isinstance(entry, dict):
            raise ValueError("phase2a_owner_review_output_entry_invalid")
        file_name = str(entry.get("file_name") or "")
        if entry.get("ordinal") != ordinal or not _SAFE_FILE.fullmatch(file_name):
            raise ValueError("phase2a_owner_review_output_entry_binding_invalid")
        path = package_root / file_name
        raw = path.read_bytes() if path.is_file() and not path.is_symlink() else b""
        if _sha256(raw) != entry.get("file_sha256") or len(raw) != entry.get("bytes"):
            raise ValueError("phase2a_owner_review_output_artifact_digest_invalid")
        artifact = json.loads(raw)
        if not isinstance(artifact, dict):
            raise ValueError("phase2a_owner_review_output_artifact_invalid")
        artifact_material = dict(artifact)
        artifact_seal = str(artifact_material.pop("artifact_sha256", ""))
        if artifact_seal != _sealed(artifact_material):
            raise ValueError("phase2a_owner_review_output_artifact_seal_invalid")
        category = str(entry.get("category") or "")
        if category == "control":
            if entry.get("item_count") != 0:
                raise ValueError("phase2a_owner_review_control_item_count_invalid")
            continue
        if category not in category_counts or artifact.get("category") != category:
            raise ValueError("phase2a_owner_review_output_category_invalid")
        items = _require_sequence(
            artifact.get("items"), code="phase2a_owner_review_output_items_invalid"
        )
        if (
            len(items) != artifact.get("item_count")
            or len(items) != entry.get("item_count")
            or artifact.get("owner_decisions_recorded") != 0
            or artifact.get("advisory_ai_reviews_recorded") != 0
            or artifact.get("authoritative") is not False
            or artifact.get("phase2b_authority") is not False
        ):
            raise ValueError("phase2a_owner_review_output_batch_boundary_invalid")
        for item in items:
            if not isinstance(item, dict):
                raise ValueError("phase2a_owner_review_output_item_invalid")
            item_material = dict(item)
            item_seal = str(item_material.pop("item_sha256", ""))
            if (
                item_seal != _sealed(item_material)
                or item.get("category") != category
                or item.get("owner_decision_required") is not True
                or item.get("advisory_ai_review") is not None
                or item.get("owner_decision") is not None
                or item.get("authoritative") is not False
                or item.get("phase2b_authority") is not False
            ):
                raise ValueError("phase2a_owner_review_output_item_boundary_invalid")
        category_counts[category] += len(items)
        category_batch_counts[category] += 1
    if (
        category_counts != index.get("category_item_counts")
        or category_batch_counts != index.get("category_batch_counts")
        or index.get("owner_decisions_recorded") != 0
        or index.get("advisory_ai_reviews_recorded") != 0
    ):
        raise ValueError("phase2a_owner_review_output_counts_invalid")
    return index


def _review_item(
    *, category: str, item_id: str, source_record: Mapping[str, Any], evidence: Mapping[str, Any]
) -> dict[str, Any]:
    if category not in ALLOWED_OWNER_OUTCOMES:
        raise ValueError("phase2a_owner_review_category_invalid")
    source_record_sha256 = str(
        source_record.get("row_evidence_sha256")
        or source_record.get("record_sha256")
        or source_record.get("sha256")
        or _sealed(source_record)
    )
    if not _SHA256.fullmatch(source_record_sha256):
        source_record_sha256 = _sealed(source_record)
    item: dict[str, Any] = {
        "schema": "legalbot.v111-phase2a-owner-review-item.v1",
        "category": category,
        "item_id": item_id,
        "source_record_sha256": source_record_sha256,
        "deterministic_evidence": dict(evidence),
        "owner_decision_required": True,
        "advisory_ai_review": None,
        "owner_decision": None,
        "allowed_owner_outcomes": sorted(ALLOWED_OWNER_OUTCOMES[category]),
        "authoritative": False,
        "phase2b_authority": False,
    }
    item["item_sha256"] = _sealed(item)
    return item


def _issue_items(rows: Sequence[Mapping[str, Any]]) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    for row in rows:
        row_id = str(row.get("row_id") or "")
        if not row_id:
            raise ValueError("phase2a_owner_review_issue_identity_missing")
        items.append(
            _review_item(
                category="issue",
                item_id=row_id,
                source_record=row,
                evidence={
                    "ordinal": row.get("ordinal"),
                    "case_id": row.get("case_id"),
                    "issue_id": row.get("issue_id"),
                    "issue_label": row.get("issue_label"),
                    "issue_label_sha256": row.get("issue_label_sha256"),
                    "legal_domain": row.get("legal_domain"),
                    "baseline_primary_status": row.get("baseline_primary_status"),
                    "baseline_official_finding_ids": row.get("baseline_official_finding_ids", []),
                    "determined_defects": row.get("determined_defects", []),
                    "candidate_evidence_candidates": row.get("candidate_evidence_candidates", []),
                    "remediation_result": row.get("remediation_result"),
                    "technical_status": row.get("technical_status"),
                },
            )
        )
    return items


def _effect_items(rows: Sequence[Mapping[str, Any]]) -> list[dict[str, Any]]:
    return [
        _review_item(
            category="legislative_effect",
            item_id=f"effect-{int(row.get('ordinal') or 0):04d}",
            source_record=row,
            evidence=row,
        )
        for row in rows
    ]


def _judgment_items(rows: Sequence[Mapping[str, Any]]) -> list[dict[str, Any]]:
    return [
        _review_item(
            category="judgment",
            item_id=f"judgment-{int(row.get('ordinal') or 0):02d}",
            source_record=row,
            evidence=row,
        )
        for row in rows
    ]


def _source_version_items(records: Sequence[Mapping[str, Any]]) -> list[dict[str, Any]]:
    selected = [
        row
        for row in records
        if (
            row.get("target_type") == "candidate_legislation"
            and row.get("result") == "DOWNLOADED_QUARANTINED"
            and row.get("matches_expected_version_sha256") is False
        )
        or row.get("result") == "OFFICIAL_SOURCE_UNAVAILABLE"
    ]
    if len(selected) != 68:
        raise ValueError("phase2a_owner_review_source_version_inventory_invalid")
    return [
        _review_item(
            category="source_version",
            item_id=f"source-version-review-{ordinal:03d}",
            source_record=row,
            evidence=row,
        )
        for ordinal, row in enumerate(selected, start=1)
    ]


def _chunks(values: Sequence[dict[str, Any]], size: int) -> list[list[dict[str, Any]]]:
    return [list(values[start : start + size]) for start in range(0, len(values), size)]


def _decision_field_contracts() -> dict[str, Any]:
    return {
        "common_required_fields": [
            "schema",
            "category",
            "item_id",
            "item_sha256",
            "owner_typed_name",
            "owner_decision_date",
            "owner_outcome",
            "owner_rationale",
            "decision_basis_sha256s",
            "advisory_ai_disposition",
            "advisory_ai_review_sha256",
            "findings",
        ],
        "outcome_required_findings": {
            "APPROVE_PROPOSITION_BINDINGS": [
                "proposition",
                "official_source_version_ids",
                "exact_span_binding_sha256s",
                "currentness_conclusion",
                "candidate_change_required",
            ],
            "APPROVE_EFFECT_DISPOSITION": [
                "effect_disposition",
                "commencement_conclusion",
                "extent_conclusion",
                "transition_saving_conclusion",
                "proposition_materiality",
            ],
            "APPROVE_JUDGMENT_TREATMENT": [
                "proposition_relied_upon",
                "binding_or_persuasive_status",
                "later_treatment_conclusion",
                "gold_proposition_current",
            ],
            "APPROVE_SOURCE_VERSION_MATERIALITY": [
                "substantive_change",
                "affected_proposition_row_ids",
                "source_admission_required",
            ],
            "APPROVE_NONMATERIAL_NOTE": ["nonmaterial_note"],
            "CONFIRM_MATERIAL_GAP": ["gap_description", "required_next_evidence"],
            "REQUEST_MORE_EVIDENCE": ["evidence_request"],
        },
    }


def validate_advisory_ai_review(value: Mapping[str, Any]) -> dict[str, Any]:
    """Require complete provenance and keep the AI record non-authoritative."""

    allowed_keys = {
        "schema",
        "status",
        "item_sha256",
        "model_id",
        "model_version",
        "model_artifact_sha256",
        "prompt_sha256",
        "configuration_sha256",
        "toolchain_sha256",
        "reviewer_execution_mode",
        "model_independent",
        "pinned",
        "logged",
        "stateless",
        "official_sources_checked_first",
        "recommendation",
        "concise_findings",
        "evidence_reference_sha256s",
        "unavailable_reason",
        "authoritative",
        "can_decide_or_adopt",
        "can_admit_sources",
        "can_authorize_gates",
        "may_raise_fail_closed_owner_review_hold",
        "owner_decision_applied",
        "hidden_reasoning_persisted",
        "review_sha256",
    }
    if set(value) - allowed_keys:
        raise ValueError("phase2a_advisory_ai_review_unknown_fields")
    material = dict(value)
    supplied_seal = str(material.pop("review_sha256", ""))
    if (
        material.get("schema") != ADVISORY_AI_SCHEMA
        or material.get("status") not in {"AVAILABLE", "UNAVAILABLE"}
        or material.get("authoritative") is not False
        or material.get("can_decide_or_adopt") is not False
        or material.get("can_admit_sources") is not False
        or material.get("can_authorize_gates") is not False
        or material.get("may_raise_fail_closed_owner_review_hold") is not True
        or material.get("owner_decision_applied") is not False
        or material.get("hidden_reasoning_persisted") is not False
        or not _SHA256.fullmatch(str(material.get("item_sha256") or ""))
    ):
        raise ValueError("phase2a_advisory_ai_review_boundary_invalid")
    if material["status"] == "AVAILABLE":
        if (
            material.get("reviewer_execution_mode")
            != "separate_verification_pass_same_model_adapter"
            or material.get("model_independent") is not False
        ):
            raise ValueError("phase2a_advisory_ai_review_execution_mode_invalid")
        for field in (
            "model_id",
            "model_version",
            "model_artifact_sha256",
            "prompt_sha256",
            "configuration_sha256",
            "toolchain_sha256",
            "recommendation",
        ):
            if not material.get(field):
                raise ValueError("phase2a_advisory_ai_review_pin_incomplete")
        for field in (
            "model_artifact_sha256",
            "prompt_sha256",
            "configuration_sha256",
            "toolchain_sha256",
        ):
            if not _SHA256.fullmatch(str(material[field])):
                raise ValueError("phase2a_advisory_ai_review_pin_invalid")
        if any(
            material.get(field) is not True
            for field in (
                "pinned",
                "logged",
                "stateless",
                "official_sources_checked_first",
            )
        ):
            raise ValueError("phase2a_advisory_ai_review_controls_invalid")
    elif not str(material.get("unavailable_reason") or "").strip():
        raise ValueError("phase2a_advisory_ai_review_unavailable_reason_missing")
    if supplied_seal != _sealed(material):
        raise ValueError("phase2a_advisory_ai_review_seal_invalid")
    return {**material, "review_sha256": supplied_seal}


def validate_owner_decision(
    *,
    item: Mapping[str, Any],
    decision: Mapping[str, Any],
    advisory_ai_review: Mapping[str, Any] | None,
) -> dict[str, Any]:
    """Validate one explicit owner decision without elevating AI authority."""

    contracts = _decision_field_contracts()
    required = set(contracts["common_required_fields"])
    if set(decision) != required:
        raise ValueError("phase2a_owner_decision_fields_invalid")
    category = str(item.get("category") or "")
    outcome = str(decision.get("owner_outcome") or "")
    if (
        decision.get("schema") != DECISION_SCHEMA
        or decision.get("category") != category
        or decision.get("item_id") != item.get("item_id")
        or decision.get("item_sha256") != item.get("item_sha256")
        or outcome not in ALLOWED_OWNER_OUTCOMES.get(category, frozenset())
    ):
        raise ValueError("phase2a_owner_decision_binding_invalid")
    owner_name = str(decision.get("owner_typed_name") or "").strip()
    if len(owner_name) < 2 or "[" in owner_name or "]" in owner_name:
        raise ValueError("phase2a_owner_decision_identity_missing")
    try:
        date.fromisoformat(str(decision.get("owner_decision_date") or ""))
    except ValueError as exc:
        raise ValueError("phase2a_owner_decision_date_invalid") from exc
    if not str(decision.get("owner_rationale") or "").strip():
        raise ValueError("phase2a_owner_decision_rationale_missing")
    basis = decision.get("decision_basis_sha256s")
    if (
        not isinstance(basis, list)
        or not basis
        or any(not _SHA256.fullmatch(str(value)) for value in basis)
    ):
        raise ValueError("phase2a_owner_decision_basis_invalid")
    findings = decision.get("findings")
    if not isinstance(findings, dict):
        raise ValueError("phase2a_owner_decision_findings_invalid")
    required_findings = set(contracts["outcome_required_findings"][outcome])
    if set(findings) != required_findings:
        raise ValueError("phase2a_owner_decision_findings_incomplete")
    if any(value is None or value == "" or value == [] for value in findings.values()):
        raise ValueError("phase2a_owner_decision_findings_blank")

    ai_disposition = decision.get("advisory_ai_disposition")
    ai_sha = decision.get("advisory_ai_review_sha256")
    if ai_disposition == "USED":
        if advisory_ai_review is None:
            raise ValueError("phase2a_owner_decision_ai_review_missing")
        verified_ai = validate_advisory_ai_review(advisory_ai_review)
        if (
            verified_ai["status"] != "AVAILABLE"
            or verified_ai["item_sha256"] != item.get("item_sha256")
            or ai_sha != verified_ai["review_sha256"]
        ):
            raise ValueError("phase2a_owner_decision_ai_binding_invalid")
    elif ai_disposition == "UNAVAILABLE_OWNER_PROCEEDED_WITH_DETERMINISTIC_EVIDENCE":
        if advisory_ai_review is None:
            raise ValueError("phase2a_owner_decision_ai_unavailability_record_missing")
        verified_ai = validate_advisory_ai_review(advisory_ai_review)
        if (
            verified_ai["status"] != "UNAVAILABLE"
            or verified_ai["item_sha256"] != item.get("item_sha256")
            or ai_sha != verified_ai["review_sha256"]
        ):
            raise ValueError("phase2a_owner_decision_ai_unavailability_binding_invalid")
    else:
        raise ValueError("phase2a_owner_decision_ai_disposition_invalid")
    return dict(decision)


def build(
    *,
    remediation_package_root: Path,
    output_root: Path,
    run_id: str,
    recorded_at: datetime,
    batch_size: int,
) -> dict[str, Any]:
    """Create a sealed, non-authorizing review package for explicit owner decisions."""

    if not _SAFE_RUN_ID.fullmatch(run_id):
        raise ValueError("phase2a_owner_review_run_id_invalid")
    if recorded_at.tzinfo is None:
        raise ValueError("phase2a_owner_review_recorded_at_must_be_timezone_aware")
    if batch_size < 1 or batch_size > 100:
        raise ValueError("phase2a_owner_review_batch_size_invalid")
    source = verify_remediation_package(remediation_package_root)
    if output_root.exists():
        raise ValueError("phase2a_owner_review_output_already_exists")
    output_root.mkdir(parents=True, mode=0o700)
    if stat.S_IMODE(output_root.stat().st_mode) != 0o700:
        raise ValueError("phase2a_owner_review_output_mode_invalid")

    artifacts = source["artifacts"]
    categories = {
        "issue": _issue_items(artifacts["remediation-matrix-585"]["rows"]),
        "legislative_effect": _effect_items(
            artifacts["legislative-effects-register-1896"]["effects"]
        ),
        "judgment": _judgment_items(artifacts["judgment-later-treatment-register-20"]["records"]),
        "source_version": _source_version_items(source["provenance_records"]),
    }
    if {key: len(value) for key, value in categories.items()} != {
        "issue": 585,
        "legislative_effect": 1896,
        "judgment": 20,
        "source_version": 68,
    }:
        raise ValueError("phase2a_owner_review_category_counts_invalid")

    policy = _artifact(
        REVIEW_POLICY_SCHEMA,
        {
            "route": OWNER_ROUTE,
            "professional_legal_certification": False,
            "private_local_owner_only": True,
            "owner_is_substantive_decision_maker": True,
            "owner_typed_name_and_date_required_per_decision": True,
            "official_primary_sources_and_deterministic_checks_first": True,
            "advisory_ai_allowed": True,
            "advisory_ai_must_be_pinned_logged_stateless": True,
            "advisory_ai_may_propose_findings": True,
            "advisory_ai_can_decide_or_adopt": False,
            "advisory_ai_can_admit_source_or_rebuild": False,
            "advisory_ai_can_authorize_gate": False,
            "advisory_ai_may_raise_fail_closed_owner_review_hold": True,
            "current_ai_reviewer_execution_mode": (
                "separate_verification_pass_same_model_adapter"
            ),
            "current_ai_reviewer_is_model_independent": False,
            "answer_generation_allowed": False,
            "automatic_source_admission": False,
            "automatic_indexing_or_embedding": False,
            "phase2b_allowed": False,
            "development_30_allowed": False,
            "promotion_validation_or_live_allowed": False,
            "stop_after_phase2a": True,
        },
    )
    decision_contract = _artifact(
        "legalbot.v111-phase2a-owner-decision-contract.v1",
        {
            "owner_decision_schema": DECISION_SCHEMA,
            "advisory_ai_schema": ADVISORY_AI_SCHEMA,
            "allowed_owner_outcomes": {
                key: sorted(value) for key, value in ALLOWED_OWNER_OUTCOMES.items()
            },
            **_decision_field_contracts(),
            "ai_output_never_substitutes_for_owner_record": True,
            "silence_never_counts_as_owner_approval": True,
            "phase2b_authority": False,
        },
    )

    entries: list[dict[str, Any]] = []
    for artifact_id, artifact in (
        ("owner-review-policy", policy),
        ("owner-decision-contract", decision_contract),
    ):
        file_name = f"{artifact_id}.json"
        raw = _canonical_json(artifact)
        _write_exclusive(output_root / file_name, raw)
        entries.append(
            {
                "ordinal": len(entries) + 1,
                "artifact_id": artifact_id,
                "category": "control",
                "file_name": file_name,
                "file_sha256": _sha256(raw),
                "bytes": len(raw),
                "item_count": 0,
            }
        )

    category_batch_counts: dict[str, int] = {}
    for category, items in categories.items():
        batches = _chunks(items, batch_size)
        category_batch_counts[category] = len(batches)
        for batch_number, batch_items in enumerate(batches, start=1):
            artifact_id = f"{category.replace('_', '-')}-batch-{batch_number:03d}"
            batch = _artifact(
                "legalbot.v111-phase2a-owner-review-batch.v1",
                {
                    "run_id": run_id,
                    "route": OWNER_ROUTE,
                    "category": category,
                    "batch_number": batch_number,
                    "batch_count": len(batches),
                    "item_count": len(batch_items),
                    "items": batch_items,
                    "owner_decisions_recorded": 0,
                    "advisory_ai_reviews_recorded": 0,
                    "authoritative": False,
                    "phase2b_authority": False,
                },
            )
            file_name = f"{artifact_id}.json"
            raw = _canonical_json(batch)
            _write_exclusive(output_root / file_name, raw)
            entries.append(
                {
                    "ordinal": len(entries) + 1,
                    "artifact_id": artifact_id,
                    "category": category,
                    "file_name": file_name,
                    "file_sha256": _sha256(raw),
                    "bytes": len(raw),
                    "item_count": len(batch_items),
                }
            )

    companion_file_name = "LegalBot-Phase2A-Owner-Review-Control.docx"
    companion_raw = _owner_review_companion_docx(
        run_id=run_id,
        recorded_at=recorded_at,
        source_package_digest=str(source["index"]["package_digest"]),
        category_counts={key: len(value) for key, value in categories.items()},
        category_batch_counts=category_batch_counts,
    )
    _write_exclusive(output_root / companion_file_name, companion_raw)
    companion_document = {
        "file_name": companion_file_name,
        "file_sha256": _sha256(companion_raw),
        "bytes": len(companion_raw),
        "role": "READABLE_OWNER_REVIEW_COMPANION",
        "authorizing": False,
        "owner_editable": True,
    }

    index: dict[str, Any] = {
        "schema": REVIEW_INDEX_SCHEMA,
        "run_id": run_id,
        "recorded_at": recorded_at.astimezone(UTC).isoformat(timespec="seconds"),
        "route": OWNER_ROUTE,
        "source_remediation_package_digest": source["index"]["package_digest"],
        "source_package_index_sha256": source["index_file_sha256"],
        "category_item_counts": {key: len(value) for key, value in categories.items()},
        "category_batch_counts": category_batch_counts,
        "batch_size": batch_size,
        "artifact_count": len(entries),
        "entries": entries,
        "companion_document": companion_document,
        "owner_decisions_recorded": 0,
        "advisory_ai_reviews_recorded": 0,
        "status": "AWAITING_PINNED_ADVISORY_AI_AND_EXPLICIT_OWNER_DECISIONS",
        "authorizing": False,
        "phase2b_allowed": False,
        "development_30_allowed": False,
        "answer_model_invoked": False,
        "source_admitted": False,
        "candidate_mutated": False,
        "terminal_verdict": TERMINAL_VERDICT,
    }
    index["package_digest"] = _sealed(index)
    _write_exclusive(output_root / "OWNER-REVIEW-INDEX.json", _canonical_json(index))
    _write_exclusive(
        output_root / "OUTCOME.txt",
        (
            "PHASE 2A OWNER REVIEW PREPARED — EXPLICIT OWNER DECISIONS REQUIRED; "
            "PHASE 2B AND DEVELOPMENT 30 NOT AUTHORIZED\n"
        ).encode(),
    )
    return index


def _parse_datetime(value: str) -> datetime:
    parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    if parsed.tzinfo is None:
        raise argparse.ArgumentTypeError("recorded-at must include a UTC offset")
    return parsed


def _parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--remediation-package-root", type=Path, required=True)
    parser.add_argument("--output-root", type=Path, required=True)
    parser.add_argument("--run-id", required=True)
    parser.add_argument("--recorded-at", type=_parse_datetime, required=True)
    parser.add_argument("--batch-size", type=int, default=25)
    return parser


def _safe_error_code(exc: BaseException) -> str:
    candidate = re.sub(r"(?<!^)(?=[A-Z])", "_", type(exc).__name__).casefold()
    if isinstance(exc, ValueError) and exc.args:
        explicit = str(exc.args[0]).strip().casefold().replace("-", "_")
        if re.fullmatch(r"[a-z0-9][a-z0-9._:]{0,127}", explicit):
            return explicit
    return candidate if re.fullmatch(r"[a-z0-9][a-z0-9._:]{0,127}", candidate) else "failure"


def _persist_failure(output_root: Path, *, run_id: str, exc: BaseException) -> None:
    try:
        if not output_root.exists():
            output_root.mkdir(parents=True, mode=0o700)
        if not output_root.is_dir() or output_root.is_symlink():
            return
        code = _safe_error_code(exc)
        completed = sorted(path.name for path in output_root.iterdir() if path.is_file())
        report: dict[str, Any] = {
            "schema": "legalbot.v111-phase2a-owner-review-failure.v1",
            "run_id": run_id,
            "stage": "OWNER_REVIEW_BATCH_BUILD",
            "error_code": code,
            "failure_fingerprint": _sealed(
                {"stage": "OWNER_REVIEW_BATCH_BUILD", "error_code": code}
            ),
            "completed_files": completed,
            "root_cause_status": "REQUIRES_EXECUTION_PLAN_CHANGE_OR_INPUT_CORRECTION",
            "automatic_retry_attempted": False,
            "phase2b_allowed": False,
            "development_30_allowed": False,
        }
        report["report_sha256"] = _sealed(report)
        _write_exclusive(output_root / "FAILURE.json", _canonical_json(report))
    except Exception:
        return


def main(argv: Sequence[str] | None = None) -> int:
    args = _parser().parse_args(argv)
    output_root = args.output_root.resolve()
    try:
        index = build(
            remediation_package_root=args.remediation_package_root.resolve(strict=True),
            output_root=output_root,
            run_id=str(args.run_id),
            recorded_at=args.recorded_at,
            batch_size=int(args.batch_size),
        )
    except Exception as exc:
        _persist_failure(output_root, run_id=str(args.run_id), exc=exc)
        print(
            json.dumps(
                {
                    "status": "safely_stopped",
                    "error_code": _safe_error_code(exc),
                    "phase2b_allowed": False,
                    "development_30_allowed": False,
                },
                sort_keys=True,
            )
        )
        return 2
    print(
        json.dumps(
            {
                "status": "owner_review_prepared",
                "package_digest": index["package_digest"],
                "category_item_counts": index["category_item_counts"],
                "owner_decisions_recorded": 0,
                "phase2b_allowed": False,
                "development_30_allowed": False,
            },
            sort_keys=True,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
