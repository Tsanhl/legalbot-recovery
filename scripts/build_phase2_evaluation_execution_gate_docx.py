from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "output" / "docx" / "LegalBot-Phase2-Evaluation-Execution-Approval.docx"

INK = "202124"
BLACK = "000000"
NAVY = "17365D"
BLUE = "2E74B5"
MUTED = "5F6368"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E8F3E8"
PALE_AMBER = "FFF4CE"
PALE_RED = "FCE8E6"
BORDER = "C9D2DD"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def _prevent_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def _cell_margins(
    cell, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _cell_borders(cell, *, color: str = BORDER, size: int = 5) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError("table widths must total 9360 DXA")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for index, width in enumerate(widths_dxa):
        table._tbl.tblGrid.gridCol_lst[index].set(qn("w:w"), str(width))
    for row in table.rows:
        _prevent_split(row)
        for index, width in enumerate(widths_dxa):
            cell = row.cells[index]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _cell_margins(cell)
            _cell_borders(cell)


def _set_font(run, *, size: float = 11, bold: bool = False, color: str = INK) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _page_field(paragraph) -> None:
    run = paragraph.add_run()
    for kind, _value in (("begin", None), ("separate", None), ("end", None)):
        if kind == "separate":
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = " PAGE "
            run._r.append(instruction)
        node = OxmlElement("w:fldChar")
        node.set(qn("w:fldCharType"), kind)
        run._r.append(node)


def _heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    _set_font(
        paragraph.add_run(text),
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color=BLUE if level < 3 else NAVY,
    )
    return paragraph


def _body(document: Document, text: str = "", *, lead: str | None = None):
    paragraph = document.add_paragraph(style="Normal")
    if lead and text.startswith(lead):
        _set_font(paragraph.add_run(lead), bold=True)
        _set_font(paragraph.add_run(text[len(lead) :]))
    else:
        _set_font(paragraph.add_run(text))
    return paragraph


def _bullet(document: Document, text: str, *, numbered: bool = False):
    paragraph = document.add_paragraph(style="List Number" if numbered else "List Bullet")
    _set_font(paragraph.add_run(text), size=10.5)
    return paragraph


def _callout(document: Document, title: str, text: str, *, fill: str = PALE_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    _table_geometry(table, [CONTENT_DXA])
    _set_repeat_header(table.rows[0])
    cell = table.cell(0, 0)
    _shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    _set_font(paragraph.add_run(title), size=11, bold=True, color=NAVY)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    _set_font(paragraph.add_run(text), size=10.5)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def _metadata(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    _set_font(paragraph.add_run(f"{label}: "), bold=True, color=BLACK)
    _set_font(paragraph.add_run(value), color=BLACK)


def _form_line(document: Document, label: str, hint: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    _set_font(paragraph.add_run(f"{label}: "), bold=True)
    _set_font(paragraph.add_run("_" * 62), color=MUTED)
    note = document.add_paragraph()
    note.paragraph_format.left_indent = Inches(0.25)
    note.paragraph_format.space_after = Pt(6)
    _set_font(note.add_run(hint), size=9, color=MUTED)


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    for level, before, after, size, color in (
        (1, 12, 6, 16, BLUE),
        (2, 10, 5, 13, BLUE),
        (3, 8, 4, 12, NAVY),
    ):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.paragraph_format.space_after = Pt(0)
    header_paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    _set_font(
        header_paragraph.add_run("LEGALBOT v1.11  |  PHASE 2"), size=8.5, bold=True, color=MUTED
    )
    _set_font(
        header_paragraph.add_run("\tOWNER EVALUATION EXECUTION GATE"),
        size=8.5,
        bold=True,
        color=MUTED,
    )

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(footer_paragraph.add_run("Private owner review  |  Page "), size=8.5, color=MUTED)
    _page_field(footer_paragraph)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    _set_font(
        title.add_run("Phase 2 evaluation execution approval"), size=24, bold=True, color=BLACK
    )
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_font(
        subtitle.add_run(
            "Factual-first General Enquiry baseline, controlled knowledge-gap repair loop and protected unseen"
        ),
        size=13,
        color=MUTED,
    )

    _metadata(document, "Owner", "Agnes")
    _metadata(document, "Date prepared", "1 September 2026")
    _metadata(
        document,
        "Current status",
        "GE loop controls implemented; real execution remains blocked by exact owner-gated inputs",
    )
    _metadata(
        document,
        "Scope",
        "331 fixed visible GE cases + 32 separate system scenarios + separate visible diagnostics",
    )
    _metadata(
        document,
        "Excluded",
        "Unseen disclosure, evaluation-data training, promotion, live activation and premature GitHub publication",
    )

    _callout(
        document,
        "No further design approval is requested",
        "The system design and the three Phase-2 preparation recommendations are already accepted. "
        "The GE improvement loop requested by the owner is now part of that living design. This document "
        "records the completed controls and asks only for the missing exact inputs and later candidate-bound "
        "authority needed to run the real visible evaluation.",
        fill=PALE_GREEN,
    )

    _heading(document, "1. What has been completed", 1)
    for text in (
        "The accepted visible set is exactly 331 cases: 306 core and 25 stress. All case IDs and ordinals are unique; no visible case is missing or deleted.",
        "All 331 cases remain visible-development material only. They are ineligible for unseen custody and training export. The 306 unseen drafts remain unopened by the Phase-2 harness.",
        "The factual/legal hard gate runs before quality. A factual hold receives no quality score.",
        "The General Enquiry quality gate totals 100 points, requires at least 70, and also enforces critical floors for accuracy, authority/currentness and practical help.",
        "Selected-schema validation, canonical digests, strict deletion authorization, encrypted immutable chain persistence and atomic chain-to-outbox binding are implemented technical foundations; runner/live activation remains disabled.",
        "The visible harness reconciles and encrypts all 331 selected case results, preserves every terminal outcome, and requires an externally pinned ten-capability execution admission before any run.",
        "The 32 system scenarios and every later diagnostic supplement have separate immutable identities, denominators and encrypted result records. They never change the fixed 331 score denominator.",
        "Every cycle replays an opaque verifier-issued owner authorization for one exact ordered GE coverage topology. Its mandatory floor is 23 distinct domains: the 17 current topics plus separate housing, employment, family, immigration, benefits/debt and consumer domains. Empty public-domain assignments remain missing cells and require separate diagnostics.",
        "A proved knowledge gap may create a hashed research intent for a registered official source. Downloaded bytes remain quarantined and create-only; source rights, identity, currentness and qualified legal review must pass before any new chunks can be admitted.",
        "Approved gap material creates a versioned, non-ACTIVE successor index. It cannot mutate the evaluated index, become ACTIVE, enter unseen custody or become training data through the GE loop.",
        "After any answer, retrieval, prompt or knowledge change, the complete 331 visible set and all 32 system scenarios must be rerun. New missing-area questions remain a separate visible diagnostic supplement and are also rerun.",
        "The same stable failure fingerprint stops after its second failed repaired run. Earlier cycles, answers, diagnoses, downloads and prose remain attributable; no deletion is authorized.",
        "Future normal-live WebSocket completion is bound to the actual persisted VerifiedRelease digest and selected terminal identity; missing bindings fail closed.",
        "A fresh catalogue backup and isolated restore drill passed. The retained non-ACTIVE retrieval build completed its bounded second attempt: 85 sources and 149,855 vectors. Its sealed tree, full ordered checkpoint and exact DB-to-Lance content, source, dimension and lane parity were independently replayed with zero mismatches; no ACTIVE/PREVIOUS pointer exists.",
    ):
        _bullet(document, text)

    document.add_page_break()
    _heading(document, "2. Current gate status", 1)
    status_table = document.add_table(rows=1, cols=3)
    _table_geometry(status_table, [2880, 1800, 4680])
    for index, label in enumerate(("Gate", "Status", "Meaning")):
        _shade(status_table.cell(0, index), NAVY)
        _set_font(
            status_table.cell(0, index).paragraphs[0].add_run(label), size=9, bold=True, color=WHITE
        )
    _set_repeat_header(status_table.rows[0])
    rows = (
        (
            "Visible set and worksheet",
            "READY",
            "331/331 bound; factual and quality policies sealed",
        ),
        (
            "Selected run integrity",
            "READY",
            "331 reconciliation, atomic chain binding and external admission implemented",
        ),
        (
            "GE cycle and missing-area controls",
            "READY",
            "Separate 32/system and diagnostic records; full rerun and zero-missing-area exit enforced",
        ),
        (
            "Official-source intake controls",
            "READY",
            "Exact proved-gap provenance and strict non-ACTIVE successor expansion implemented",
        ),
        (
            "Retrieval candidate",
            "HELD",
            "149,855-vector sealed non-ACTIVE build; exact tree and row parity verified",
        ),
        (
            "Answer model + transport",
            "BLOCKED",
            "Exact artifact and private UDS capability not supplied",
        ),
        ("Legal gold/currentness", "BLOCKED", "Qualified case-level decisions not supplied"),
        ("Development review root", "BLOCKED", "Distinct private owner root not supplied"),
        (
            "Ten-capability evaluation admission",
            "BLOCKED",
            "Cannot bind until candidate, model, custody, resources and exact execution inputs pass",
        ),
        (
            "Unseen / evaluation-data training / live / GitHub",
            "CLOSED",
            "No action before GE_COMPLETE_OWNER_ACCEPTED and its applicable owner gate",
        ),
    )
    for gate, state, meaning in rows:
        cells = status_table.add_row().cells
        for index, value in enumerate((gate, state, meaning)):
            _set_font(cells[index].paragraphs[0].add_run(value), size=9)
        _shade(
            cells[1],
            PALE_GREEN
            if state == "READY"
            else PALE_AMBER
            if state in {"RUNNING", "HELD"}
            else PALE_RED,
        )
    _table_geometry(status_table, [2880, 1800, 4680])

    document.add_page_break()
    _heading(document, "3. Exact visible-evaluation identity", 1)
    identity_table = document.add_table(rows=1, cols=2)
    for index, label in enumerate(("Identity field", "Bound value")):
        _shade(identity_table.cell(0, index), NAVY)
        _set_font(
            identity_table.cell(0, index).paragraphs[0].add_run(label),
            size=9,
            bold=True,
            color=WHITE,
        )
    _set_repeat_header(identity_table.rows[0])
    identity_rows = (
        ("Content version", "GE-visible-r3"),
        ("Visible denominator", "331 = 306 core + 25 stress"),
        ("Separate scenarios", "32 system scenarios outside legal-quality denominator"),
        (
            "Case manifest SHA-256",
            "dae4592f4c355d0b84ec80e29689185e396603114b4b4588d08ecc511991d080",
        ),
        ("Case order SHA-256", "ce6d64c5ddcee2a73a139ea40b048c36a9ff895b7a72ee2e4f15938402e450c4"),
        (
            "System manifest SHA-256",
            "4f9094a4157d6dd9c8e5728cbc38a6c326e995b7f9064550c0df6ec2e5266ce1",
        ),
        (
            "System order SHA-256",
            "0a7779dc046fc350d324210dc10a8fad1711e10762cfd07e2047c8bfb442b165",
        ),
        (
            "Factual policy SHA-256",
            "84bbc2b5525eb0aeedb903b5cee9990f2b8af6d2e6e71b249ba1bce77ea971ee",
        ),
        (
            "Quality policy SHA-256",
            "6237385f2244153847d10d339c2f9ea978bba24cc942b85f1ae320b45ae1db25",
        ),
        (
            "Schema selection SHA-256",
            "ca93229c3a446801424024c3074509ef589597aa3afc36f74bceade865f429d7",
        ),
        ("Coverage breadth", "23 domains = 17 current topics + 6 separate public-access domains"),
        (
            "Held retrieval build",
            "current-law-ew-full-fp16-v111-20260829-recovery-b",
        ),
        (
            "Held build seal SHA-256",
            "356475aab2b01845f5753e4b183642268a108a215dcdd087f6508c8e0560d217",
        ),
        (
            "Held source-manifest SHA-256",
            "1ab9e139e2d97e2f4b935fb8619a46c98ee257f855ce8f9a99ec309905f7623b",
        ),
        (
            "Candidate-bound execution preflight",
            "NOT ISSUED — regenerate only after all ten exact capability inputs pass",
        ),
    )
    for label, value in identity_rows:
        cells = identity_table.add_row().cells
        _shade(cells[0], LIGHT)
        _set_font(cells[0].paragraphs[0].add_run(label), size=9, bold=True)
        _set_font(cells[1].paragraphs[0].add_run(value), size=8.5)
    _table_geometry(identity_table, [2700, 6660])

    _heading(document, "4. Factual-first decision rule", 1)
    _body(
        document,
        "Every case first receives FACTUAL_PASS, FACTUAL_HOLD or SYSTEM_ERROR. "
        "FACTUAL_PASS requires all applicable identity, evidence, user-fact, jurisdiction, "
        "date/currentness, amount/deadline, citation, contradiction, safety and privacy checks. "
        "Any material failure stops quality scoring.",
    )
    _callout(
        document,
        "Legal review remains human-authorized",
        "Automated checks can prove identity, traceability and deterministic support. They cannot "
        "supply the qualified legal-currentness or gold decision required for a material legal claim.",
        fill=PALE_AMBER,
    )

    document.add_page_break()
    _heading(document, "5. Practical General Enquiry quality gate", 1)
    quality_table = document.add_table(rows=1, cols=2)
    _table_geometry(quality_table, [6660, 2700])
    for index, label in enumerate(("Dimension", "Maximum")):
        _shade(quality_table.cell(0, index), NAVY)
        _set_font(
            quality_table.cell(0, index).paragraphs[0].add_run(label),
            size=9,
            bold=True,
            color=WHITE,
        )
    _set_repeat_header(quality_table.rows[0])
    for dimension, points in (
        ("Legal and factual accuracy", "25"),
        ("Issue coverage and reasoning", "15"),
        ("Authority and currentness", "15"),
        ("Practical steps and urgency", "15"),
        ("Uncertainty, limits and clarification", "10"),
        ("Organisation and plain language", "10"),
        ("Traceability and citations", "10"),
    ):
        cells = quality_table.add_row().cells
        _set_font(cells[0].paragraphs[0].add_run(dimension), size=9)
        _set_font(cells[1].paragraphs[0].add_run(points), size=9, bold=True)
    _table_geometry(quality_table, [6660, 2700])
    _body(
        document,
        "Pass rule: at least 70/100 and all critical floors. A nominal 70 cannot pass if accuracy, "
        "authority/currentness or practical help is materially weak.",
        lead="Pass rule:",
    )

    document.add_page_break()
    _heading(document, "6. Accepted GE improvement loop", 1)
    for text in (
        "Run the fixed 331 visible cases and the separate 32 system scenarios only under one exact, verified execution admission.",
        "Before coverage audit or closure, replay the exact stored owner request and resolution and issue an opaque authorization over all 23 required breadth domains and their ordered cells. Reject arbitrary hashes, narrow manifests, omitted or aliased public domains, substitutions, reordering and stale decisions.",
        "Apply the factual/legal gate first. The selected v2 result seals every ordered factual check and quality dimension; the aggregate outcome is recomputed from that evidence. The older v1 result remains audit-only. Quality scoring is permitted only for FACTUAL_PASS results; every hold and system error remains in the denominator evidence.",
        "Diagnose each failure against exact answer, retrieval and EvidenceSpan identities. A knowledge gap is proved only when the sealed retrieval attempt and available approved sources cannot support the required issue.",
        "For a proved gap, create a hashed official-research task. Keep raw user questions out of task records and prompts. Discovery is limited to an explicitly registered official source and unsupported dispatchers fail closed.",
        "Keep downloaded material in quarantine. Verify byte digest, rights, source identity, jurisdiction and currentness; require the qualified legal review decision before it can support a legal claim.",
        "Replay the exact diagnosed-result, research, review, quarantine, source-row and chunk chain. Preserve one completed sealed predecessor source set byte-for-byte and append at least one separately qualified source; shrink, replacement, equal-set or substitution fails.",
        "Chunk and embed that strict expansion into a new versioned, non-ACTIVE successor. Bind its predecessor, additions, source manifest, parser, chunker, embedding model, lane labels and owner decision. Only the capability-bound GE evaluator may read the held tree.",
        "After every material change, rerun all 331 visible cases, all 32 system scenarios and the cumulative visible diagnostic supplement. Never replace a full rerun with only failed-case retesting.",
        "Audit the approved coverage topology after the full run. Add separate visible diagnostic questions for every missing area, then repeat the controlled repair and full-rerun cycle.",
        "Stop a repair path after the same stable fingerprint fails twice. Exit only with complete reconciled results, zero missing approved coverage cells, no unresolved factual holds or system errors, all quality thresholds and critical floors met, no custody leakage, and explicit owner acceptance.",
    ):
        _bullet(document, text, numbered=True)

    _callout(
        document,
        "Training remains a separate decision",
        "Visible GE questions, answers, gold, diagnostic supplements, user history and protected unseen "
        "are excluded from weight training. A future clean-corpus training experiment requires its own "
        "rights, privacy, leakage and owner approval contract, followed by fresh visible and unseen evaluation.",
        fill=PALE_AMBER,
    )

    document.add_page_break()
    _heading(document, "7. Owner inputs required before the visible run", 1)
    _body(
        document,
        "Complete these fields or provide the exact same information in your reply. Codex will verify "
        "each item and return one final, candidate-bound execution record before any answer-model call.",
    )

    _heading(document, "7.1 Answer model and private transport", 2)
    _form_line(
        document, "Model artifact path", "Local approved artifact; no cloud or network fallback"
    )
    _form_line(document, "Model artifact SHA-256", "Exact 64-character digest")
    _form_line(
        document,
        "Model/runtime identity",
        "Format, quantization, context limit and runtime version",
    )
    _form_line(
        document, "Private UDS capability", "Exact owner-approved socket parent/path capability"
    )

    _heading(document, "7.2 Qualified legal gold and currentness", 2)
    _form_line(
        document,
        "Reviewer role and identity reference",
        "Qualified legal reviewer; AI cannot fill this role",
    )
    _form_line(
        document,
        "Gold/currentness package path",
        "Exact private package covering all 331 cases or explicit holds",
    )
    _form_line(document, "Package SHA-256", "Exact 64-character digest")
    _form_line(
        document,
        "Legal currentness cutoff/decision",
        "Owner/reviewer decision; do not infer from source dates",
    )

    document.add_page_break()
    _heading(document, "7.3 Private Development review root", 2)
    _form_line(document, "Development root", "Distinct owner-approved mode-0700 private root")
    _form_line(
        document, "Root capability/identity", "Must not be Validation or live review custody"
    )

    _heading(document, "7.4 Final candidate and ten-capability run authority", 2)
    _metadata(
        document,
        "Candidate build ID",
        "current-law-ew-full-fp16-v111-20260829-recovery-b (sealed non-ACTIVE)",
    )
    _metadata(
        document,
        "Candidate manifest/seal SHA-256",
        "356475aab2b01845f5753e4b183642268a108a215dcdd087f6508c8e0560d217",
    )
    _form_line(
        document,
        "Evaluation authority SHA-256",
        "Binds candidate, model, policies, denominator, root and resources",
    )
    _form_line(
        document,
        "Resource envelope/policy SHA-256",
        "Exact local capacity, timeout and cancellation limits",
    )
    _form_line(
        document, "Unseen custody ledger SHA-256", "Proves unseen remains sealed and unexposed"
    )
    _form_line(
        document,
        "System-suite manifest/order SHA-256",
        "Exact 32-scenario identity and execution order",
    )
    _form_line(
        document,
        "GE iteration-plan SHA-256",
        "Exact cycle, predecessor, repair and full-rerun bindings",
    )
    _form_line(
        document,
        "Visible diagnostic custody SHA-256",
        "Proves diagnostics remain visible-only and outside unseen/training",
    )

    _heading(document, "7.5 Exact GE coverage-topology authority", 2)
    _form_line(
        document,
        "Coverage topology manifest SHA-256",
        "Ordered 23-domain topology; empty public domains remain explicit gaps",
    )
    _form_line(
        document,
        "Coverage decision ID + request/resolution SHA-256",
        "Exact stored trusted decision pair; self-sealed JSON is insufficient",
    )

    decision_heading = _heading(document, "8. Requested owner decision", 1)
    decision_heading.paragraph_format.page_break_before = True
    _callout(
        document,
        "Recommended decision after all exact fields pass",
        "Authorize one visible baseline over all 331 GE cases and the 32 separately reported system "
        "scenarios, factual gate first and quality gate second. Require every hold/error in the review. "
        "Allow the accepted create-only official-source gap repair loop, non-ACTIVE successor builds, "
        "the exact authorized 23-domain coverage topology, separate visible diagnostics and complete "
        "successor reruns. Return the complete improved GE "
        "evidence and owner-acceptance record before any unseen, training, live or GitHub action.",
        fill=PALE_GREEN,
    )
    for item in (
        "[ ] APPROVE the exact visible GE loop execution only after every field in section 7 is filled and technically verified.",
        "[ ] AMEND the proposed evaluation authority as written below.",
        "[ ] DECLINE the visible run; retain the prepared set and evidence without execution.",
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        _set_font(paragraph.add_run(item), size=11, bold=item.startswith("[ ] APPROVE"))
    _form_line(document, "Owner amendments", "Leave blank if approving without amendment")
    _form_line(document, "Owner typed name", "Owner decision identity")
    _form_line(document, "Decision date", "ISO date preferred")

    document.add_page_break()
    _heading(document, "9. Authority explicitly excluded", 1)
    for text in (
        "No private unseen prompt may be opened, projected, scored or disclosed.",
        "No evaluation question, answer, gold, reviewer note, user history or unseen material may become training data.",
        "No weight training is authorized. Any future training experiment requires a separate rights/privacy-reviewed corpus and approval.",
        "No promotion, ACTIVE/PREVIOUS pointer write, Owner Certification 60, Validation 30, live activation, public access, sharing or cloud storage is authorized.",
        "No deletion, overwrite, in-place index mutation or history removal is authorized by this document.",
        "No Git commit, push or GitHub publication occurs before GE_COMPLETE_OWNER_ACCEPTED and the applicable exact publication gate.",
    ):
        _bullet(document, text)

    _callout(
        document,
        "Stop rule",
        "If any exact identity, capability, denominator, legal decision or candidate binding is missing or changes, "
        "the run remains blocked. An answer-model or evaluation run must never be used to debug the missing gate.",
        fill=PALE_RED,
    )

    document.core_properties.title = "LegalBot Phase 2 GE Improvement Loop Execution Approval"
    document.core_properties.subject = "Owner gate for the visible GE factual-first evaluation and controlled repair loop"
    document.core_properties.author = "LegalBot v1.11"
    document.core_properties.keywords = "LegalBot, Phase 2, evaluation, owner approval"
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
