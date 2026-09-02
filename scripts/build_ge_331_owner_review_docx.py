#!/usr/bin/env python3
"""Build the complete 331-case General Enquiries owner-review DOCX."""

from __future__ import annotations

import hashlib
import json
import os
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
RUN_ID = "LegalBot-GE-2026-09-01-owner-draft-r4"
RUN_ROOT = PROJECT_ROOT / "data/evaluations/general-enquiries" / RUN_ID
OUTPUT_PATH = PROJECT_ROOT / "output/docx/LegalBot-GE-331-Full-Answer-Review.docx"
DIAGNOSTIC_R2 = (
    PROJECT_ROOT
    / "data/evaluations/general-enquiries/LegalBot-GE-2026-09-01-owner-draft-r2/cases"
)
DIAGNOSTIC_R3 = (
    PROJECT_ROOT
    / "data/evaluations/general-enquiries/LegalBot-GE-2026-09-01-owner-draft-r3/cases"
)

NAVY = "17365D"
BLUE = "2F75B5"
TEAL = "008C95"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E8F5F4"
PALE_RED = "FDECEC"
RED = "A61B1B"
ORANGE = "C65911"
PALE_ORANGE = "FFF2CC"
INK = "1F2937"
MID = "5B6573"
LIGHT = "D9E2F3"
WHITE = "FFFFFF"

TOPIC_LABELS = {
    "administrative-law": "Administrative law",
    "ai-and-data-protection": "AI and data protection",
    "business-and-company-law": "Business and company law",
    "commercial-law": "Commercial law",
    "competition-law": "Competition law",
    "contemporary-biolaw-and-regulation": "Contemporary biolaw and regulation",
    "contract-law": "Contract law",
    "criminal-law": "Criminal law",
    "eu-internal-market-law": "EU internal market law",
    "international-commercial-mediation": "International commercial mediation",
    "land-law": "Land law",
    "law-and-medicine": "Law and medicine",
    "pensions-law": "Pensions law",
    "private-international-law": "Private international law",
    "tort-law": "Tort law",
    "trusts-law": "Trusts law",
    "wills-and-estates": "Wills and estates",
}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        while block := handle.read(1024 * 1024):
            digest.update(block)
    return digest.hexdigest()


def load_cases() -> tuple[dict[str, Any], list[dict[str, Any]]]:
    manifest_path = RUN_ROOT / "RUN-MANIFEST.json"
    if not manifest_path.is_file():
        raise RuntimeError("the complete 331-case run manifest is not available")
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    if manifest.get("case_count") != 331 or len(manifest.get("cases") or []) != 331:
        raise RuntimeError("owner-review manifest does not contain exactly 331 cases")
    rows: list[dict[str, Any]] = []
    for item in manifest["cases"]:
        path = RUN_ROOT / str(item["path"])
        if not path.is_file() or sha256_file(path) != item["sha256"]:
            raise RuntimeError(f"owner-review case artifact changed: {item['ordinal']}")
        value = json.loads(path.read_text(encoding="utf-8"))
        if value.get("ordinal") != item["ordinal"] or value.get("case_id") != item["case_id"]:
            raise RuntimeError("owner-review case identity differs from the manifest")
        rows.append(value)
    if [row["ordinal"] for row in rows] != list(range(1, 332)):
        raise RuntimeError("owner-review case order is incomplete")
    return manifest, rows


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell: Any, **edges: dict[str, str]) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, values in edges.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in values.items():
            element.set(qn(f"w:{key}"), value)


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _keep_with_next(paragraph: Any) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char, instr, sep, text, end))


def configure_document(doc: DocumentType) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.70)
    section.left_margin = Inches(0.80)
    section.right_margin = Inches(0.80)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 12),
        ("Subtitle", 13, MID, 0, 8),
        ("Heading 1", 18, NAVY, 12, 8),
        ("Heading 2", 14, BLUE, 10, 5),
        ("Heading 3", 11.5, NAVY, 7, 3),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Case Question" not in styles:
        style = styles.add_style("Case Question", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY)
        style.paragraph_format.left_indent = Inches(0.12)
        style.paragraph_format.right_indent = Inches(0.12)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_together = True

    if "Small Note" not in styles:
        style = styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor.from_string(MID)
        style.paragraph_format.space_after = Pt(3)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(2.1)
    left = table.cell(0, 0).paragraphs[0]
    run = left.add_run("LEGALBOT  /  GENERAL ENQUIRIES")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right.add_run("OWNER REVIEW  ·  331 CASES")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    for cell in table.rows[0].cells:
        _set_cell_border(cell, bottom={"val": "single", "sz": "12", "color": LIGHT})

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Owner-review draft  •  Non-authorizing  •  Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID)
    _add_field(p, "PAGE")


def add_cover(doc: DocumentType, manifest: dict[str, Any]) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(60)
    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = label.add_run("PHASE 2  /  OWNER REVIEW")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("General Enquiries\n331-Case Full Review")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Factual gate first  •  Law-folder 70+ quality second  •  Owner decision last")

    line = doc.add_table(rows=1, cols=1)
    line.alignment = WD_TABLE_ALIGNMENT.CENTER
    line.columns[0].width = Inches(2.4)
    _set_cell_shading(line.cell(0, 0), TEAL)
    line.cell(0, 0).height = Pt(4)

    doc.add_paragraph()
    card = doc.add_table(rows=4, cols=2)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    card.autofit = False
    labels = (
        ("Run", manifest["run_id"]),
        ("Visible cases", "331 (306 core + 25 stress)"),
        ("Legal cutoff", "28 August 2026"),
        ("Status", "Owner-review draft; not legal gold or live output"),
    )
    for index, (key, value) in enumerate(labels):
        left, right = card.rows[index].cells
        left.width = Inches(1.5)
        right.width = Inches(4.8)
        _set_cell_shading(left, NAVY)
        _set_cell_shading(right, PALE_BLUE)
        left.text = key
        right.text = str(value)
        for run in left.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size = Pt(9)
        for run in right.paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(INK)
            run.font.size = Pt(9)

    doc.add_paragraph()
    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = notice.add_run(
        "This document deliberately shows every factual hold, system error, and sub-70 result. "
        "A hold is evidence about the present system; it is not a conclusion about the user's rights."
    )
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor.from_string(MID)
    doc.add_page_break()


def add_summary(doc: DocumentType, manifest: dict[str, Any], rows: list[dict[str, Any]]) -> None:
    doc.add_heading("How to use this review pack", level=1)
    doc.add_paragraph(
        "Review each question and the user-facing response. The factual gate is decisive: if any "
        "material factual check fails, the case is held and receives no 70+ score. For each case, "
        "select one owner decision: Approve, Re-evaluate, Tune or train, or Hold."
    )
    steps = (
        "Check whether the response addresses the user's real problem in plain English.",
        "Check each failed factual item and the source list before considering prose quality.",
        "Where a factual pass exists, review the seven quality dimensions and critical floors.",
        "Record one owner decision; later tuning or training must create a new evaluated version.",
    )
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)

    factual = Counter(row["factual_gate"]["outcome"] for row in rows)
    quality = Counter(
        (row.get("quality_review") or {}).get("quality_outcome", "NOT_SCORED") for row in rows
    )
    evidence = Counter(
        "WITH_EVIDENCE" if row["retrieval"]["evidence_count"] else "NO_QUALIFIED_EVIDENCE"
        for row in rows
    )
    doc.add_heading("Run result at a glance", level=2)
    table = doc.add_table(rows=2, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headings = ("Factual pass", "Factual hold", "70+ scored", "No qualified evidence")
    values = (
        factual.get("FACTUAL_PASS", 0),
        factual.get("FACTUAL_HOLD", 0),
        sum(value for key, value in quality.items() if key in {"MEETS_70_STANDARD", "EXCEEDS_70_STANDARD"}),
        evidence.get("NO_QUALIFIED_EVIDENCE", 0),
    )
    for col, heading in enumerate(headings):
        cell = table.cell(0, col)
        _set_cell_shading(cell, NAVY)
        cell.text = heading
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size = Pt(8.5)
        value_cell = table.cell(1, col)
        value_cell.text = str(values[col])
        value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(value_cell, PALE_TEAL if col in {0, 2} else PALE_RED)
        for run in value_cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor.from_string(TEAL if col in {0, 2} else RED)

    doc.add_heading("Important execution limitation", level=2)
    p = doc.add_paragraph()
    run = p.add_run("Release attestation stopped at the clean-Git-tree gate. ")
    run.bold = True
    p.add_run(
        "The exact sealed index was verified read-only for this pack, but its catalogue status "
        "remained built_unscored and no ACTIVE or PREVIOUS pointer was written. The production "
        "semantic threshold is also pending calibration, so ordinary semantic hits do not qualify."
    )

    by_topic: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        by_topic[str(row["topic_id"])].append(row)
    doc.add_heading("Coverage and outcome by topic", level=2)
    topic_table = doc.add_table(rows=1, cols=5)
    topic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Topic", "Cases", "Pass", "Hold", "70+ scored")
    for col, value in enumerate(headers):
        cell = topic_table.cell(0, col)
        _set_cell_shading(cell, NAVY)
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size = Pt(8)
    _set_repeat_table_header(topic_table.rows[0])
    for topic, topic_rows in by_topic.items():
        cells = topic_table.add_row().cells
        topic_factual = Counter(item["factual_gate"]["outcome"] for item in topic_rows)
        topic_quality = Counter(
            (item.get("quality_review") or {}).get("quality_outcome", "NOT_SCORED")
            for item in topic_rows
        )
        values = (
            TOPIC_LABELS.get(topic, topic.replace("-", " ").title()),
            len(topic_rows),
            topic_factual.get("FACTUAL_PASS", 0),
            topic_factual.get("FACTUAL_HOLD", 0),
            sum(
                value
                for key, value in topic_quality.items()
                if key in {"MEETS_70_STANDARD", "EXCEEDS_70_STANDARD"}
            ),
        )
        for col, value in enumerate(values):
            cells[col].text = str(value)
            cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cells[col].paragraphs[0].runs:
                run.font.size = Pt(8.5)
        if topic_factual.get("FACTUAL_HOLD"):
            _set_cell_shading(cells[3], PALE_RED)
    doc.add_page_break()


def _add_status_card(doc: DocumentType, row: dict[str, Any]) -> None:
    factual = row["factual_gate"]["outcome"]
    quality = (row.get("quality_review") or {}).get("quality_outcome", "NOT_SCORED")
    evidence = int(row["retrieval"]["evidence_count"])
    card = doc.add_table(rows=2, cols=3)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    card.autofit = False
    for col, label in enumerate(("FACTUAL", "70+ QUALITY", "QUALIFIED EVIDENCE")):
        cell = card.cell(0, col)
        _set_cell_shading(cell, NAVY)
        cell.text = label
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
            run.font.size = Pt(7.5)
    values = (factual.replace("_", " "), quality.replace("_", " "), str(evidence))
    for col, value in enumerate(values):
        cell = card.cell(1, col)
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        passing = value in {"FACTUAL PASS", "MEETS 70 STANDARD", "EXCEEDS 70 STANDARD"}
        _set_cell_shading(cell, PALE_TEAL if passing else PALE_RED if col < 2 else PALE_BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(TEAL if passing else RED if col < 2 else NAVY)


def _add_answer_text(doc: DocumentType, text: str) -> None:
    heading_names = {
        "Short answer",
        "What the law requires",
        "What to do now",
        "Uncertainty and next questions",
        "Evidence limitations",
        "Evidence hold",
        "Questions that must be answered",
        "Review reason",
    }
    for raw in text.splitlines():
        line = raw.strip().replace("\u200b", "")
        if not line:
            continue
        if line in heading_names:
            doc.add_heading(line, level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.keep_together = True


def _add_case(doc: DocumentType, row: dict[str, Any], first_in_topic: bool) -> None:
    topic = str(row["topic_id"])
    if first_in_topic:
        doc.add_heading(TOPIC_LABELS.get(topic, topic.replace("-", " ").title()), level=1)
        p = doc.add_paragraph(style="Small Note")
        p.add_run(f"Topic section  •  {topic}")

    heading = doc.add_heading(
        f"Case {row['ordinal']:03d}  ·  {row['case_id']}", level=2
    )
    _keep_with_next(heading)
    meta = doc.add_paragraph(style="Small Note")
    meta.add_run(
        f"{row['lane'].replace('_', ' ').title()}  •  Scenario {row['scenario_family_id']}  •  "
        f"Answer type: {row['answer_kind'].replace('_', ' ')}"
    )
    q = doc.add_paragraph(style="Case Question")
    q.add_run(row["question"])
    _add_status_card(doc, row)

    doc.add_heading("Response shown to the user", level=3)
    _add_answer_text(doc, str(row["answer"]))

    failed = [
        name for name, outcome in row["factual_gate"]["checks"].items() if outcome == "FAIL"
    ]
    doc.add_heading("Factual review", level=3)
    if failed:
        for name in failed:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(name.replace("_", " ").title() + ": ")
            run.bold = True
            p.add_run(row["factual_gate"]["reasons"][name])
    else:
        doc.add_paragraph("All ten material factual checks passed.")

    quality = row.get("quality_review")
    doc.add_heading("70+ quality review", level=3)
    if quality is None:
        doc.add_paragraph(
            str(row.get("quality_not_scored_reason") or "Quality was not scored."),
            style="Small Note",
        )
    else:
        p = doc.add_paragraph()
        p.add_run(f"Overall: {quality['quality_score']:.2f}/100  •  ").bold = True
        p.add_run(str(quality["quality_outcome"]).replace("_", " ").title())
        score_table = doc.add_table(rows=1, cols=3)
        score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for col, label in enumerate(("Dimension", "Score", "Maximum")):
            cell = score_table.cell(0, col)
            _set_cell_shading(cell, NAVY)
            cell.text = label
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)
                run.font.size = Pt(8)
        maxima = {
            "legal_and_factual_accuracy": 25,
            "issue_coverage_and_reasoning": 15,
            "authority_and_currentness": 15,
            "practical_steps_and_urgency": 15,
            "uncertainty_limits_and_clarification": 10,
            "organisation_and_plain_language": 10,
            "traceability_and_citations": 10,
        }
        for dimension, score in quality["quality_dimensions"].items():
            cells = score_table.add_row().cells
            cells[0].text = dimension.replace("_", " ").title()
            cells[1].text = f"{float(score):.2f}"
            cells[2].text = str(maxima[dimension])
            for cell in cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)

    doc.add_heading("Sources retrieved", level=3)
    evidence = row["retrieval"]["evidence"]
    if not evidence:
        doc.add_paragraph("No production-qualified evidence was admitted for this question.")
    else:
        seen: set[tuple[str, str]] = set()
        for item in evidence:
            citation = str(
                item.get("canonical_citation")
                or (item.get("citation_data") or {}).get("title")
                or item["source_version_id"]
            )
            locator = str(item.get("locator") or "")
            key = (citation, locator)
            if key in seen:
                continue
            seen.add(key)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(citation).bold = True
            if locator:
                p.add_run(f", {locator}")
            route = item.get("retrieval_route")
            qualified = item.get("retrieval_threshold_qualified")
            p.add_run(f"  [{route}; qualified={qualified}]")

    doc.add_heading("Owner decision", level=3)
    decisions = doc.add_table(rows=1, cols=4)
    decisions.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, label in enumerate(("☐ Approve", "☐ Re-evaluate", "☐ Tune / train", "☐ Hold")):
        cell = decisions.cell(0, col)
        cell.text = label
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_shading(cell, PALE_BLUE if col != 3 else PALE_ORANGE)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(NAVY if col != 3 else ORANGE)
    note = doc.add_paragraph(style="Small Note")
    note.add_run("Owner note: ______________________________________________________________")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_cases(doc: DocumentType, rows: list[dict[str, Any]]) -> None:
    doc.add_heading("Full 331-case review", level=1)
    doc.add_paragraph(
        "Cases remain in the exact visible-pack order. Each case starts on a new page so an "
        "owner decision can be recorded without ambiguity."
    )
    doc.add_page_break()
    previous_topic: str | None = None
    for row in rows:
        topic = str(row["topic_id"])
        _add_case(doc, row, first_in_topic=topic != previous_topic)
        previous_topic = topic


def _load_diagnostic(path: Path, ordinal: int) -> dict[str, Any] | None:
    matches = list(path.glob(f"{ordinal:03d}-*.json")) if path.is_dir() else []
    return json.loads(matches[0].read_text(encoding="utf-8")) if len(matches) == 1 else None


def add_diagnostics(doc: DocumentType) -> None:
    doc.add_heading("Appendix A  ·  Forced-semantic diagnostics", level=1)
    doc.add_paragraph(
        "These preserved diagnostics are outside the scored 331-case r4 run. They show why the "
        "production factual gate must not be bypassed while semantic relevance calibration is pending."
    )
    r2 = _load_diagnostic(DIAGNOSTIC_R2, 232)
    if r2:
        doc.add_heading("Diagnostic 1  ·  Pensions case with unqualified semantic context", level=2)
        doc.add_paragraph(r2["question"], style="Case Question")
        _add_answer_text(doc, r2["answer"])
        doc.add_paragraph(
            "Observed issue: the model produced readable prose but drew on employment and trade-union "
            "provisions that were not reliably targeted to missing workplace pension contributions. "
            "The factual gate held the answer at claim-evidence support.",
            style="Small Note",
        )
    r3 = _load_diagnostic(DIAGNOSTIC_R3, 232)
    if r3:
        doc.add_heading("Diagnostic 2  ·  Bounded semantic attempt", level=2)
        doc.add_paragraph(
            "The reduced-evidence attempt reached the five-minute gateway timeout. Its fallback was "
            "preserved, no result was treated as verified, and the full run returned to the production gate.",
            style="Small Note",
        )
    doc.add_page_break()


def add_policy_appendix(doc: DocumentType, manifest: dict[str, Any]) -> None:
    doc.add_heading("Appendix B  ·  Review rules and integrity", level=1)
    doc.add_heading("Factual-first rule", level=2)
    doc.add_paragraph(
        "All ten material checks must pass before quality is scored: integrity chain; claim-evidence "
        "support; user-fact provenance; jurisdiction; cutoff/currentness; dates, amounts and deadlines; "
        "citation/quotation identity; contradictions/counterauthority; safety/urgency; and privacy/instruction isolation."
    )
    doc.add_heading("70+ quality dimensions", level=2)
    dimensions = (
        ("Legal and factual accuracy", "25", "Critical floor 17.5"),
        ("Issue coverage and reasoning", "15", "—"),
        ("Authority and currentness", "15", "Critical floor 10.5"),
        ("Practical steps and urgency", "15", "Critical floor 9.0"),
        ("Uncertainty, limits and clarification", "10", "—"),
        ("Organisation and plain language", "10", "—"),
        ("Traceability and citations", "10", "—"),
    )
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, value in enumerate(("Dimension", "Maximum", "Floor")):
        cell = table.cell(0, col)
        _set_cell_shading(cell, NAVY)
        cell.text = value
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)
    for values in dimensions:
        cells = table.add_row().cells
        for col, value in enumerate(values):
            cells[col].text = value
    doc.add_paragraph(
        "70–79 meets the target; 80+ exceeds it. A critical-floor failure requires material "
        "improvement regardless of the total. Automated results remain advisory until qualified legal review."
    )
    doc.add_heading("Bound identities", level=2)
    identities = (
        ("Visible pack manifest", manifest["visible_pack_content_sha256"]),
        ("Visible case manifest", manifest["visible_case_manifest_sha256"]),
        ("Visible case order", manifest["visible_case_order_sha256"]),
        ("Retrieval build", manifest["build_id"]),
        ("Source manifest", manifest["source_manifest_sha256"]),
        ("Model", manifest["model_version"]),
        ("Assessment bundle", manifest["assessment_bundle_sha256"]),
    )
    for label, value in identities:
        p = doc.add_paragraph(style="Small Note")
        p.add_run(label + ": ").bold = True
        p.add_run(str(value))
    doc.add_paragraph(
        "No question, evidence record, diagnostic, runtime artifact, or prior package was deleted "
        "or overwritten while preparing this review."
    )


def core_properties(doc: DocumentType) -> None:
    props = doc.core_properties
    props.title = "LegalBot General Enquiries 331-Case Full Owner Review"
    props.subject = "Factual-first and 70+ owner review of the visible General Enquiries set"
    props.author = "LegalBot recovery workspace"
    props.keywords = "LegalBot, General Enquiries, factual gate, 70+, owner review"
    props.comments = "Non-authorizing owner-review artifact; not legal gold, unseen, training, promotion, or live output."


def main() -> int:
    manifest, rows = load_cases()
    if OUTPUT_PATH.exists():
        raise FileExistsError(f"refusing to replace existing DOCX: {OUTPUT_PATH}")
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    core_properties(document)
    add_cover(document, manifest)
    add_summary(document, manifest, rows)
    add_cases(document, rows)
    add_diagnostics(document)
    add_policy_appendix(document, manifest)
    document.save(OUTPUT_PATH)
    os.chmod(OUTPUT_PATH, 0o600)
    print(
        json.dumps(
            {
                "path": str(OUTPUT_PATH),
                "sha256": sha256_file(OUTPUT_PATH),
                "case_count": len(rows),
                "paragraphs": len(document.paragraphs),
                "tables": len(document.tables),
            },
            indent=2,
            sort_keys=True,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
