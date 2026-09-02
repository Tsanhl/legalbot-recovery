from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "output" / "docx" / "LegalBot-Phase2-Owner-Approval.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = '17365D'
BLUE = '2F5597'
LIGHT_BLUE = 'D9EAF7'
PALE_BLUE = 'EEF5FB'
PALE_GREY = 'F2F2F2'
MID_GREY = '666666'
DARK = '1F1F1F'
WHITE = 'FFFFFF'
AMBER = 'FFF2CC'
GREEN = 'E2F0D9'
RED = 'FCE4D6'
BORDER = 'B7C9DD'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, **edges):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge not in edges:
            continue
        tag = 'w:' + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edges[edge].items():
            element.set(qn('w:' + key), str(value))


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for margin, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn('w:' + margin))
        if node is None:
            node = OxmlElement('w:' + margin)
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_table_fixed(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in('w:tblLayout')
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tbl_pr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    indent = tbl_pr.first_child_found_in('w:tblInd')
    if indent is None:
        indent = OxmlElement('w:tblInd')
        tbl_pr.append(indent)
    indent.set(qn('w:w'), '120')
    indent.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for idx, width in enumerate(widths):
        grid_col = grid.gridCol_lst[idx]
        grid_col.set(qn('w:w'), str(int(width * 1440)))
    for row in table.rows:
        prevent_row_split(row)
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(int(width * 1440)))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    repeat_table_header(row)


def format_table_text(table, size=9.1, has_header=True):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = 'Arial'
                    if has_header and row_index == 0:
                        run.font.size = Pt(9)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)
                    else:
                        run.font.size = Pt(size)
                        run.font.color.rgb = RGBColor.from_string(DARK)


def add_table_header(table, labels):
    row = table.rows[0]
    set_repeat_table_header(row)
    for idx, label in enumerate(labels):
        cell = row.cells[idx]
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(WHITE)


def set_all_table_borders(table, color=BORDER, size=6):
    edge = {'val': 'single', 'sz': size, 'space': '0', 'color': color}
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=edge, left=edge, bottom=edge, right=edge)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text='', bold_lead=None):
    p = doc.add_paragraph(style='Body Text')
    if bold_lead and text.startswith(bold_lead):
        first, rest = text[:len(bold_lead)], text[len(bold_lead):]
        p.add_run(first).bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_list(doc, items, numbered=False):
    style = 'List Number' if numbered else 'List Bullet'
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_callout(doc, title, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={'val': 'single', 'sz': '20', 'space': '0', 'color': BLUE})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.05
    for run in p2.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


doc = Document()
doc.settings.odd_and_even_pages_header_footer = False
section = doc.sections[0]
section.different_first_page_header_footer = False
section.orientation = WD_ORIENT.PORTRAIT
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.82)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.28)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(DARK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

body = styles['Body Text']
body.font.name = 'Arial'
body.font.size = Pt(11)
body.font.color.rgb = RGBColor.from_string(DARK)
body.paragraph_format.space_after = Pt(6)
body.paragraph_format.line_spacing = 1.05

for style_name in ('List Bullet', 'List Number'):
    st = styles[style_name]
    st.font.name = 'Arial'
    st.font.size = Pt(10.5)
    st.font.color.rgb = RGBColor.from_string(DARK)
    st.paragraph_format.space_after = Pt(3)

h1 = styles['Heading 1']
h1.font.name = 'Arial'
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor.from_string(BLUE)
h1.paragraph_format.space_before = Pt(13)
h1.paragraph_format.space_after = Pt(5)
h1.paragraph_format.keep_with_next = True

h2 = styles['Heading 2']
h2.font.name = 'Arial'
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor.from_string(BLUE)
h2.paragraph_format.space_before = Pt(10)
h2.paragraph_format.space_after = Pt(4)
h2.paragraph_format.keep_with_next = True

h3 = styles['Heading 3']
h3.font.name = 'Arial'
h3.font.size = Pt(11.5)
h3.font.bold = True
h3.font.color.rgb = RGBColor.from_string(NAVY)
h3.paragraph_format.space_before = Pt(8)
h3.paragraph_format.space_after = Pt(3)
h3.paragraph_format.keep_with_next = True

# Header masthead
header = section.header
header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
set_table_fixed(header_table, [3.35, 3.0])
header_table.cell(0, 0).text = 'LEGALBOT v1.11'
header_table.cell(0, 1).text = 'OWNER DECISION MEMO'
header_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
for cell in header_table.rows[0].cells:
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, top=55, bottom=55, start=100, end=100)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(8.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)

# Footer
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(2)
r = p.add_run('Owner review • Phase 2 not started')
r.font.name = 'Arial'
r.font.size = Pt(8)
r.font.color.rgb = RGBColor.from_string(MID_GREY)
r2 = p.add_run(' ' * 44 + 'Page ')
r2.font.name = 'Arial'
r2.font.size = Pt(8)
r2.font.color.rgb = RGBColor.from_string(MID_GREY)
add_page_field(p)
for run in p.runs:
    if run.font.size is None:
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MID_GREY)

# Title area
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(4)
r = p.add_run('LegalBot Phase 2')
r.font.name = 'Arial'
r.font.size = Pt(25)
r.font.bold = True
r.font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(13)
r = p.add_run('Owner approval for factual-first evaluation, quality review, improvement and unseen testing')
r.font.name = 'Arial'
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(BLUE)

meta = doc.add_table(rows=4, cols=2)
set_table_fixed(meta, [1.55, 4.80])
meta_data = [
    ('Date', '1 September 2026'),
    ('Status', 'Prepared for owner decision; no evaluation, training, unseen or live run started'),
    ('Design', 'Accepted and editable; no further design approval requested'),
    ('Decision', 'Approve, amend or decline the three recommended Phase-2 preparation decisions'),
]
for idx, (label, value) in enumerate(meta_data):
    set_cell_shading(meta.cell(idx, 0), PALE_GREY)
    p1 = meta.cell(idx, 0).paragraphs[0]
    p1.add_run(label).bold = True
    p2 = meta.cell(idx, 1).paragraphs[0]
    p2.add_run(value)
format_table_text(meta, 9.5, has_header=False)
set_all_table_borders(meta)

doc.add_paragraph().paragraph_format.space_after = Pt(0)
add_callout(
    doc,
    'Decision requested',
    'Approve the scope, review standard and controlled improvement sequence below. Approval prepares Phase 2; it does not by itself authorize a model run, source-currentness judgment, weight training, unseen execution, promotion or live activation.',
    fill=AMBER,
)

add_heading(doc, '1. What this approval will give you', 1)
add_body(doc, 'You will review the current model through readable answers, not raw gate logs. Every answer will first receive a material factual/legal check. Only factually eligible answers receive a quality result. The final owner pack will show every case, including holds, failures and system errors, and will identify the layer that needs repair.')
add_list(doc, [
    'Readable question-and-answer reviews grouped by topic.',
    'A factual result with the exact material problem when an answer is held.',
    'A mode-specific quality result and a concrete improvement action.',
    'Full denominators: no failed, held or system-error case disappears from the report.',
    'One final owner approval DOCX summarising results, risks, proposed repairs and the next exact decision.',
])

add_heading(doc, '2. Proposed Phase-2 sequence', 1)
add_list(doc, [
    'Prepare the exact visible evaluation inputs and verify the technical baseline, legal evidence, currentness, gold, model capability, resources and private review roots.',
    'Run the authorized visible General Enquiries baseline on all 331 accepted cases; report the 32 system scenarios separately.',
    'Apply the Layer 1 factual/legal eligibility gate to every answer.',
    'Apply the Layer 2 mode-quality gate only to FACTUAL_PASS answers.',
    'Return readable answers and the final owner review document before making model changes.',
    'Repair the responsible layer. Use weight training only if later separately justified and approved.',
    'Re-evaluate all visible cases on the exact improved candidate and return the comparison for owner approval.',
    'After that approval, run the separately protected unseen scope once; never tune the tested candidate from unseen prompts or findings.',
], numbered=True)

doc.add_page_break()
add_heading(doc, '3. Layer 1 — factual and legal eligibility', 1)
add_body(doc, 'This is a hard gate. Strong writing cannot cure a material factual or legal defect. Each answer receives one of three outcomes: FACTUAL_PASS, FACTUAL_HOLD or SYSTEM_ERROR. Only FACTUAL_PASS continues to quality review.')

factual = doc.add_table(rows=1, cols=3)
add_table_header(factual, ['Check', 'Required evidence', 'Hold when'])
factual_rows = [
    ('Identity and scope', 'Exact case/question, attributed user facts, jurisdiction and requested date', 'The answer changes the question, invents facts or answers outside scope'),
    ('Current authority', 'Current, applicable primary authority and properly used secondary material', 'A material rule is unsupported, outdated, inapplicable or treated as settled when disputed'),
    ('Material precision', 'Support for quotations, dates, sums, deadlines, calculations, procedures and remedies', 'A material detail cannot be traced or a calculation/procedure is wrong'),
    ('Balance and uncertainty', 'Contrary authority, exceptions, missing facts and genuine uncertainty', 'The answer overstates certainty or omits a material counterposition'),
    ('Citation and provenance', 'Reviewed source metadata, locators and deterministic OSCOLA output', 'A source or locator is invented, mismatched or not reviewable'),
    ('Safety and privacy', 'No guarantee, fabricated action/referral, cross-user leakage or protected-evaluation leakage', 'Any material safety, privacy or custody breach occurs'),
]
for i, rowdata in enumerate(factual_rows, start=1):
    cells = factual.add_row().cells
    for j, value in enumerate(rowdata):
        cells[j].text = value
    if i % 2 == 0:
        for cell in cells:
            set_cell_shading(cell, PALE_BLUE)
set_table_fixed(factual, [1.55, 2.75, 2.05])
format_table_text(factual, 8.7)
set_all_table_borders(factual)

add_heading(doc, '4. Layer 2 — quality standard', 1)
add_body(doc, 'The quality target adapts the official Levels 2 and 3 first-class criteria and the owner’s Y2/Y3 feedback. It is a diagnostic product standard, not a university award. The target is accurate, detailed and comprehensive knowledge; insight; complete issue identification; persuasive application; relevant primary and secondary research; synthesis and critical evaluation; independent judgment; and clear, economical, confident communication.')

mode_table = doc.add_table(rows=1, cols=3)
add_table_header(mode_table, ['Mode', '70+ quality target', 'Practical emphasis'])
mode_rows = [
    ('General Enquiry', 'Correct direct answer; sound legal basis; material qualifications; clear and economical explanation', 'Plain language, only necessary clarification, usable options/next steps, proportionate limits and urgent/safety handling. Academic length is not required.'),
    ('Problem Based', 'Complete issue spotting; accurate rules; fact-sensitive application; competing arguments; supported conclusions', 'Identify missing facts, remedies, defences, deadlines, procedural choices and outcome uncertainty.'),
    ('Essay', 'Responsive thesis; accurate authority; structured synthesis; competing views; critical evaluation; independent judgment', 'Use wider academic material, explain rationales, engage with dissent/reform and reach a supported conclusion.'),
]
for idx, rowdata in enumerate(mode_rows, start=1):
    cells = mode_table.add_row().cells
    for j, value in enumerate(rowdata):
        cells[j].text = value
    if idx == 1:
        for cell in cells:
            set_cell_shading(cell, GREEN)
    elif idx % 2 == 0:
        for cell in cells:
            set_cell_shading(cell, PALE_BLUE)
set_table_fixed(mode_table, [1.20, 2.55, 2.60])
format_table_text(mode_table, 8.9)
set_all_table_borders(mode_table)

doc.add_page_break()
add_heading(doc, 'Diagnostic bands', 2)
bands = doc.add_table(rows=1, cols=3)
add_table_header(bands, ['Band', 'Meaning', 'Owner interpretation'])
band_rows = [
    ('80–100', 'Exceeds 70+ standard', 'Factually eligible and consistently strong; still review material limitations'),
    ('70–79', 'Meets 70+ standard', 'Factually eligible and reaches the approved quality target'),
    ('60–69', 'Below 70+ standard', 'Substantially competent but needs a defined quality improvement'),
    ('Below 60', 'Material improvement required', 'Significant quality repair required'),
    ('Any factual hold', 'No quality approval', 'Repair and re-check the material factual/legal defect first'),
]
for idx, rowdata in enumerate(band_rows, start=1):
    cells = bands.add_row().cells
    for j, value in enumerate(rowdata):
        cells[j].text = value
    if idx == 2:
        for cell in cells:
            set_cell_shading(cell, GREEN)
    elif idx == 5:
        for cell in cells:
            set_cell_shading(cell, RED)
    elif idx % 2 == 0:
        for cell in cells:
            set_cell_shading(cell, PALE_BLUE)
set_table_fixed(bands, [1.05, 2.10, 3.20])
format_table_text(bands, 8.9)
set_all_table_borders(bands)

add_callout(
    doc,
    'Quality rule',
    'A polished answer with a material legal error remains FACTUAL_HOLD. A correct answer can still fall below the quality target when it is incomplete, unclear, poorly applied, weakly researched or insufficiently critical for its mode.',
)

add_heading(doc, '5. Current General Enquiries scope', 1)
scope = doc.add_table(rows=1, cols=4)
add_table_header(scope, ['Set', 'Count', 'Use now', 'Treatment'])
scope_rows = [
    ('Accepted visible GE', '331', 'Yes, after exact run approval', 'All cases; no sampling; factual and quality results retained'),
    ('System scenarios', '32', 'Yes, separately', 'Reliability/safety behaviour; excluded from legal-quality denominator'),
    ('Private unseen GE', '306', 'No', 'Separate custody; no development, training or visible disclosure'),
    ('PB and Essay', 'Present', 'Later when selected', 'Same factual-first sequence; separate mode/topic results'),
]
for idx, rowdata in enumerate(scope_rows, start=1):
    cells = scope.add_row().cells
    for j, value in enumerate(rowdata):
        cells[j].text = value
    if idx % 2 == 0:
        for cell in cells:
            set_cell_shading(cell, PALE_BLUE)
set_table_fixed(scope, [1.55, 0.72, 1.55, 2.53])
format_table_text(scope, 8.8)
set_all_table_borders(scope)
add_body(doc, 'The question sets are review/evaluation inputs. They are not completed legal gold, legal authority or an approved training corpus.')

doc.add_page_break()
add_heading(doc, '6. Improvement and training guardrails', 1)
add_body(doc, 'The first response to a failure is diagnosis. The repair must target the responsible layer instead of treating every problem as a model-weight problem.')

repair = doc.add_table(rows=1, cols=2)
add_table_header(repair, ['Failure layer', 'Correct first response'])
repair_rows = [
    ('Source or currentness', 'Admit/review the correct authority and date; rebuild attributable evidence'),
    ('Retrieval or reranking', 'Repair chunking, exact-reference parsing, recall, fusion, reranking, qualification or evidence allocation'),
    ('Matter facts', 'Repair attribution, missing-fact clarification, snapshot identity or FactRef linkage'),
    ('Prompt, code or validation', 'Repair the contract, output checker, fallback or deterministic renderer'),
    ('Gold or rubric', 'Independently correct the expected evidence/answer or evaluation rule before rerun'),
    ('Model capability', 'Consider a model/configuration change; weight training only under a later separate approved scope'),
]
for idx, rowdata in enumerate(repair_rows, start=1):
    cells = repair.add_row().cells
    cells[0].text, cells[1].text = rowdata
    if idx % 2 == 0:
        for cell in cells:
            set_cell_shading(cell, PALE_BLUE)
set_table_fixed(repair, [2.00, 4.35])
format_table_text(repair, 9.0)
set_all_table_borders(repair)

add_body(doc, 'Training exclusions:', bold_lead='Training exclusions:')
add_list(doc, [
    'No evaluation questions/answers, reviewer notes or quality corrections as training data.',
    'No user conversation histories or uploads.',
    'No private unseen prompts, expected answers, metadata or findings.',
    'No teaching/feedback material treated as independent legal authority.',
    'No weight change until a separate rights, privacy, dataset, resource and execution decision is approved.',
])

doc.add_page_break()
add_heading(doc, '7. Owner decisions', 1)
add_body(doc, 'The design amendments are already accepted. Record only the Phase-2 preparation decisions below.')

decisions = doc.add_table(rows=1, cols=4)
add_table_header(decisions, ['No.', 'Recommended decision', 'Owner choice', 'Owner comment / amendment'])
decision_rows = [
    ('1', 'Use all 331 accepted visible GE cases for the baseline; report the 32 system scenarios separately; do not use unseen.', 'APPROVE / AMEND / DECLINE', '________________________________\n________________________________'),
    ('2', 'Use the factual/legal hard gate first, then practical GE and adapted 70+ PB/Essay quality; deliver readable answers plus one final owner approval DOCX.', 'APPROVE / AMEND / DECLINE', '________________________________\n________________________________'),
    ('3', 'After baseline review, allow diagnosis and authorized non-weight repairs. Keep weight training separate. Return improved visible results before any unseen run.', 'APPROVE / AMEND / DECLINE', '________________________________\n________________________________'),
]
for idx, rowdata in enumerate(decision_rows, start=1):
    cells = decisions.add_row().cells
    for j, value in enumerate(rowdata):
        cells[j].text = value
    for cell in cells:
        set_cell_shading(cell, GREEN if idx in (1, 2, 3) else WHITE)
set_table_fixed(decisions, [0.45, 3.05, 1.25, 1.60])
format_table_text(decisions, 8.4)
set_all_table_borders(decisions)

add_heading(doc, 'Owner decision record', 2)
record = doc.add_table(rows=4, cols=2)
set_table_fixed(record, [1.65, 4.70])
record_rows = [
    ('Overall decision', 'APPROVE ALL / APPROVE WITH AMENDMENTS / DO NOT APPROVE'),
    ('Owner name', '____________________________________________________________'),
    ('Decision date', '____________________________________________________________'),
    ('Additional conditions', '____________________________________________________________\n____________________________________________________________'),
]
for idx, (label, value) in enumerate(record_rows):
    set_cell_shading(record.cell(idx, 0), PALE_GREY)
    record.cell(idx, 0).text = label
    record.cell(idx, 1).text = value
format_table_text(record, 9.2, has_header=False)
set_all_table_borders(record)

add_callout(
    doc,
    'Fast approval reply',
    'If the three recommendations are acceptable, reply: “Approve the three recommended Phase-2 decisions.”',
    fill=GREEN,
)

doc.add_page_break()
add_heading(doc, '8. What this document does not approve', 1)
add_list(doc, [
    'No source admission or owner legal-currentness judgment.',
    'No gold-answer acceptance or change to an existing certification/split contract.',
    'No model artifact, transport, resource envelope, signing key or private review root.',
    'No actual evaluation, training, unseen disclosure/run, promotion or live activation.',
    'No Git mutation, public accounts, cloud storage, sharing or external human-referral integration.',
])
add_body(doc, 'Those decisions are requested only when the exact, reviewable input or proposed action is ready. This keeps the owner workflow short without weakening legal, privacy or release controls.')

add_heading(doc, '9. Reference basis for the 70+ target', 1)
add_body(doc, 'The quality framework was adapted from permitted owner reference materials in the Law folder. These materials guide answer quality; they do not establish legal truth and are not instructions to execute the system.')
add_list(doc, [
    'Official Levels 2 and 3 Assessment Criteria: first-class (70–100) descriptors for knowledge, problem solving, research, synthesis, evaluation, independence and communication.',
    'Y3 Pensions Law general feedback: current authority, counterarguments, missing/ambiguous facts, independent research, sound calculations, SPaG and OSCOLA precision.',
    'Selected Y3 feedback examples: stronger criticism of judgments, rationale and dissent, wider academic writing, secondary sources and pinpoint citations; strong work combined legal/policy analysis, academic views and evidence while still improving counterarguments.',
])
add_body(doc, 'For GE, the same intellectual discipline is expressed through a concise public-facing answer. The product should be legally careful and useful without imitating an academic essay.')

# Document metadata
props = doc.core_properties
props.title = 'LegalBot Phase 2 — Owner Approval'
props.subject = 'Factual-first evaluation, quality review, improvement and unseen testing'
props.author = 'LegalBot project'
props.keywords = 'LegalBot, Phase 2, evaluation, quality, owner approval'
props.comments = 'Prepared for owner decision; no run started.'

doc.save(OUTPUT)
print(OUTPUT)
