#!/usr/bin/env python3
"""Render the r94 digest-bound Phase-2A owner review as a polished DOCX."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
from collections import Counter, defaultdict
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
R94_ROOT = (
    PROJECT_ROOT / "data/evaluations/phase2a-owner-review/"
    "LegalBot-Phase2AB-2026-08-25-r94-consolidated-substantive-owner-batch"
)
DEFAULT_BATCH = R94_ROOT / "OWNER-SUBSTANTIVE-DECISION-BATCH.json"
DEFAULT_TRIAGE = (
    PROJECT_ROOT / "data/evaluations/phase2a-owner-review/"
    "LegalBot-Phase2AB-2026-08-25-r71-gap-triage/ISSUE-GAP-TRIAGE-448.json"
)
DEFAULT_OUTPUT = R94_ROOT / "LegalBot-Phase2A-Owner-Decision-Batch-Agnes-2026-08-25.docx"
EXPECTED_BATCH_DIGEST = "496c36b665b8114f6ba44169d23f911ec31fc2718aecb1747b2e26962bd889f7"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
GOLD_FILL = "FFF4CE"
GOLD_TEXT = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "000000"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def _canonical_json(value: Any) -> bytes:
    return (
        json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")) + "\n"
    ).encode()


def _load_object(path: Path) -> dict[str, Any]:
    if path.is_symlink() or not path.is_file():
        raise ValueError("phase2a_owner_docx_input_not_regular")
    value = json.loads(path.read_bytes())
    if not isinstance(value, dict):
        raise ValueError("phase2a_owner_docx_input_not_object")
    return value


def _verify_batch(value: dict[str, Any]) -> None:
    material = dict(value)
    supplied = str(material.pop("artifact_content_sha256", ""))
    calculated = hashlib.sha256(_canonical_json(material)).hexdigest()
    if supplied != calculated or supplied != EXPECTED_BATCH_DIGEST:
        raise ValueError("phase2a_owner_docx_batch_digest_invalid")
    if (
        value.get("owner_approved") is not False
        or value.get("source_admission_authorized") is not False
        or value.get("candidate_mutated") is not False
        or value.get("phase2b_authorized") is not False
        or value.get("development30_authorized") is not False
    ):
        raise ValueError("phase2a_owner_docx_batch_boundary_invalid")


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _prevent_row_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(
    cell: Any, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(
    table: Any, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA
) -> None:
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError("phase2a_owner_docx_table_width_invalid")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    width.set(qn("w:type"), "dxa")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for cell, value in zip(row.cells, widths_dxa, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(value))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(value / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)


def _set_run(
    run: Any, *, size: float = 10.0, bold: bool = False, color: str = BLACK, italic: bool = False
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_text(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    color: str = BLACK,
    size: float = 9.0,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    _set_run(paragraph.add_run(text), size=size, bold=bold, color=color)


def _add_table(
    doc: Document,
    headers: list[str],
    widths: list[int],
    rows: Iterable[list[str]],
    *,
    font_size: float = 8.6,
) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        _set_cell_shading(table.rows[0].cells[index], BLUE_GRAY)
        _set_cell_text(table.rows[0].cells[index], header, bold=True, color=INK, size=9.0)
    _set_repeat_table_header(table.rows[0])
    _prevent_row_split(table.rows[0])
    for values in rows:
        row = table.add_row()
        _prevent_row_split(row)
        for index, value in enumerate(values):
            _set_cell_text(row.cells[index], value, size=font_size)
    _set_table_geometry(table, widths)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    return table


def _add_heading(doc: Document, text: str, level: int = 1, *, new_page: bool = False) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = new_page


def _add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        _set_run(paragraph.add_run(bold_lead), size=11, bold=True)
        _set_run(paragraph.add_run(text[len(bold_lead) :]), size=11)
    else:
        _set_run(paragraph.add_run(text), size=11)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    _set_run(paragraph.add_run(text), size=11)


def _add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    _set_run(paragraph.add_run(text), size=11)


def _add_callout(
    doc: Document, label: str, text: str, *, fill: str = GOLD_FILL, color: str = GOLD_TEXT
) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    _set_run(paragraph.add_run(f"{label}: "), size=10.5, bold=True, color=color)
    _set_run(paragraph.add_run(text), size=10.5, color=BLACK)
    _set_table_geometry(table, [TABLE_WIDTH_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _add_rule(paragraph: Any, color: str = BLUE, size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def _add_page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)
    _set_run(run, size=9, color=MUTED)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    for header in (section.header, section.even_page_header):
        table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_cell_text(
            table.cell(0, 0),
            "LEGALBOT v1.11 | PHASE 2A OWNER GATE",
            bold=True,
            color=MUTED,
            size=8.5,
        )
        _set_cell_text(
            table.cell(0, 1),
            "CONFIDENTIAL - OWNER REVIEW",
            bold=True,
            color=MUTED,
            size=8.5,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        _set_table_geometry(table, [5700, 3660], indent_dxa=0)
    for footer_container in (section.footer, section.even_page_footer):
        footer = footer_container.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run(footer.add_run("LegalBot v1.11 | Page "), size=9, color=MUTED)
        _add_page_field(footer)


def _read_source(batch: dict[str, Any], source_name: str) -> dict[str, Any]:
    reference = next(row for row in batch["source_artifacts"] if row["source_name"] == source_name)
    return _load_object(PROJECT_ROOT / reference["relative_path"])


def _short(value: str, length: int = 12) -> str:
    return value[:length]


def _source_admission_rows(batch: dict[str, Any]) -> list[list[str]]:
    rows: list[list[str]] = []
    for item in batch["judgment_source_admission_decisions"]:
        use = "; ".join(
            f"{relation['target_neutral_citation']}: {relation['advisory_relationship']}"
            for source in [_read_source(batch, "judgments")]
            for proposal in source["source_admission_proposals"]
            if proposal["proposal_id"] == item["proposal_id"]
            for relation in proposal["proposition_level_relationships"]
        )
        rows.append(
            [
                "Judgment",
                item["neutral_citation"],
                use,
                _short(item["source_proposal_content_sha256"]),
            ]
        )
    supplemental = batch["supplemental_binding_and_source_decisions"]
    grouped: dict[tuple[str, str], list[dict[str, Any]]] = defaultdict(list)
    for item in supplemental:
        if item["owner_source_admission_required"]:
            grouped[(item["source_batch"], item["source_target_id"])].append(item)
    for (_group, source_id), values in sorted(grouped.items()):
        rows.append(
            [
                "Official legislation",
                source_id,
                ", ".join(sorted({row for value in values for row in value["row_ids"]})),
                _short(values[0]["proposal_content_sha256"]),
            ]
        )
    for item in batch["uksc_source_review_decisions"]:
        if item["recommended_owner_outcome"] == "APPROVE_PROPOSITION_BINDINGS_AND_SOURCE_ADMISSION":
            rows.append(
                [
                    "UKSC judgment",
                    item["neutral_citation"],
                    f"{len(item['supported_claim_binding_content_sha256s'])} exact claim binding(s)",
                    _short(item["source_review_content_sha256"]),
                ]
            )
    if len(rows) != 20:
        raise ValueError("phase2a_owner_docx_source_admission_count_invalid")
    return rows


def render(*, batch_path: Path, triage_path: Path, output_path: Path) -> dict[str, Any]:
    if output_path.exists() or output_path.is_symlink():
        raise ValueError("phase2a_owner_docx_output_exists")
    batch = _load_object(batch_path)
    _verify_batch(batch)
    triage = _load_object(triage_path)
    triage_by_id = {str(row["row_id"]): row for row in triage["rows"]}
    if len(triage_by_id) != 448:
        raise ValueError("phase2a_owner_docx_triage_inventory_invalid")

    doc = Document()
    _configure_styles(doc)
    _configure_page(doc)
    doc.core_properties.title = "LegalBot v1.11 Phase 2A Consolidated Owner Decision Package"
    doc.core_properties.subject = "Digest-bound Phase 2A owner review"
    doc.core_properties.author = "LegalBot v1.11"
    doc.core_properties.keywords = "Phase 2A, owner decision, evidence, source admission"

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(4)
    _set_run(title.add_run("PHASE 2A OWNER DECISION PACKAGE"), size=23, bold=True, color=BLACK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_run(
        subtitle.add_run("LegalBot v1.11 - Consolidated substantive batch r94"),
        size=14,
        color=MUTED,
    )
    for label, value in (
        ("Owner", "Agnes"),
        ("Decision date", "2026-08-25"),
        ("Qualification route", "Owner-adopted internal research tool"),
        ("Status", "Exact owner approval required - continued Phase 2A only"),
        ("Batch digest", EXPECTED_BATCH_DIGEST),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        _set_run(paragraph.add_run(f"{label}: "), size=10.5, bold=True)
        _set_run(paragraph.add_run(value), size=10.5, color=(RED if label == "Status" else BLACK))
    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    _add_rule(rule)
    _add_callout(
        doc,
        "Decision gate",
        "Approval applies only to the exact listed recommendations and 20 official-source admissions. It does not qualify the 364 unresolved material-gap rows and does not authorize Phase 2B or Development 30.",
    )

    _add_heading(doc, "1. Decision requested")
    _add_body(
        doc,
        "Approve the exact r94 batch digest so Codex may apply the evidence-ready Phase-2A decisions and admit the 20 identified official sources into the later consolidated successor-source scope. No indexing, embedding, or candidate rebuild occurs merely because this batch is approved.",
    )
    summary = batch["decision_summary"]
    _add_table(
        doc,
        ["Decision group", "Count", "Effect of approval"],
        [3000, 900, 5460],
        [
            [
                "Candidate exact-span bindings",
                str(summary["candidate_exact_binding_decision_count"]),
                "Accept proposition/span use for continued remediation; no row is technically qualified yet.",
            ],
            [
                "Judgment later-treatment dispositions",
                str(summary["judgment_later_treatment_decision_count"]),
                "Apply the listed recommended relationship or no-reliance outcome.",
            ],
            [
                "Unique official-source admissions",
                str(summary["total_unique_source_admission_count"]),
                "Admit only for the exact proposition-level uses listed in the batch.",
            ],
            [
                "Supplemental binding decisions",
                str(summary["supplemental_binding_decision_count"]),
                "Accept exact official material/currentness spans.",
            ],
            [
                "UKSC source reviews",
                str(summary["uksc_source_review_decision_count"]),
                "Accept direct mappings and reject the listed unrelated mappings.",
            ],
            [
                "XML byte mismatch",
                "1",
                "Treat as serialization-only because canonical XML is identical.",
            ],
            [
                "Patents s 60(7) delta",
                "1",
                "Defer substantive outcome until final Patents bindings are available.",
            ],
            [
                "Unresolved material-gap rows",
                str(summary["unresolved_material_gap_row_count"]),
                "Remain blocked and outside this approval.",
            ],
        ],
        font_size=9.0,
    )

    _add_heading(doc, "2. Boundaries")
    _add_body(
        doc, "This is an evidence-remediation decision, not legal certification or legal advice."
    )
    for text in (
        "AI recommendations remain advisory; Agnes is the substantive owner decision-maker.",
        "No source outside the exact 20-source admission inventory is approved.",
        "Crawler output and quarantined bytes remain staging evidence until the approved decisions are applied.",
        "The sealed predecessor candidate remains immutable and cannot be used as a silent fallback.",
        "Phase 2B still requires a successful final Phase-2A package and explicit adoption of its exact digest.",
        "Development 30 still requires Phase 2B success and explicit approval of its exact authorization payload.",
    ):
        _add_bullet(doc, text)

    _add_heading(doc, "3. Candidate proposition and exact-span decisions", new_page=True)
    _add_body(
        doc,
        "All 84 rows below have a sealed candidate-bound exact span. Approval accepts the proposition/span pairing for continued Phase-2A remediation only. Currentness or later-treatment work shown in the machine-readable record still remains a final qualification dependency.",
    )
    issue_rows: list[list[str]] = []
    for item in batch["candidate_exact_binding_decisions"]:
        triage_row = triage_by_id[item["row_id"]]
        currentness = item.get("source_currentness") or {}
        issue_rows.append(
            [
                item["row_id"],
                "Direct" if item["assessment"].startswith("DIRECT") else "Partial",
                item["atomic_proposition"],
                f"{triage_row['issue_label']} | {_short(item['exact_span_binding_content_sha256'])} | {currentness.get('currentness_status', 'review required')}",
            ]
        )
    _add_table(
        doc,
        ["Row", "Scope", "Atomic proposition", "Evidence/currentness"],
        [1550, 850, 4300, 2660],
        issue_rows,
        font_size=7.8,
    )

    _add_heading(doc, "4. Judgment later-treatment decisions", new_page=True)
    judgments = _read_source(batch, "judgments")
    judgment_by_citation = {row["neutral_citation"]: row for row in judgments["records"]}
    _add_table(
        doc,
        ["Citation", "Case", "Recommended outcome", "Category"],
        [1400, 2200, 3400, 2360],
        [
            [
                item["neutral_citation"],
                judgment_by_citation[item["neutral_citation"]]["title"],
                item["recommended_owner_outcome"].replace("_", " "),
                judgment_by_citation[item["neutral_citation"]]["advisory_category"].replace(
                    "_", " "
                ),
            ]
            for item in batch["judgment_later_treatment_decisions"]
        ],
        font_size=8.2,
    )
    _add_callout(
        doc,
        "Search limitation",
        "The judgment review records bounded official searches. They expressly do not treat absence of another hit as proof that no other later treatment exists.",
        fill=LIGHT_GRAY,
        color=DARK_BLUE,
    )

    _add_heading(doc, "5. Exact official-source admission inventory", new_page=True)
    _add_body(
        doc,
        "Approval admits exactly these 20 official primary sources for the proposition-level uses shown. Admission does not itself index, embed, or rebuild a candidate.",
    )
    _add_table(
        doc,
        ["Group", "Source", "Approved use", "Record digest"],
        [1500, 2200, 4400, 1260],
        _source_admission_rows(batch),
        font_size=8.0,
    )

    _add_heading(doc, "6. UKSC mapping corrections")
    uksc_source = _read_source(batch, "uksc_rehoming")
    uksc_by_id = {row["proposal_id"]: row for row in uksc_source["source_reviews"]}
    _add_table(
        doc,
        ["Citation", "Decision", "Supported rows", "Rejected rows"],
        [1550, 2350, 2700, 2760],
        [
            [
                item["neutral_citation"],
                item["recommended_owner_outcome"].replace("_", " "),
                ", ".join(
                    sorted(
                        {
                            row
                            for claim in uksc_by_id[item["proposal_id"]]["supported_claims"]
                            for row in claim["row_ids"]
                        }
                    )
                )
                or "None",
                ", ".join(
                    row["row_id"] for row in uksc_by_id[item["proposal_id"]]["rejected_mappings"]
                )
                or "None",
            ]
            for item in batch["uksc_source_review_decisions"]
        ],
        font_size=8.0,
    )
    _add_body(
        doc,
        "Key correction: Byers v Saudi National Bank is proposed for the knowing-receipt row; Lifestyle Equities is rejected for the unrelated restitution/tracing mappings.",
    )

    _add_heading(doc, "7. Statute and byte-level dispositions")
    _add_table(
        doc,
        ["Item", "Verified finding", "Recommended decision"],
        [2100, 3800, 3460],
        [
            [
                "Data (Use and Access) Act 2025 XML",
                "Raw bytes differ, but canonical XML digest is identical and no legal XML infoset change was detected.",
                "Approve nonmaterial serialization-only disposition; preserve both raw provenance records.",
            ],
            [
                "Patents Act 1977 s 60(7)",
                "Fresh text removes the obsolete Civil Aviation Act reference; no final exact proposition binding yet resolves the benchmark use.",
                "Defer owner outcome until final Patents proposition bindings; if retained, use a fresh successor version, never an in-place patch.",
            ],
        ],
        font_size=8.6,
    )

    _add_heading(doc, "8. Rows that remain blocked", new_page=True)
    _add_body(
        doc,
        "The 364 rows below are not approved or qualified by this package. They require exact official-source research, gold/issue repair, or owner review before Phase 2A can pass.",
    )
    by_case: dict[str, list[str]] = defaultdict(list)
    class_counts: Counter[str] = Counter()
    for item in batch["unresolved_material_gap_rows"]:
        row = triage_by_id[item["row_id"]]
        by_case[str(row["case_id"])].append(str(row["row_id"]).split(":", 1)[1])
        class_counts[str(row["triage_class"])] += 1
    _add_table(
        doc,
        ["Deterministic triage class", "Rows"],
        [6200, 3160],
        [[name.replace("_", " "), str(count)] for name, count in sorted(class_counts.items())],
        font_size=9.0,
    )
    _add_heading(doc, "Blocked-row inventory by case", 2)
    case_rows = [
        [case_id, str(len(values)), ", ".join(values)]
        for case_id, values in sorted(by_case.items())
    ]
    _add_table(doc, ["Case", "Count", "Issue rows"], [1650, 800, 6910], case_rows, font_size=8.2)

    _add_heading(doc, "9. What happens after approval")
    for text in (
        "Apply only the exact r94 decisions and preserve an owner-approval receipt bound to the digest.",
        "Continue official-source research and exact-span remediation for the 364 blocked rows.",
        "Consolidate any approved missing/stale sources into one successor-source manifest; do not patch the predecessor.",
        "Build and verify one successor candidate only when the final source scope is proven.",
        "Run retrieval re-attestation and all-585 technical qualification.",
        "Return the final Phase-2A package and exact digest for separate owner adoption.",
        "Only after that adoption may the conditional Phase-2B authorization activate.",
    ):
        _add_number(doc, text)

    _add_heading(doc, "10. Exact approval text", new_page=True)
    _add_callout(
        doc,
        "Copy exactly",
        "The digest must remain unchanged. A general OK does not approve a different package.",
        fill=BLUE_GRAY,
        color=DARK_BLUE,
    )
    prompt = (R94_ROOT / "OWNER-APPROVAL-PROMPT.txt").read_text(encoding="utf-8").strip()
    prompt_table = doc.add_table(rows=1, cols=1)
    prompt_table.style = "Table Grid"
    prompt_cell = prompt_table.cell(0, 0)
    _set_cell_shading(prompt_cell, LIGHT_GRAY)
    _set_cell_margins(prompt_cell, top=160, bottom=160, start=180, end=180)
    prompt_cell.text = ""
    for index, line in enumerate(prompt.splitlines()):
        paragraph = prompt_cell.paragraphs[0] if index == 0 else prompt_cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)
        _set_run(
            paragraph.add_run(line or " "),
            size=9.2,
            bold=(index == 0),
            color=(DARK_BLUE if index == 0 else BLACK),
        )
    _set_table_geometry(prompt_table, [TABLE_WIDTH_DXA])

    _add_heading(doc, "11. Audit identities")
    _add_body(doc, f"Owner batch digest: {EXPECTED_BATCH_DIGEST}")
    for reference in batch["source_artifacts"]:
        _add_body(
            doc,
            f"{reference['source_name']}: {reference['artifact_content_sha256']} (file {_short(reference['file_sha256'], 16)})",
        )
    _add_body(
        doc,
        "Reviewer status: advisory review uses the same pinned model adapter as drafting and is not a genuinely model-independent reviewer. Deterministic evidence checks remain controlling.",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = output_path.with_suffix(".tmp.docx")
    doc.save(temp_path)
    descriptor = os.open(output_path, os.O_WRONLY | os.O_CREAT | os.O_EXCL | os.O_NOFOLLOW, 0o600)
    try:
        with os.fdopen(descriptor, "wb") as target, temp_path.open("rb") as source:
            for block in iter(lambda: source.read(1024 * 1024), b""):
                target.write(block)
            target.flush()
            os.fsync(target.fileno())
    except BaseException:
        output_path.unlink(missing_ok=True)
        raise
    finally:
        temp_path.unlink(missing_ok=True)
    return {
        "output_path": str(output_path),
        "output_sha256": hashlib.sha256(output_path.read_bytes()).hexdigest(),
        "owner_batch_content_sha256": EXPECTED_BATCH_DIGEST,
        "owner_approved": False,
        "phase2b_authorized": False,
        "development30_authorized": False,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--batch", type=Path, default=DEFAULT_BATCH)
    parser.add_argument("--triage", type=Path, default=DEFAULT_TRIAGE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    result = render(
        batch_path=args.batch.resolve(strict=True),
        triage_path=args.triage.resolve(strict=True),
        output_path=args.output.resolve(),
    )
    print(json.dumps(result, ensure_ascii=False, sort_keys=True, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
