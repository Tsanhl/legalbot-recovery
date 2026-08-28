#!/usr/bin/env python3
"""Render the exact post-r110 Phase-2A owner decision packet as DOCX."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import stat
import sys
from collections import defaultdict
from collections.abc import Mapping, Sequence
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

PROJECT_ROOT = Path(__file__).resolve().parents[1]
REVIEW_ROOT = PROJECT_ROOT / "data/evaluations/phase2a-owner-review"
R111_ROOT = REVIEW_ROOT / "LegalBot-Phase2AB-2026-08-26-r111-source-currentness-owner-batch"
BATCH_PATH = R111_ROOT / "OWNER-SOURCE-CURRENTNESS-DECISION-BATCH.json"
PROMPT_PATH = R111_ROOT / "OWNER-APPROVAL-PROMPT.txt"
BATCH_CONTENT_SHA256 = "6c9eda0de5c9c921b99127cac9c6e41bb3ae87151178e250b9f4abcf4a0d7fa1"
BATCH_FILE_SHA256 = "e657bffbf77e264fd658bb4c61bb3932f8e6e65fa10923b7eebcf20a8739cf20"
PROMPT_FILE_SHA256 = "46013e08350cc3a54d568cf3c7ab86a12ca3cfdde6b39320e48f009d71e21374"
DEFAULT_OUTPUT_ROOT = REVIEW_ROOT / "LegalBot-Phase2AB-2026-08-26-r112b-post-r110-owner-review-docx"
DOCX_NAME = "LegalBot-Phase2A-Post-r110-Owner-Decision-Agnes-2026-08-26.docx"
BUILD_MANIFEST_NAME = "DOCX-BUILD-MANIFEST.json"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x5F, 0x6B, 0x7A)
RISK = RGBColor(0x9B, 0x1C, 0x1C)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
PALE_BLUE = "E8EEF5"


def _canonical_json(value: Any) -> bytes:
    return (
        json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")) + "\n"
    ).encode("utf-8")


def _pretty_json(value: Any) -> bytes:
    return (json.dumps(value, ensure_ascii=False, sort_keys=True, indent=2) + "\n").encode("utf-8")


def _sha256(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _sealed(value: Any) -> str:
    return _sha256(_canonical_json(value))


def _write_exclusive(path: Path, raw: bytes) -> None:
    descriptor = os.open(
        path,
        os.O_WRONLY | os.O_CREAT | os.O_EXCL | os.O_NOFOLLOW,
        0o600,
    )
    try:
        with os.fdopen(descriptor, "wb") as handle:
            handle.write(raw)
            handle.flush()
            os.fsync(handle.fileno())
    except BaseException:
        path.unlink(missing_ok=True)
        raise


def _load_inputs() -> tuple[dict[str, Any], str]:
    if BATCH_PATH.is_symlink() or not BATCH_PATH.is_file():
        raise ValueError("phase2a_r112_batch_not_regular")
    if PROMPT_PATH.is_symlink() or not PROMPT_PATH.is_file():
        raise ValueError("phase2a_r112_prompt_not_regular")
    if _sha256_file(BATCH_PATH) != BATCH_FILE_SHA256:
        raise ValueError("phase2a_r112_batch_file_digest_invalid")
    if _sha256_file(PROMPT_PATH) != PROMPT_FILE_SHA256:
        raise ValueError("phase2a_r112_prompt_file_digest_invalid")
    batch = json.loads(BATCH_PATH.read_bytes())
    if not isinstance(batch, dict):
        raise ValueError("phase2a_r112_batch_not_object")
    material = dict(batch)
    supplied = str(material.pop("artifact_content_sha256", ""))
    if supplied != BATCH_CONTENT_SHA256 or supplied != _sealed(material):
        raise ValueError("phase2a_r112_batch_content_seal_invalid")
    for field in (
        "owner_approved",
        "owner_decisions_applied",
        "source_admission_authorized",
        "automatic_indexing",
        "automatic_embedding",
        "candidate_mutated",
        "technical_qualification_assigned",
        "phase2b_authorized",
        "development30_authorized",
    ):
        if batch.get(field) is not False:
            raise ValueError("phase2a_r112_batch_boundary_invalid")
    return batch, PROMPT_PATH.read_text(encoding="utf-8")


def _set_run_font(
    run,
    *,
    name: str = "Arial",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, *, name: str, size: float, color: RGBColor) -> None:
    style.font.name = name
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = color


def _set_spacing(style, *, before: float, after: float, line: float) -> None:
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line


def _keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    _set_style_font(normal, name="Arial", size=11, color=INK)
    _set_spacing(normal, before=0, after=6, line=1.10)

    title = doc.styles["Title"]
    _set_style_font(title, name="Arial", size=23, color=RGBColor(0, 0, 0))
    title.font.bold = True
    _set_spacing(title, before=0, after=4, line=1.0)

    subtitle = doc.styles["Subtitle"]
    _set_style_font(subtitle, name="Arial", size=14, color=MUTED)
    _set_spacing(subtitle, before=0, after=16, line=1.1)

    heading_tokens = {
        "Heading 1": (16, BLUE, 12, 6),
        "Heading 2": (13, BLUE, 10, 5),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        _set_style_font(style, name="Arial", size=size, color=color)
        style.font.bold = True
        _set_spacing(style, before=before, after=after, line=1.05)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        _set_style_font(style, name="Arial", size=11, color=INK)
        _set_spacing(style, before=0, after=8, line=1.167)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.tab_stops.add_tab_stop(Inches(0.5))

    styles = doc.styles
    digest_style = styles.add_style("Digest", 1)
    _set_style_font(digest_style, name="Courier New", size=9, color=DARK_BLUE)
    _set_spacing(digest_style, before=2, after=7, line=1.0)
    digest_style.paragraph_format.keep_together = True

    quote_style = styles.add_style("Evidence Quote", 1)
    _set_style_font(quote_style, name="Arial", size=9.5, color=INK)
    quote_style.font.italic = True
    _set_spacing(quote_style, before=4, after=6, line=1.10)
    quote_style.paragraph_format.left_indent = Inches(0.25)
    quote_style.paragraph_format.right_indent = Inches(0.15)

    table_text = styles.add_style("Table Text", 1)
    _set_style_font(table_text, name="Arial", size=9.2, color=INK)
    _set_spacing(table_text, before=0, after=0, line=1.05)

    table_citation = styles.add_style("Table Citation", 1)
    _set_style_font(table_citation, name="Arial", size=8.5, color=MUTED)
    _set_spacing(table_citation, before=4, after=4, line=1.0)

    props = doc.core_properties
    props.title = "LegalBot v1.11 Phase 2A Post-r110 Owner Decision Packet"
    props.subject = "Exact digest-bound source and currentness owner gate"
    props.author = "LegalBot Phase 2A Remediation"
    props.keywords = "LegalBot; Phase 2A; owner decision; source admission"
    fixed = datetime(2026, 8, 26, 0, 0, 0, tzinfo=UTC)
    props.created = fixed
    props.modified = fixed


def _header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("LEGALBOT v1.11  |  PHASE 2A OWNER REVIEW")
    _set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("Owner review  |  Page ")
    _set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    _set_run_font(field_run, size=8.5, color=MUTED)
    field_run._r.extend((begin, instruction, separate, display, end))


def _add_key_value(doc: Document, label: str, value: str, *, after: float = 2) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    label_run = paragraph.add_run(f"{label}: ")
    _set_run_font(label_run, size=10.5, bold=True, color=INK)
    value_run = paragraph.add_run(value)
    _set_run_font(value_run, size=10.5, color=INK)


def _shade_paragraph(paragraph, fill: str, *, left_border: str = "2E74B5") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), left_border)
    borders.append(left)
    p_pr.append(borders)


def _box_paragraph(paragraph, fill: str, *, border_color: str = "1F2937") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "8")
        border.set(qn("w:color"), border_color)
        borders.append(border)
    p_pr.append(borders)


def _add_callout(
    doc: Document,
    label: str,
    text: str,
    *,
    fill: str = CALLOUT,
    color: RGBColor = DARK_BLUE,
) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    _shade_paragraph(paragraph, fill)
    lead = paragraph.add_run(f"{label}  ")
    _set_run_font(lead, size=10.5, color=color, bold=True)
    body = paragraph.add_run(text)
    _set_run_font(body, size=10.5, color=INK)


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_width(parent, tag: str, width: int) -> None:
    element = _ensure_child(parent, tag)
    element.set(qn("w:type"), "dxa")
    element.set(qn("w:w"), str(width))


def _apply_table_geometry(table, widths: Sequence[int]) -> None:
    widths = [int(value) for value in widths]
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("phase2a_r112_table_width_invalid")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    _set_width(tbl_pr, "w:tblW", CONTENT_WIDTH_DXA)
    indent = _ensure_child(tbl_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = _ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for column_index, cell in enumerate(row.cells):
            width = widths[column_index]
            cell.width = Twips(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            _set_width(tc_pr, "w:tcW", width)
            margins = _ensure_child(tc_pr, "w:tcMar")
            for side, value in CELL_MARGINS_DXA.items():
                margin = _ensure_child(margins, f"w:{side}")
                margin.set(qn("w:type"), "dxa")
                margin.set(qn("w:w"), str(value))


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _keep_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = _ensure_child(tc_pr, "w:shd")
    shading.set(qn("w:fill"), fill)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: RGBColor = INK,
    size: float = 9.2,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "Table Text"
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    _set_run_font(run, size=size, color=color, bold=bold)


def _add_summary_table(doc: Document, batch: Mapping[str, Any]) -> None:
    summary = batch["decision_summary"]
    rows = [
        ("Source-mapping decisions", str(summary["row_source_link_decision_count"])),
        ("Affected Phase-2A rows", str(summary["affected_unique_row_count"])),
        ("Unrelated mappings rejected", "21"),
        ("Mappings superseded", "2"),
        ("Partial new-source bindings", "2"),
        ("Partial existing-source binding", "1"),
        (
            "Proposition-level source admissions requested",
            str(summary["proposition_level_source_admission_count"]),
        ),
        (
            "Corrected same-adapter false negatives",
            str(summary["same_adapter_false_negative_count"]),
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _set_cell_text(table.rows[0].cells[0], "Decision inventory", bold=True)
    _set_cell_text(
        table.rows[0].cells[1],
        "Count",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    for cell in table.rows[0].cells:
        _shade_cell(cell, LIGHT_GRAY)
    _repeat_header(table.rows[0])
    for label, value in rows:
        row = table.add_row()
        _set_cell_text(row.cells[0], label)
        _set_cell_text(row.cells[1], value, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _keep_row(row)
    _apply_table_geometry(table, (7440, 1920))


def _add_source_table(doc: Document, batch: Mapping[str, Any]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ("Official source", "Date", "Affected row", "Proposed use")
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True)
        _shade_cell(table.rows[0].cells[index], PALE_BLUE)
    _repeat_header(table.rows[0])
    for source in batch["source_admission_decisions"]:
        row = table.add_row()
        title = f"{source['source_title']}\n{source['authority_identity_id']}"
        _set_cell_text(row.cells[0], title)
        _set_cell_text(
            row.cells[1],
            str(source["source_date"]),
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _set_cell_text(row.cells[2], "\n".join(source["affected_row_ids"]))
        _set_cell_text(row.cells[3], source["proposed_candidate_use"])
        _keep_row(row)
    _apply_table_geometry(table, (3000, 1200, 1950, 3210))
    citation = doc.add_paragraph(style="Table Citation")
    citation.add_run(
        "Scope: proposition-level admission only. No source is indexed or embedded "
        "by this decision packet."
    )


def _add_evidence_binding(doc: Document, binding: Mapping[str, Any]) -> None:
    _add_key_value(doc, "Atomic proposition", binding["atomic_proposition"], after=2)
    _add_key_value(doc, "Exact locator", binding["locator"], after=2)
    paragraph = doc.add_paragraph(style="Evidence Quote")
    _shade_paragraph(paragraph, "F8FAFC", left_border="7A8CA5")
    run = paragraph.add_run(f"“{binding['quote']}”")
    _set_run_font(run, size=9.5, color=INK, italic=True)
    digest = doc.add_paragraph(style="Digest")
    digest.add_run(
        "Binding SHA-256: "
        + str(binding.get("binding_content_sha256") or binding.get("quote_sha256"))
    )


def _add_masthead(doc: Document, batch: Mapping[str, Any]) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)
    title = doc.add_paragraph("OWNER DECISION PACKET", style="Title")
    _keep_with_next(title)
    subtitle = doc.add_paragraph(
        "LegalBot v1.11 - Phase 2A Post-r110 Source and Currentness Gate",
        style="Subtitle",
    )
    _keep_with_next(subtitle)
    for label, value in (
        ("Owner", "Agnes"),
        ("Decision date", "26 August 2026"),
        ("Status", "Exact owner approval required"),
        ("Route", "Owner-adopted internal research tool"),
        ("Professional legal certification", "No"),
    ):
        _add_key_value(doc, label, value)
    digest_label = doc.add_paragraph()
    digest_label.paragraph_format.space_before = Pt(8)
    digest_label.paragraph_format.space_after = Pt(1)
    run = digest_label.add_run("Exact approval-bound artifact digest")
    _set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
    digest = doc.add_paragraph(batch["artifact_content_sha256"], style="Digest")
    _shade_paragraph(digest, "EEF4FA", left_border="2E74B5")
    _add_callout(
        doc,
        "Gate status",
        "No source has been admitted, indexed, or embedded. The sealed predecessor "
        "candidate is unchanged. Phase 2B and Development 30 remain closed.",
        color=RISK,
        fill="FFF4F4",
    )


def _add_decision_overview(doc: Document, batch: Mapping[str, Any]) -> None:
    doc.add_heading("Decision overview", level=1)
    paragraph = doc.add_paragraph()
    paragraph.add_run("What approval does. ").bold = True
    paragraph.add_run(
        "It applies the 26 listed mapping dispositions and records proposition-level "
        "admission of exactly five official sources for continued Phase 2A."
    )
    paragraph = doc.add_paragraph()
    paragraph.add_run("What approval does not do. ").bold = True
    paragraph.add_run(
        "It does not technically qualify any row, index or embed any material, mutate "
        "a candidate, start Phase 2B, or authorize Development 30."
    )
    _add_summary_table(doc, batch)
    _add_callout(
        doc,
        "Why owner approval is required",
        "These are substantive source-to-proposition decisions. Deterministic checks "
        "and AI review may recommend an outcome, but only the owner may adopt it.",
    )


def _add_source_admissions(doc: Document, batch: Mapping[str, Any]) -> None:
    doc.add_heading("Five proposed official-source admissions", level=1)
    doc.add_paragraph(
        "Approval is limited to the exact source identities, representations, affected "
        "rows, uses, and evidence spans sealed in the batch."
    )
    _add_source_table(doc, batch)
    for index, source in enumerate(batch["source_admission_decisions"], start=1):
        heading = doc.add_heading(f"Source {index}: {source['authority_identity_id']}", level=2)
        _keep_with_next(heading)
        _add_key_value(doc, "Title", source["source_title"])
        _add_key_value(doc, "Source date", str(source["source_date"]))
        _add_key_value(doc, "Currentness status", source["currentness_status"])
        _add_key_value(doc, "Proposed candidate use", source["proposed_candidate_use"])
        _add_key_value(
            doc,
            "Source representation SHA-256",
            source["source_representation_sha256"],
        )
        _add_key_value(
            doc,
            "Canonical XML SHA-256",
            source["source_canonical_xml_sha256"],
        )
        for binding in source["exact_proposition_bindings"]:
            _add_evidence_binding(doc, binding)


def _mapping_label(outcome: str) -> str:
    return {
        "RECOMMEND_SUPERSEDE_WITH_CURRENT_AUTHORITY": ("Supersede with current authority"),
        "RECOMMEND_PARTIAL_BINDING_AND_SOURCE_ADMISSION": ("Partial binding and source admission"),
        "RECOMMEND_PARTIAL_EXISTING_SOURCE_BINDING": ("Partial existing-source binding"),
        "RECOMMEND_REJECT_MAPPING": "Reject unrelated mapping",
    }[outcome]


def _add_mapping_decisions(doc: Document, batch: Mapping[str, Any]) -> None:
    doc.add_heading("Detailed mapping decisions", level=1)
    doc.add_paragraph(
        "The groups below contain every one of the 26 decisions. A positive or "
        "superseding decision includes its exact bound evidence; a rejected mapping "
        "states the deterministic subject-matter mismatch."
    )
    grouped: dict[str, list[Mapping[str, Any]]] = defaultdict(list)
    for row in batch["mapping_decisions"]:
        grouped[row["recommended_owner_outcome"]].append(row)
    order = (
        "RECOMMEND_SUPERSEDE_WITH_CURRENT_AUTHORITY",
        "RECOMMEND_PARTIAL_BINDING_AND_SOURCE_ADMISSION",
        "RECOMMEND_PARTIAL_EXISTING_SOURCE_BINDING",
        "RECOMMEND_REJECT_MAPPING",
    )
    for outcome in order:
        rows = grouped[outcome]
        doc.add_heading(f"{_mapping_label(outcome)} ({len(rows)})", level=2)
        for row in rows:
            heading = doc.add_heading(
                f"Mapping {int(row['ordinal']):02d} | {row['row_id']}", level=3
            )
            _keep_with_next(heading)
            _add_key_value(doc, "Issue", row["issue_label"])
            _add_key_value(doc, "Mapped authority", row["mapped_authority_identity_id"])
            _add_key_value(doc, "Source title", row["source_title"])
            _add_key_value(doc, "Reason code", row["deterministic_reason_code"])
            _add_key_value(doc, "Rationale", row["deterministic_rationale"], after=4)
            replacements = row["replacement_authority_identity_ids"]
            if replacements:
                _add_key_value(doc, "Replacement authority", ", ".join(replacements))
            if row["same_adapter_false_negative"]:
                _add_callout(
                    doc,
                    "Reviewer correction",
                    "The same-adapter advisory pass labelled this source unrelated; "
                    "the deterministic official-source review found the listed exact "
                    "support. The advisory result is not used as a gate.",
                    fill="FFF8E8",
                    color=RGBColor(0x7A, 0x5A, 0x00),
                )
            for binding in row["exact_proposition_bindings"]:
                _add_evidence_binding(doc, binding)
            digest = doc.add_paragraph(style="Digest")
            digest.add_run("Decision SHA-256: " + row["decision_content_sha256"])


def _add_currentness_metadata(doc: Document, batch: Mapping[str, Any]) -> None:
    currentness = batch["currentness_metadata_only_decision"]
    doc.add_heading("Currentness metadata-only decision", level=1)
    _add_key_value(doc, "Authority", currentness["authority_identity_id"])
    _add_key_value(doc, "Affected row", ", ".join(currentness["affected_row_ids"]))
    _add_key_value(doc, "Relationship", currentness["treatment_relationship"])
    _add_key_value(
        doc,
        "Recommended outcome",
        currentness["recommended_owner_outcome"],
    )
    _add_callout(
        doc,
        "Candidate scope",
        "This later same-case judgment is proposed as currentness metadata only. "
        "Candidate-source admission is not recommended.",
    )
    exact_span = currentness["exact_span"]
    _add_key_value(doc, "Exact locator", exact_span["locator"])
    paragraph = doc.add_paragraph(style="Evidence Quote")
    _shade_paragraph(paragraph, "F8FAFC", left_border="7A8CA5")
    paragraph.add_run(f"“{exact_span['quote']}”")
    digest = doc.add_paragraph(style="Digest")
    digest.add_run("Decision SHA-256: " + currentness["decision_content_sha256"])


def _add_controls_and_provenance(doc: Document, batch: Mapping[str, Any]) -> None:
    doc.add_heading("Controls and provenance", level=1)
    controls = (
        "Owner decision not yet applied",
        "Source admission not yet authorized",
        "Automatic indexing disabled",
        "Automatic embedding disabled",
        "Candidate mutation disabled",
        "Technical qualification not assigned",
        "Phase 2B not authorized by this batch",
        "Development 30 not authorized by this batch",
    )
    for control in controls:
        doc.add_paragraph(control, style="List Bullet")
    doc.add_heading("Bound artifact identities", level=2)
    _add_key_value(
        doc,
        "r110 reconciliation content SHA-256",
        batch["source_r110_artifact_content_sha256"],
    )
    _add_key_value(
        doc,
        "r110 reconciliation file SHA-256",
        batch["source_r110_file_sha256"],
    )
    _add_key_value(
        doc,
        "r111 owner batch content SHA-256",
        batch["artifact_content_sha256"],
    )
    _add_key_value(doc, "r111 owner batch file SHA-256", BATCH_FILE_SHA256)


def _add_exact_prompt(doc: Document, prompt: str, digest: str) -> None:
    doc.add_page_break()
    doc.add_heading("Exact owner approval text", level=1)
    _add_callout(
        doc,
        "Action required",
        "Review this packet, then send the exact text below. Approval must retain "
        f"the full digest {digest}.",
    )
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_together = True
    _box_paragraph(paragraph, "FAFAFA")
    run = paragraph.add_run(prompt.rstrip("\n"))
    _set_run_font(run, name="Courier New", size=8.6, color=INK)
    first_run = paragraph.runs[0]
    first_run.bold = False
    footer = doc.add_paragraph(style="Table Citation")
    footer.add_run("This text is copied from the sealed r111 OWNER-APPROVAL-PROMPT.txt.")


def _save_exclusive(doc: Document, path: Path) -> None:
    temporary = path.parent / f".{path.name}.tmp-{os.getpid()}"
    if path.exists() or path.is_symlink() or temporary.exists():
        raise ValueError("phase2a_r112_docx_output_exists")
    try:
        doc.save(temporary)
        os.chmod(temporary, 0o600)
        os.link(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)


def build(output_root: Path = DEFAULT_OUTPUT_ROOT) -> dict[str, Any]:
    if output_root.exists() or output_root.is_symlink():
        raise ValueError("phase2a_r112_output_already_exists")
    batch, prompt = _load_inputs()
    output_root.mkdir(parents=True, mode=0o700)
    if stat.S_IMODE(output_root.stat().st_mode) != 0o700:
        raise ValueError("phase2a_r112_output_mode_invalid")

    doc = Document()
    _configure_document(doc)
    for section in doc.sections:
        _header_footer(section)
    _add_masthead(doc, batch)
    _add_decision_overview(doc, batch)
    _add_source_admissions(doc, batch)
    _add_mapping_decisions(doc, batch)
    _add_currentness_metadata(doc, batch)
    _add_controls_and_provenance(doc, batch)
    _add_exact_prompt(doc, prompt, batch["artifact_content_sha256"])

    docx_path = output_root / DOCX_NAME
    _save_exclusive(doc, docx_path)
    manifest_material = {
        "schema": "legalbot.v111.phase2a.post-r110-owner-docx-build.v1",
        "status": "DOCX_CREATED_VISUAL_QA_REQUIRED",
        "design_preset": "decision_memo",
        "header_pattern": "memo_masthead",
        "source_owner_batch_content_sha256": BATCH_CONTENT_SHA256,
        "source_owner_batch_file_sha256": BATCH_FILE_SHA256,
        "source_owner_prompt_file_sha256": PROMPT_FILE_SHA256,
        "docx_member": DOCX_NAME,
        "docx_file_sha256": _sha256_file(docx_path),
        "visual_qa_completed": False,
        "owner_approved": False,
        "source_admission_authorized": False,
        "candidate_mutated": False,
        "phase2b_authorized": False,
        "development30_authorized": False,
    }
    manifest = {
        **manifest_material,
        "manifest_content_sha256": _sealed(manifest_material),
    }
    _write_exclusive(
        output_root / BUILD_MANIFEST_NAME,
        _pretty_json(manifest),
    )
    _write_exclusive(
        output_root / "OUTCOME.txt",
        b"OWNER REVIEW DOCX CREATED; RENDER AND VISUAL QA REQUIRED BEFORE DELIVERY; "
        b"NO OWNER DECISION APPLIED; PHASE 2B CLOSED.\n",
    )
    return manifest


def _persist_failure(output_root: Path, exc: BaseException) -> None:
    try:
        output_root.mkdir(parents=True, mode=0o700, exist_ok=True)
        path = output_root / "FAILURE.json"
        if path.exists() or path.is_symlink():
            return
        fingerprint_material = {
            "exception_type": type(exc).__name__,
            "error": str(exc),
            "affected_stage": "PHASE2A_POST_R110_OWNER_DOCX_BUILD",
        }
        material = {
            "schema": "legalbot.v111.phase2a.post-r110-owner-docx-failure.v1",
            "failure_fingerprint": _sealed(fingerprint_material),
            **fingerprint_material,
            "affected_rows": "26_SOURCE_LINKS_ACROSS_22_ROWS",
            "completed_work": "PRESERVED_BEFORE_EXCEPTION",
            "root_cause_status": "DEBUG_REQUIRED",
            "required_execution_plan_change": (
                "INSPECT_FAILURE_AND_BOUND_R111_INPUTS_BEFORE_RETRY"
            ),
            "source_admission_authorized": False,
            "candidate_mutated": False,
            "phase2b_authorized": False,
            "development30_authorized": False,
        }
        _write_exclusive(
            path,
            _pretty_json({**material, "failure_content_sha256": _sealed(material)}),
        )
    except Exception:
        return


def main(argv: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-root", type=Path, default=DEFAULT_OUTPUT_ROOT)
    args = parser.parse_args(argv)
    output_root = args.output_root.resolve()
    try:
        manifest = build(output_root)
    except Exception as exc:
        _persist_failure(output_root, exc)
        print(f"{type(exc).__name__}: {exc}", file=sys.stderr)
        return 1
    print(json.dumps(manifest, ensure_ascii=False, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
