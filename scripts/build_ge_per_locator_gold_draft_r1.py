#!/usr/bin/env python3
"""Create-only per-locator gold DRAFT for the 43 PASS locators plus 008/174/312.

Unsigned. Every row is PENDING. legal_gold stays false until the owner ticks
APPROVE on a later signed receipt pack.
"""

from __future__ import annotations

import hashlib
import json
import os
import stat
from collections.abc import Mapping
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PACK = (
    ROOT
    / "data/evaluations/general-enquiries"
    / "LegalBot-GE-2026-09-02-per-locator-gold-draft-r1"
)
RESULTS = (
    ROOT
    / "data/evaluations/general-enquiries"
    / "LegalBot-GE-2026-09-02-visible-331-diagnostic-r1"
    / "visible/RESULTS.jsonl"
)
MANIFEST = (
    ROOT
    / "data/indexes/builds/current-law-ew-full-fp16-v111-20260829-recovery-b"
    / "approved-source-manifest.json"
)
MECHANICAL = (
    ROOT
    / "data/evaluations/general-enquiries"
    / "LegalBot-GE-2026-09-02-authorized-next-work-r1"
    / "MECHANICAL-CURRENTNESS-REVIEW.json"
)
DOCX_PATH = ROOT / "output/docx/LegalBot-GE-2026-09-02-Per-Locator-Gold-Draft-r1.docx"
TULRCA_ID = "ukpga-1992-52"
EVALUATION_AS_OF = "2026-08-28"

EXTRA_ROWS: tuple[dict[str, Any], ...] = (
    {
        "case_ids": ["administrative-law:cp-d08"],
        "title": "Equality Act 2010",
        "locator": "Schedule 2",
        "source_class": "legislation",
        "notes": "Overlay route for case 008; not selected in the 331 r1 evidence rows.",
    },
    {
        "case_ids": ["administrative-law:cp-d08"],
        "title": "The Public Sector Bodies (Websites and Mobile Applications) (No. 2) Accessibility Regulations 2018",
        "locator": "regulation 12",
        "source_class": "legislation",
        "notes": "Staged as SI 2018/952; not in the 85-source recovery-b index.",
        "staging_id": "uksi-2018-952",
    },
    {
        "case_ids": ["international-commercial-mediation:cp-d01"],
        "title": "ICC Mediation Rules (contractually incorporated edition)",
        "locator": "article 5",
        "source_class": "official_secondary",
        "notes": "Staged HTML. Owner must bind the contractually incorporated edition. Not gold.",
        "staging_id": "icc-mediation-rules-html",
    },
    {
        "case_ids": ["international-commercial-mediation:cp-d01"],
        "title": "Ohpen Operations UK Ltd v Invesco Fund Managers Ltd",
        "locator": "judgment",
        "source_class": "judgment",
        "notes": "Find Case Law XML captured. Proposition currentness remains HOLD unless later treatment is confirmed.",
        "staging_id": "ewhc-tcc-2019-2246",
    },
    {
        "case_ids": ["international-commercial-mediation:cp-d01"],
        "title": "Kajima Construction Europe (UK) Ltd v Children's Ark Partnership Ltd",
        "locator": "judgment",
        "source_class": "judgment",
        "notes": "Find Case Law XML captured. Proposition currentness remains HOLD.",
        "staging_id": "ewca-civ-2023-292",
    },
    {
        "case_ids": ["international-commercial-mediation:cp-d01"],
        "title": "Churchill v Merthyr Tydfil County Borough Council",
        "locator": "judgment",
        "source_class": "judgment",
        "notes": "Find Case Law XML captured. Proposition currentness remains HOLD.",
        "staging_id": "ewca-civ-2023-1416",
    },
    {
        "case_ids": ["international-commercial-mediation:cp-d01"],
        "title": "Cable & Wireless plc v IBM United Kingdom Ltd",
        "locator": "judgment",
        "source_class": "judgment",
        "notes": "Find Case Law data.xml HTTP 404. Remain fail-closed. Do not use BAILII unless later named.",
        "fail_closed": True,
    },
    {
        "case_ids": ["wills-and-estates:cp-d02"],
        "title": "Wills Act 1837 (as at 2024-01-15)",
        "locator": "section 9",
        "source_class": "legislation",
        "notes": "Point-in-time snapshot required. Latest 2026 Wills Act text is not this locator. Validity remains HOLD.",
        "point_in_time_as_at": "2024-01-15",
    },
    {
        "case_ids": ["wills-and-estates:cp-d02"],
        "title": "The Wills Act 1837 (Electronic Communications) (Amendment) (Coronavirus) Order 2020",
        "locator": "article 2",
        "source_class": "legislation",
        "notes": "Staged SI 2020/952. Validity of the video-will remains HOLD even if formality text is approved.",
        "staging_id": "uksi-2020-952",
        "point_in_time_as_at": "2024-01-15",
    },
    {
        "case_ids": ["wills-and-estates:cp-d02"],
        "title": "The Wills Act 1837 (Electronic Communications) (Amendment) Order 2022",
        "locator": "article 2",
        "source_class": "legislation",
        "notes": "Staged SI 2022/18. Extends the video-will window to 31 January 2024.",
        "staging_id": "uksi-2022-18",
        "point_in_time_as_at": "2024-01-15",
    },
)


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _canonical_bytes(value: Any) -> bytes:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()


def _digest(value: dict[str, Any]) -> dict[str, Any]:
    body = {key: item for key, item in value.items() if key != "content_sha256"}
    result = dict(value)
    result["content_sha256"] = _sha256_bytes(_canonical_bytes(body))
    return result


def _write_bytes(path: Path, data: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd = os.open(path, os.O_WRONLY | os.O_CREAT | os.O_EXCL, 0o600)
    with os.fdopen(fd, "wb") as handle:
        handle.write(data)
        handle.flush()
        os.fsync(handle.fileno())


def _write_text(path: Path, text: str) -> None:
    data = text.encode("utf-8")
    if not data.endswith(b"\n"):
        data += b"\n"
    _write_bytes(path, data)


def _write_json(path: Path, value: Any) -> None:
    _write_text(path, json.dumps(value, ensure_ascii=False, indent=2))


def _review_id(authority_identity_id: str) -> str | None:
    text = str(authority_identity_id or "")
    if text.startswith(("ukpga:", "uksi:", "eur:", "eut:")):
        return text.replace(":", "-")
    return None


def _load_results() -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    with RESULTS.open(encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                rows.append(json.loads(line))
    return rows


def _recommend(
    *,
    source_class: str,
    review_row: Mapping[str, Any] | None,
    authority_identity_id: str,
    fail_closed: bool,
) -> tuple[str, str]:
    if fail_closed:
        return "HOLD", "Official bytes are missing. Fail-closed."
    if TULRCA_ID in {_review_id(authority_identity_id) or "", str((review_row or {}).get("id") or "")}:
        return "HOLD", "14 vs 28 August normalized body changed (TULRCA 1992)."
    if source_class == "judgment":
        return (
            "HOLD",
            "Judgment identity may be historically citable; proposition currentness stays HOLD pending later-treatment confirmation.",
        )
    if source_class == "official_secondary":
        return "HOLD", "Official-secondary candidate. Bind edition before APPROVE."
    if review_row and review_row.get("normalized_body_equal") is True:
        return (
            "APPROVE",
            "Mechanical 14 vs 28 August body identity after URI normalisation. Effects still need owner review; this is not gold.",
        )
    if review_row is None:
        return "HOLD", "No 14 vs 28 August pair is bound to this locator yet."
    return "HOLD", "Normalized body was not identical between 14 and 28 August."


def _shade(cell, fill: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): fill, qn("w:val"): "clear"})
    tc_pr.append(shading)


def _build_docx(locators: list[dict[str, Any]]) -> bytes:
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    title = document.add_paragraph()
    run = title.add_run("LegalBot GE per-locator gold draft")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x17, 0x36, 0x5D)
    intro = document.add_paragraph()
    intro.add_run(
        "This is still evaluation. Tick APPROVE, HOLD or REJECT for each locator. "
        "Nothing is gold, admitted, or full-current-law eligible until a later signed "
        "receipt records your ticks. AI is not the qualified England-and-Wales legal reviewer. "
        "Effects may be reviewed even when the official XML lists unapplied effects. "
        "One evaluation as-of date is 28 August 2026."
    )
    document.add_paragraph(
        "Cable & Wireless remains fail-closed after Find Case Law 404. "
        "Case 312 validity remains HOLD even if formality locators are later approved. "
        "Do not admit an unidentified Mediation Act 2025."
    )
    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = (
        "Locator",
        "Quote / note",
        "14 vs 28 Aug",
        "Recommend",
        "Owner",
        "Why still not gold",
    )
    for index, label in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = label
        _shade(cell, "17365D")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
                run.font.size = Pt(9)
    for row in locators:
        cells = table.add_row().cells
        cells[0].text = f"{row['title']}, {row['locator']}"
        quote = str(row.get("quote") or row.get("notes") or "")[:280]
        cells[1].text = quote
        cells[2].text = str(row.get("pit_14_vs_28") or "uncompared")
        cells[3].text = str(row.get("recommended_owner_action") or "HOLD")
        cells[4].text = "PENDING"
        cells[5].text = str(row.get("recommend_reason") or "")
        fill = "E8F5F4" if row.get("recommended_owner_action") == "APPROVE" else "FFF2CC"
        if row.get("fail_closed"):
            fill = "FDECEC"
        for cell in cells:
            _shade(cell, fill)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    from io import BytesIO

    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def main() -> int:
    if PACK.exists() or PACK.is_symlink():
        raise FileExistsError(f"create-only pack exists: {PACK}")
    results = _load_results()
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    sources = {str(row["source_version_id"]): row for row in manifest["sources"]}
    mechanical = json.loads(MECHANICAL.read_text(encoding="utf-8"))
    review_by_id = {str(row["id"]): row for row in mechanical.get("rows") or []}

    grouped: dict[tuple[str, str, str], dict[str, Any]] = {}
    for case in results:
        if str((case.get("factual_result") or {}).get("checks", {}).get("claim_evidence_support")) != "PASS":
            continue
        for evidence in case.get("evidence") or []:
            source_id = str(evidence.get("source_version_id") or "")
            locator = str(evidence.get("locator") or "")
            title = str(evidence.get("title") or "")
            key = (source_id, title, locator)
            bucket = grouped.setdefault(
                key,
                {
                    "source_version_id": source_id,
                    "title": title,
                    "locator": locator,
                    "chunk_id": evidence.get("chunk_id"),
                    "quote": evidence.get("quote"),
                    "quote_sha256": evidence.get("evidence_span_sha256"),
                    "case_ids": [],
                    "unapplied_effect_count": evidence.get("unapplied_effect_count"),
                    "provision_extent_status": evidence.get("provision_extent_status"),
                    "from_pass43": True,
                },
            )
            case_id = str(case.get("case_id") or "")
            if case_id not in bucket["case_ids"]:
                bucket["case_ids"].append(case_id)

    locators: list[dict[str, Any]] = []
    for item in grouped.values():
        source = sources.get(str(item["source_version_id"])) or {}
        authority = str(source.get("authority_identity_id") or "")
        review_id = _review_id(authority)
        review_row = review_by_id.get(review_id or "")
        source_class = "judgment" if authority.startswith("neutral-citation:") else "legislation"
        recommend, reason = _recommend(
            source_class=source_class,
            review_row=review_row,
            authority_identity_id=authority,
            fail_closed=False,
        )
        locators.append(
            {
                "source_version_id": item["source_version_id"],
                "title": item["title"],
                "locator": item["locator"],
                "chunk_id": item.get("chunk_id"),
                "quote": item.get("quote"),
                "quote_sha256": item.get("quote_sha256"),
                "case_ids": item["case_ids"],
                "authority_identity_id": authority,
                "source_class": source_class,
                "unapplied_effect_count": item.get("unapplied_effect_count")
                if item.get("unapplied_effect_count") is not None
                else source.get("unapplied_effect_count"),
                "provision_extent_status": "unverified",
                "pit_14_vs_28": (
                    "body_identical"
                    if review_row and review_row.get("normalized_body_equal") is True
                    else "body_changed"
                    if review_row and review_row.get("normalized_body_equal") is False
                    else "uncompared"
                ),
                "recommended_owner_action": recommend,
                "recommend_reason": reason,
                "owner_decision": "PENDING",
                "owner_signed": False,
                "effects_reviewed": False,
                "legal_gold": False,
                "admitted": False,
                "full_current_law_eligible": False,
                "qualified_legal_review": False,
                "from_pass43": True,
            }
        )

    seen = {(row["title"], row["locator"]) for row in locators}
    for extra in EXTRA_ROWS:
        key = (str(extra["title"]), str(extra["locator"]))
        if key in seen:
            continue
        seen.add(key)
        review_row = review_by_id.get(str(extra.get("staging_id") or ""))
        recommend, reason = _recommend(
            source_class=str(extra["source_class"]),
            review_row=review_row,
            authority_identity_id="",
            fail_closed=bool(extra.get("fail_closed")),
        )
        locators.append(
            {
                "source_version_id": "",
                "title": extra["title"],
                "locator": extra["locator"],
                "chunk_id": None,
                "quote": None,
                "quote_sha256": None,
                "case_ids": list(extra["case_ids"]),
                "authority_identity_id": extra.get("staging_id"),
                "source_class": extra["source_class"],
                "unapplied_effect_count": None,
                "provision_extent_status": "unverified",
                "pit_14_vs_28": "uncompared" if extra.get("fail_closed") else (
                    "body_identical"
                    if review_row and review_row.get("normalized_body_equal") is True
                    else "staged_not_in_85"
                ),
                "recommended_owner_action": recommend,
                "recommend_reason": extra["notes"] + " " + reason,
                "owner_decision": "PENDING",
                "owner_signed": False,
                "effects_reviewed": False,
                "legal_gold": False,
                "admitted": False,
                "full_current_law_eligible": False,
                "qualified_legal_review": False,
                "from_pass43": False,
                "fail_closed": bool(extra.get("fail_closed")),
                "point_in_time_as_at": extra.get("point_in_time_as_at"),
                "notes": extra["notes"],
            }
        )

    PACK.mkdir(parents=True, mode=0o700)
    os.chmod(PACK, stat.S_IRWXU)
    payload = _digest(
        {
            "schema": "legalbot.ge-per-locator-gold-draft.v1",
            "recorded_at_utc": datetime.now(UTC).isoformat(),
            "evaluation_state": True,
            "evaluation_as_of_date": EVALUATION_AS_OF,
            "owner_pack_signed": False,
            "predecessor_run_id": "LegalBot-GE-2026-09-02-visible-331-diagnostic-r1",
            "source_manifest_sha256": manifest["manifest_sha256"],
            "mechanical_review_content_sha256": mechanical.get("content_sha256"),
            "locator_count": len(locators),
            "pass43_locator_count": sum(1 for row in locators if row.get("from_pass43")),
            "recommended_approve_count": sum(
                1 for row in locators if row.get("recommended_owner_action") == "APPROVE"
            ),
            "locators": locators,
            "owner_decisions_still_required": [
                "APPROVE_HOLD_or_REJECT_each_locator",
                "cable_and_wireless_keep_fail_closed_or_supply_official_bytes",
                "icc_bind_staged_html_edition_or_hold",
                "case_312_validity_remains_hold",
                "whether_signed_pack_also_authorizes_diagnostic_331_r2",
            ],
            "legal_gold": False,
            "admitted": False,
            "full_current_law_eligible": False,
            "qualified_legal_review": False,
            "answer_weight_training": False,
            "sealed_unseen": False,
            "promotion": False,
            "live": False,
            "ai_may_act_as_legal_reviewer": False,
        }
    )
    _write_json(PACK / "LOCATOR-GOLD-DRAFT.json", payload)
    _write_text(
        PACK / "README.md",
        """# Per-locator gold draft r1

Unsigned owner draft. Every locator is PENDING. This is evaluation, not gold.

Tick APPROVE / HOLD / REJECT in the DOCX. A later create-only receipt pack
will record signed ticks. Until then the diagnostic overlay is a no-op.
""",
    )
    docx_bytes = _build_docx(locators)
    _write_bytes(PACK / "LegalBot-GE-2026-09-02-Per-Locator-Gold-Draft-r1.docx", docx_bytes)
    if not DOCX_PATH.parent.exists():
        DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    _write_bytes(DOCX_PATH, docx_bytes)
    print(
        json.dumps(
            {
                "pack": str(PACK),
                "content_sha256": payload["content_sha256"],
                "locator_count": payload["locator_count"],
                "recommended_approve_count": payload["recommended_approve_count"],
                "docx_sha256": _sha256_bytes(docx_bytes),
            },
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
