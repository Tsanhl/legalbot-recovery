#!/usr/bin/env python3
"""Render the r47 consolidated Phase-2A owner gate as a polished DOCX."""

from __future__ import annotations

import argparse
import json
import os
import sys
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE_ROOT = (
    PROJECT_ROOT
    / "data/evaluations/phase2a-owner-review/LegalBot-Phase2AB-2026-08-24-r47-consolidated-owner-gate"
)
DEFAULT_OUTPUT = (
    PROJECT_ROOT
    / "data/evaluations/phase2a-owner-review/LegalBot-Phase2AB-2026-08-24-r48-owner-review-delivery"
    / "LegalBot-v111-Phase2A-Consolidated-Owner-Review-Agnes-2026-08-24-rev2.docx"
)
EXPECTED_DECISION_DIGEST = "7a471bed936bf901cca49413f1abb8e27db54157862a1f369136a0704e811414"
EXPECTED_MACHINE_DIGEST = "3ba8de75875cd2192a0707450c206fbb91220fbf3d3ac2704b1fd18046d1227c"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 101, 116)
GREEN = RGBColor(26, 109, 74)
AMBER = RGBColor(122, 90, 0)
RED = RGBColor(155, 28, 28)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_AMBER = "FFF8E1"
LIGHT_RED = "FDECEC"
WHITE = RGBColor(255, 255, 255)
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def _load(path: Path) -> dict[str, Any]:
    value = json.loads(path.read_bytes())
    if not isinstance(value, dict):
        raise ValueError("phase2a_docx_input_not_object")
    return value


def _set_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def _field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, display, end))
    _set_font(run, size=8.5, color=MUTED)


def _set_cell_text(cell, text: str, *, header: bool = False, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align or WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(str(text))
    _set_font(
        run,
        size=9.2 if header else 9.5,
        color=INK,
        bold=header,
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _add_table(doc, headers, rows, widths, apply_table_geometry):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, header=True)
        _shade(table.rows[0].cells[index], LIGHT_GRAY)
    for values in rows:
        row = table.add_row()
        _cant_split(row)
        for index, value in enumerate(values):
            align = (
                WD_ALIGN_PARAGRAPH.CENTER
                if index == 0 and len(headers) > 2
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            _set_cell_text(row.cells[index], str(value), align=align)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
        cell_margins_dxa=CELL_MARGINS_DXA,
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def _add_callout(doc, label: str, body: str, *, fill: str, color: RGBColor, apply_table_geometry):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.08
    lead = paragraph.add_run(f"{label}: ")
    _set_font(lead, size=10.5, color=color, bold=True)
    run = paragraph.add_run(body)
    _set_font(run, size=10.5, color=INK)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    apply_table_geometry(
        table,
        [CONTENT_WIDTH_DXA],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=180,
        cell_margins_dxa={"top": 130, "bottom": 130, "start": 180, "end": 180},
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _add_bullet(doc, text: str, *, level: int = 0) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    _set_font(run, size=11, color=INK)


def _add_number(doc, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    _set_font(run, size=11, color=INK)


def _add_body(doc, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        _set_font(first, size=11, color=INK, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        _set_font(rest, size=11, color=INK)
    else:
        run = paragraph.add_run(text)
        _set_font(run, size=11, color=INK)


def _add_heading(doc, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True


def _set_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (12, DARK_BLUE, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _set_page(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def _set_header_footer(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
    left = paragraph.add_run("LEGALBOT v1.11  |  PHASE 2A OWNER REVIEW")
    _set_font(left, size=8.5, color=MUTED, bold=True)
    right = paragraph.add_run("\tINTERNAL")
    _set_font(right, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run("Page ")
    _set_font(run, size=8.5, color=MUTED)
    _field(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    _set_font(run, size=8.5, color=MUTED)
    _field(paragraph, "NUMPAGES")


def render(source_root: Path, output_path: Path, doc_skill_scripts: Path) -> dict[str, Any]:
    if output_path.exists() or output_path.is_symlink():
        raise ValueError("phase2a_docx_output_already_exists")
    output_path.parent.mkdir(parents=True, mode=0o700, exist_ok=True)
    if not doc_skill_scripts.is_dir():
        raise ValueError("phase2a_docx_skill_scripts_missing")
    sys.path.insert(0, str(doc_skill_scripts))
    from table_geometry import apply_table_geometry

    overview = _load(source_root / "PHASE2A-CONSOLIDATED-OWNER-GATE.json")
    index = _load(source_root / "MACHINE-PACKAGE-INDEX.json")
    decisions = _load(source_root / "OWNER-DECISION-BATCH-1058.json")
    mismatch = _load(source_root / "COMPLETE-LEGISLATION-BYTE-MISMATCH-REGISTER-65.json")
    if (
        overview.get("owner_decision_batch_content_sha256") != EXPECTED_DECISION_DIGEST
        or index.get("machine_package_content_sha256") != EXPECTED_MACHINE_DIGEST
        or decisions.get("owner_decision_batch_content_sha256") != EXPECTED_DECISION_DIGEST
    ):
        raise ValueError("phase2a_docx_source_digest_invalid")
    changed = [
        row
        for row in mismatch["records"]
        if row["byte_mismatch_record"]["comparison"]["classification"]
        != "SEMANTIC_PROVISION_TEXT_IDENTICAL_BYTE_MISMATCH_ONLY"
    ]
    if len(changed) != 1:
        raise ValueError("phase2a_docx_changed_mismatch_count_invalid")
    patents = changed[0]

    document = Document()
    _set_page(document)
    _set_styles(document)
    _set_header_footer(document.sections[0])

    core = document.core_properties
    core.title = "LegalBot v1.11 Phase 2A Consolidated Owner Review"
    core.subject = "Non-authorizing owner review gate"
    core.author = ""
    core.last_modified_by = ""
    core.keywords = ""
    core.comments = ""
    core.created = datetime(2000, 1, 1, tzinfo=UTC)
    core.modified = datetime(2000, 1, 1, tzinfo=UTC)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("PHASE 2A CONSOLIDATED OWNER REVIEW")
    _set_font(run, name="Arial", size=23, color=INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("LegalBot v1.11 - evidence closure and owner decision gate")
    _set_font(run, name="Arial", size=13.5, color=MUTED)

    for label, value in (
        ("Owner", "Agnes"),
        ("Decision date", "24 August 2026"),
        ("Route", "Owner-adopted internal research tool; not professional legal certification"),
        ("Current phase", "Phase 2A only"),
        ("Status", "Stopped at explicit owner-decision gate"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        lead = paragraph.add_run(f"{label}: ")
        _set_font(lead, size=10.5, color=INK, bold=True)
        body = paragraph.add_run(value)
        _set_font(body, size=10.5, color=INK)

    _add_callout(
        document,
        "Gate verdict",
        "Phase 2A is not yet complete. Phase 2B and Development 30 remain unauthorized. "
        "No candidate was mutated, no new source was admitted, and no answer model was invoked.",
        fill=LIGHT_RED,
        color=RED,
        apply_table_geometry=apply_table_geometry,
    )

    _add_heading(document, "Decision snapshot", 1)
    _add_table(
        document,
        ("Area", "Recorded / verified", "Still needs owner decision"),
        (
            ("585 issue rows", "137", "448"),
            ("1,896 legislative effects", "1,380", "516"),
            ("20 judgments", "20 historical snapshots verified", "20 later-treatment outcomes"),
            (
                "65 legislation byte mismatches",
                "64 text-identical",
                "1 text delta + owner disposition",
            ),
        ),
        (2520, 3240, 3600),
        apply_table_geometry,
    )

    _add_heading(document, "What is already complete", 1)
    _add_bullet(
        document,
        "All 585 issue rows are present in one sealed matrix; none is silently dropped or averaged away.",
    )
    _add_bullet(
        document,
        "The independent advisory pass ranked all 448 pending rows using pinned Qwen/Qwen3-Reranker-0.6B at revision e61197ed45024b0ed8a2d74b80b4d909f1255473.",
    )
    _add_bullet(
        document,
        "All 516 held legislative effects were remapped against same-authority evidence; none had an exact affected-provision intersection, so all remain owner decisions.",
    )
    _add_bullet(
        document,
        "All 20 historical judgment snapshots passed byte-integrity checks; nine targeted official later-treatment leads were quarantined for review.",
    )
    _add_bullet(
        document,
        "The 65 legislation byte mismatches were compared at stable provision anchors: 64 are semantic-text identical and one is a real section 60(7) delta.",
    )

    document.add_page_break()
    _add_heading(document, "Immediate owner action", 1)
    _add_body(
        document,
        "A deterministic subset of 580 recommendations is ready for one digest-bound owner approval: all 516 metadata/currentness-only legislative-effect dispositions and all 64 semantic-text-identical byte-mismatch dispositions.",
    )
    _add_heading(document, "Exact approval wording", 2)
    approval = (
        "I, Agnes, approve the exact 516 legislative-effect metadata/currentness-only "
        "recommendations and the exact 64 semantic-text-identical byte-mismatch "
        "recommendations bound to Phase-2A owner-decision batch digest "
        f"{EXPECTED_DECISION_DIGEST}. This approval is for continued Phase 2A only. "
        "It does not approve the 448 issue proposition/span selections, the 20 judgment "
        "later-treatment decisions, the Patents Act 1977 section 60(7) text delta, or "
        "the 9 new source admissions. It does not authorize Phase 2B or Development 30."
    )
    _add_callout(
        document,
        "Reply exactly",
        approval,
        fill=LIGHT_AMBER,
        color=AMBER,
        apply_table_geometry=apply_table_geometry,
    )
    _add_body(document, "Owner-decision batch digest:", bold_lead="Owner-decision batch digest:")
    digest_p = document.add_paragraph()
    digest_p.paragraph_format.space_after = Pt(8)
    digest_run = digest_p.add_run(EXPECTED_DECISION_DIGEST)
    _set_font(digest_run, name="Courier New", size=8.5, color=INK, bold=True)

    _add_heading(document, "Why a blanket approval of all 1,058 items is not enough", 1)
    _add_bullet(
        document,
        "The 448 unresolved issue records still identify topics and ranked spans, but do not yet contain owner-confirmed atomic proposition text for every row.",
    )
    _add_bullet(
        document,
        "A relevance score ranks evidence; it is not a legal-currentness or release threshold and cannot substitute for owner judgment.",
    )
    _add_bullet(
        document,
        "Twenty judgment records still require an explicit affirmed, limited, distinguished, displaced, no-material-treatment, exclusion, or more-evidence outcome.",
    )
    _add_bullet(
        document,
        "The nine new official cases remain quarantine evidence. They cannot be admitted, indexed, or embedded until their proposition-level role is expressly approved.",
    )

    _add_heading(document, "Outstanding decision map", 1)
    _add_table(
        document,
        ("Count", "Decision class", "Current safe state"),
        (
            (
                "448",
                "Issue proposition + exact span",
                "Evidence packet ready; owner selection required",
            ),
            (
                "20",
                "Judgment later treatment",
                "Historical bytes verified; substantive relationship unresolved",
            ),
            ("1", "Patents Act text delta", "Fresh bytes quarantined; materiality unresolved"),
            ("9", "New source admissions", "Quarantined; not indexed or embedded"),
        ),
        (900, 3000, 5460),
        apply_table_geometry,
    )

    document.add_page_break()
    _add_heading(document, "Evidence findings", 1)
    _add_heading(document, "585 issue rows", 2)
    _add_body(
        document,
        "The matrix records 107 prior owner decisions in the 509 gold/case group and 30 in the 76 candidate-impact group. The remaining split is 402 gold/case rows and 46 candidate-impact rows. Technical qualification has not been rerun, so recorded owner scope decisions are not mislabeled as completed qualification.",
    )
    _add_body(
        document,
        "The independent reviewer is honestly recorded as a separate, model-independent reranking pass. It is non-generative, does not use the drafting adapter, applies no release threshold, and cannot decide or authorize any row.",
    )

    _add_heading(document, "1,896 legislative effects", 2)
    _add_body(
        document,
        "The previous owner package records 1,380 effect dispositions. The remaining 516 all have same-authority retrieval candidates, including 472 found through cross-subject recovery, but zero exact affected-provision intersections. The advisory recommendation is metadata/currentness-only pending final proposition-binding confirmation.",
    )

    _add_heading(document, "20 judgments and three unavailable pages", 2)
    _add_body(
        document,
        "All 20 sealed historical judgment representations passed integrity checks. Seventeen fresh official identities were downloaded. Three Twinsectra Parliament pages returned HTTP 403; their historical official bytes and provenance remain verified, but present-law currentness is not inferred from those bytes.",
    )
    _add_body(
        document,
        "Nine targeted UKSC/JCPC leads were verified to contain both the later case citation and the target citation. The searches are expressly non-exhaustive, and absence of another lead proves nothing.",
    )

    _add_heading(document, "65 legislation byte mismatches", 2)
    _add_body(
        document,
        "After normalizing point-in-time and representation selectors in provision anchors, 64 records have identical canonical provision text. Their byte differences are representation-only and are ready for the safe subset approval above.",
    )
    patent_record = patents["byte_mismatch_record"]
    patent_comparison = patent_record["comparison"]
    _add_callout(
        document,
        "One genuine text delta",
        f"{patent_record['title']} {patent_comparison['changed'][0]['locator']} at "
        f"{patent_comparison['changed'][0]['source_anchor']} changed between the sealed "
        "and fresh official versions. Seventeen pending issue packets reference the Patents "
        "Act source somewhere, but none uses the exact changed locator as its selected locator. "
        "That reduces immediate concern but does not decide materiality.",
        fill=LIGHT_AMBER,
        color=AMBER,
        apply_table_geometry=apply_table_geometry,
    )

    document.add_page_break()
    _add_heading(document, "Answer-safety and system hardening", 1)
    _add_table(
        document,
        ("Layer", "Implemented fail-closed control"),
        (
            (
                "Prompt engineering",
                "Pinned prompt/model/configuration identities; atomic claims; explicit evidence IDs; no silent context truncation.",
            ),
            (
                "Quality",
                "Material dates, amounts, percentages, durations and provision IDs must bind to exact evidence spans; unsupported facts and unrelated citations block release.",
            ),
            (
                "Output validation",
                "Strict JSON/schema checks, clause atomicity, quotation/source/version/jurisdiction/currentness checks, and sealed retrieval metadata.",
            ),
            (
                "Fallback",
                "No evidence, below-threshold evidence, index/reranker failure or unavailable index prevents model invocation; repeated fingerprints trigger debug before a third attempt.",
            ),
        ),
        (2100, 7260),
        apply_table_geometry,
    )

    _add_heading(document, "Local support-agent architecture choices", 2)
    _add_bullet(
        document,
        "Conversation memory uses encrypted durable SQLite plus a bounded in-memory hot cache, rather than adding a Redis service to an owner-only local deployment. Retention is 30 days, the hot-cache window is seven days, and sliding-window quotas are enforced.",
    )
    _add_bullet(
        document,
        "Knowledge freshness is event driven: official material enters quarantine, then a durable update event can drive chunking, embedding and ingestion only after owner admission. Source date and last_updated are retained, and retrieval collapses to the latest approved version.",
    )
    _add_bullet(
        document,
        "WebSocket job events now require an exact subprotocol. A gRPC streaming contract records time-to-first-token and gap-free sentence diagnostics, but the real Unix-domain-socket transport remains a Phase-2B provisioning step.",
    )
    _add_bullet(
        document,
        "The code keeps a single local deployable system while defining service boundaries. A full microservice split, Redis deployment and network gRPC activation are deferred until actual load or operational evidence justifies them.",
    )

    _add_heading(document, "Controls designed but not activated", 2)
    _add_bullet(
        document,
        "Pinned Ed25519 verification, three private roots, session/CSRF secrets, literal 127.0.0.1 access, UDS model transport, 12 GiB memory ceiling and 3 GiB free-memory admission remain Phase-2B work.",
    )
    _add_bullet(
        document,
        "No split secret, Development/Validation roots, 30/30 allocation, Development authorization payload, Stage A, answer generation, promotion, Validation or live activation exists yet.",
    )

    document.add_page_break()
    _add_heading(document, "What happens next", 1)
    _add_number(
        document,
        "Owner approves or rejects the exact 580-item deterministic subset using the digest-bound wording in this report.",
    )
    _add_number(
        document,
        "Phase 2A continues with explicit proposition/span decisions for 448 issues, later-treatment decisions for 20 judgments, the Patents Act section 60(7) disposition and source-admission decisions for nine leads.",
    )
    _add_number(
        document,
        "Only after all material dependencies close may one consolidated successor candidate be justified, built, sealed and retrieval re-attested. The existing sealed candidate is never patched.",
    )
    _add_number(
        document,
        "Codex returns the successful Phase-2A package and exact adoption digest. Owner adoption of that exact digest is required before Phase 2B begins.",
    )
    _add_number(
        document,
        "After Phase 2B passes, Codex returns the exact Development authorization payload. Development 30 remains unauthorized until the owner approves that payload.",
    )

    _add_heading(document, "Machine-readable companion", 1)
    _add_body(
        document,
        "The companion package contains full JSON registers plus CSV review sheets with blank owner_outcome and owner_comments columns. The DOCX intentionally summarizes rather than reproducing thousands of rows.",
    )
    _add_table(
        document,
        ("Artifact", "Purpose"),
        (
            (
                "COMPLETE-REMEDIATION-MATRIX-585.json",
                "Every issue, recorded decision, advisory ranking and deep-recovery comparison",
            ),
            (
                "COMPLETE-LEGISLATIVE-EFFECT-REGISTER-1896.json",
                "All recorded and pending effect dispositions",
            ),
            (
                "COMPLETE-JUDGMENT-LATER-TREATMENT-REGISTER-20.json",
                "Judgment custody and targeted leads",
            ),
            ("OWNER-DECISION-BATCH-1058.json", "Exact unresolved decision payload"),
            ("OWNER-REVIEW-*.csv", "Owner-editable outcomes and comments"),
            ("SHA256SUMS.txt", "File integrity verification"),
        ),
        (3960, 5400),
        apply_table_geometry,
    )

    _add_heading(document, "Exact package identities", 1)
    for label, digest in (
        ("Owner-decision batch", EXPECTED_DECISION_DIGEST),
        ("Machine package", EXPECTED_MACHINE_DIGEST),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        lead = paragraph.add_run(f"{label}: ")
        _set_font(lead, size=9.5, color=INK, bold=True)
        run = paragraph.add_run(digest)
        _set_font(run, name="Courier New", size=8, color=INK)

    _add_callout(
        document,
        "Legal scope",
        "This is an owner-adopted internal research-tool review package. It is not "
        "professional England-and-Wales legal certification and is not legal advice.",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
        apply_table_geometry=apply_table_geometry,
    )

    for section in document.sections:
        _set_page(document)
        _set_header_footer(section)
    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    document.save(output_path)
    os.chmod(output_path, 0o600)
    return {"output": str(output_path), "bytes": output_path.stat().st_size}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-root", type=Path, default=DEFAULT_SOURCE_ROOT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--doc-skill-scripts", type=Path, required=True)
    args = parser.parse_args()
    result = render(
        args.source_root.resolve(strict=True),
        args.output.resolve(),
        args.doc_skill_scripts.resolve(strict=True),
    )
    print(json.dumps(result, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
