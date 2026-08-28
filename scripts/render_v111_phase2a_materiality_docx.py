#!/usr/bin/env python3
"""Render the 54-row Phase-2A materiality batch as an owner-review DOCX."""

from __future__ import annotations

import argparse
import json
from collections import OrderedDict
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "FFF4CE"
RISK = "FDECEC"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120


def _load_batch(path: Path) -> dict[str, Any]:
    value = json.loads(path.read_bytes())
    if value.get("schema") != "legalbot.v111.phase2a.owner-materiality-batch-54.v1":
        raise ValueError("phase2a_materiality_docx_batch_schema_invalid")
    if value.get("row_count") != 54 or value.get("owner_decisions_applied") is not False:
        raise ValueError("phase2a_materiality_docx_batch_boundary_invalid")
    return value


def _set_run(
    run: Any,
    *,
    size: float = 11,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str = "000000",
    font: str = "Calibri",
) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("end", CELL_START_END_DXA),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError("phase2a_materiality_docx_table_width_invalid")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", TABLE_WIDTH_DXA), ("tblInd", TABLE_INDENT_DXA)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _shade_cell(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _shade_paragraph(paragraph: Any, *, fill: str, border_color: str = "D8DEE8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "bottom", "left", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), border_color)
        borders.append(node)
    p_pr.append(borders)


def _add_field(paragraph: Any, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, end))


def _display_code(value: Any) -> str:
    return str(value or "").replace("_", " ").strip().capitalize()


def _add_label_value(
    doc: Document,
    label: str,
    value: str,
    *,
    after: float = 2,
    keep_with_next: bool = False,
) -> Any:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    label_run = paragraph.add_run(f"{label}: ")
    _set_run(label_run, bold=True, color=INK)
    value_run = paragraph.add_run(value)
    _set_run(value_run)
    return paragraph


def _add_callout(doc: Document, label: str, text: str, *, fill: str = CALLOUT) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.08)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    _shade_paragraph(paragraph, fill=fill)
    label_run = paragraph.add_run(f"{label}  ")
    _set_run(label_run, bold=True, color=INK)
    text_run = paragraph.add_run(text)
    _set_run(text_run, color=INK)


def _add_checkbox_line(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(f"☐ {text}")
    _set_run(run, size=10.5)


def _add_comment_lines(doc: Document, count: int = 2) -> None:
    label = doc.add_paragraph()
    label.paragraph_format.space_before = Pt(4)
    label.paragraph_format.space_after = Pt(2)
    label.paragraph_format.keep_with_next = True
    _set_run(label.add_run("Owner comments:"), bold=True, color=INK)
    for index in range(count):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = index < count - 1
        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "B8C1CC")
        borders.append(bottom)
        p_pr.append(borders)
        paragraph.add_run(" ")


def _component_span(component: dict[str, Any]) -> tuple[str, str, str, str]:
    if "proposed_bound_span_text" in component:
        return (
            str(component["proposed_bound_span_text"]),
            str(component["proposed_bound_span_sha256"]),
            str(component.get("anchor_id") or ""),
            str(component.get("official_url") or ""),
        )
    return (
        str(component["exact_normalized_span_text"]),
        str(component["exact_normalized_span_sha256"]),
        str(component.get("anchor_id") or ""),
        str(component.get("official_url") or ""),
    )


def _collect_spans(batch: dict[str, Any]) -> OrderedDict[str, dict[str, str]]:
    spans: OrderedDict[str, dict[str, str]] = OrderedDict()
    for row in batch["rows"]:
        for component in row.get("component_evidence", []):
            text, digest, anchor, url = _component_span(component)
            spans.setdefault(
                digest,
                {
                    "text": text,
                    "anchor": anchor,
                    "url": url,
                    "source": str(row["official_source_title"]),
                },
            )
        for component in row.get("currentness_evidence", []):
            text, digest, anchor, url = _component_span(component)
            spans.setdefault(
                digest,
                {
                    "text": text,
                    "anchor": anchor,
                    "url": url,
                    "source": str(row["official_source_title"]),
                },
            )
    return spans


def _configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(23)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(13)
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _set_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(0)
        _set_run(
            header.add_run("LEGALBOT v1.11  |  PHASE 2A OWNER REVIEW"),
            size=8.5,
            bold=True,
            color=MUTED,
        )
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.space_before = Pt(0)
        run = footer.add_run("Confidential owner-review draft  |  Page ")
        _set_run(run, size=8.5, color=MUTED)
        _add_field(footer, "PAGE")


def _add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        _shade_cell(cells[0], LIGHT_GRAY)
        p0 = cells[0].paragraphs[0]
        p1 = cells[1].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        _set_run(p0.add_run(label), size=9.5, bold=True, color=INK)
        _set_run(p1.add_run(value), size=9.5)
    _set_table_geometry(table, [2200, 7160])
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _add_source_table(doc: Document, authorities: list[str]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, text in enumerate(("#", "Proposed new authority")):
        cell = table.rows[0].cells[index]
        _shade_cell(cell, LIGHT_GRAY)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _set_run(p.add_run(text), size=9.5, bold=True, color=INK)
    for ordinal, authority in enumerate(authorities, start=1):
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        p0 = cells[0].paragraphs[0]
        p1 = cells[1].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        _set_run(p0.add_run(str(ordinal)), size=9.5)
        _set_run(p1.add_run(authority), size=9.5, font="Courier New")
    _set_table_geometry(table, [720, 8640])


def _row_candidate_summary(row: dict[str, Any]) -> str:
    if "candidate_coverage" in row:
        return str(row["candidate_coverage"]["assessment"])
    assessments = sorted(
        {
            str(component["candidate_coverage"]["assessment"])
            for component in row.get("component_evidence", [])
        }
    )
    return "; ".join(assessments)


def build_docx(batch_path: Path, output_path: Path) -> None:
    batch = _load_batch(batch_path)
    spans = _collect_spans(batch)
    doc = Document()
    _configure_styles(doc)
    _set_header_footer(doc)
    doc.core_properties.title = "LegalBot v1.11 Phase 2A Owner Materiality Review"
    doc.core_properties.subject = "Exact 54-row official-source decision batch"
    doc.core_properties.author = "LegalBot Phase 2A remediation"
    doc.core_properties.comments = "Owner-review draft; no gate authorization"

    doc.add_paragraph("PHASE 2A OWNER DECISION MEMO", style="Title")
    doc.add_paragraph(
        "LegalBot v1.11 — 54-row corrected materiality and source-admission batch",
        style="Subtitle",
    )
    _add_metadata_table(
        doc,
        [
            ("Owner", "Agnes (typed-name confirmation required)"),
            ("Decision date", "2026-08-24"),
            ("Status", "Owner decisions required — none applied"),
            ("Rows", "54 corrected/repaired issue rows"),
            ("New authorities", str(batch["proposed_source_admission_authority_count"])),
            ("Exact batch digest", batch["artifact_content_sha256"]),
        ],
    )
    _add_callout(
        doc,
        "Gate status",
        "This package authorizes nothing by itself. Phase 2B and Development 30 remain "
        "unauthorized. Approval applies only when the owner confirms the exact digest above.",
        fill=CAUTION,
    )

    doc.add_heading("Decision requested", level=1)
    paragraph = doc.add_paragraph()
    _set_run(
        paragraph.add_run(
            "Review the 54 row recommendations and the 16 proposed source authorities. "
            "Approve, hold, or reject each row. A package-wide approval is permitted only "
            "after reviewing this document and must quote the exact batch digest."
        )
    )
    _add_checkbox_line(doc, "Approve all 54 recommendations exactly as sealed.")
    _add_checkbox_line(doc, "Hold one or more rows; identify each row in owner comments.")
    _add_checkbox_line(doc, "Reject one or more rows; identify each row and reason.")

    doc.add_heading("What approval permits", level=2)
    for text in (
        "Apply the approved proposition bindings to versioned successor gold/case artifacts.",
        "Admit only the proposition-level official authorities listed in this batch.",
        "Continue Phase 2A remediation and, if needed, build one consolidated successor candidate.",
    ):
        _add_checkbox_line(doc, text)
    doc.add_heading("What approval does not permit", level=2)
    for text in (
        "Phase 2B provisioning or split freeze.",
        "Development 30, Validation, promotion, ACTIVE/PREVIOUS writes, or live activation.",
        "Automatic indexing, embedding, or any model-generated unsupported proposition.",
    ):
        _add_checkbox_line(doc, text)

    doc.add_page_break()
    doc.add_heading("Proposed new source authorities", level=1)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(4)
    _set_run(
        note.add_run(
            "These authorities are absent from the sealed predecessor. Approval is "
            "proposition-level and does not allow unrelated material from the same source."
        ),
        size=9.5,
        italic=True,
        color=MUTED,
    )
    _add_source_table(doc, batch["proposed_source_admission_authorities"])

    doc.add_page_break()
    doc.add_heading("54-row owner decision register", level=1)
    intro = doc.add_paragraph()
    _set_run(
        intro.add_run(
            "Each span preview is an excerpt for navigation. The complete, untruncated, "
            "hash-bound span appears once in Appendix A."
        ),
        italic=True,
        color=MUTED,
    )

    for ordinal, row in enumerate(batch["rows"], start=1):
        doc.add_heading(
            f"{ordinal}. {row['row_id']} — {row['official_source_title']}",
            level=2,
        )
        _add_label_value(
            doc,
            "Locator",
            str(row.get("stated_official_legal_locator") or "See bound anchors"),
            keep_with_next=True,
        )
        _add_label_value(
            doc,
            "Recommendation",
            _display_code(row["advisory_recommendation"]),
            keep_with_next=True,
        )
        _add_label_value(
            doc,
            "Candidate",
            _display_code(_row_candidate_summary(row)),
            after=4,
            keep_with_next=True,
        )
        for index, component in enumerate(row.get("component_evidence", []), start=1):
            text, digest, anchor, _ = _component_span(component)
            action = str(
                component.get("component_action")
                or row.get("proposed_action")
                or "BIND_EXACT_OFFICIAL_SPAN"
            )
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            _set_run(
                p.add_run(f"Component {index}: {_display_code(action)}"),
                size=10,
                bold=True,
                color=DARK_BLUE,
            )
            preview = text if len(text) <= 420 else text[:417].rstrip() + "…"
            span_p = doc.add_paragraph()
            span_p.paragraph_format.left_indent = Inches(0.08)
            span_p.paragraph_format.right_indent = Inches(0.08)
            span_p.paragraph_format.space_after = Pt(3)
            span_p.paragraph_format.keep_together = True
            _shade_paragraph(span_p, fill=CALLOUT)
            _set_run(span_p.add_run(preview), size=9.25, color=INK)
            _add_label_value(doc, "Anchor", anchor, after=1)
            _add_label_value(doc, "Full-span SHA-256", digest, after=3)
        _add_checkbox_line(doc, "APPROVE this row recommendation")
        _add_checkbox_line(doc, "HOLD this row for further evidence")
        _add_checkbox_line(doc, "REJECT this row recommendation")
        if row["owner_source_admission_required"]:
            _add_checkbox_line(doc, "APPROVE proposition-level source admission for this row")
        _add_comment_lines(doc, count=1)

    doc.add_page_break()
    doc.add_heading("Appendix A — complete exact bound spans", level=1)
    appendix_intro = doc.add_paragraph()
    _set_run(
        appendix_intro.add_run(
            "Every span below is complete and untruncated. Repeated spans are listed once "
            "and referenced by SHA-256 in the decision register."
        )
    )
    for ordinal, (digest, span) in enumerate(spans.items(), start=1):
        doc.add_heading(f"A{ordinal}. {span['source']} — {span['anchor']}", level=2)
        _add_label_value(doc, "Span SHA-256", digest, after=2, keep_with_next=True)
        if span["url"]:
            _add_label_value(
                doc,
                "Official URL",
                span["url"],
                after=3,
                keep_with_next=True,
            )
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Keep ordinary evidence spans intact, but allow unusually long official
        # pages to flow immediately after their heading instead of stranding the
        # heading on an otherwise empty page.
        paragraph.paragraph_format.keep_together = len(span["text"]) <= 4_000
        _shade_paragraph(paragraph, fill=CALLOUT)
        _set_run(
            paragraph.add_run(span["text"]),
            size=9.25,
            color=INK,
            font="Calibri",
        )

    doc.add_page_break()
    doc.add_heading("Appendix B — exact package-wide approval wording", level=1)
    approval = (
        "OWNER DECISION — APPROVE EXACT PHASE-2A MATERIALITY BATCH ONLY\n\n"
        "I, Agnes, approve all 54 row-level materiality and binding recommendations "
        f"in artifact digest {batch['artifact_content_sha256']}.\n\n"
        "I expressly approve proposition-level admission of only the new official "
        "authorities listed in that artifact.\n\n"
        "This approval authorizes continued Phase 2A remediation only. It does not "
        "authorize Phase 2B, Development 30, Validation, promotion, or live activation.\n\n"
        "Owner typed name: Agnes\n"
        "Owner decision date: 2026-08-24\n\n"
        "I APPROVE THIS EXACT DIGEST."
    )
    approval_p = doc.add_paragraph()
    approval_p.paragraph_format.space_before = Pt(4)
    approval_p.paragraph_format.space_after = Pt(10)
    _shade_paragraph(approval_p, fill=CAUTION, border_color="D7B84A")
    _set_run(approval_p.add_run(approval), size=10, color=INK)
    _add_comment_lines(doc, count=2)

    final = doc.add_paragraph()
    final.paragraph_format.space_before = Pt(10)
    final.paragraph_format.space_after = Pt(0)
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(
        final.add_run(
            "PHASE 2A MATERIALITY DECISION GATE — PHASE 2B AND DEVELOPMENT 30 NOT AUTHORIZED"
        ),
        size=10,
        bold=True,
        color="9B1C1C",
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--batch", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    if args.output.exists() or args.output.is_symlink():
        raise ValueError("phase2a_materiality_docx_output_already_exists")
    build_docx(args.batch.resolve(strict=True), args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
