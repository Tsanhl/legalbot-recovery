#!/usr/bin/env python3
"""Create-only r1-versus-r2 visible 331 comparison. One owner-facing DOCX."""

from __future__ import annotations

import hashlib
import json
import os
import stat
from collections import Counter
from collections.abc import Mapping
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
R1 = (
    ROOT
    / "data/evaluations/general-enquiries"
    / "LegalBot-GE-2026-09-02-visible-331-diagnostic-r1"
)
R2 = (
    ROOT
    / "data/evaluations/general-enquiries"
    / "LegalBot-GE-2026-09-02-visible-331-diagnostic-r2"
)
PACK = (
    ROOT
    / "data/evaluations/general-enquiries"
    / "LegalBot-GE-2026-09-02-visible-331-diagnostic-r2-comparison"
)
DOCX = ROOT / "output/docx/LegalBot-GE-2026-09-02-visible-331-diagnostic-r1-vs-r2.docx"
FOCUS = (
    "administrative-law:cp-d08",
    "international-commercial-mediation:cp-d01",
    "wills-and-estates:cp-d02",
)
NAVY = RGBColor(0x17, 0x36, 0x5D)
INK = RGBColor(0x1F, 0x29, 0x37)


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _canonical_bytes(value: Any) -> bytes:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()


def _digest(value: dict[str, Any]) -> dict[str, Any]:
    body = {key: item for key, item in value.items() if key != "content_sha256"}
    result = dict(value)
    result["content_sha256"] = hashlib.sha256(_canonical_bytes(body)).hexdigest()
    return result


def _write_text(path: Path, text: str) -> None:
    data = text.encode("utf-8")
    if not data.endswith(b"\n"):
        data += b"\n"
    fd = os.open(path, os.O_WRONLY | os.O_CREAT | os.O_EXCL, 0o600)
    with os.fdopen(fd, "wb") as handle:
        handle.write(data)
        handle.flush()
        os.fsync(handle.fileno())


def _write_json(path: Path, value: Any) -> None:
    _write_text(path, json.dumps(value, ensure_ascii=False, indent=2))


def _load_results(path: Path) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    with path.open(encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if line:
                rows.append(json.loads(line))
    return rows


def _locators(row: dict[str, Any]) -> list[str]:
    evidence = row.get("evidence") or []
    names: list[str] = []
    if isinstance(evidence, list):
        for item in evidence:
            if isinstance(item, dict):
                title = str(item.get("title") or "")
                locator = str(item.get("locator") or "")
                names.append(f"{title} {locator}".strip())
    return names


def _claim(row: Mapping[str, Any]) -> str:
    checks = ((row.get("factual_result") or {}).get("checks") or {})
    return str(checks.get("claim_evidence_support") or "")


def _counts(rows: list[dict[str, Any]]) -> dict[str, Any]:
    factual = Counter(str((row.get("factual_result") or {}).get("outcome") or "") for row in rows)
    quality = Counter(str((row.get("quality_70_plus") or {}).get("outcome") or "") for row in rows)
    return {
        "total": len(rows),
        "factual": dict(sorted(factual.items())),
        "quality": dict(sorted(quality.items())),
        "evidence_present": sum(bool(row.get("evidence")) for row in rows),
        "claim_support_pass": sum(_claim(row) == "PASS" for row in rows),
        "factual_pass": factual.get("FACTUAL_PASS", 0),
        "factual_hold": factual.get("FACTUAL_HOLD", 0),
    }


def _set_run(paragraph, size: int = 11, *, bold: bool = False, color: RGBColor | None = None) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color or INK
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "Calibri")


def _heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    _set_run(paragraph, 16 if level == 1 else 13, bold=True, color=NAVY)


def build_docx(comparison: dict[str, Any]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    title = document.add_paragraph("Visible 331 diagnostic: r1 versus r2")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run(title, 22, bold=True, color=NAVY)
    intro = document.add_paragraph(
        "This is the single owner-facing evaluation report after the resolved "
        "67-locator package. It is locator-level evaluation gold and a diagnostic "
        "rerun only. It is not qualified England-and-Wales legal review, answer gold, "
        "runtime admission, weight training, sealed unseen, promotion or live."
    )
    _set_run(intro, 11)
    _heading(document, "Progress rule")
    progress = comparison.get("progress") or {}
    body = document.add_paragraph(
        f"overall_progress = {progress.get('overall_progress')}; "
        f"overall_state = {progress.get('overall_state')}. "
        "One held or fail-closed case does not stop the rest of Phase 2."
    )
    _set_run(body, 11)
    _heading(document, "Headline counts")
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Measure"
    hdr[1].text = "r1"
    hdr[2].text = "r2"
    r1c = comparison["counts"]["r1"]
    r2c = comparison["counts"]["r2"]
    rows = [
        ("Cases", r1c["total"], r2c["total"]),
        ("FACTUAL_PASS", r1c.get("factual_pass", 0), r2c.get("factual_pass", 0)),
        ("FACTUAL_HOLD", r1c.get("factual_hold", 0), r2c.get("factual_hold", 0)),
        ("Evidence present", r1c["evidence_present"], r2c["evidence_present"]),
        ("Claim-support PASS", r1c["claim_support_pass"], r2c["claim_support_pass"]),
    ]
    for label, left, right in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(left)
        cells[2].text = str(right)
    _heading(document, "Cases 008, 174 and 312")
    for key, row in (comparison.get("focus_cases") or {}).items():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{key} ({row.get('id')})")
        run.bold = True
        _set_run(paragraph, 11, bold=True, color=NAVY)
        detail = document.add_paragraph(
            f"r1 locators: {', '.join(row.get('r1_locators') or ['(none)'])}. "
            f"r2 locators: {', '.join(row.get('r2_locators') or ['(none)'])}. "
            f"r1 claim-support {row.get('r1_claim_support')}; "
            f"r2 claim-support {row.get('r2_claim_support')}; "
            f"r2 factual {row.get('r2_factual')}."
        )
        _set_run(detail, 11)
    _heading(document, "Boundaries that remain false")
    false_flags = document.add_paragraph(
        "qualified_legal_review, answer_legal_gold, legal_gold, admitted, "
        "full_current_law_eligible, answer_weight_training, sealed_unseen_execution, "
        "promotion and live remain false."
    )
    _set_run(false_flags, 11)
    if DOCX.exists():
        raise FileExistsError(f"create-only DOCX exists: {DOCX}")
    document.save(DOCX)
    os.chmod(DOCX, 0o600)


def main() -> int:
    if PACK.exists() or PACK.is_symlink():
        raise FileExistsError(f"create-only comparison pack exists: {PACK}")
    r1_rows = _load_results(R1 / "visible/RESULTS.jsonl")
    r2_rows = _load_results(R2 / "visible/RESULTS.jsonl")
    if len(r1_rows) != 331 or len(r2_rows) != 331:
        raise RuntimeError("visible result counts are not 331")
    r1_by_id = {str(row.get("case_id")): row for row in r1_rows}
    r2_by_id = {str(row.get("case_id")): row for row in r2_rows}
    focus: dict[str, Any] = {}
    labels = {
        "administrative-law:cp-d08": "case_008",
        "international-commercial-mediation:cp-d01": "case_174",
        "wills-and-estates:cp-d02": "case_312",
    }
    for case_id, label in labels.items():
        left = r1_by_id[case_id]
        right = r2_by_id[case_id]
        focus[label] = {
            "id": case_id,
            "r1_locators": _locators(left),
            "r2_locators": _locators(right),
            "r1_claim_support": _claim(left),
            "r2_claim_support": _claim(right),
            "r1_factual": (left.get("factual_result") or {}).get("outcome"),
            "r2_factual": (right.get("factual_result") or {}).get("outcome"),
        }
    r2_manifest = json.loads((R2 / "RUN-MANIFEST.json").read_text(encoding="utf-8"))
    progress_path = R2 / "PROGRESS-AND-BLOCKER-LEDGER.json"
    progress = json.loads(progress_path.read_text(encoding="utf-8")) if progress_path.is_file() else {}
    comparison = _digest(
        {
            "schema": "legalbot.ge-visible-331-diagnostic-comparison.v2",
            "recorded_at_utc": datetime.now(UTC).isoformat(),
            "predecessor_run_id": R1.name,
            "rerun_id": R2.name,
            "rerun_manifest_content_sha256": r2_manifest.get("content_sha256"),
            "rerun_visible_results_sha256": _sha256_file(R2 / "visible/RESULTS.jsonl"),
            "evaluation_state": True,
            "counts": {"r1": _counts(r1_rows), "r2": _counts(r2_rows)},
            "focus_cases": focus,
            "progress": {
                "overall_progress": progress.get("overall_progress"),
                "overall_state": progress.get("overall_state"),
                "held_or_fail_closed_cases": progress.get("held_or_fail_closed_cases"),
            },
            "locator_resolution": {
                "APPROVE": 66,
                "HOLD": 0,
                "REJECT": 1,
                "PENDING": 0,
            },
            "unseen": {"probe": "omit", "fresh_unseen": False, "sealed_306": False},
            "non_authorizing": {
                "qualified_legal_review": False,
                "legal_gold": False,
                "answer_legal_gold": False,
                "admitted": False,
                "full_current_law_eligible": False,
                "answer_weight_training": False,
                "sealed_unseen": False,
                "promotion": False,
                "live": False,
            },
        }
    )
    PACK.mkdir(parents=True, mode=0o700)
    os.chmod(PACK, stat.S_IRWXU)
    _write_json(PACK / "COMPARISON.json", comparison)
    _write_text(
        PACK / "README.md",
        "# Visible 331 diagnostic r1 versus r2\n\n"
        "Create-only comparison after the owner-adopted locator-evaluation-gold "
        "resolution. This remains Phase 2 evaluation.\n",
    )
    build_docx(comparison)
    print(
        json.dumps(
            {
                "comparison": comparison["content_sha256"],
                "docx": _sha256_file(DOCX),
            }
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
