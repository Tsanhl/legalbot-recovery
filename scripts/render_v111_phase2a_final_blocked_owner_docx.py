#!/usr/bin/env python3
"""Render the final blocked Phase-2A owner-review brief as a create-only DOCX."""

from __future__ import annotations

import argparse
import json
import os
import sys
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
RUN_NAME = "LegalBot-Phase2A-2026-08-27-final-owner-review-blocked"
DEFAULT_SOURCE_ROOT = (
    PROJECT_ROOT / "data/evaluations/phase2a-owner-review" / RUN_NAME / "machine"
)
DEFAULT_OUTPUT = (
    PROJECT_ROOT
    / "data/evaluations/phase2a-owner-review"
    / RUN_NAME
    / "LegalBot-v111-Phase2A-Final-Blocked-Owner-Review-2026-08-27-rev4.docx"
)
EXPECTED_MACHINE_DIGEST = (
    "24490cd8ae21fa9eb2f0217096a7d8556f1910514d6b9c7e5889c243d21d91d8"
)
EXPECTED_QUALIFICATION_SHA = (
    "4170aa192181c7b9a368af01cf4f813eb6b3417c0c57c58bb7b4f03257727df8"
)
EXPECTED_RETRIEVAL_SHA = (
    "e8345b3e5c2dd8be164dc7e7dfaed90712fa59c01dbbc1c499fbe1a4e5997224"
)
EXPECTED_VERDICT = (
    "PHASE 2A SAFELY STOPPED - PHASE 2B AND DEVELOPMENT 30 NOT AUTHORIZED"
)

INK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 101, 116)
DARK_BLUE = RGBColor(31, 77, 120)
RED = RGBColor(155, 28, 28)
AMBER = RGBColor(122, 90, 0)
GREEN = RGBColor(26, 109, 74)
WHITE = RGBColor(255, 255, 255)
LIGHT_RED = "FDECEC"
LIGHT_AMBER = "FFF8E1"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E8F5EE"
LIGHT_GRAY = "F2F4F7"
CONTENT_WIDTH_DXA = 9360


def _load(path: Path) -> dict[str, Any]:
    value = json.loads(path.read_bytes())
    if not isinstance(value, dict):
        raise ValueError(f"owner_review_input_not_object:{path.name}")
    return value


def _set_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def _mark_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    properties.append(marker)


def _field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, separate, display, end))
    _set_font(run, size=8.5, color=MUTED)


def _header_footer(document: Document) -> None:
    def populate_header(header) -> None:
        paragraph = header.paragraphs[0]
        paragraph.text = ""
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Pt(468), alignment=WD_TAB_ALIGNMENT.RIGHT
        )
        left = paragraph.add_run("LEGALBOT v1.11  |  PHASE 2A FINAL REVIEW")
        _set_font(left, size=8.5, color=MUTED, bold=True)
        right = paragraph.add_run("\tINTERNAL / NON-AUTHORIZING")
        _set_font(right, size=8.5, color=MUTED, bold=True)

    def populate_footer(footer) -> None:
        paragraph = footer.paragraphs[0]
        paragraph.text = ""
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("Page ")
        _set_font(run, size=8.5, color=MUTED)
        _field(paragraph, "PAGE")
        run = paragraph.add_run(" of ")
        _set_font(run, size=8.5, color=MUTED)
        _field(paragraph, "NUMPAGES")

    document.settings.odd_and_even_pages_header_footer = True
    for section in document.sections:
        section.different_first_page_header_footer = True
        for header in (
            section.header,
            section.even_page_header,
            section.first_page_header,
        ):
            populate_header(header)
        for footer in (
            section.footer,
            section.even_page_footer,
            section.first_page_footer,
        ):
            populate_footer(footer)


def _cell(cell, text: str, *, header: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.02
    run = paragraph.add_run(str(text))
    _set_font(run, size=9.1 if header else 9.4, color=INK, bold=header)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        _shade(cell, fill)


def _table(document, headers, rows, widths, apply_table_geometry):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _mark_header(table.rows[0])
    for index, header in enumerate(headers):
        _cell(table.rows[0].cells[index], header, header=True, fill=LIGHT_GRAY)
    for values in rows:
        row = table.add_row()
        properties = row._tr.get_or_add_trPr()
        properties.append(OxmlElement("w:cantSplit"))
        for index, value in enumerate(values):
            _cell(row.cells[index], value)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    document.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def _callout(document, label: str, body: str, *, fill: str, color, apply_table_geometry):
    table = document.add_table(rows=1, cols=1)
    _mark_header(table.rows[0])
    cell = table.cell(0, 0)
    _shade(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    lead = paragraph.add_run(f"{label}: ")
    _set_font(lead, size=10.3, color=color, bold=True)
    run = paragraph.add_run(body)
    _set_font(run, size=10.3, color=INK)
    apply_table_geometry(
        table,
        [CONTENT_WIDTH_DXA],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=180,
        cell_margins_dxa={"top": 130, "bottom": 130, "start": 180, "end": 180},
    )
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def _heading(document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True


def _body(document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.1
    _set_font(paragraph.add_run(text), size=10.7, color=INK)


def _bullet(document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    _set_font(paragraph.add_run(text), size=10.5, color=INK)


def _number(document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.08
    _set_font(paragraph.add_run(text), size=10.5, color=INK)


def _masthead(document, apply_table_geometry) -> None:
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    _mark_header(table.rows[0])
    _cell(table.cell(0, 0), "FINAL OWNER REVIEW", header=True, fill="1F4D78")
    _cell(table.cell(0, 1), "27 AUGUST 2026", header=True, fill="1F4D78")
    for cell in table.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = WHITE
    _cell(table.cell(1, 0), "LegalBot v1.11 · Phase 2A")
    _cell(table.cell(1, 1), "Owner-adopted internal research-tool route")
    apply_table_geometry(
        table,
        [4680, 4680],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=0,
        cell_margins_dxa={"top": 110, "bottom": 110, "start": 150, "end": 150},
    )


def _configure_document(document: Document) -> None:
    for section in document.sections:
        section.page_width = Pt(612)
        section.page_height = Pt(792)
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        section.header_distance = Pt(35.4)
        section.footer_distance = Pt(35.4)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.7)
    normal.font.color.rgb = INK
    for level, tokens in {1: (16, 14, 7), 2: (12.5, 10, 5)}.items():
        size, before, after = tokens
        style = document.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = DARK_BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    _header_footer(document)


def _validate_inputs(source_root: Path) -> tuple[dict[str, Any], ...]:
    package = _load(source_root / "MACHINE-PACKAGE-INDEX.json")
    verdict = _load(source_root / "FINAL-PHASE2A-VERDICT.json")
    qualification = _load(
        source_root / "qualification/DETERMINISTIC-ALL585-QUALIFICATION.json"
    )
    retrieval = _load(source_root / "retrieval/HELD-RETRIEVAL-REATTESTATION.json")
    debug = _load(source_root / "DEBUG-AND-ANTI-LOOP-REGISTER.json")
    code = _load(source_root / "CODE-IDENTITY-AND-WORKTREE-STATUS.json")
    if (
        package.get("machine_package_content_sha256") != EXPECTED_MACHINE_DIGEST
        or package.get("terminal_verdict") != EXPECTED_VERDICT
        or package.get("phase2a_technical_qualification_passed") is not False
        or package.get("files", {})
        .get("qualification/DETERMINISTIC-ALL585-QUALIFICATION.json", {})
        .get("sha256")
        != EXPECTED_QUALIFICATION_SHA
        or package.get("files", {})
        .get("retrieval/HELD-RETRIEVAL-REATTESTATION.json", {})
        .get("sha256")
        != EXPECTED_RETRIEVAL_SHA
        or verdict.get("terminal_verdict") != EXPECTED_VERDICT
        or qualification.get("issue_count") != 585
        or retrieval.get("retrieval_quality_passed") is not True
        or debug.get("same_failure_third_attempts") != 0
    ):
        raise ValueError("phase2a_final_docx_bound_input_invalid")
    return package, verdict, qualification, retrieval, debug, code


def render(source_root: Path, output_path: Path, doc_skill_scripts: Path) -> dict[str, Any]:
    if output_path.exists() or output_path.is_symlink():
        raise ValueError("phase2a_final_docx_output_already_exists")
    if not doc_skill_scripts.is_dir():
        raise ValueError("phase2a_docx_skill_scripts_missing")
    package, verdict, qualification, retrieval, debug, code = _validate_inputs(source_root)
    sys.path.insert(0, str(doc_skill_scripts))
    from table_geometry import apply_table_geometry

    document = Document()
    _configure_document(document)
    core = document.core_properties
    core.title = "LegalBot v1.11 Final Phase 2A Owner Review"
    core.subject = "Safely stopped, non-authorizing Phase-2A evidence report"
    core.author = ""
    core.last_modified_by = ""
    core.keywords = ""
    core.comments = ""
    core.created = datetime(2000, 1, 1, tzinfo=UTC)
    core.modified = datetime(2000, 1, 1, tzinfo=UTC)

    _masthead(document, apply_table_geometry)
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(3)
    title.paragraph_format.keep_with_next = True
    _set_font(
        title.add_run("PHASE 2A FINAL OWNER REVIEW"),
        name="Arial",
        size=22,
        color=INK,
        bold=True,
    )
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(13)
    _set_font(
        subtitle.add_run("Successor retrieval passed; legal-evidence qualification remains blocked"),
        name="Arial",
        size=12.5,
        color=MUTED,
    )
    _callout(
        document,
        "Final verdict",
        "Phase 2A safely stopped. Phase 2B and Development 30 are not authorized. "
        "The successor candidate remains non-ACTIVE and answer-release ineligible.",
        fill=LIGHT_RED,
        color=RED,
        apply_table_geometry=apply_table_geometry,
    )

    _heading(document, "Executive result")
    _body(
        document,
        "The single frozen successor candidate was built, sealed and retrieval "
        "re-attested successfully. The final deterministic review of all 585 issues "
        "then found unresolved legal-evidence work. Retrieval quality cannot substitute "
        "for proposition-level support, owner legal judgment, source currentness or later "
        "treatment review.",
    )
    _table(
        document,
        ("Control", "Result", "Release consequence"),
        (
            ("Successor candidate", "Built and sealed", "Held; non-ACTIVE"),
            ("24-query retrieval", "PASS", "Quality gate passed"),
            ("All-585 qualification", "FAIL-CLOSED", "Phase 2A not complete"),
            ("Common legal cutoff", "None supportable", "Cannot freeze currentness"),
            ("Phase 2B", "LOCKED", "No provisioning or split"),
            ("Development 30", "LOCKED", "No answer run"),
        ),
        (2600, 2200, 4560),
        apply_table_geometry,
    )

    _heading(document, "All-585 outcome")
    statuses = qualification["status_counts"]
    _table(
        document,
        ("Issue status", "Count", "Meaning"),
        (
            (
                "Technically evidence-ready for owner adoption",
                statuses["TECHNICALLY_EVIDENCE_READY_FOR_OWNER_ADOPTION"],
                "Evidence-ready only; not owner-adopted",
            ),
            (
                "Owner decision required",
                statuses["OWNER_DECISION_REQUIRED"],
                "Substantive selection remains unresolved",
            ),
            (
                "Blocked material gap",
                statuses["BLOCKED_MATERIAL_GAP"],
                "Exact-span evidence is not yet sufficient",
            ),
        ),
        (4400, 900, 4060),
        apply_table_geometry,
    )
    _callout(
        document,
        "Important",
        "The blocker counts below overlap. Do not add 98, 263, 186 and 135 as though "
        "they were separate issue totals.",
        fill=LIGHT_AMBER,
        color=AMBER,
        apply_table_geometry=apply_table_geometry,
    )

    document.add_page_break()
    _heading(document, "What passed")
    candidate = verdict["candidate_identity"]
    metrics = retrieval["metrics"]
    _bullet(
        document,
        "The one permitted source scan completed 3,848 of 3,848 files and was not "
        "repeated. Its manifest digest is "
        "fb6e0d82ff205e74052aa0f536049702e84c8af86624305f5fa03e19eb6e820d.",
    )
    _bullet(
        document,
        "The frozen successor scope contains 251 approved source versions and 222,200 "
        "chunks. No unseen or held source was admitted.",
    )
    _bullet(
        document,
        "The successor contains 222,200 finite vectors at 1,024 dimensions; the sealed "
        "predecessor was not patched and no ACTIVE/PREVIOUS pointer was written.",
    )
    _table(
        document,
        ("Retrieval measure", "Result", "Frozen gate"),
        (
            ("Binding", f"{metrics['binding_count']}/24", "24/24"),
            ("Recall@5", f"{metrics['positive_recall_at_5']:.2f}", "1.00"),
            ("Recall@10", f"{metrics['positive_recall_at_10']:.2f}", ">= 0.95"),
            ("MRR", f"{metrics['mrr']:.4f}", ">= 0.80"),
            ("Teaching contamination", metrics["teaching_assessment_hits"], "0"),
            ("Private-path contamination", metrics["private_path_hits"], "0"),
            ("Wrong-version contamination", metrics["wrong_version_count"], "0"),
        ),
        (4200, 2100, 3060),
        apply_table_geometry,
    )
    _callout(
        document,
        "Advisory diagnostic",
        "Mean exact-span recall@5 was 0.6293 across 18 diagnostic cases. It was not a "
        "frozen retrieval release gate, but it reinforces the need to resolve the 98 "
        "material exact-span gaps.",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
        apply_table_geometry=apply_table_geometry,
    )

    _heading(document, "What remains blocked")
    holds = verdict["source_holds"]
    _table(
        document,
        ("Required work", "Count", "Current state"),
        (
            ("Issue exact-span remediation", 98, "Blocked material gap"),
            ("Substantive owner decisions", 263, "Explicit owner selection required"),
            (
                "Source currentness verification",
                holds["currentness_unverified_source_count"],
                "Held; 65 sources verified",
            ),
            (
                "Judgment later-treatment review",
                holds["later_treatment_required_source_count"],
                "Held; 0 verified",
            ),
        ),
        (4300, 900, 4160),
        apply_table_geometry,
    )
    _body(
        document,
        "Until those material dependencies are resolved, there is no supportable common "
        "legal-currentness cutoff and no successful Phase-2A adoption payload to approve.",
    )

    document.add_page_break()
    _heading(document, "Anti-loop and debugging record")
    incidents = {item["stage"]: item for item in debug["incidents"]}
    _table(
        document,
        ("Path", "Attempt record", "Debug / disposition"),
        (
            (
                "Prior planner flow",
                "38 rows exceeded cap",
                "Sealed and excluded; deterministic route used",
            ),
            (
                "Successor build",
                "Attempt 1 failed; attempt 2 passed",
                incidents["successor_build"]["debug"],
            ),
            (
                "Post-build audit",
                "No rebuild",
                "Metadata-only counter reconciliation; candidate bytes unchanged",
            ),
            (
                "Retrieval re-attestation",
                "Attempt 1 stopped; attempt 2 passed",
                incidents["held_retrieval_reattestation"]["debug"],
            ),
        ),
        (2600, 2600, 4160),
        apply_table_geometry,
    )
    _callout(
        document,
        "Loop control",
        "There were zero unchanged third attempts. No planner, advisory reviewer or "
        "answer-generation model was used in the deterministic route.",
        fill=LIGHT_GREEN,
        color=GREEN,
        apply_table_geometry=apply_table_geometry,
    )
    _heading(document, "Candidate identity")
    _table(
        document,
        ("Field", "Recorded identity"),
        (
            ("Build", candidate["build_id"]),
            ("Corpus", candidate["corpus_id"]),
            ("Stage / status", f"{candidate['stage']} / {candidate['status']}"),
            ("Lance tree SHA-256", candidate["lance_tree_sha256"]),
            ("Build seal file SHA-256", candidate["build_seal_sha256"]),
            ("Source-manifest SHA-256", candidate["source_manifest_sha256"]),
        ),
        (2800, 6560),
        apply_table_geometry,
    )
    _heading(document, "Code identity limitation")
    _body(
        document,
        f"Recorded Git HEAD: {code['head_sha256']}. The branch was "
        f"{code['branch']} with {code['worktree_status_entry_count']} status entries. "
        "The evidence therefore does not claim a clean-HEAD candidate build, and "
        "exact-HEAD verification is recorded as failed.",
    )

    document.add_page_break()
    _heading(document, "Required path to Phase 2B")
    _number(
        document,
        "Prepare deterministic exact-span evidence for the 98 blocked issue rows. New "
        "official sources, if any, must remain quarantined until exact digest-bound owner "
        "admission.",
    )
    _number(
        document,
        "Present the 263 substantive owner selections in clear digest-bound review "
        "batches and record explicit decisions.",
    )
    _number(
        document,
        "Complete source-currentness review for 186 held sources and later-treatment "
        "review for 135 sources; preserve exclusions and holds.",
    )
    _number(
        document,
        "Rerun deterministic all-585 qualification after the evidence changes. The "
        "retrieval benchmark must remain passed against the exact final candidate.",
    )
    _number(
        document,
        "Only if every material dependency passes may Codex present a successful exact "
        "Phase-2A digest for owner adoption. That explicit adoption—not a blanket advance "
        "authorization—can unlock Phase 2B.",
    )
    _callout(
        document,
        "No approval requested now",
        "This package can only be acknowledged as received. It cannot be adopted as a "
        "successful Phase-2A package, and no owner wording can bypass the unresolved "
        "evidence and currentness gates.",
        fill=LIGHT_RED,
        color=RED,
        apply_table_geometry=apply_table_geometry,
    )

    _heading(document, "Owner-review companion files")
    _table(
        document,
        ("Artifact", "Use"),
        (
            ("UNRESOLVED-MATERIAL-GAPS-98.csv", "Exact-span remediation queue"),
            ("OWNER-DECISIONS-REQUIRED-263.csv", "Substantive owner decision queue"),
            ("SOURCE-CURRENTNESS-AND-LATER-TREATMENT-HOLDS-251.csv", "Source hold review"),
            ("CASE-QUALIFICATION-SUMMARY-60.csv", "Case-level fail-closed summary"),
            ("DETERMINISTIC-ALL585-QUALIFICATION.json", "Authoritative issue outcomes"),
            ("DEBUG-AND-ANTI-LOOP-REGISTER.json", "Attempts, fixes and stop evidence"),
        ),
        (4580, 4780),
        apply_table_geometry,
    )
    _body(
        document,
        f"The machine package indexes {package['file_count']} evidence files, plus its "
        "package index and checksum list. It includes no source blobs, vectors, private "
        "keys, secrets, private review roots, split secret, Development payload, "
        "Validation material or live state.",
    )

    document.add_page_break()
    _heading(document, "Exact evidence digests")
    _table(
        document,
        ("Evidence", "SHA-256"),
        (
            ("Machine package", EXPECTED_MACHINE_DIGEST),
            ("All-585 qualification", EXPECTED_QUALIFICATION_SHA),
            ("Retrieval re-attestation", EXPECTED_RETRIEVAL_SHA),
            ("Final verdict record", verdict["verdict_content_sha256"]),
        ),
        (3000, 6360),
        apply_table_geometry,
    )
    _heading(document, "Package status", 2)
    _table(
        document,
        ("Item", "Recorded state"),
        (
            ("Machine evidence", "Sealed and checksummed"),
            ("Retrieval", "Passed against the held successor"),
            ("Phase 2A", "Safely stopped; material blockers remain"),
            ("Phase 2B / Development 30", "Not authorized"),
            ("Owner signature", "Not requested for this blocked package"),
        ),
        (3800, 5560),
        apply_table_geometry,
    )
    _callout(
        document,
        "Optional acknowledgement",
        "I acknowledge receipt of the exact blocked Phase-2A owner-review package. "
        "This acknowledgement does not adopt Phase 2A, authorize Phase 2B, or authorize "
        "Development 30.",
        fill=LIGHT_AMBER,
        color=AMBER,
        apply_table_geometry=apply_table_geometry,
    )
    _callout(
        document,
        "Legal scope",
        "This is an owner-adopted internal research-tool evidence report. It is not "
        "professional England-and-Wales legal certification and is not legal advice.",
        fill=LIGHT_BLUE,
        color=DARK_BLUE,
        apply_table_geometry=apply_table_geometry,
    )

    settings = document.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_path.parent.mkdir(parents=True, mode=0o700, exist_ok=True)
    document.save(output_path)
    os.chmod(output_path, 0o600)
    return {"output": str(output_path), "bytes": output_path.stat().st_size}


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-root", type=Path, default=DEFAULT_SOURCE_ROOT)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--doc-skill-scripts", type=Path, required=True)
    args = parser.parse_args()
    result = render(
        args.source_root.resolve(strict=True),
        args.output.resolve(),
        args.doc_skill_scripts.resolve(strict=True),
    )
    print(json.dumps(result, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
